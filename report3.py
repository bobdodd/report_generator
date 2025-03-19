from pymongo import MongoClient
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement, ns
from docx.oxml import parse_xml
from docx.oxml.ns import qn
import os
import json
from bson import ObjectId
import click
from docx.enum.style import WD_STYLE_TYPE

#############################################
# Accessible names Database Access
#############################################

class AccessibilityDB:
    def __init__(self):
        try:
            self.client = MongoClient('mongodb://localhost:27017/',
                                    serverSelectionTimeoutMS=5000)
            self.client.server_info()
            self.db = self.client['accessibility_tests']
            
            # Separate collections for test runs and page results
            self.test_runs = self.db['test_runs']
            self.page_results = self.db['page_results']
            
            # Create indexes
            self.page_results.create_index([('url', 1), ('test_run_id', 1)])
            self.page_results.create_index('timestamp')
            self.test_runs.create_index('timestamp')
        except Exception as e:
            print(f"Failed to connect to MongoDB: {e}")
            raise

    def get_latest_test_run(self):
        """Get the most recent test run"""
        return self.test_runs.find_one(
            sort=[('timestamp_start', -1)]
        )
    
    def get_all_test_runs(self):
        """Get all test runs"""
        return list(self.test_runs.find(
           sort=[('timestamp_start', -1)]
        ))

    def get_page_results(self, test_run_id):
        """Get all page results for a specific test run"""
        try:
            return list(self.page_results.find(
                {'test_run_id': str(test_run_id)},
                {'_id': 0}
            ))
        except Exception as e:
            print(f"Error getting page results: {e}")
            return []

    def __del__(self):
        if hasattr(self, 'client'):
            self.client.close()

####################################################
# Helper Functions
####################################################

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)



"""
    affected_domains = set()
    for page in pages_without_lang:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        affected_domains.add(domain)

    # Get total number of domains in the test run for comparison
    all_urls = db_connection.page_results.distinct('url', {'test_run_id': str(latest_test_run['_id'])})
    total_domains = set()
    for url in all_urls:
        domain = url.replace('http://', '').replace('https://', '').split('/')[0]
        total_domains.add(domain)

    percentage = (len(affected_domains) / len(total_domains)) * 100 if total_domains else 0
"""

#################################################
# Word Styling / TOC functions
#################################################

def add_table_of_contents(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    # Add hyperlink and web options to the TOC field code
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u \\w'  # \h adds hyperlinks, \w preserves tab spacing

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def format_toc_styles(doc):
    """Format TOC styles to properly handle numbered headings"""
    for i in range(1, 4):  # TOC levels 1-3
        style_name = f'TOC {i}'
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.paragraph_format.tab_stops.clear_all()
            
            # Calculate indentation based on level
            base_indent = Inches(0.25)
            level_indent = Inches(0.25 * (i-1))
            
            # Set paragraph indentation
            style.paragraph_format.left_indent = level_indent
            style.paragraph_format.first_line_indent = -base_indent
            
            # Add tab stops:
            # First tab for after the number
            style.paragraph_format.tab_stops.add_tab_stop(level_indent, WD_TAB_ALIGNMENT.LEFT)
            # Final tab for page number
            style.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            
            # Set font
            style.font.name = 'Arial'
            style.font.size = Pt(14)

def set_document_styles(doc):
    """Set up document styles with specified font sizes and spacing"""
    # Default paragraph text
    style = doc.styles['Normal']
    style.font.size = Pt(14)
    style.font.name = 'Arial'
    
    # Heading Styles
    title_style = doc.styles['Title']  # Level 0
    title_style.font.size = Pt(26)
    title_style.font.name = 'Arial'
    title_style.font.bold = True
    title_style.paragraph_format.space_after = Pt(14)  # One line height
    
    h1_style = doc.styles['Heading 1']
    h1_style.font.size = Pt(22)
    h1_style.font.name = 'Arial'
    h1_style.font.bold = True
    h1_style.paragraph_format.space_after = Pt(14)  # One line height
    
    h2_style = doc.styles['Heading 2']
    h2_style.font.size = Pt(18)
    h2_style.font.name = 'Arial'
    h2_style.font.bold = True
    h2_style.paragraph_format.space_after = Pt(14)  # One line height
    
    h3_style = doc.styles['Heading 3']
    h3_style.font.size = Pt(16)
    h3_style.font.name = 'Arial'
    h3_style.font.bold = True
    h3_style.paragraph_format.space_after = Pt(14)  # One line height
    
    # List styles
    list_style = doc.styles['List Bullet']
    list_style.font.size = Pt(14)
    list_style.font.name = 'Arial'
    
    # Caption style
    caption_style = doc.styles['Caption']
    caption_style.font.size = Pt(14)
    caption_style.font.name = 'Arial'
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # TOC styles are automatically created when the TOC is inserted

def format_table_text(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(14)
                    run.font.name = 'Arial'

##########################################################
# Fn to Create the Report template
##########################################################

def create_report_template(db_connection, title, author, date):
    print("Starting report creation...")

    ####################################################
    # Get list of URLs and domains used by the reporting
    # Gets used everywhere
    ####################################################

    # Get the latest test run data
    #latest_test_run = db_connection.get_latest_test_run()
    #all_urls = db_connection.page_results.distinct('url', {'test_run_id': str(latest_test_run['_id'])})

    all_test_runs = db_connection.get_all_test_runs()
    test_run_ids = [str(run['_id']) for run in all_test_runs]
    all_urls = db_connection.page_results.distinct('url', {'test_run_id': {'$in': test_run_ids}})    


    total_domains = set()
    for url in all_urls:
        domain = url.replace('http://', '').replace('https://', '').split('/')[0]
        total_domains.add(domain)

    doc = Document()
    
    # Set up document styles
    set_document_styles(doc)

    #################################### 
    # Add sections for header and footer
    ####################################
    section = doc.sections[0]
    
    # Set different first page for header/footer
    section.different_first_page_header_footer = True
    
    # Header (will only appear from second page onwards)
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.text = title
    for run in header_para.runs:
        run.font.size = Pt(14)

    # Footer - Create a table for better alignment control
    footer = section.footer
    footer_table = footer.add_table(1, 3, width=Inches(6.5))
    footer_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set equal widths for all columns
    for cell in footer_table.columns:
        cell.width = Inches(2.17)
    
    # Set table properties
    footer_table.style = 'Table Grid'
    footer_table.autofit = False
    
    # Remove all borders and set font size
    for cell in footer_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(14)
                run.font.name = 'Arial'
        cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>'))
    
    # Footer content
    space_paragraph = footer_table.rows[0].cells[0].paragraphs[0]
    space_paragraph.add_run(" ")
    
    page_num_paragraph = footer_table.rows[0].cells[1].paragraphs[0]
    page_num_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(page_num_paragraph)
    
    logo_paragraph = footer_table.rows[0].cells[2].paragraphs[0]
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists('logo.png'):
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture('logo.png', width=Inches(1))
    
    # Set table cell vertical alignment to center
    for cell in footer_table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:vAlign')
        tcVAlign.set(qn('w:val'), "center")
        tcPr.append(tcVAlign)

    ##########################################
    # Title Page
    ##########################################

    title_heading = doc.add_heading(title, level=0)
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add image - calculate width based on page width
    doc.add_paragraph()  # Add space before image
    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Get page width (accounting for margins)
    section = doc.sections[0]
    page_width = section.page_width - section.left_margin - section.right_margin
    
    if os.path.exists('access labs.jpg'):
        image_run = image_paragraph.add_run()
        image_run.add_picture('access labs.jpg', width=page_width)
    else:
        print("Warning: 'access labs.jpg' not found in the current directory")
    
    doc.add_paragraph()  # Add space after image
    
    # Add subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("CNIB Access Labs")
    subtitle_run.font.size = Pt(24)  # Large text
    subtitle_run.font.bold = True

    doc.add_paragraph()  # Add space after image
        
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_para.style.font.size = Pt(14)
    author_para.add_run(f"Author: {author}")

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"Date: {date}")
    
    # Get the latest test run data
    latest_test_run = db_connection.get_latest_test_run()



    if latest_test_run:
        doc.add_paragraph()
        test_run_para = doc.add_paragraph()
        test_run_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        test_run_para.add_run(f"Test Run Date: {latest_test_run['timestamp_start']}")
        status_para = doc.add_paragraph()
        status_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        status_para.add_run(f"Status: {latest_test_run['status']}")

    ################################
    # Table of Contents
    ################################

    doc.add_page_break()
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_heading.style = doc.styles['Heading 1']
    add_table_of_contents(doc)
    format_toc_styles(doc)  

    #################################
    # Executive Summary
    #################################

    doc.add_page_break()
    h1 = doc.add_heading('Executive Summary', level=1)
    h1.style = doc.styles['Heading 1']
    
    h2 = doc.add_heading('Disclaimer', level=2)
    h2.style = doc.styles['Heading 2']
    
    # Get the page and domain counts directly from the database
    all_test_runs = db_connection.get_all_test_runs()
    test_run_ids = [str(run['_id']) for run in all_test_runs]
    all_urls = db_connection.page_results.distinct('url', {'test_run_id': {'$in': test_run_ids}})

    if all_test_runs:
        test_run_ids = [str(run['_id']) for run in all_test_runs]
        unique_urls = db_connection.page_results.distinct('url', {'test_run_id': {'$in': test_run_ids}})
        page_count = len(unique_urls)


    #if latest_test_run:
    #    unique_urls = db_connection.page_results.distinct('url', {'test_run_id': str(latest_test_run['_id'])})
    #    page_count = len(unique_urls)
        
        # Extract unique domains
        domains = set()
        for url in unique_urls:
            domain = url.replace('http://', '').replace('https://', '').split('/')[0]
            domains.add(domain)
        domain_count = len(domains)    
    overview_disclaimer = doc.add_paragraph()
    overview_disclaimer.add_run(f"""
Disclaimer: This accessibility review represents a digital accessiblilty inspection of {page_count} pages from {domain_count} websites, looking at page properties that help indicate the accessibility health of the page and site. It is not designed to be a comprehensive report on every accessibility issue on each of the pages inspected, merely indicitave of potential issues. A formal manual inspection as part of a digital accessibility audit that includes an element of lived experience user testing is required to make any claim on conformance to accessibility standards.
""".strip())

    h2 = doc.add_heading('Subject areas', level=2)
    h2.style = doc.styles['Heading 2']

    overview_gen_1 = doc.add_paragraph()
    overview_gen_1.add_run("""
The report sets out to identify features of websites that impact on web accessibility across a random selection of pages from each site, and to consider what potential issues those pages and sites may have, and t compare this across sites to see which sites are at most risk of failing accessibility.
""".strip())

    overview_gen_2 = doc.add_paragraph()
    overview_gen_2.add_run("""
A number of subjects that imact on web accessibility were considered:
""".strip())
    
    doc.add_paragraph("Basic HTML structure (page language, title etc.)", style='List Number')
    doc.add_paragraph("Content styling (fonts, lists, animation, title attribute)", style='List Number')
    doc.add_paragraph("Multmedia content (images, videos, maps)", style='List Number')
    doc.add_paragraph("Non-text content (adjacent blocks)", style='List Number')
    doc.add_paragraph("Underlying semantic structure (headings, landmarks, tables, timers)", style='List Number')  
    doc.add_paragraph("Navigation aids (menus, forms, dialogs, links, buttons, 'more' controls, floating dialogs), ", style='List Number')
    doc.add_paragraph("Navigation order (tab order, tabindex), ", style='List Number')
    doc.add_paragraph("Accessility supported (event handling, focus management, floating content, accessible names) ", style='List Number')
    doc.add_paragraph("Use of Colour (text/non-text contrast,user as indicator) ", style='List Number')
    doc.add_paragraph("Local electronic documents (Linked PDFs)", style='List Number')

    #############################################
    # Summary Findings
    # XXXXXXXXXXX
    #############################################

    h2 = doc.add_heading('Summary findings', level=1)
    h2.style = doc.styles['Heading 1']

    #############################
    # Accessible names
    #############################

    h2 = doc.add_heading('Accessible names', level=2)
    h2.style = doc.styles['Heading 2']

    # Original query to get pages with missing accessible names (for total counts)
    pages_with_name_issues = list(db_connection.page_results.find(
        {"results.accessibility.tests.accessible_names.accessible_names.details.summary.missingNames": {"$gt": 0}},
        {
            "url": 1,
            "results.accessibility.tests.accessible_names.accessible_names.details.summary.missingNames": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Count affected domains (needed for overall statistics)
    affected_domains = set()
    total_missing_names = 0
    for page in pages_with_name_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        affected_domains.add(domain)
        total_missing_names += page['results']['accessibility']['tests']['accessible_names']['accessible_names']['details']['summary']['missingNames']

    # Query for pages with violations for tag-specific analysis
    pages_with_violations = list(db_connection.page_results.find(
        {"results.accessibility.tests.accessible_names.accessible_names.details.violations": {"$exists": True}},
        {
            "url": 1,
            "results.accessibility.tests.accessible_names.accessible_names.details.violations": 1,
            "_id": 0
        }
    ))

    # Process violations to count by tag
    tag_statistics = {}
    
    for page in pages_with_violations:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        
        # Get the violations array and parse it if it's a string
        violations = page['results']['accessibility']['tests']['accessible_names']['accessible_names']['details']['violations']
        if isinstance(violations, str):
            violations = json.loads(violations)
        
        # Track unique tags for this page
        page_tags = set()
        
        for violation in violations:
            tag = violation['element']
            
            if tag not in tag_statistics:
                tag_statistics[tag] = {
                    'count': 0,
                    'pages': set(),
                    'domains': set()
                }
            
            tag_statistics[tag]['count'] += 1
            tag_statistics[tag]['pages'].add(page['url'])
            tag_statistics[tag]['domains'].add(domain)

    # Create results table
    table = doc.add_table(rows=len(tag_statistics) + 1, cols=4)
    table.style = 'Table Grid'

    # Set column headers
    headers = table.rows[0].cells
    headers[0].text = "Tag name"
    headers[1].text = "# of instances"
    headers[2].text = "# of sites"
    headers[3].text = "% of sites"

    # Add data for each tag
    for i, (tag, stats) in enumerate(sorted(tag_statistics.items()), 1):
        row = table.rows[i].cells
        percentage = (len(stats['domains']) / len(total_domains)) * 100 if total_domains else 0
        
        row[0].text = f"<{tag}>"
        row[1].text = str(stats['count'])
        row[2].text = str(len(stats['domains']))
        row[3].text = f"{percentage:.1f}%"

    # Format the table text
    format_table_text(table)

    # Add some space after the table
    doc.add_paragraph()    

    #############################
    # Animation
    #############################

    h2 = doc.add_heading('Animation', level=2)
    h2.style = doc.styles['Heading 2']

    # Query for pages that have animations but lack reduced motion support
    pages_lacking_motion_support = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.animations.animations.pageFlags.hasAnimations": True,
            "results.accessibility.tests.animations.animations.pageFlags.lacksReducedMotionSupport": True
        },
        {
            "url": 1,
            "results.accessibility.tests.animations.animations.details.summary": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Count affected domains
    affected_domains = set()
    for page in pages_lacking_motion_support:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        affected_domains.add(domain)

    # Calculate percentage
    percentage = (len(affected_domains) / len(total_domains)) * 100 if total_domains else 0

    # Create summary table
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'

    # Set column headers
    headers = table.rows[0].cells
    headers[0].text = "Issue"
    headers[1].text = "# of pages"
    headers[2].text = "# of sites affected"
    headers[3].text = "% of sites"

    # Add data
    row = table.rows[1].cells
    row[0].text = "No reduced motion media query"
    row[1].text = str(len(pages_lacking_motion_support))
    row[2].text = str(len(affected_domains))
    row[3].text = f"{percentage:.1f}%"

    # Format the table text
    format_table_text(table)

    # Add some space after the table
    doc.add_paragraph()

    ###################################
    # Colour Contrast
    ###################################

    h2 = doc.add_heading('Colour Contrast', level=2)
    h2.style = doc.styles['Heading 2']

    # Define the contrast issues to be analyzed
    contrast_issues = [
        {
            'name': 'Text Contrast Issues',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.hasContrastIssues',
            'details_field': 'results.accessibility.tests.colors.colors.details.summary.contrastViolations'
        },
        {
            'name': 'Non-Text Contrast Issues',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.hasNonTextContrastIssues',
            'details_field': 'results.accessibility.tests.colors.colors.details.summary.nonTextContrastViolations'
        },
        {
            'name': 'Adjacent Contrast Issues',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.hasAdjacentContrastIssues',
            'details_field': 'results.accessibility.tests.colors.colors.details.summary.adjacentContrastViolations'
        },
        {
            'name': 'Contrast Preferences Support',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.supportsContrastPreferences',
            'details_field': None  # This is a boolean field, not a count
        }
    ]

    # Gather the data for each issue type
    issue_data = {}

    for issue in contrast_issues:
        # For the Contrast Preferences Support, we want sites that DO support it
        # For other issues, we want sites that have problems
        if issue['name'] == 'Contrast Preferences Support':
            query = {issue['db_field']: True}
        else:
            query = {issue['db_field']: True}
        
        # Prepare projection
        projection = {"url": 1, "_id": 0}
        if issue['details_field']:
            projection[issue['details_field']] = 1
        
        # Query the database to find pages with this issue
        pages_with_issue = list(db_connection.page_results.find(query, projection))
        
        # Count affected domains and total issue instances
        affected_domains = set()
        total_instances = 0
        
        for page in pages_with_issue:
            domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
            affected_domains.add(domain)
            
            # Count instances if applicable
            if issue['details_field']:
                # Navigate the nested structure to get the count
                parts = issue['details_field'].split('.')
                value = page
                try:
                    for part in parts:
                        if part in value:
                            value = value[part]
                        else:
                            value = 0
                            break
                    
                    if isinstance(value, (int, float)):
                        total_instances += value
                except:
                    pass  # Handle any issues with nested access
        
        # Store the data
        issue_data[issue['name']] = {
            'pages': pages_with_issue,
            'domains': affected_domains,
            'instances': total_instances
        }

    # Create summary table
    last_para = doc.add_paragraph()
    last_para._element.get_or_add_pPr().append(
        parse_xml(r'<w:keepNext xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    )

    summary_table = doc.add_table(rows=len(contrast_issues) + 1, cols=4)
    summary_table.style = 'Table Grid'

    # Keep table together
    for row in summary_table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcPr.append(parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>'))
            tcPr.append(parse_xml(r'<w:noWrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

    # Set column headers
    headers = summary_table.rows[0].cells
    headers[0].text = "Color Accessibility Issue"
    headers[1].text = "Pages Affected"
    headers[2].text = "Sites Affected"
    headers[3].text = "% of Total Sites"

    # Add data
    for i, issue in enumerate(contrast_issues, 1):
        row = summary_table.rows[i].cells
        data = issue_data[issue['name']]
        
        row[0].text = issue['name']
        row[1].text = str(len(data['pages']))
        row[2].text = str(len(data['domains']))
        
        percentage = (len(data['domains']) / len(total_domains)) * 100 if total_domains else 0
        row[3].text = f"{percentage:.1f}%"

    # Format the table text
    format_table_text(summary_table)

    ###################################
    # Colour as indicator
    ###################################

    doc.add_paragraph()
    h2 = doc.add_heading('Colour as Indicator', level=2)
    h2.style = doc.styles['Heading 2']

    # Define the indicator issues to be analyzed
    indicator_issues = [
        {
            'name': 'Color-Only Links',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.hasColorOnlyLinks',
            'details_field': 'results.accessibility.tests.colors.colors.details.summary.colorOnlyLinks'
        },
        {
            'name': 'Color References',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.hasColorReferences',
            'details_field': 'results.accessibility.tests.colors.colors.details.summary.colorReferenceCount'
        },
        {
            'name': 'Color Scheme Preferences Support',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.supportsColorSchemePreferences',
            'details_field': None  # This is a boolean field, not a count
        }
    ]

    # Gather the data for each issue type
    indicator_data = {}

    for issue in indicator_issues:
        # For the Color Scheme Preferences Support, we want sites that DO support it
        # For other issues, we want sites that have problems
        if issue['name'] == 'Color Scheme Preferences Support':
            query = {issue['db_field']: True}
        else:
            query = {issue['db_field']: True}
        
        # Prepare projection
        projection = {"url": 1, "_id": 0}
        if issue['details_field']:
            projection[issue['details_field']] = 1
        
        # Query the database to find pages with this issue
        pages_with_issue = list(db_connection.page_results.find(query, projection))
        
        # Count affected domains and total issue instances
        affected_domains = set()
        total_instances = 0
        
        for page in pages_with_issue:
            domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
            affected_domains.add(domain)
            
            # Count instances if applicable
            if issue['details_field']:
                # Navigate the nested structure to get the count
                parts = issue['details_field'].split('.')
                value = page
                try:
                    for part in parts:
                        if part in value:
                            value = value[part]
                        else:
                            value = 0
                            break
                    
                    if isinstance(value, (int, float)):
                        total_instances += value
                except:
                    pass  # Handle any issues with nested access
        
        # Store the data
        indicator_data[issue['name']] = {
            'pages': pages_with_issue,
            'domains': affected_domains,
            'instances': total_instances
        }

    # Create summary table
    last_para = doc.add_paragraph()
    last_para._element.get_or_add_pPr().append(
        parse_xml(r'<w:keepNext xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    )

    summary_table = doc.add_table(rows=len(indicator_issues) + 1, cols=4)
    summary_table.style = 'Table Grid'

    # Keep table together
    for row in summary_table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcPr.append(parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>'))
            tcPr.append(parse_xml(r'<w:noWrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

    # Set column headers
    headers = summary_table.rows[0].cells
    headers[0].text = "Color Accessibility Issue"
    headers[1].text = "Pages Affected"
    headers[2].text = "Sites Affected"
    headers[3].text = "% of Total Sites"

    # Add data
    for i, issue in enumerate(indicator_issues, 1):
        row = summary_table.rows[i].cells
        data = indicator_data[issue['name']]
        
        row[0].text = issue['name']
        row[1].text = str(len(data['pages']))
        row[2].text = str(len(data['domains']))
        
        percentage = (len(data['domains']) / len(total_domains)) * 100 if total_domains else 0
        row[3].text = f"{percentage:.1f}%"

    # Format the table text
    format_table_text(summary_table)

    ################################################
    # Dialogs
    ################################################

    doc.add_paragraph()
    h2 = doc.add_heading('Dialogs', level=2)
    h2.style = doc.styles['Heading 2']

 
    # Query for pages with modal issues
    pages_with_modal_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.modals.modals.pageFlags.hasModals": True,
            "results.accessibility.tests.modals.modals.pageFlags.hasModalViolations": True
        },
        {
            "url": 1,
            "results.accessibility.tests.modals.modals.pageFlags": 1,
            "results.accessibility.tests.modals.modals.details.summary": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    modal_issues = {
        "modalsWithoutClose": {"name": "Missing close mechanism", "pages": set(), "domains": set()},
        "modalsWithoutFocusManagement": {"name": "Improper focus management", "pages": set(), "domains": set()},
        "modalsWithoutProperHeading": {"name": "Missing/improper heading", "pages": set(), "domains": set()},
        "modalsWithoutTriggers": {"name": "Missing/improper triggers", "pages": set(), "domains": set()}
    }

    # Count issues
    for page in pages_with_modal_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        summary = page['results']['accessibility']['tests']['modals']['modals']['details']['summary']
        
        for flag in modal_issues:
            if summary.get(flag, 0) > 0:
                modal_issues[flag]['pages'].add(page['url'])
                modal_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in modal_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Modal Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

    else:
        doc.add_paragraph("No dialog accessibility issues were found.")


    ####################################################
    # Event handling
    ####################################################

    h2 = doc.add_heading('Event Handling and Keyboard Interaction', level=2)
    h2.style = doc.styles['Heading 2']

 
    # Query for pages with event information
    pages_with_events = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.events.events": {"$exists": True}
        },
        {
            "url": 1,
            "results.accessibility.tests.events.events": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize tracking structures
    property_data = {
        # Event Types
        "event_mouse": {"name": "Mouse Events", "pages": set(), "domains": set(), "count": 0},
        "event_keyboard": {"name": "Keyboard Events", "pages": set(), "domains": set(), "count": 0},
        "event_focus": {"name": "Focus Events", "pages": set(), "domains": set(), "count": 0},
        "event_touch": {"name": "Touch Events", "pages": set(), "domains": set(), "count": 0},
        "event_timer": {"name": "Timer Events", "pages": set(), "domains": set(), "count": 0},
        "event_lifecycle": {"name": "Lifecycle Events", "pages": set(), "domains": set(), "count": 0},
        "event_other": {"name": "Other Events", "pages": set(), "domains": set(), "count": 0},
        
        # Tab Order
        "explicit_tabindex": {"name": "Explicit tabindex Usage", "pages": set(), "domains": set(), "count": 0},
        "visual_violations": {"name": "Visual Order Violations", "pages": set(), "domains": set(), "count": 0},
        "column_violations": {"name": "Column Order Violations", "pages": set(), "domains": set(), "count": 0},
        "negative_tabindex": {"name": "Negative Tabindex", "pages": set(), "domains": set(), "count": 0},
        "high_tabindex": {"name": "High Tabindex Values", "pages": set(), "domains": set(), "count": 0},
        
        # Interactive Elements
        "mouse_only": {"name": "Mouse-only Elements", "pages": set(), "domains": set(), "count": 0},
        "missing_tabindex": {"name": "Missing tabindex", "pages": set(), "domains": set(), "count": 0},
        "non_interactive": {"name": "Non-interactive with Handlers", "pages": set(), "domains": set(), "count": 0},
        
        # Modal Support
        "modals_no_escape": {"name": "Modals Missing Escape", "pages": set(), "domains": set(), "count": 0}
    }

    # Create detailed violation tracking organized by domain and URL
    domain_data = {}

    # Process each page
    for page in pages_with_events:
        try:
            url = page['url']
            
            domain = url.replace('http://', '').replace('https://', '').split('/')[0]
            event_data = page['results']['accessibility']['tests']['events']['events']
            
            # Initialize domain and URL tracking if needed
            if domain not in domain_data:
                domain_data[domain] = {
                    'urls': {}
                }
            
            # Initialize URL data
            domain_data[domain]['urls'][url] = {
                'event_types': {},
                'violations': {},
                'handlers_count': 0,
                'focusable_elements': 0,
                'total_violations': 0
            }
            
            # Get pageFlags data for the most reliable summary information
            pageFlags = event_data.get('pageFlags', {})
            details = pageFlags.get('details', {})
            
            # Track total handlers and violations
            total_handlers = details.get('totalHandlers', 0)
            total_violations = details.get('totalViolations', 0)
            domain_data[domain]['urls'][url]['handlers_count'] = total_handlers
            domain_data[domain]['urls'][url]['total_violations'] = total_violations

            # Track total focusable elements
            tab_order_data = details.get('tabOrder', {})
            focusable_elements = tab_order_data.get('totalFocusableElements', 0)
            domain_data[domain]['urls'][url]['focusable_elements'] = focusable_elements
            
            # Process event types using the updated structure
            by_type = details.get('byType', {})
            
            for event_type in ['mouse', 'keyboard', 'focus', 'touch', 'timer', 'lifecycle', 'other']:
                count = by_type.get(event_type, 0)
                
                if isinstance(count, list):
                    count = len(count)
                elif not isinstance(count, (int, float)):
                    try:
                        count = int(count or 0)
                    except (ValueError, TypeError):
                        count = 0
                
                # Track event type for this URL
                domain_data[domain]['urls'][url]['event_types'][event_type] = count
                
                if count > 0:
                    key = f"event_{event_type}"
                    property_data[key]['pages'].add(url)
                    property_data[key]['domains'].add(domain)
                    property_data[key]['count'] += count

            # Process violation counts by type
            violation_counts = details.get('violationCounts', {})
            
            # Tab order violations
            explicit_count = tab_order_data.get('elementsWithExplicitTabIndex', 0)
            visual_violations = violation_counts.get('visual-order', 0) or tab_order_data.get('visualOrderViolations', 0)
            column_violations = violation_counts.get('column-order', 0) or tab_order_data.get('columnOrderViolations', 0)
            
            # Track negative and high tabindex
            negative_tabindex = 1 if pageFlags.get('hasNegativeTabindex', False) else 0
            high_tabindex = 1 if pageFlags.get('hasHighTabindex', False) else 0
            
            # Track violations for this URL
            domain_data[domain]['urls'][url]['violations']['explicit_tabindex'] = explicit_count
            domain_data[domain]['urls'][url]['violations']['visual_order'] = visual_violations
            domain_data[domain]['urls'][url]['violations']['column_order'] = column_violations
            domain_data[domain]['urls'][url]['violations']['negative_tabindex'] = negative_tabindex
            domain_data[domain]['urls'][url]['violations']['high_tabindex'] = high_tabindex
            
            if explicit_count > 0:
                property_data['explicit_tabindex']['pages'].add(url)
                property_data['explicit_tabindex']['domains'].add(domain)
                property_data['explicit_tabindex']['count'] += explicit_count
                
            if visual_violations > 0:
                property_data['visual_violations']['pages'].add(url)
                property_data['visual_violations']['domains'].add(domain)
                property_data['visual_violations']['count'] += visual_violations
                
            if column_violations > 0:
                property_data['column_violations']['pages'].add(url)
                property_data['column_violations']['domains'].add(domain)
                property_data['column_violations']['count'] += column_violations
                
            if negative_tabindex > 0:
                property_data['negative_tabindex']['pages'].add(url)
                property_data['negative_tabindex']['domains'].add(domain)
                property_data['negative_tabindex']['count'] += negative_tabindex
                
            if high_tabindex > 0:
                property_data['high_tabindex']['pages'].add(url)
                property_data['high_tabindex']['domains'].add(domain)
                property_data['high_tabindex']['count'] += high_tabindex

            # Process element violations
            mouse_only = violation_counts.get('mouse-only', 0) or details.get('mouseOnlyElements', {}).get('count', 0)
            missing_tabindex = violation_counts.get('missing-tabindex', 0) or details.get('missingTabindex', 0)
            non_interactive = details.get('nonInteractiveWithHandlers', 0)
            modals_without_escape = violation_counts.get('modal-without-escape', 0)
            
            # Track violations for this URL
            domain_data[domain]['urls'][url]['violations']['mouse_only'] = mouse_only
            domain_data[domain]['urls'][url]['violations']['missing_tabindex'] = missing_tabindex
            domain_data[domain]['urls'][url]['violations']['non_interactive'] = non_interactive
            domain_data[domain]['urls'][url]['violations']['modals_no_escape'] = modals_without_escape
            
            if mouse_only > 0:
                property_data['mouse_only']['pages'].add(url)
                property_data['mouse_only']['domains'].add(domain)
                property_data['mouse_only']['count'] += mouse_only
                
            if missing_tabindex > 0:
                property_data['missing_tabindex']['pages'].add(url)
                property_data['missing_tabindex']['domains'].add(domain)
                property_data['missing_tabindex']['count'] += missing_tabindex
                
            if non_interactive > 0:
                property_data['non_interactive']['pages'].add(url)
                property_data['non_interactive']['domains'].add(domain)
                property_data['non_interactive']['count'] += non_interactive
                
            if modals_without_escape > 0:
                property_data['modals_no_escape']['pages'].add(url)
                property_data['modals_no_escape']['domains'].add(domain)
                property_data['modals_no_escape']['count'] += modals_without_escape

        except Exception as e:
            print(f"Error processing page {url}:")
            print("Exception:", str(e))
            traceback.print_exc()
            continue

    if pages_with_events:
        # Part 1: Overall Summary Table
        doc.add_heading('Event Handling Summary', level=3)
        
        # Calculate number of rows needed
        rows_needed = 1  # Header row
        
        # Event Types section (header + all event types)
        rows_needed += 1  # Section header
        rows_needed += len([k for k in property_data.keys() if k.startswith('event_')])
        
        # Tab Order section (header + 5 items)
        rows_needed += 1  # Section header
        rows_needed += 5  # explicit_tabindex, visual_violations, column_violations, negative_tabindex, high_tabindex
        
        # Interactive Elements section (header + 3 items)
        rows_needed += 1  # Section header
        rows_needed += 3  # mouse_only, missing_tabindex, non_interactive
        
        # Modal Support section (header + 1 item)
        rows_needed += 1  # Section header
        rows_needed += 1  # modals_no_escape

        # Create table with correct number of rows
        table = doc.add_table(rows=rows_needed, cols=4)
        table.style = 'Table Grid'
        
        # Add headers
        headers = table.rows[0].cells
        headers[0].text = "Property"
        headers[1].text = "Occurrences"
        headers[2].text = "Pages Affected"
        headers[3].text = "% of Sites"
        
        current_row = 1
        
        # Add Event Types section
        row = table.rows[current_row].cells
        row[0].text = "Event Types:"
        current_row += 1
        
        for key, data in sorted([(k, v) for k, v in property_data.items() if k.startswith('event_')], 
                            key=lambda x: x[1]['count'], reverse=True):
            row = table.rows[current_row].cells
            row[0].text = "  " + data['name']
            row[1].text = str(data['count'])
            row[2].text = str(len(data['pages']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"
            current_row += 1
        
        # Add Tab Order section
        row = table.rows[current_row].cells
        row[0].text = "Tab Order:"
        current_row += 1
        
        for key in ['explicit_tabindex', 'visual_violations', 'column_violations', 'negative_tabindex', 'high_tabindex']:
            row = table.rows[current_row].cells
            row[0].text = "  " + property_data[key]['name']
            row[1].text = str(property_data[key]['count'])
            row[2].text = str(len(property_data[key]['pages']))
            row[3].text = f"{(len(property_data[key]['domains']) / len(total_domains) * 100):.1f}%"
            current_row += 1
        
        # Add Interactive Elements section
        row = table.rows[current_row].cells
        row[0].text = "Interactive Elements:"
        current_row += 1
        
        for key in ['mouse_only', 'missing_tabindex', 'non_interactive']:
            row = table.rows[current_row].cells
            row[0].text = "  " + property_data[key]['name']
            row[1].text = str(property_data[key]['count'])
            row[2].text = str(len(property_data[key]['pages']))
            row[3].text = f"{(len(property_data[key]['domains']) / len(total_domains) * 100):.1f}%"
            current_row += 1
        
        # Add Modal Support section
        row = table.rows[current_row].cells
        row[0].text = "Modal Support:"
        current_row += 1
        
        key = 'modals_no_escape'
        row = table.rows[current_row].cells
        row[0].text = "  " + property_data[key]['name']
        row[1].text = str(property_data[key]['count'])
        row[2].text = str(len(property_data[key]['pages']))
        row[3].text = f"{(len(property_data[key]['domains']) / len(total_domains) * 100):.1f}%"

        format_table_text(table)
        
    else:
        doc.add_paragraph("No event handling data was found.")


    #################################
    # Floating dialogs
    #################################

    doc.add_paragraph()
    h3 = doc.add_heading('Floating Dialogs', level=2)
    h3.style = doc.styles['Heading 2']

    # Query for documentation to use in our explanation
    dialog_docs = list(db_connection.page_results.find(
        {"results.accessibility.tests.floating_dialogs.dialogs.documentation": {"$exists": True}},
        {"results.accessibility.tests.floating_dialogs.dialogs.documentation": 1, "_id": 0}
    ).limit(1))

    # Query for pages with dialog issues - using the consolidated results field
    pages_with_dialog_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.floating_dialogs.dialogs.consolidated": {"$exists": True},
            "results.accessibility.tests.floating_dialogs.dialogs.consolidated.summary.totalIssues": {"$gt": 0}
        },
        {
            "url": 1,
            "results.accessibility.tests.floating_dialogs.dialogs.consolidated": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type by severity
    dialog_issues = {
        "violations": {
            "hiddenInteractiveContent": {"name": "Hidden interactive content", "pages": set(), "domains": set(), "severity": "critical"},
            "incorrectHeadingLevel": {"name": "Incorrect heading structure", "pages": set(), "domains": set(), "severity": "high"},
            "missingCloseButton": {"name": "Missing close button", "pages": set(), "domains": set(), "severity": "high"},
            "improperFocusManagement": {"name": "Improper focus management", "pages": set(), "domains": set(), "severity": "high"}
        },
        "warnings": {
            "contentOverlap": {"name": "Content overlap issues", "pages": set(), "domains": set(), "severity": "moderate"}
        }
    }

    # Count issues and store URLs by domain
    domain_to_urls = {}

    for page in pages_with_dialog_issues:
        url = page['url']
        domain = url.replace('http://', '').replace('https://', '').split('/')[0]
        consolidated = page['results']['accessibility']['tests']['floating_dialogs']['dialogs']['consolidated']
        
        # Initialize domain entry if it doesn't exist
        if domain not in domain_to_urls:
            domain_to_urls[domain] = {}
        
        # Process violations
        if 'issuesByType' in consolidated:
            issues_by_type = consolidated['issuesByType']
            
            # Process violations
            for violation_type, violation_data in issues_by_type.get('violations', {}).items():
                if violation_type in dialog_issues['violations'] and violation_data.get('count', 0) > 0:
                    dialog_issues['violations'][violation_type]['pages'].add(url)
                    dialog_issues['violations'][violation_type]['domains'].add(domain)
                    
                    # Store the severity if available
                    if 'severity' in violation_data:
                        dialog_issues['violations'][violation_type]['severity'] = violation_data['severity']
                    
                    # Store URL by issue type for this domain
                    if violation_type not in domain_to_urls[domain]:
                        domain_to_urls[domain][violation_type] = []
                    domain_to_urls[domain][violation_type].append(url)
            
            # Process warnings
            for warning_type, warning_data in issues_by_type.get('warnings', {}).items():
                if warning_type in dialog_issues['warnings'] and warning_data.get('count', 0) > 0:
                    dialog_issues['warnings'][warning_type]['pages'].add(url)
                    dialog_issues['warnings'][warning_type]['domains'].add(domain)
                    
                    # Store the severity if available
                    if 'severity' in warning_data:
                        dialog_issues['warnings'][warning_type]['severity'] = warning_data['severity']
                    
                    # Store URL by issue type for this domain
                    if warning_type not in domain_to_urls[domain]:
                        domain_to_urls[domain][warning_type] = []
                    domain_to_urls[domain][warning_type].append(url)

    # Create filtered list of issues that have affected pages
    all_active_issues = []

    for category in ['violations', 'warnings']:
        for issue_type, data in dialog_issues[category].items():
            if len(data['pages']) > 0:
                all_active_issues.append({
                    'category': category,
                    'type': issue_type,
                    'name': data['name'],
                    'severity': data['severity'],
                    'pages': data['pages'],
                    'domains': data['domains']
                })

    # Sort issues by severity - critical first, then high, then moderate
    severity_order = {'critical': 0, 'high': 1, 'moderate': 2, 'low': 3}
    all_active_issues.sort(key=lambda x: severity_order.get(x['severity'], 4))

    if all_active_issues:
        # Create summary table
        
        summary_table = doc.add_table(rows=len(all_active_issues) + 1, cols=5)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue"
        headers[1].text = "Severity"
        headers[2].text = "Pages Affected"
        headers[3].text = "Sites Affected"
        headers[4].text = "% of Total Sites"

        # Add data
        for i, issue in enumerate(all_active_issues, 1):
            row = summary_table.rows[i].cells
            row[0].text = issue['name']
            row[1].text = issue['severity'].capitalize()
            row[2].text = str(len(issue['pages']))
            row[3].text = str(len(issue['domains']))
            row[4].text = f"{(len(issue['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

    else:
        doc.add_paragraph("No floating dialog accessibility issues were found.")


    #################################
    # Focus Management (general)
    #################################

    doc.add_paragraph()
    h2 = doc.add_heading('Focus Management (General)', level=2)
    h2.style = doc.styles['Heading 2']

    # Query for pages with focus management information
    pages_with_focus = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.focus_management.focus_management": {"$exists": True}
        },
        {
            "url": 1,
            "results.accessibility.tests.focus_management.focus_management": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize tracking
    site_data = {}
    url_data = {}  # Track data by individual URL
    total_interactive_elements = 0
    total_violations = 0
    total_breakpoints_tested = 0

    # First, add a description of the test based on the documentation
    if pages_with_focus:
        first_page = pages_with_focus[0]
        '''
        test_doc = first_page.get('results', {}).get('accessibility', {}).get('tests', {}).get(
            'focus_management', {}).get('focus_management', {}).get('test_documentation', {})
        '''
        
    # Process each page
    for page in pages_with_focus:
        try:
            url = page['url']
            domain = url.replace('http://', '').replace('https://', '').split('/')[0]
            focus_data = page['results']['accessibility']['tests']['focus_management']['focus_management']
            
            # Initialize domain tracking
            if domain not in site_data:
                site_data[domain] = {
                    "total_violations": 0,
                    "breakpoints_tested": 0,
                    "urls": set(),
                    "tests": {
                        "focus_outline_presence": {"violations": 0, "elements": set()},
                        "focus_outline_contrast": {"violations": 0, "elements": set()},
                        "focus_outline_offset": {"violations": 0, "elements": set()},
                        "hover_feedback": {"violations": 0, "elements": set()},
                        "focus_obscurement": {"violations": 0, "elements": set()},
                        "anchor_target_tabindex": {"violations": 0, "elements": set()}
                    }
                }
            
            # Initialize URL tracking
            if url not in url_data:
                url_data[url] = {
                    "domain": domain,
                    "total_violations": 0,
                    "breakpoints_tested": 0,
                    "tests": {
                        "focus_outline_presence": {"violations": 0, "elements": set()},
                        "focus_outline_contrast": {"violations": 0, "elements": set()},
                        "focus_outline_offset": {"violations": 0, "elements": set()},
                        "hover_feedback": {"violations": 0, "elements": set()},
                        "focus_obscurement": {"violations": 0, "elements": set()},
                        "anchor_target_tabindex": {"violations": 0, "elements": set()}
                    }
                }

            # Get metadata
            metadata = focus_data.get('metadata', {})
            url_violations = metadata.get('total_violations_found', 0)
            
            total_violations += url_violations
            site_data[domain]["total_violations"] += url_violations
            url_data[url]["total_violations"] = url_violations
            
            breakpoints_tested = metadata.get('total_breakpoints_tested', 0)
            total_breakpoints_tested += breakpoints_tested
            site_data[domain]["breakpoints_tested"] = max(site_data[domain]["breakpoints_tested"], breakpoints_tested)
            url_data[url]["breakpoints_tested"] = breakpoints_tested
            
            # Add URL to domain list
            site_data[domain]["urls"].add(url)

            # Process each test
            tests = focus_data.get('tests', {})
            for test_name, test_data in tests.items():
                if test_name in site_data[domain]["tests"]:
                    # Get summary data
                    summary = test_data.get('summary', {})
                    violations = summary.get('total_violations', 0)
                    
                    # Update site data
                    site_data[domain]["tests"][test_name]["violations"] += violations
                    
                    # Update URL data
                    url_data[url]["tests"][test_name]["violations"] = violations
                    
                    # Track affected elements
                    elements = test_data.get('elements_affected', [])
                    if isinstance(elements, list):
                        site_data[domain]["tests"][test_name]["elements"].update(elements)
                        url_data[url]["tests"][test_name]["elements"].update(elements)

        except Exception as e:
            print(f"Error processing page {page.get('url', 'unknown')}: {str(e)}")
            continue

    if pages_with_focus:
        # Overall Statistics
        #doc.add_paragraph("Focus Management Statistics:", style='Normal')
        stats_table = doc.add_table(rows=3, cols=2)
        stats_table.style = 'Table Grid'
        
        rows = stats_table.rows
        rows[0].cells[0].text = "Pages Tested"
        rows[0].cells[1].text = str(len(pages_with_focus))
        rows[1].cells[0].text = "Total Breakpoints Tested (across all pages)"
        rows[1].cells[1].text = str(total_breakpoints_tested)
        rows[2].cells[0].text = "Total Violations Found"
        rows[2].cells[1].text = str(total_violations)
        
        format_table_text(stats_table)

        # Tests information
        #doc.add_paragraph()
        #doc.add_paragraph("Tests Performed:", style='Normal')
        
        if first_page:
            '''
            tests_performed = test_doc.get('tests_performed', [])
            test_table = doc.add_table(rows=len(tests_performed) + 1, cols=3)
            test_table.style = 'Table Grid'
            
            # Add headers
            headers = test_table.rows[0].cells
            headers[0].text = "Test Name"
            headers[1].text = "Description"
            headers[2].text = "Success Criteria"
            
            # Add data for each test
            for i, test_info in enumerate(tests_performed, 1):
                row = test_table.rows[i].cells
                row[0].text = test_info.get('name', '')
                row[1].text = test_info.get('description', '')
                row[2].text = test_info.get('success_criteria', '')
                
            format_table_text(test_table)
            '''


        # Map test IDs to more readable names
        test_name_map = {
            "focus_outline_presence": "Missing Focus Outlines",
            "focus_outline_contrast": "Insufficient Outline Contrast",
            "focus_outline_offset": "Insufficient Outline Offset/Width",
            "hover_feedback": "Insufficient Hover Feedback",
            "focus_obscurement": "Obscured Focus Outlines",
            "anchor_target_tabindex": "Improper Local Target Configuration"
        }

        # Detailed Issues Summary
        doc.add_paragraph()
        doc.add_paragraph("Focus Management Issues by Test Type:", style='Normal')
        
        # Calculate totals for each test type
        test_totals = {}
        for test_id, display_name in test_name_map.items():
            total_violations = sum(site["tests"][test_id]["violations"] for site in site_data.values())
            affected_sites = sum(1 for site in site_data.values() if site["tests"][test_id]["violations"] > 0)
            affected_pages = sum(1 for url in url_data.values() if url["tests"][test_id]["violations"] > 0)
            
            # Count unique elements across all sites
            all_elements = set()
            for site in site_data.values():
                all_elements.update(site["tests"][test_id]["elements"])
            
            test_totals[test_id] = {
                "violations": total_violations,
                "affected_sites": affected_sites,
                "affected_pages": affected_pages,
                "unique_elements": len(all_elements)
            }
        
        # Create table for test summaries
        summary_table = doc.add_table(rows=len(test_name_map) + 1, cols=5)
        summary_table.style = 'Table Grid'
        
        # Add headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue Type"
        headers[1].text = "Total Violations"
        headers[2].text = "Pages Affected"
        headers[3].text = "Sites Affected"
        headers[4].text = "% of Total Sites"
        
        # Add data for each test type
        row_idx = 1
        for test_id, display_name in test_name_map.items():
            totals = test_totals[test_id]
            row = summary_table.rows[row_idx].cells
            row[0].text = display_name
            row[1].text = str(totals["violations"])
            row[2].text = str(totals["affected_pages"])
            row[3].text = str(totals["affected_sites"])
            row[4].text = f"{(totals['affected_sites'] / len(site_data) * 100):.1f}%" if site_data else "0%"
            row_idx += 1
        
        format_table_text(summary_table)
        

    else:
        doc.add_paragraph("No focus management data available in the database.", style='Normal')

    #################################
    # Fonts
    #################################

    doc.add_paragraph()
    h2 = doc.add_heading('Fonts', level=2)
    h2.style = doc.styles['Heading 2']

    # System fonts definition
    SYSTEM_FONTS = [
        "Arial", "Helvetica", "Times New Roman", "Times", "Courier New", 
        "Courier", "Verdana", "Georgia", "Palatino", "Garamond", "Bookman",
        "Tahoma", "Trebuchet MS", "Impact", "Comic Sans MS", "Webdings", 
        "Symbol", "Calibri", "Cambria", "Segoe UI"
    ]



 
    # Query for pages with font information
    pages_with_fonts = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.fonts.font_analysis": {"$exists": True}
        },
        {
            "url": 1,
            "results.accessibility.tests.fonts.font_analysis": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize tracking structures
    font_usage = {}
    site_fonts = {}
    total_font_count = 0
    heading_sizes = []

    # Typography issue tracking
    typography_issues = {
        "small_text": {"name": "Small text", "pages": set(), "domains": set()},
        "small_line_height": {"name": "Small line height", "pages": set(), "domains": set()},
        "justified_text": {"name": "Justified text", "pages": set(), "domains": set()},
        "right_aligned": {"name": "Right-aligned text", "pages": set(), "domains": set()},
        "italic_text": {"name": "Italic text usage", "pages": set(), "domains": set()},
        "bold_larger_than_headings": {"name": "Bold text larger than headings", "pages": set(), "domains": set()}
    }

    # CSS Variable tracking
    css_var_usage = {
        'awb-text-font-family': set(),
        'body_typography-font-family': set(),
        'fontsBaseFamily': set(),
        'h1_typography-font-family': set(),
        'h2_typography-font-family': set(),
        'h3_typography-font-family': set(),
        'homepage-title-font': set()
    }

    # Process each page
    for page in pages_with_fonts:
        try:
            domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
            font_data = page['results']['accessibility']['tests']['fonts']['font_analysis']
            
            # Initialize domain in site_fonts if not present
            if domain not in site_fonts:
                site_fonts[domain] = {
                    'fonts': set(),
                    'css_vars': set(),
                    'system_fonts': set(),
                    'web_fonts': set(),
                    'smallest_heading': None,
                    'typography_issues': set()
                }

            # Process accessibility data
            accessibility_data = font_data.get('accessibility', {})
            tests = accessibility_data.get('tests', {})
            
            # Track smallest heading size
            smallest_heading = accessibility_data.get('smallestHeadingSize')
            if smallest_heading:
                heading_sizes.append(smallest_heading)
                site_fonts[domain]['smallest_heading'] = smallest_heading

            # Check typography issues
            if tests.get('hasSmallText'):
                typography_issues['small_text']['pages'].add(page['url'])
                typography_issues['small_text']['domains'].add(domain)
                site_fonts[domain]['typography_issues'].add('small_text')

            if tests.get('hasSmallLineHeight'):
                typography_issues['small_line_height']['pages'].add(page['url'])
                typography_issues['small_line_height']['domains'].add(domain)
                site_fonts[domain]['typography_issues'].add('small_line_height')

            if tests.get('hasJustifiedText'):
                typography_issues['justified_text']['pages'].add(page['url'])
                typography_issues['justified_text']['domains'].add(domain)
                site_fonts[domain]['typography_issues'].add('justified_text')

            if tests.get('hasRightAlignedText'):
                typography_issues['right_aligned']['pages'].add(page['url'])
                typography_issues['right_aligned']['domains'].add(domain)
                site_fonts[domain]['typography_issues'].add('right_aligned')

            if tests.get('hasItalicText'):
                typography_issues['italic_text']['pages'].add(page['url'])
                typography_issues['italic_text']['domains'].add(domain)
                site_fonts[domain]['typography_issues'].add('italic_text')

            if tests.get('hasBoldNonHeadingLargerThanHeadings'):
                typography_issues['bold_larger_than_headings']['pages'].add(page['url'])
                typography_issues['bold_larger_than_headings']['domains'].add(domain)
                site_fonts[domain]['typography_issues'].add('bold_larger_than_headings')

            # Process fonts
            fonts = font_data.get('fonts', {})
            for font_name in fonts.keys():
                # Skip generic families
                if font_name.lower() in ['inherit', 'sans-serif', 'monospace']:
                    continue
                
                # Handle CSS variables
                if font_name.startswith('var('):
                    site_fonts[domain]['css_vars'].add(font_name)
                    for var_name in css_var_usage.keys():
                        if var_name in font_name:
                            css_var_usage[var_name].add(domain)
                    continue
                
                # Track regular fonts
                if font_name not in font_usage:
                    font_usage[font_name] = {"domains": set()}
                
                font_usage[font_name]["domains"].add(domain)
                site_fonts[domain]['fonts'].add(font_name)
                
                # Categorize as system or web font
                if any(sf.lower() in font_name.lower() for sf in SYSTEM_FONTS):
                    site_fonts[domain]['system_fonts'].add(font_name)
                else:
                    site_fonts[domain]['web_fonts'].add(font_name)

            total_font_count += font_data.get('totalFonts', 0)
                
        except Exception as e:
            print(f"Error processing page {page.get('url', 'unknown')}: {str(e)}")
            continue

    if pages_with_fonts:
        
        active_issues = {flag: data for flag, data in typography_issues.items() 
                        if len(data['pages']) > 0}
        
        if active_issues:
            issues_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
            issues_table.style = 'Table Grid'

            headers = issues_table.rows[0].cells
            headers[0].text = "Issue"
            headers[1].text = "Pages Affected"
            headers[2].text = "Sites Affected"
            headers[3].text = "% of Total Sites"

            for i, (flag, data) in enumerate(active_issues.items(), 1):
                row = issues_table.rows[i].cells
                row[0].text = data['name']
                row[1].text = str(len(data['pages']))
                row[2].text = str(len(data['domains']))
                row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

            format_table_text(issues_table)

        # 2. Heading Size Analysis
        doc.add_paragraph()
        doc.add_paragraph("Heading Size Analysis:", style='Normal')
        if heading_sizes:
            doc.add_paragraph(f"Smallest heading size detected: {min(heading_sizes)}px")
            doc.add_paragraph(f"Average smallest heading size: {sum(heading_sizes) / len(heading_sizes):.1f}px")


    #################################
    # Forms
    #################################

    doc.add_paragraph()
    h2 = doc.add_heading('Forms', level=2)
    h2.style = doc.styles['Heading 2']

    doc.add_paragraph()

    # Query for pages with form issues
    pages_with_form_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.forms.forms.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.forms.forms.pageFlags.hasInputsWithoutLabels": True},
                {"results.accessibility.tests.forms.forms.pageFlags.hasPlaceholderOnlyInputs": True},
                {"results.accessibility.tests.forms.forms.pageFlags.hasFormsWithoutHeadings": True},
                {"results.accessibility.tests.forms.forms.pageFlags.hasFormsOutsideLandmarks": True},
                {"results.accessibility.tests.forms.forms.pageFlags.hasContrastIssues": True},
                {"results.accessibility.tests.forms.forms.pageFlags.hasLayoutIssues": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.forms.forms": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for different form issues
    form_issues = {
        "missing_labels": {
            "name": "Inputs without labels",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "placeholder_only": {
            "name": "Placeholder-only inputs",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "no_headings": {
            "name": "Forms without headings",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "outside_landmarks": {
            "name": "Forms outside landmarks",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "contrast_issues": {
            "name": "Input contrast issues",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "layout_issues": {
            "name": "Form layout issues",
            "pages": set(),
            "domains": set(),
            "count": 0
        }
    }

    # Process each page
    total_forms = 0
    for page in pages_with_form_issues:
        try:
            domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
            form_data = page['results']['accessibility']['tests']['forms']['forms']
            flags = form_data.get('pageFlags', {})
            summary = form_data.get('details', {}).get('summary', {})
            
            # Update total forms count
            total_forms += summary.get('totalForms', 0)
            
            # Check inputs without labels
            if flags.get('hasInputsWithoutLabels'):
                form_issues['missing_labels']['pages'].add(page['url'])
                form_issues['missing_labels']['domains'].add(domain)
                form_issues['missing_labels']['count'] += summary.get('inputsWithoutLabels', 0)
            
            # Check placeholder-only inputs
            if flags.get('hasPlaceholderOnlyInputs'):
                form_issues['placeholder_only']['pages'].add(page['url'])
                form_issues['placeholder_only']['domains'].add(domain)
                form_issues['placeholder_only']['count'] += summary.get('inputsWithPlaceholderOnly', 0)
            
            # Check forms without headings
            if flags.get('hasFormsWithoutHeadings'):
                form_issues['no_headings']['pages'].add(page['url'])
                form_issues['no_headings']['domains'].add(domain)
                form_issues['no_headings']['count'] += summary.get('formsWithoutHeadings', 0)
            
            # Check forms outside landmarks
            if flags.get('hasFormsOutsideLandmarks'):
                form_issues['outside_landmarks']['pages'].add(page['url'])
                form_issues['outside_landmarks']['domains'].add(domain)
                form_issues['outside_landmarks']['count'] += summary.get('formsOutsideLandmarks', 0)
            
            # Check contrast issues
            if flags.get('hasContrastIssues'):
                form_issues['contrast_issues']['pages'].add(page['url'])
                form_issues['contrast_issues']['domains'].add(domain)
                form_issues['contrast_issues']['count'] += summary.get('inputsWithContrastIssues', 0)
            
            # Check layout issues
            if flags.get('hasLayoutIssues'):
                form_issues['layout_issues']['pages'].add(page['url'])
                form_issues['layout_issues']['domains'].add(domain)
                form_issues['layout_issues']['count'] += summary.get('inputsWithLayoutIssues', 0)
                
        except Exception as e:
            print(f"Error processing page {page.get('url', 'unknown')}: {str(e)}")
            continue

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in form_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=5)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue Type"
        headers[1].text = "Number of Occurrences"
        headers[2].text = "Pages Affected"
        headers[3].text = "Sites Affected"
        headers[4].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(data['count'])
            row[2].text = str(len(data['pages']))
            row[3].text = str(len(data['domains']))
            row[4].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add statistics
        doc.add_paragraph()
        doc.add_paragraph("Form Statistics:", style='Normal')
        doc.add_paragraph(f"Total number of forms across all pages: {total_forms}")

 
        # Add domain details for each issue type
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)

    else:
        doc.add_paragraph("No form accessibility issues were found.")


    #######################################
    # Headings
    #######################################

    doc.add_paragraph()
    h3 = doc.add_heading('Headings', level=2)
    h3.style = doc.styles['Heading 2']

    # Query for pages with heading issues
    pages_with_heading_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.headings.headings.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.headings.headings.pageFlags.missingH1": True},
                {"results.accessibility.tests.headings.headings.pageFlags.multipleH1s": True},
                {"results.accessibility.tests.headings.headings.pageFlags.hasHierarchyGaps": True},
                {"results.accessibility.tests.headings.headings.pageFlags.hasHeadingsBeforeMain": True},
                {"results.accessibility.tests.headings.headings.pageFlags.hasVisualHierarchyIssues": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.headings.headings": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for different heading issues
    heading_issues = {
        "missing_h1": {"name": "Missing H1", "pages": set(), "domains": set()},
        "multiple_h1": {"name": "Multiple H1s", "pages": set(), "domains": set()},
        "hierarchy_gaps": {"name": "Hierarchy gaps", "pages": set(), "domains": set(), "count": 0},
        "headings_before_main": {"name": "Headings before main", "pages": set(), "domains": set(), "count": 0},
        "visual_hierarchy": {"name": "Visual hierarchy issues", "pages": set(), "domains": set(), "count": 0}
    }

    # Process each page
    total_headings = 0
    for page in pages_with_heading_issues:
        try:
            domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
            heading_data = page['results']['accessibility']['tests']['headings']['headings']
            flags = heading_data.get('pageFlags', {})
            
            # The issue is here - details might be in a different place or structure than expected
            details = flags.get('details', {})
            if not details:  # If details not found in flags, try the main heading_data
                details = heading_data.get('details', {})
            
            summary = heading_data.get('details', {}).get('summary', {})
            
            # Update total headings count
            headings_count = summary.get('totalHeadings', 0)
            if isinstance(headings_count, (int, float)):
                total_headings += headings_count
            
            # Check missing H1
            if flags.get('missingH1'):
                heading_issues['missing_h1']['pages'].add(page['url'])
                heading_issues['missing_h1']['domains'].add(domain)
            
            # Check multiple H1s
            if flags.get('multipleH1s'):
                heading_issues['multiple_h1']['pages'].add(page['url'])
                heading_issues['multiple_h1']['domains'].add(domain)
            
            # Check hierarchy gaps
            if flags.get('hasHierarchyGaps'):
                heading_issues['hierarchy_gaps']['pages'].add(page['url'])
                heading_issues['hierarchy_gaps']['domains'].add(domain)
                
                # Fix for the potential list issue
                hierarchy_gaps = details.get('hierarchyGaps', 0)
                if isinstance(hierarchy_gaps, list):
                    heading_issues['hierarchy_gaps']['count'] += len(hierarchy_gaps)
                elif isinstance(hierarchy_gaps, (int, float)):
                    heading_issues['hierarchy_gaps']['count'] += hierarchy_gaps
            
            # Check headings before main
            if flags.get('hasHeadingsBeforeMain'):
                heading_issues['headings_before_main']['pages'].add(page['url'])
                heading_issues['headings_before_main']['domains'].add(domain)
                
                # Fix for the potential list issue
                headings_before_main = details.get('headingsBeforeMain', 0)
                if isinstance(headings_before_main, list):
                    heading_issues['headings_before_main']['count'] += len(headings_before_main)
                elif isinstance(headings_before_main, (int, float)):
                    heading_issues['headings_before_main']['count'] += headings_before_main
            
            # Check visual hierarchy issues
            if flags.get('hasVisualHierarchyIssues'):
                heading_issues['visual_hierarchy']['pages'].add(page['url'])
                heading_issues['visual_hierarchy']['domains'].add(domain)
                
                # Fix for the potential list issue
                visual_hierarchy_issues = details.get('visualHierarchyIssues', 0)
                if isinstance(visual_hierarchy_issues, list):
                    heading_issues['visual_hierarchy']['count'] += len(visual_hierarchy_issues)
                elif isinstance(visual_hierarchy_issues, (int, float)):
                    heading_issues['visual_hierarchy']['count'] += visual_hierarchy_issues
                
        except Exception as e:
            print(f"Error processing page {page.get('url', 'unknown')}: {str(e)}")
            continue

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in heading_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=5)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue Type"
        headers[1].text = "Number of Occurrences"
        headers[2].text = "Pages Affected"
        headers[3].text = "Sites Affected"
        headers[4].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(data.get('count', len(data['pages'])))
            row[2].text = str(len(data['pages']))
            row[3].text = str(len(data['domains']))
            row[4].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

    ##################################
    # Images
    ##################################
    
    doc.add_paragraph()
    h2 = doc.add_heading('Images', level=2)
    h2.style = doc.styles['Heading 2']
    doc.add_paragraph()

 
    # Query for pages with image issues
    pages_with_image_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.images.images.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.images.images.pageFlags.hasImagesWithoutAlt": True},
                {"results.accessibility.tests.images.images.pageFlags.hasImagesWithInvalidAlt": True},
                {"results.accessibility.tests.images.images.pageFlags.hasSVGWithoutRole": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.images.images": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for different image issues
    image_issues = {
        "missing_alt": {
            "name": "Missing alt text",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "invalid_alt": {
            "name": "Invalid alt text",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "missing_role": {
            "name": "SVGs missing role",
            "pages": set(),
            "domains": set(),
            "count": 0
        }
    }

    # Process each page
    total_images = 0
    total_decorative = 0

    for page in pages_with_image_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        image_data = page['results']['accessibility']['tests']['images']['images']
        flags = image_data['pageFlags']
        details = flags['details']
        
        # Count total and decorative images
        total_images += details.get('totalImages', 0)
        total_decorative += details.get('decorativeImages', 0)
        
        # Check missing alt text
        if flags.get('hasImagesWithoutAlt'):
            image_issues['missing_alt']['pages'].add(page['url'])
            image_issues['missing_alt']['domains'].add(domain)
            image_issues['missing_alt']['count'] += details.get('missingAlt', 0)
        
        # Check invalid alt text
        if flags.get('hasImagesWithInvalidAlt'):
            image_issues['invalid_alt']['pages'].add(page['url'])
            image_issues['invalid_alt']['domains'].add(domain)
            image_issues['invalid_alt']['count'] += details.get('invalidAlt', 0)
        
        # Check missing SVG roles
        if flags.get('hasSVGWithoutRole'):
            image_issues['missing_role']['pages'].add(page['url'])
            image_issues['missing_role']['domains'].add(domain)
            image_issues['missing_role']['count'] += details.get('missingRole', 0)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in image_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=5)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue Type"
        headers[1].text = "Number of Images"
        headers[2].text = "Pages Affected"
        headers[3].text = "Sites Affected"
        headers[4].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(data['count'])
            row[2].text = str(len(data['pages']))
            row[3].text = str(len(data['domains']))
            row[4].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

    else:
        doc.add_paragraph("No image accessibility issues were found.")   

    #######################################
    # Landmarks
    #######################################

    doc.add_paragraph()
    h2 = doc.add_heading('Landmarks', level=2)
    h2.style = doc.styles['Heading 2']

    # Query for pages with landmark issues
    pages_with_landmark_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.landmarks.landmarks.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.landmarks.landmarks.pageFlags.missingRequiredLandmarks": True},
                {"results.accessibility.tests.landmarks.landmarks.pageFlags.hasDuplicateLandmarksWithoutNames": True},
                {"results.accessibility.tests.landmarks.landmarks.pageFlags.hasNestedTopLevelLandmarks": True},
                {"results.accessibility.tests.landmarks.landmarks.pageFlags.hasContentOutsideLandmarks": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.landmarks.landmarks": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for different landmark issues
    landmark_issues = {
        "missing": {
            "name": "Missing required landmarks",
            "pages": set(),
            "domains": set(),
            "details": {
                "banner": 0,
                "main": 0,
                "contentinfo": 0,
                "search": 0
            }
        },
        "duplicate": {
            "name": "Duplicate landmarks without unique names",
            "pages": set(),
            "domains": set(),
            "details": {
                "banner": 0,
                "main": 0,
                "navigation": 0,
                "complementary": 0,
                "contentinfo": 0,
                "search": 0,
                "form": 0,
                "region": 0
            }
        },
        "nested": {
            "name": "Nested top-level landmarks",
            "pages": set(),
            "domains": set()
        },
        "outside": {
            "name": "Content outside landmarks",
            "pages": set(),
            "domains": set(),
            "count": 0
        }
    }

    # Process each page
    total_landmarks = 0
    for page in pages_with_landmark_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        landmark_data = page['results']['accessibility']['tests']['landmarks']['landmarks']
        flags = landmark_data['pageFlags']
        details = flags['details']
        
        # Count total landmarks
        if 'totalLandmarks' in landmark_data.get('details', {}).get('summary', {}):
            total_landmarks += landmark_data['details']['summary']['totalLandmarks']
        
        # Check missing landmarks
        if flags.get('missingRequiredLandmarks'):
            landmark_issues['missing']['pages'].add(page['url'])
            landmark_issues['missing']['domains'].add(domain)
            missing = details.get('missingLandmarks', {})
            for landmark in ['banner', 'main', 'contentinfo', 'search']:
                if missing.get(landmark):
                    landmark_issues['missing']['details'][landmark] += 1

        # Check duplicate landmarks
        if flags.get('hasDuplicateLandmarksWithoutNames'):
            landmark_issues['duplicate']['pages'].add(page['url'])
            landmark_issues['duplicate']['domains'].add(domain)
            duplicates = details.get('duplicateLandmarks', {})
            for landmark in landmark_issues['duplicate']['details'].keys():
                if landmark in duplicates:
                    landmark_issues['duplicate']['details'][landmark] += duplicates[landmark].get('count', 0)

        # Check nested landmarks
        if flags.get('hasNestedTopLevelLandmarks'):
            landmark_issues['nested']['pages'].add(page['url'])
            landmark_issues['nested']['domains'].add(domain)

        # Check content outside landmarks
        if flags.get('hasContentOutsideLandmarks'):
            landmark_issues['outside']['pages'].add(page['url'])
            landmark_issues['outside']['domains'].add(domain)
            landmark_issues['outside']['count'] += details.get('contentOutsideLandmarksCount', 0)

    # Create summary table
    if any(len(issue['pages']) > 0 for issue in landmark_issues.values()):
        # Create main issues summary table
        summary_table = doc.add_table(rows=len(landmark_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue Type"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        row_idx = 1
        for issue_type, data in landmark_issues.items():
            if len(data['pages']) > 0:
                row = summary_table.rows[row_idx].cells
                row[0].text = data['name']
                row[1].text = str(len(data['pages']))
                row[2].text = str(len(data['domains']))
                row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"
                row_idx += 1

        # Format the table text
        format_table_text(summary_table)

        # Add specific details for missing landmarks
        if landmark_issues['missing']['pages']:
            doc.add_paragraph()
            doc.add_paragraph("Missing Required Landmarks Breakdown:", style='Normal')
            
            missing_table = doc.add_table(rows=5, cols=2)
            missing_table.style = 'Table Grid'
            
            headers = missing_table.rows[0].cells
            headers[0].text = "Landmark Type"
            headers[1].text = "Number of Pages Missing"
            
            landmarks = [("Banner", "banner"), ("Main", "main"), 
                        ("Footer", "contentinfo"), ("Search", "search")]
            
            for idx, (name, key) in enumerate(landmarks, 1):
                row = missing_table.rows[idx].cells
                row[0].text = name
                row[1].text = str(landmark_issues['missing']['details'][key])
            
            format_table_text(missing_table)

    else:
        doc.add_paragraph("No landmark structure issues were found.")

    #######################################################
    # Language of page
    #######################################################

    doc.add_paragraph()
    h2 = doc.add_heading('Language of Page', level=2)
    h2.style = doc.styles['Heading 2']    
        
    findings = doc.add_paragraph()

    # If there are pages without lang attribute, list them
    pages_without_lang = list(db_connection.page_results.find(
        {"results.accessibility.tests.html_structure.html_structure.tests.hasValidLang": False},
        {"url": 1, "_id": 0}
    ).sort("url", 1))

    if pages_without_lang:
        doc.add_paragraph(f"{len(pages_without_lang)} pages found without valid language attribute:".strip())

        for page in pages_without_lang:
            doc.add_paragraph(page['url'], style='List Bullet')

    else:
        doc.add_paragraph("""
All pages have a valid lang attribute.
                        """.strip())

 

    #################################
    # Lists
    #################################

    doc.add_paragraph()
    h2 = doc.add_heading('Lists', level=2)
    h2.style = doc.styles['Heading 2']
    
    # Query for pages with list issues
    pages_with_list_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.lists.lists.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.lists.lists.pageFlags.hasEmptyLists": True},
                {"results.accessibility.tests.lists.lists.pageFlags.hasFakeLists": True},
                {"results.accessibility.tests.lists.lists.pageFlags.hasCustomBullets": True},
                {"results.accessibility.tests.lists.lists.pageFlags.hasDeepNesting": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.lists.lists.pageFlags": 1,
            "results.accessibility.tests.lists.lists.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    list_issues = {
        "hasEmptyLists": {"name": "Empty lists", "pages": set(), "domains": set()},
        "hasFakeLists": {"name": "Fake lists (not using proper HTML)", "pages": set(), "domains": set()},
        "hasCustomBullets": {"name": "Custom bullet implementations", "pages": set(), "domains": set()},
        "hasDeepNesting": {"name": "Excessively nested lists", "pages": set(), "domains": set()}
    }

    # Count issues
    for page in pages_with_list_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        flags = page['results']['accessibility']['tests']['lists']['lists']['pageFlags']
        
        for flag in list_issues:
            if flags.get(flag, False):
                list_issues[flag]['pages'].add(page['url'])
                list_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in list_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "List Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

    else:
        doc.add_paragraph(r'''
No list issues found.                          
                          '''.strip())      



    #########################################
    # Maps
    #########################################

    doc.add_paragraph()
    h2 = doc.add_heading('Maps', level=2)
    h2.style = doc.styles['Heading 2']

    # Query for pages with map issues
    pages_with_map_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.maps.maps.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.maps.maps.pageFlags.hasMaps": True},
                {"results.accessibility.tests.maps.maps.pageFlags.hasMapsWithoutTitle": True},
                {"results.accessibility.tests.maps.maps.pageFlags.hasMapsWithAriaHidden": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.maps.maps.pageFlags": 1,
            "results.accessibility.tests.maps.maps.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    map_issues = {
        "hasMaps": {"name": "Pages containing maps", "pages": set(), "domains": set()},
        "hasMapsWithoutTitle": {"name": "Maps without proper titles", "pages": set(), "domains": set()},
        "hasMapsWithAriaHidden": {"name": "Maps hidden from screen readers", "pages": set(), "domains": set()}
    }

    # Count issues
    for page in pages_with_map_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        flags = page['results']['accessibility']['tests']['maps']['maps']['pageFlags']
        
        for flag in map_issues:
            if flags.get(flag, False):
                map_issues[flag]['pages'].add(page['url'])
                map_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in map_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

    else:
        doc.add_paragraph(r'''
No interactive mqps found.                   
        '''.strip())

    #####################################
    # Menus
    #####################################

    doc.add_paragraph()
    h2 = doc.add_heading('Menus', level=2)
    h2.style = doc.styles['Heading 2']

    # Query for pages with menu issues
    pages_with_menu_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.menus.menus.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.menus.menus.pageFlags.hasInvalidMenuRoles": True},
                {"results.accessibility.tests.menus.menus.pageFlags.hasMenusWithoutCurrent": True},
                {"results.accessibility.tests.menus.menus.pageFlags.hasUnnamedMenus": True},
                {"results.accessibility.tests.menus.menus.pageFlags.hasDuplicateMenuNames": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.menus.menus": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    menu_issues = {
        "invalidRoles": {"name": "Invalid menu roles", "pages": set(), "domains": set(), "count": 0},
        "menusWithoutCurrent": {"name": "Missing current page indicators", "pages": set(), "domains": set(), "count": 0},
        "unnamedMenus": {"name": "Unnamed menus", "pages": set(), "domains": set(), "count": 0},
        "duplicateNames": {"name": "Duplicate menu names", "pages": set(), "domains": set(), "count": 0}
    }

    # Count issues
    total_menus = 0
    for page in pages_with_menu_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        menu_data = page['results']['accessibility']['tests']['menus']['menus']
        flags = menu_data['pageFlags']
        details = menu_data['pageFlags']['details']
        
        total_menus += details.get('totalMenus', 0)
        
        # Check each type of issue
        if flags.get('hasInvalidMenuRoles'):
            menu_issues['invalidRoles']['pages'].add(page['url'])
            menu_issues['invalidRoles']['domains'].add(domain)
            menu_issues['invalidRoles']['count'] += details.get('invalidRoles', 0)
            
        if flags.get('hasMenusWithoutCurrent'):
            menu_issues['menusWithoutCurrent']['pages'].add(page['url'])
            menu_issues['menusWithoutCurrent']['domains'].add(domain)
            menu_issues['menusWithoutCurrent']['count'] += details.get('menusWithoutCurrent', 0)
            
        if flags.get('hasUnnamedMenus'):
            menu_issues['unnamedMenus']['pages'].add(page['url'])
            menu_issues['unnamedMenus']['domains'].add(domain)
            menu_issues['unnamedMenus']['count'] += details.get('unnamedMenus', 0)
            
        if flags.get('hasDuplicateMenuNames'):
            menu_issues['duplicateNames']['pages'].add(page['url'])
            menu_issues['duplicateNames']['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in menu_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=5)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Menu Issue"
        headers[1].text = "Number of Occurrences"
        headers[2].text = "Pages Affected"
        headers[3].text = "Sites Affected"
        headers[4].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(data['count']) if flag != 'duplicateNames' else 'N/A'
            row[2].text = str(len(data['pages']))
            row[3].text = str(len(data['domains']))
            row[4].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

    else:
        doc.add_paragraph("No navigation menu accessibility issues were found.")

    #########################################
    # 'More' Controls
    #########################################
    
    doc.add_page_break()
    h2 = doc.add_heading('"More" Controls', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
Generic "Read More" or "Learn More" style links can create barriers for screen reader users who rely on link and button text to understand where a link/button will take them. When they are taken out of context:
""".strip())

    doc.add_paragraph("Users can't determine the link's purpose from the link text alone", style='List Bullet')
    doc.add_paragraph("Screen reader users may get a list of identical 'read more' links", style='List Bullet')
    doc.add_paragraph("The destination of the link isn't clear without surrounding context slowing down reading for screen-reader and screen-magnifier users", style='List Bullet')

    # Add recommendations
    doc.add_paragraph()
    doc.add_paragraph("Recommendations for improving link text:")
    doc.add_paragraph("Make link and button text descriptive of its destination or purpose", style='List Bullet')
    doc.add_paragraph("Use aria-label or visually hidden text if additional context is needed", style='List Bullet')
    doc.add_paragraph("Ensure link text makes sense when read out of context", style='List Bullet')

    doc.add_paragraph()

    # Query for pages with read more link issues
    pages_with_readmore_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.read_more_links.read_more_links.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.read_more_links.read_more_links.pageFlags.hasGenericReadMoreLinks": True},
                {"results.accessibility.tests.read_more_links.read_more_links.pageFlags.hasInvalidReadMoreLinks": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.read_more_links.read_more_links.pageFlags": 1,
            "results.accessibility.tests.read_more_links.read_more_links.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    readmore_issues = {
        "hasGenericReadMoreLinks": {"name": "Generic 'Read More' links", "pages": set(), "domains": set()},
        "hasInvalidReadMoreLinks": {"name": "Invalid implementation of 'Read More' links", "pages": set(), "domains": set()}
    }

    # Count issues
    for page in pages_with_readmore_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        flags = page['results']['accessibility']['tests']['read_more_links']['read_more_links']['pageFlags']
        
        for flag in readmore_issues:
            if flags.get(flag, False):
                readmore_issues[flag]['pages'].add(page['url'])
                readmore_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in readmore_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add domain details for each issue
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)

        # Add examples if available
        for page in pages_with_readmore_issues:
            details = page['results']['accessibility']['tests']['read_more_links']['read_more_links']['details']
            if 'items' in details and details['items']:
                doc.add_paragraph()
                doc.add_paragraph("Examples of problematic link text found:")
                for item in details['items'][:5]:  # Show up to 5 examples
                    doc.add_paragraph(item, style='List Bullet')
                break  # Only show examples from first page with issues

    else:
        doc.add_paragraph("No issues with generic 'Read More' links were found.")

    #########################################
    # Tabindex
    #########################################

    doc.add_paragraph()
    h3 = doc.add_heading('Tabindex', level=2)
    h3.style = doc.styles['Heading 2']

    # Query for pages with tabindex issues
    pages_with_tabindex_issues = list(db_connection.page_results.find(
        {"results.accessibility.tests.tabindex.tabindex.pageFlags": {"$exists": True}},
        {
            "url": 1,
            "results.accessibility.tests.tabindex.tabindex.pageFlags": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    tabindex_issues = {
        "hasPositiveTabindex": {"name": "Elements with positive tabindex", "pages": set(), "domains": set()},
        "hasNonInteractiveZeroTabindex": {"name": "Non-interactive elements with tabindex=0", "pages": set(), "domains": set()},
        "hasMissingRequiredTabindex": {"name": "Interactive elements missing required tabindex", "pages": set(), "domains": set()},
        "hasSvgTabindexWarnings": {"name": "SVG elements with tabindex warnings", "pages": set(), "domains": set()}
    }    


    # Count issues
    for page in pages_with_tabindex_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        flags = page['results']['accessibility']['tests']['tabindex']['tabindex']['pageFlags']
        
        for flag in tabindex_issues:
            if flags.get(flag, False):  # If issue exists (True)
                tabindex_issues[flag]['pages'].add(page['url'])
                tabindex_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in tabindex_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        tiheaders = summary_table.rows[0].cells
        tiheaders[0].text = "Issue"
        tiheaders[1].text = "Pages Affected"
        tiheaders[2].text = "Sites Affected"
        tiheaders[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

    #################################
    # Title Attribute
    #################################

    doc.add_paragraph()
    h3 = doc.add_heading('Title Attribute', level=2)
    h3.style = doc.styles['Heading 2']

    # Query for pages with title attribute issues
    pages_with_title_issues = list(db_connection.page_results.find(
        {"results.accessibility.tests.title.titleAttribute.pageFlags.hasImproperTitleAttributes": True},
        {
            "url": 1,
            "results.accessibility.tests.title.titleAttribute.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Count affected domains
    affected_domains = set()
    total_improper_uses = 0
    domain_counts = {}

    for page in pages_with_title_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        affected_domains.add(domain)
        
        # Count improper uses from the details
        improper_uses = len(page['results']['accessibility']['tests']['title']['titleAttribute']['details']['improperUse'])
        total_improper_uses += improper_uses
        
        # Track counts by domain
        if domain not in domain_counts:
            domain_counts[domain] = 0
        domain_counts[domain] += improper_uses

    # Calculate percentage
    percentage = (len(affected_domains) / len(total_domains)) * 100 if total_domains else 0

    # Create summary table
    summary_table = doc.add_table(rows=2, cols=4)
    summary_table.style = 'Table Grid'

    # Set column headers
    headers = summary_table.rows[0].cells
    headers[0].text = "Issue"
    headers[1].text = "Total Occurrences"
    headers[2].text = "Sites Affected"
    headers[3].text = "% of Total Sites"

    # Add data
    row = summary_table.rows[1].cells
    row[0].text = "Improper use of title attribute"
    row[1].text = str(total_improper_uses)
    row[2].text = str(len(affected_domains))
    row[3].text = f"{percentage:.1f}%"

    # Format the table text
    format_table_text(summary_table)

    # Add some space after the table
    doc.add_paragraph()

    ########################################
    # Tables
    ########################################
            
    doc.add_page_break()
    h2 = doc.add_heading('Tables', level=2)
    h2.style = doc.styles['Heading 2']
 
    # Query for pages with table issues
    pages_with_table_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.tables.tables.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.tables.tables.pageFlags.hasMissingHeaders": True},
                {"results.accessibility.tests.tables.tables.pageFlags.hasNoScope": True},
                {"results.accessibility.tests.tables.tables.pageFlags.hasMissingCaption": True},
                {"results.accessibility.tests.tables.tables.pageFlags.hasLayoutTables": True},
                {"results.accessibility.tests.tables.tables.pageFlags.hasComplexTables": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.tables.tables.pageFlags": 1,
            "results.accessibility.tests.tables.tables.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    table_issues = {
        "hasMissingHeaders": {"name": "Missing table headers", "pages": set(), "domains": set()},
        "hasNoScope": {"name": "Missing scope attributes", "pages": set(), "domains": set()},
        "hasMissingCaption": {"name": "Missing table captions", "pages": set(), "domains": set()},
        "hasLayoutTables": {"name": "Layout tables", "pages": set(), "domains": set()},
        "hasComplexTables": {"name": "Complex tables without proper structure", "pages": set(), "domains": set()}
    }

    # Count issues
    for page in pages_with_table_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        flags = page['results']['accessibility']['tests']['tables']['tables']['pageFlags']
        
        for flag in table_issues:
            if flags.get(flag, False):
                table_issues[flag]['pages'].add(page['url'])
                table_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in table_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Table Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)
    else:
        doc.add_paragraph('No tables found.')    

    #########################################
    # Timers
    #########################################

    doc.add_paragraph()
    h2 = doc.add_heading('Timers', level=2)
    h2.style = doc.styles['Heading 2']

    # Query for pages with timer issues
    pages_with_timer_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.timers.timers.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.timers.timers.pageFlags.hasTimers": True},
                {"results.accessibility.tests.timers.timers.pageFlags.hasAutoStartTimers": True},
                {"results.accessibility.tests.timers.timers.pageFlags.hasTimersWithoutControls": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.timers.timers.pageFlags": 1,
            "results.accessibility.tests.timers.timers.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    timer_issues = {
        "hasTimers": {"name": "Pages with timers", "pages": set(), "domains": set()},
        "hasAutoStartTimers": {"name": "Auto-starting timers", "pages": set(), "domains": set()},
        "hasTimersWithoutControls": {"name": "Timers without adequate controls", "pages": set(), "domains": set()}
    }

    # Count issues
    for page in pages_with_timer_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        flags = page['results']['accessibility']['tests']['timers']['timers']['pageFlags']
        
        for flag in timer_issues:
            if flags.get(flag, False):
                timer_issues[flag]['pages'].add(page['url'])
                timer_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in timer_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Timer Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

    else:
        doc.add_paragraph("No timer-related issues were found.")

    ##################################
    # Videos
    ##################################

    doc.add_page_break()
    h2 = doc.add_heading('Videos', level=2)
    h2.style = doc.styles['Heading 2']
    findings = doc.add_paragraph()

 
    # Query for pages with video issues
    pages_with_video_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.video.video.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.video.video.pageFlags.missingCaptions": True},
                {"results.accessibility.tests.video.video.pageFlags.missingAudioDescription": True},
                {"results.accessibility.tests.video.video.pageFlags.inaccessibleControls": True},
                {"results.accessibility.tests.video.video.pageFlags.missingTranscript": True},
                {"results.accessibility.tests.video.video.pageFlags.hasAutoplay": True},
                {"results.accessibility.tests.video.video.pageFlags.missingLabels": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.video.video.pageFlags": 1,
            "results.accessibility.tests.video.video.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    video_issues = {
        "missingCaptions": {"name": "Missing closed captions", "pages": set(), "domains": set()},
        "missingAudioDescription": {"name": "Missing audio descriptions", "pages": set(), "domains": set()},
        "inaccessibleControls": {"name": "Inaccessible video controls", "pages": set(), "domains": set()},
        "missingTranscript": {"name": "Missing transcripts", "pages": set(), "domains": set()},
        "hasAutoplay": {"name": "Autoplay without user control", "pages": set(), "domains": set()},
        "missingLabels": {"name": "Missing video labels/titles", "pages": set(), "domains": set()}
    }

    # Count issues
    if (len(pages_with_video_issues) > 0):

        for page in pages_with_video_issues:
            domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
            flags = page['results']['accessibility']['tests']['video']['video']['pageFlags']
            
            for flag in video_issues:
                if flags.get(flag, False):
                    video_issues[flag]['pages'].add(page['url'])
                    video_issues[flag]['domains'].add(domain)

        # Create filtered list of issues that have affected pages
        active_issues = {flag: data for flag, data in video_issues.items() 
                        if len(data['pages']) > 0}

        if active_issues:
            # Create summary table
            summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
            summary_table.style = 'Table Grid'

            # Set column headers
            headers = summary_table.rows[0].cells
            headers[0].text = "Video Issue"
            headers[1].text = "Pages Affected"
            headers[2].text = "Sites Affected"
            headers[3].text = "% of Total Sites"

            # Add data
            for i, (flag, data) in enumerate(active_issues.items(), 1):
                row = summary_table.rows[i].cells
                row[0].text = data['name']
                row[1].text = str(len(data['pages']))
                row[2].text = str(len(data['domains']))
                row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

            # Format the table text
            format_table_text(summary_table)

        else:
            doc.add_paragraph("No video accessibility issues were found.")
    else:
        doc.add_paragraph("No videos were found.")


    #############################################
    # Detailed Findings
    # YYYYYYYYYYYY
    #############################################
    h2 = doc.add_heading('Detailed findings', level=1)
    h2.style = doc.styles['Heading 1']

#####################################
    # Helper Functions for Structure Report
    #####################################
    
    def count_descendants(element):
        """Count the total number of descendant elements"""
        count = 0
        if element and 'children' in element:
            count += len(element['children'])
            for child in element['children']:
                count += count_descendants(child)
        return count

    def count_element_type(element, tag_name):
        """Count elements of a specific tag type within a parent element"""
        count = 0
        if element:
            if element.get('tag', '').lower() == tag_name.lower():
                count += 1
            if 'children' in element:
                for child in element['children']:
                    count += count_element_type(child, tag_name)
        return count

    def element_contains_tag(element, tag_name):
        """Check if element contains a specific tag anywhere in its descendants"""
        if element:
            if element.get('tag', '').lower() == tag_name.lower():
                return True
            if 'children' in element:
                for child in element['children']:
                    if element_contains_tag(child, tag_name):
                        return True
        return False
    
    #########################################
    # Page Structure Analysis
    #########################################

    h2 = doc.add_heading('Page Structure Analysis', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
    Understanding the structure of web pages is fundamental to accessibility. This section analyzes the common elements found across pages, such as headers, footers, and navigation components. Consistent structure helps users understand and navigate content efficiently.
    """.strip())

    # Query for structure analysis results
    structure_analysis = list(db_connection.db.structure_analysis.find(
        {},
        {"_id": 0}
    ).sort("timestamp", -1).limit(1))

    if structure_analysis:
        analysis = structure_analysis[0]
        
        # Get the overall summary and domain analyses
        overall_summary = analysis.get('overall_summary', {})
        domain_analyses = analysis.get('domain_analyses', {})
        
        # Calculate total pages from domain analyses if not in overall summary
        total_pages = overall_summary.get('total_pages', 0)
        if total_pages == 0 and domain_analyses:
            total_pages = sum(domain_data.get('page_count', 0) for domain_data in domain_analyses.values())
        
        # Add summary statistics
        doc.add_paragraph()
        doc.add_paragraph("Structure Analysis Summary:", style='Normal')
        
        stats_table = doc.add_table(rows=6, cols=2)  # Expanded to include main and complementary content
        stats_table.style = 'Table Grid'
        
        # Add summary data using overall_summary
        rows = stats_table.rows
        rows[0].cells[0].text = "Pages Analyzed"
        rows[0].cells[1].text = str(total_pages)
        
        rows[1].cells[0].text = "Header Consistency"
        header_score = overall_summary.get('average_header_score', 0) * 100
        rows[1].cells[1].text = f"{header_score:.1f}%"
        
        rows[2].cells[0].text = "Footer Consistency"
        footer_score = overall_summary.get('average_footer_score', 0) * 100
        rows[2].cells[1].text = f"{footer_score:.1f}%"
        
        rows[3].cells[0].text = "Navigation Consistency"
        nav_score = overall_summary.get('average_navigation_score', 0) * 100
        rows[3].cells[1].text = f"{nav_score:.1f}%"
        
        rows[4].cells[0].text = "Main Content Consistency"
        main_score = overall_summary.get('average_main_content_score', 0) * 100
        rows[4].cells[1].text = f"{main_score:.1f}%"
        
        rows[5].cells[0].text = "Complementary Content Consistency"
        comp_score = overall_summary.get('average_complementary_score', 0) * 100
        rows[5].cells[1].text = f"{comp_score:.1f}%"
        
        format_table_text(stats_table)
        
        # Component Presence
        doc.add_paragraph()
        doc.add_paragraph("Common UI Components Presence:", style='Normal')
        
        # Calculate component presence across all domains
        header_pages = 0
        footer_pages = 0
        nav_pages = 0
        main_pages = 0
        complementary_pages = 0
        search_pages = 0
        
        for domain_data in domain_analyses.values():
            header_pages += domain_data.get('header_analysis', {}).get('pages_with_component', 0)
            footer_pages += domain_data.get('footer_analysis', {}).get('pages_with_component', 0)
            nav_pages += domain_data.get('navigation_analysis', {}).get('pages_with_component', 0)
            main_pages += domain_data.get('main_content_analysis', {}).get('pages_with_component', 0)
            complementary_pages += domain_data.get('complementary_analysis', {}).get('pages_with_component', 0)
            
            # Search might be in component_presence or directly in the domain data
            search_count = domain_data.get('component_presence', {}).get('search', 0)
            if search_count == 0:  # Try alternate location
                search_count = domain_data.get('search_components', 0)
            search_pages += search_count
        
        component_table = doc.add_table(rows=7, cols=3)  # Expanded to include main and complementary content
        component_table.style = 'Table Grid'
        
        # Headers
        component_headers = component_table.rows[0].cells
        component_headers[0].text = "Component"
        component_headers[1].text = "Pages"
        component_headers[2].text = "% of Total"
        
        # Component data
        components = [
            ("Header", header_pages),
            ("Footer", footer_pages),
            ("Navigation", nav_pages),
            ("Main Content", main_pages),
            ("Complementary Content", complementary_pages),
            ("Search", search_pages)
        ]
        
        for i, (component, count) in enumerate(components, 1):
            row = component_table.rows[i].cells
            row[0].text = component
            row[1].text = str(count)
            percentage = (count / total_pages) * 100 if total_pages > 0 else 0
            row[2].text = f"{percentage:.1f}%"
        
        format_table_text(component_table)
        
        # Header Analysis
        doc.add_paragraph()
        h3 = doc.add_heading('Header Analysis', level=3)
        h3.style = doc.styles['Heading 3']
        
        # Find a domain with header data to use as an example
        example_domain = None
        for domain, data in domain_analyses.items():
            if data.get('header_analysis', {}).get('pages_with_component', 0) > 0:
                example_domain = domain
                break
        
        if example_domain:
            header_analysis = domain_analyses[example_domain]['header_analysis']
            
            # Common patterns
            patterns = header_analysis.get('common_patterns', {})
            
            doc.add_paragraph(f"Most common header tag: <{patterns.get('tag', 'unknown')}>")
            
            # Header structure details
            doc.add_paragraph("Header structure details:")
            
            # If we have sample data, use it to provide more information
            sample_pages = domain_analyses[example_domain].get('sample_pages', {})
            if sample_pages:
                sample_url = next(iter(sample_pages.keys()))
                sample_data = sample_pages[sample_url]
                
                header_element = sample_data.get('keyElements', {}).get('header', {})
                if header_element:
                    # Count total descendants (not just direct children)
                    descendants = count_descendants(header_element)
                    doc.add_paragraph(f"Average header complexity: {descendants} total elements", style='List Bullet')
                    
                    # Look for specific elements within header
                    links = count_element_type(header_element, 'a')
                    buttons = count_element_type(header_element, 'button')
                    images = count_element_type(header_element, 'img')
                    
                    doc.add_paragraph(f"Typical header contains: {links} links, {buttons} buttons, {images} images", style='List Bullet')
                    
                    # Determine if header likely contains site navigation
                    has_nav = element_contains_tag(header_element, 'nav')
                    doc.add_paragraph(f"Header contains navigation menu: {'Yes' if has_nav else 'No'}", style='List Bullet')
            
            if patterns.get('common_classes'):
                doc.add_paragraph("Common header CSS classes:")
                for cls in patterns.get('common_classes', []):
                    doc.add_paragraph(cls, style='List Bullet')
        else:
            doc.add_paragraph("No consistent header structure was identified across pages.")
        
        # Footer Analysis
        doc.add_paragraph()
        h3 = doc.add_heading('Footer Analysis', level=3)
        h3.style = doc.styles['Heading 3']
        
        # Find a domain with footer data to use as an example
        example_domain = None
        for domain, data in domain_analyses.items():
            if data.get('footer_analysis', {}).get('pages_with_component', 0) > 0:
                example_domain = domain
                break
        
        if example_domain:
            footer_analysis = domain_analyses[example_domain]['footer_analysis']
            
            # Common patterns
            patterns = footer_analysis.get('common_patterns', {})
            
            doc.add_paragraph(f"Most common footer tag: <{patterns.get('tag', 'unknown')}>")
            
            # Footer structure details
            doc.add_paragraph("Footer structure details:")
            
            # If we have sample data, use it to provide more information
            sample_pages = domain_analyses[example_domain].get('sample_pages', {})
            if sample_pages:
                sample_url = next(iter(sample_pages.keys()))
                sample_data = sample_pages[sample_url]
                
                footer_element = sample_data.get('keyElements', {}).get('footer', {})
                if footer_element:
                    # Count total descendants (not just direct children)
                    descendants = count_descendants(footer_element)
                    doc.add_paragraph(f"Average footer complexity: {descendants} total elements", style='List Bullet')
                    
                    # Look for specific elements within footer
                    links = count_element_type(footer_element, 'a')
                    buttons = count_element_type(footer_element, 'button')
                    images = count_element_type(footer_element, 'img')
                    
                    doc.add_paragraph(f"Typical footer contains: {links} links, {buttons} buttons, {images} images", style='List Bullet')
            
            if patterns.get('common_classes'):
                doc.add_paragraph("Common footer CSS classes:")
                for cls in patterns.get('common_classes', []):
                    doc.add_paragraph(cls, style='List Bullet')
        else:
            doc.add_paragraph("No consistent footer structure was identified across pages.")
        
        # Navigation Analysis
        doc.add_paragraph()
        h3 = doc.add_heading('Navigation Analysis', level=3)
        h3.style = doc.styles['Heading 3']
        
        # Find a domain with navigation data to use as an example
        example_domain = None
        for domain, data in domain_analyses.items():
            if data.get('navigation_analysis', {}).get('pages_with_component', 0) > 0:
                example_domain = domain
                break
        
        if example_domain:
            navigation_analysis = domain_analyses[example_domain]['navigation_analysis']
            
            # Common patterns
            patterns = navigation_analysis.get('common_patterns', {})
            
            doc.add_paragraph(f"Most common navigation tag: <{patterns.get('tag', 'unknown')}>")
            
            # Navigation structure details
            doc.add_paragraph("Navigation structure details:")
            
            # If we have sample data, use it to provide more information
            sample_pages = domain_analyses[example_domain].get('sample_pages', {})
            if sample_pages:
                sample_url = next(iter(sample_pages.keys()))
                sample_data = sample_pages[sample_url]
                
                nav_element = sample_data.get('keyElements', {}).get('navigation', {})
                if nav_element:
                    # Count total links
                    links = count_element_type(nav_element, 'a')
                    doc.add_paragraph(f"Average navigation contains: {links} links", style='List Bullet')
            
            if patterns.get('common_classes'):
                doc.add_paragraph("Common navigation CSS classes:")
                for cls in patterns.get('common_classes', []):
                    doc.add_paragraph(cls, style='List Bullet')
        else:
            doc.add_paragraph("No consistent navigation structure was identified across pages.")
            
        # Main Content Analysis
        doc.add_paragraph()
        h3 = doc.add_heading('Main Content Analysis', level=3)
        h3.style = doc.styles['Heading 3']
        
        # Find a domain with main content data to use as an example
        example_domain = None
        for domain, data in domain_analyses.items():
            if data.get('main_content_analysis', {}).get('pages_with_component', 0) > 0:
                example_domain = domain
                break
        
        if example_domain:
            main_content_analysis = domain_analyses[example_domain]['main_content_analysis']
            
            # Common patterns
            patterns = main_content_analysis.get('common_patterns', {})
            
            doc.add_paragraph(f"Most common main content tag: <{patterns.get('tag', 'unknown')}>")
            
            # Main content structure details
            doc.add_paragraph("Main content structure details:")
            
            # If we have sample data, use it to provide more information
            sample_pages = domain_analyses[example_domain].get('sample_pages', {})
            if sample_pages:
                sample_url = next(iter(sample_pages.keys()))
                sample_data = sample_pages[sample_url]
                
                main_element = sample_data.get('keyElements', {}).get('mainContent', {})
                if main_element:
                    # Check for important content elements
                    has_headings = element_contains_tag(main_element, 'h1') or element_contains_tag(main_element, 'h2')
                    doc.add_paragraph(f"Main content contains headings: {'Yes' if has_headings else 'No'}", style='List Bullet')
            
            if patterns.get('common_classes'):
                doc.add_paragraph("Common main content CSS classes:")
                for cls in patterns.get('common_classes', []):
                    doc.add_paragraph(cls, style='List Bullet')
        else:
            doc.add_paragraph("No consistent main content structure was identified across pages.")
        
        # Complementary Content Analysis
        doc.add_paragraph()
        h3 = doc.add_heading('Complementary Content Analysis', level=3)
        h3.style = doc.styles['Heading 3']
        
        # Find a domain with complementary content data to use as an example
        example_domain = None
        for domain, data in domain_analyses.items():
            if data.get('complementary_analysis', {}).get('pages_with_component', 0) > 0:
                example_domain = domain
                break
        
        if example_domain:
            complementary_analysis = domain_analyses[example_domain]['complementary_analysis']
            
            # Common patterns
            patterns = complementary_analysis.get('common_patterns', {})
            
            doc.add_paragraph(f"Most common complementary content tag: <{patterns.get('tag', 'unknown')}>")
            
            # Complementary content structure details
            doc.add_paragraph("Complementary content structure details:")
            
            # If we have sample data, provide more information
            sample_pages = domain_analyses[example_domain].get('sample_pages', {})
            if sample_pages:
                sample_url = next(iter(sample_pages.keys()))
                sample_data = sample_pages[sample_url]
                
                comp_element = sample_data.get('keyElements', {}).get('complementaryContent', {})
                if comp_element:
                    # Analyze content
                    has_links = element_contains_tag(comp_element, 'a')
                    doc.add_paragraph(f"Complementary content contains links: {'Yes' if has_links else 'No'}", style='List Bullet')
            
            if patterns.get('common_classes'):
                doc.add_paragraph("Common complementary content CSS classes:")
                for cls in patterns.get('common_classes', []):
                    doc.add_paragraph(cls, style='List Bullet')
        else:
            doc.add_paragraph("No consistent complementary content structure was identified across pages.")
        
        # Forms Analysis
        doc.add_paragraph()
        h3 = doc.add_heading('Forms Analysis', level=3)
        h3.style = doc.styles['Heading 3']

        # Get forms analysis from the overall summary
        forms_analysis = overall_summary.get('forms_analysis', {})
        unique_forms = forms_analysis.get('unique_forms', {})
        forms_by_type = forms_analysis.get('forms_by_type', {})
        forms_by_domain = forms_analysis.get('forms_by_domain', {})

        if unique_forms:
            # Summary table of form types
            doc.add_paragraph(f"Found {len(unique_forms)} unique forms across all analyzed pages.")
            
            form_types_table = doc.add_table(rows=len(forms_by_type) + 1, cols=3)
            form_types_table.style = 'Table Grid'
            
            # Set headers
            headers = form_types_table.rows[0].cells
            headers[0].text = "Form Type"
            headers[1].text = "Count"
            headers[2].text = "% of All Forms"
            
            # Add data rows
            i = 1
            for form_type, forms in sorted(forms_by_type.items(), key=lambda x: len(x[1]), reverse=True):
                row = form_types_table.rows[i].cells
                row[0].text = form_type.title()
                row[1].text = str(len(forms))
                percentage = (len(forms) / len(unique_forms)) * 100
                row[2].text = f"{percentage:.1f}%"
                i += 1
            
            format_table_text(form_types_table)
            
            # Detailed form locations by domain
            doc.add_paragraph()
            doc.add_paragraph("Forms by Site and Page:", style='Normal')
            
            # For each domain
            for domain, forms in sorted(forms_by_domain.items()):
                doc.add_paragraph()
                doc.add_paragraph(f"Domain: {domain}", style='Heading 4')
                
                # Group forms by page within this domain
                forms_by_page = {}
                for form in forms:
                    page_url = form['page_url']
                    if page_url not in forms_by_page:
                        forms_by_page[page_url] = []
                    forms_by_page[page_url].append(form)
                
                # Create a table for this domain's forms
                domain_forms_table = doc.add_table(rows=len(forms_by_page) + 1, cols=4)
                domain_forms_table.style = 'Table Grid'
                
                # Set headers
                headers = domain_forms_table.rows[0].cells
                headers[0].text = "Page URL"
                headers[1].text = "Form Types"
                headers[2].text = "Count"
                headers[3].text = "Locations"
                
                # Add data rows
                i = 1
                for page_url, page_forms in sorted(forms_by_page.items()):
                    row = domain_forms_table.rows[i].cells
                    
                    # Shorten URL for display
                    short_url = page_url.replace(f"https://{domain}", "")
                    if not short_url:
                        short_url = "/"  # Homepage
                    
                    row[0].text = short_url
                    
                    # Get unique form types for this page
                    form_types = set(form['form_type'] for form in page_forms)
                    row[1].text = ", ".join(form_type.title() for form_type in form_types)
                    
                    # Count of forms
                    row[2].text = str(len(page_forms))
                    
                    # Locations of forms
                    locations = set(form['location'] for form in page_forms)
                    row[3].text = ", ".join(location.title() for location in locations)
                    
                    i += 1
                
                format_table_text(domain_forms_table)
        else:
            doc.add_paragraph("No forms were detected across the analyzed pages.")
        
        # Accessibility implications
        doc.add_paragraph()
        h3 = doc.add_heading('Accessibility Implications', level=3)
        h3.style = doc.styles['Heading 3']
        
        # Calculate overall structure score - now includes main content and complementary content
        header_score = overall_summary.get('average_header_score', 0)
        footer_score = overall_summary.get('average_footer_score', 0)
        nav_score = overall_summary.get('average_navigation_score', 0)
        main_score = overall_summary.get('average_main_content_score', 0)
        comp_score = overall_summary.get('average_complementary_score', 0)
        
        # Count non-zero scores to avoid division by zero
        score_count = sum(1 for score in [header_score, footer_score, nav_score, main_score, comp_score] if score > 0)
        
        # Calculate the average of non-zero scores
        overall_score = sum([header_score, footer_score, nav_score, main_score, comp_score]) / score_count * 100 if score_count > 0 else 0
        
        if overall_score >= 80:
            doc.add_paragraph("""
    The site demonstrates high structural consistency, which benefits users by providing a predictable and uniform experience across pages. This consistency is particularly helpful for:
            """.strip())
            doc.add_paragraph("Screen reader users who rely on consistent navigation patterns", style='List Bullet')
            doc.add_paragraph("Keyboard-only users navigating through consistent tab orders", style='List Bullet')
            doc.add_paragraph("Users with cognitive disabilities who benefit from predictable layouts", style='List Bullet')
            doc.add_paragraph("All users through reduced cognitive load and improved usability", style='List Bullet')
            doc.add_paragraph("Clear distinction between main and complementary content", style='List Bullet')
            doc.add_paragraph("Consistent placement of recurring elements like forms and popups", style='List Bullet')
        elif overall_score >= 50:
            doc.add_paragraph("""
    The site shows moderate structural consistency. While many elements remain consistent across pages, there are some variations that could impact user experience. Consider:
            """.strip())
            doc.add_paragraph("Standardizing header and footer components across all pages", style='List Bullet')
            doc.add_paragraph("Ensuring navigation remains consistent in both desktop and mobile views", style='List Bullet')
            doc.add_paragraph("Maintaining consistent placement of search and other utility functions", style='List Bullet')
            doc.add_paragraph("Improving distinction between main content and complementary content", style='List Bullet')
            doc.add_paragraph("Ensuring consistent behavior of recurring elements across the site", style='List Bullet')
        else:
            doc.add_paragraph("""
    The site demonstrates low structural consistency, which may create barriers for users. Inconsistent structure can cause:
            """.strip())
            doc.add_paragraph("Disorientation for screen reader users", style='List Bullet')
            doc.add_paragraph("Navigation challenges for keyboard users", style='List Bullet')
            doc.add_paragraph("Increased cognitive load for all users, particularly those with cognitive disabilities", style='List Bullet')
            doc.add_paragraph("Difficulty finding consistent functionality across pages", style='List Bullet')
            doc.add_paragraph("Inconsistent main content identification making page navigation difficult", style='List Bullet')
            doc.add_paragraph("Unpredictable placement of recurring elements creating confusion", style='List Bullet')
            
            doc.add_paragraph("""
    Recommendation: Implement a consistent template system with standardized header, footer, and navigation components across all pages.
            """.strip())

    else:
        doc.add_paragraph("""
    No structure analysis data was found. Please ensure the page structure analysis test is properly integrated and the analysis has been run after testing.
        """.strip())

    # Add a more detailed analysis section
    doc.add_page_break()
    h2 = doc.add_heading('Detailed Structure Analysis', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
    This section provides a detailed breakdown of structural elements found across the analyzed websites. Understanding these patterns helps identify opportunities for improving consistency and accessibility within and across sites.
    """.strip())

    # Query for structure analysis results
    structure_analysis = list(db_connection.db.structure_analysis.find(
        {},
        {"_id": 0}
    ).sort("timestamp", -1).limit(1))

    if structure_analysis and structure_analysis[0]:
        analysis = structure_analysis[0]
        
        # Get overall summary
        summary = analysis.get('overall_summary', {})
        domain_analyses = analysis.get('domain_analyses', {})
        
        # Overall cross-site consistency
        doc.add_paragraph()
        doc.add_paragraph("Cross-Site Structure Consistency:", style='Normal')
        
        cross_site_table = doc.add_table(rows=7, cols=2)  # Increased from 5 to 7 rows
        cross_site_table.style = 'Table Grid'
        
        # Add summary data
        rows = cross_site_table.rows
        rows[0].cells[0].text = "Total Domains Analyzed"
        rows[0].cells[1].text = str(summary.get('total_domains', 0))
        
        rows[1].cells[0].text = "Average Header Consistency"
        header_score = summary.get('average_header_score', 0) * 100
        rows[1].cells[1].text = f"{header_score:.1f}%"
        
        rows[2].cells[0].text = "Average Footer Consistency"
        footer_score = summary.get('average_footer_score', 0) * 100
        rows[2].cells[1].text = f"{footer_score:.1f}%"
        
        rows[3].cells[0].text = "Average Navigation Consistency"
        nav_score = summary.get('average_navigation_score', 0) * 100
        rows[3].cells[1].text = f"{nav_score:.1f}%"
        
        rows[4].cells[0].text = "Average Main Content Consistency"
        main_score = summary.get('average_main_content_score', 0) * 100
        rows[4].cells[1].text = f"{main_score:.1f}%"
        
        rows[5].cells[0].text = "Average Complementary Content Consistency"
        comp_score = summary.get('average_complementary_score', 0) * 100
        rows[5].cells[1].text = f"{comp_score:.1f}%"
        
        rows[6].cells[0].text = "Overall Structure Consistency"
        overall_score = summary.get('average_consistency_score', 0) * 100
        rows[6].cells[1].text = f"{overall_score:.1f}%"
        
        format_table_text(cross_site_table)
        
        # Display domain-by-domain overview
        doc.add_paragraph()
        doc.add_paragraph("Structure Consistency by Site:", style='Normal')
        
        if domain_analyses:
            domains_table = doc.add_table(rows=len(domain_analyses) + 1, cols=7)  # Increased columns
            domains_table.style = 'Table Grid'
            
            # Set column headers
            headers = domains_table.rows[0].cells
            headers[0].text = "Domain"
            headers[1].text = "Pages"
            headers[2].text = "Header Score"
            headers[3].text = "Footer Score"
            headers[4].text = "Navigation Score"
            headers[5].text = "Main Content Score"
            headers[6].text = "Complementary Score"
            
            # Add domain data
            i = 1
            for domain, domain_analysis in sorted(domain_analyses.items()):
                row = domains_table.rows[i].cells
                row[0].text = domain
                row[1].text = str(domain_analysis.get('page_count', 0))
                
                # Get consistency scores
                header_score = domain_analysis.get('header_analysis', {}).get('consistency_score', 0) * 100
                footer_score = domain_analysis.get('footer_analysis', {}).get('consistency_score', 0) * 100
                nav_score = domain_analysis.get('navigation_analysis', {}).get('consistency_score', 0) * 100
                main_score = domain_analysis.get('main_content_analysis', {}).get('consistency_score', 0) * 100
                comp_score = domain_analysis.get('complementary_analysis', {}).get('consistency_score', 0) * 100
                
                row[2].text = f"{header_score:.1f}%"
                row[3].text = f"{footer_score:.1f}%"
                row[4].text = f"{nav_score:.1f}%"
                row[5].text = f"{main_score:.1f}%"
                row[6].text = f"{comp_score:.1f}%"
                i += 1
            
            format_table_text(domains_table)
        else:
            doc.add_paragraph("No domain-specific analysis available.")
        
        # Detailed analysis of each domain (limited to 3 domains for brevity)
        doc.add_paragraph()
        doc.add_paragraph("Detailed Domain Structure Analysis:", style='Normal')
        
        # Process up to 3 domains
        domains_to_show = list(domain_analyses.keys())[:3]
        
        for domain in domains_to_show:
            domain_analysis = domain_analyses[domain]
            
            doc.add_paragraph()
            doc.add_paragraph(f"Domain: {domain}", style='Heading 4')
            doc.add_paragraph(f"Pages analyzed: {domain_analysis.get('page_count', 0)}", style='Normal')
            doc.add_paragraph(f"Analysis method: {domain_analysis.get('analysis_method', 'unknown')}", style='Normal')
            
            # Header analysis for this domain
            header_analysis = domain_analysis.get('header_analysis', {})
            if header_analysis and header_analysis.get('pages_with_component', 0) > 0:
                doc.add_paragraph()
                doc.add_paragraph("Header Structure:", style='List Bullet')
                header_patterns = header_analysis.get('common_patterns', {})
                
                doc.add_paragraph(f"Tag: <{header_patterns.get('tag', 'unknown')}>", style='List Bullet 2')
                doc.add_paragraph(f"Presence: {header_analysis.get('presence_ratio', 0) * 100:.1f}% of pages", style='List Bullet 2')
                doc.add_paragraph(f"Consistency: {header_analysis.get('consistency_score', 0) * 100:.1f}%", style='List Bullet 2')
                
                if header_patterns.get('common_classes'):
                    doc.add_paragraph("Common CSS classes:", style='List Bullet 2')
                    for cls in header_patterns.get('common_classes', []):
                        doc.add_paragraph(cls, style='List Bullet 3')
            
            # Footer analysis for this domain
            footer_analysis = domain_analysis.get('footer_analysis', {})
            if footer_analysis and footer_analysis.get('pages_with_component', 0) > 0:
                doc.add_paragraph()
                doc.add_paragraph("Footer Structure:", style='List Bullet')
                footer_patterns = footer_analysis.get('common_patterns', {})
                
                doc.add_paragraph(f"Tag: <{footer_patterns.get('tag', 'unknown')}>", style='List Bullet 2')
                doc.add_paragraph(f"Presence: {footer_analysis.get('presence_ratio', 0) * 100:.1f}% of pages", style='List Bullet 2')
                doc.add_paragraph(f"Consistency: {footer_analysis.get('consistency_score', 0) * 100:.1f}%", style='List Bullet 2')
                
                if footer_patterns.get('common_classes'):
                    doc.add_paragraph("Common CSS classes:", style='List Bullet 2')
                    for cls in footer_patterns.get('common_classes', []):
                        doc.add_paragraph(cls, style='List Bullet 3')
            
            # Navigation analysis for this domain
            nav_analysis = domain_analysis.get('navigation_analysis', {})
            if nav_analysis and nav_analysis.get('pages_with_component', 0) > 0:
                doc.add_paragraph()
                doc.add_paragraph("Navigation Structure:", style='List Bullet')
                nav_patterns = nav_analysis.get('common_patterns', {})
                
                doc.add_paragraph(f"Tag: <{nav_patterns.get('tag', 'unknown')}>", style='List Bullet 2')
                doc.add_paragraph(f"Presence: {nav_analysis.get('presence_ratio', 0) * 100:.1f}% of pages", style='List Bullet 2')
                doc.add_paragraph(f"Consistency: {nav_analysis.get('consistency_score', 0) * 100:.1f}%", style='List Bullet 2')
                
                if nav_patterns.get('common_classes'):
                    doc.add_paragraph("Common CSS classes:", style='List Bullet 2')
                    for cls in nav_patterns.get('common_classes', []):
                        doc.add_paragraph(cls, style='List Bullet 3')
            
            # Main content analysis for this domain
            main_content_analysis = domain_analysis.get('main_content_analysis', {})
            if main_content_analysis and main_content_analysis.get('pages_with_component', 0) > 0:
                doc.add_paragraph()
                doc.add_paragraph("Main Content Structure:", style='List Bullet')
                main_patterns = main_content_analysis.get('common_patterns', {})
                
                doc.add_paragraph(f"Tag: <{main_patterns.get('tag', 'unknown')}>", style='List Bullet 2')
                doc.add_paragraph(f"Presence: {main_content_analysis.get('presence_ratio', 0) * 100:.1f}% of pages", style='List Bullet 2')
                doc.add_paragraph(f"Consistency: {main_content_analysis.get('consistency_score', 0) * 100:.1f}%", style='List Bullet 2')
                
                if main_patterns.get('common_classes'):
                    doc.add_paragraph("Common CSS classes:", style='List Bullet 2')
                    for cls in main_patterns.get('common_classes', []):
                        doc.add_paragraph(cls, style='List Bullet 3')
            
            # Complementary content analysis for this domain
            complementary_analysis = domain_analysis.get('complementary_analysis', {})
            if complementary_analysis and complementary_analysis.get('pages_with_component', 0) > 0:
                doc.add_paragraph()
                doc.add_paragraph("Complementary Content Structure:", style='List Bullet')
                comp_patterns = complementary_analysis.get('common_patterns', {})
                
                doc.add_paragraph(f"Tag: <{comp_patterns.get('tag', 'unknown')}>", style='List Bullet 2')
                doc.add_paragraph(f"Presence: {complementary_analysis.get('presence_ratio', 0) * 100:.1f}% of pages", style='List Bullet 2')
                doc.add_paragraph(f"Consistency: {complementary_analysis.get('consistency_score', 0) * 100:.1f}%", style='List Bullet 2')
                
                if comp_patterns.get('common_classes'):
                    doc.add_paragraph("Common CSS classes:", style='List Bullet 2')
                    for cls in comp_patterns.get('common_classes', []):
                        doc.add_paragraph(cls, style='List Bullet 3')
            
            # Forms analysis for this domain
            forms_analysis = domain_analysis.get('forms_analysis', {})
            if forms_analysis and forms_analysis.get('total_forms', 0) > 0:
                doc.add_paragraph()
                doc.add_paragraph("Forms Analysis:", style='List Bullet')
                
                total_forms = forms_analysis.get('total_forms', 0)
                unique_forms = len(forms_analysis.get('unique_forms', {}))
                form_types = forms_analysis.get('form_types', {})
                form_locations = forms_analysis.get('form_locations', {})
                
                doc.add_paragraph(f"Total forms: {total_forms} ({unique_forms} unique)", style='List Bullet 2')
                
                # Form types
                if form_types:
                    doc.add_paragraph("Form types:", style='List Bullet 2')
                    for form_type, count in sorted(form_types.items(), key=lambda x: x[1], reverse=True):
                        doc.add_paragraph(f"{form_type.title()}: {count}", style='List Bullet 3')
                
                # Form locations
                if form_locations:
                    doc.add_paragraph("Form locations:", style='List Bullet 2')
                    for location, count in sorted(form_locations.items(), key=lambda x: x[1], reverse=True):
                        doc.add_paragraph(f"{location.title()}: {count}", style='List Bullet 3')
            
            # Recurring elements analysis for this domain
            recurring_elements = domain_analysis.get('recurring_elements', {})
            if recurring_elements and any(count > 0 for count in recurring_elements.values()):
                doc.add_paragraph()
                doc.add_paragraph("Recurring Elements:", style='List Bullet')
                
                recurring_table = doc.add_table(rows=len(recurring_elements) + 1, cols=3)
                recurring_table.style = 'Table Grid'
                
                # Set column headers
                headers = recurring_table.rows[0].cells
                headers[0].text = "Element Type"
                headers[1].text = "Pages with Element"
                headers[2].text = "% of Pages"
                
                # Add data
                i = 1
                for element_type, count in sorted(recurring_elements.items()):
                    row = recurring_table.rows[i].cells
                    row[0].text = element_type.replace('_', ' ').title()
                    row[1].text = str(count)
                    percentage = (count / domain_analysis.get('page_count', 1)) * 100
                    row[2].text = f"{percentage:.1f}%"
                    i += 1
                
                format_table_text(recurring_table)
            
            # Sample page example if available
            sample_pages = domain_analysis.get('sample_pages', {})
            if sample_pages:
                doc.add_paragraph()
                doc.add_paragraph("Sample Page Details:", style='List Bullet')
                
                # Get the first sample page
                sample_url = next(iter(sample_pages.keys()))
                sample_data = sample_pages[sample_url]
                
                doc.add_paragraph(f"URL: {sample_url}", style='List Bullet 2')
                
                # Show component presence
                flags = sample_data.get('pageFlags', {})
                if flags:
                    doc.add_paragraph("Component Presence:", style='List Bullet 2')
                    doc.add_paragraph(f"Header: {'Yes' if flags.get('hasHeader', False) else 'No'}", style='List Bullet 3')
                    doc.add_paragraph(f"Footer: {'Yes' if flags.get('hasFooter', False) else 'No'}", style='List Bullet 3')
                    doc.add_paragraph(f"Navigation: {'Yes' if flags.get('hasMainNavigation', False) else 'No'}", style='List Bullet 3')
                    doc.add_paragraph(f"Main Content: {'Yes' if flags.get('hasMainContent', False) else 'No'}", style='List Bullet 3')
                    doc.add_paragraph(f"Complementary Content: {'Yes' if flags.get('hasComplementaryContent', False) else 'No'}", style='List Bullet 3')
                    doc.add_paragraph(f"Search: {'Yes' if flags.get('hasSearchComponent', False) else 'No'}", style='List Bullet 3')
                    doc.add_paragraph(f"Forms: {'Yes' if flags.get('hasForms', False) else 'No'}", style='List Bullet 3')
        
        # Add improvement recommendations based on cross-site analysis
        doc.add_paragraph()
        h3 = doc.add_heading('Cross-Site Structure Recommendations', level=3)
        h3.style = doc.styles['Heading 3']
        
        if overall_score >= 80:
            doc.add_paragraph("""
    The sites demonstrate strong structural consistency both within and across domains. To maintain this high standard:
            """.strip())
            doc.add_paragraph("Continue using consistent templates across all sites and sections", style='List Bullet')
            doc.add_paragraph("Document the structure requirements in a shared style guide", style='List Bullet')
            doc.add_paragraph("Consider implementing a shared component library for headers, footers, and navigation", style='List Bullet')
            doc.add_paragraph("Ensure main content is consistently identified with proper semantic markup", style='List Bullet')
            doc.add_paragraph("Maintain clear separation between main and complementary content", style='List Bullet')
            doc.add_paragraph("Conduct regular audits to ensure consistency is maintained as sites evolve", style='List Bullet')
        elif overall_score >= 50:
            doc.add_paragraph("""
    The sites show moderate structural consistency. While internal site consistency is decent, there are cross-site variations that could impact user experience. Consider:
            """.strip())
            doc.add_paragraph("Standardizing header and footer components across all domains", style='List Bullet')
            doc.add_paragraph("Creating shared navigation patterns for a more unified experience", style='List Bullet')
            doc.add_paragraph("Implementing consistent landmark roles and accessible names across all sites", style='List Bullet')
            doc.add_paragraph("Ensuring main content is clearly identified with proper semantic markup", style='List Bullet')
            doc.add_paragraph("Establishing clear patterns for complementary content placement", style='List Bullet')
            doc.add_paragraph("Developing a cross-site style guide and component library", style='List Bullet')
        else:
            doc.add_paragraph("""
    The sites lack structural consistency both internally and across domains, which creates accessibility barriers. Priority recommendations:
            """.strip())
            doc.add_paragraph("Implement a unified structure and template system across all sites", style='List Bullet')
            doc.add_paragraph("Start with standardizing the most important components: headers, footers, and main navigation", style='List Bullet')
            doc.add_paragraph("Ensure all pages use proper semantic markup for main content with role='main'", style='List Bullet')
            doc.add_paragraph("Implement consistent patterns for complementary content with role='complementary'", style='List Bullet')
            doc.add_paragraph("Create a shared component library with accessibility built-in", style='List Bullet')
            doc.add_paragraph("Develop a comprehensive style guide with structural requirements", style='List Bullet')
            doc.add_paragraph("Conduct user testing with screen reader and keyboard-only users to validate improvements", style='List Bullet')

    else:
        doc.add_paragraph("""
    No structure analysis data was found. Please ensure the page structure analysis test is properly integrated and the analysis has been run after testing.
        """.strip())

    #############################
    # Accessible names
    #############################
        
    doc.add_page_break()
    h2 = doc.add_heading('Accessible Names', level=2)
    h2.style = doc.styles['Heading 2']

    # Original query to get pages with missing accessible names (for total counts)
    pages_with_name_issues = list(db_connection.page_results.find(
        {"results.accessibility.tests.accessible_names.accessible_names.details.summary.missingNames": {"$gt": 0}},
        {
            "url": 1,
            "results.accessibility.tests.accessible_names.accessible_names.details.summary.missingNames": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Count affected domains (needed for overall statistics)
    affected_domains = set()
    total_missing_names = 0
    for page in pages_with_name_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        affected_domains.add(domain)
        total_missing_names += page['results']['accessibility']['tests']['accessible_names']['accessible_names']['details']['summary']['missingNames']

    # Query for pages with violations for tag-specific analysis
    pages_with_violations = list(db_connection.page_results.find(
        {"results.accessibility.tests.accessible_names.accessible_names.details.violations": {"$exists": True}},
        {
            "url": 1,
            "results.accessibility.tests.accessible_names.accessible_names.details.violations": 1,
            "_id": 0
        }
    ))

    # Process violations to count by tag
    tag_statistics = {}
    
    for page in pages_with_violations:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        
        # Get the violations array and parse it if it's a string
        violations = page['results']['accessibility']['tests']['accessible_names']['accessible_names']['details']['violations']
        if isinstance(violations, str):
            violations = json.loads(violations)
        
        # Track unique tags for this page
        page_tags = set()
        
        for violation in violations:
            tag = violation['element']
            
            if tag not in tag_statistics:
                tag_statistics[tag] = {
                    'count': 0,
                    'pages': set(),
                    'domains': set()
                }
            
            tag_statistics[tag]['count'] += 1
            tag_statistics[tag]['pages'].add(page['url'])
            tag_statistics[tag]['domains'].add(domain)

    # Add explanation
    doc.add_paragraph("""
Interactive elements such as links, buttons, form fields etc. must have an accessible name that can be programmatically determined. This name is what will be announced by screen readers and other assistive technologies when the user encounters the element. Without an accessible name, users will not know the purpose or function of the element.
""".strip())

    # Create results table
    table = doc.add_table(rows=len(tag_statistics) + 1, cols=4)
    table.style = 'Table Grid'

    # Set column headers
    headers = table.rows[0].cells
    headers[0].text = "Tag name"
    headers[1].text = "# of instances"
    headers[2].text = "# of sites"
    headers[3].text = "% of sites"

    # Add data for each tag
    for i, (tag, stats) in enumerate(sorted(tag_statistics.items()), 1):
        row = table.rows[i].cells
        percentage = (len(stats['domains']) / len(total_domains)) * 100 if total_domains else 0
        
        row[0].text = f"<{tag}>"
        row[1].text = str(stats['count'])
        row[2].text = str(len(stats['domains']))
        row[3].text = f"{percentage:.1f}%"

    # Format the table text
    format_table_text(table)

    # Add some space after the table
    doc.add_paragraph()

    # Add detailed breakdown for each tag
    doc.add_heading('Detailed Breakdown by Tag', level=3)
    
    tag_domain_details = {}
    
    # Process violations again to build domain details
    for page in pages_with_violations:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        
        violations = page['results']['accessibility']['tests']['accessible_names']['accessible_names']['details']['violations']
        if isinstance(violations, str):
            violations = json.loads(violations)
        
        for violation in violations:
            tag = violation['element']
            
            # Initialize tag in domain details if not present
            if tag not in tag_domain_details:
                tag_domain_details[tag] = {}
            
            # Initialize domain for this tag if not present
            if domain not in tag_domain_details[tag]:
                tag_domain_details[tag][domain] = {
                    'count': 0,
                    'pages': set()
                }
            
            # Update statistics
            tag_domain_details[tag][domain]['count'] += 1
            tag_domain_details[tag][domain]['pages'].add(page['url'])

    # Create tables for each tag
    for tag in sorted(tag_statistics.keys()):
        # Add tag header
        doc.add_paragraph(f"Tag: <{tag}>", style='Heading 4')
        
        # Create table for this tag's domain breakdown
        domain_table = doc.add_table(rows=len(tag_domain_details[tag]) + 1, cols=3)
        domain_table.style = 'Table Grid'
        
        # Set headers
        headers = domain_table.rows[0].cells
        headers[0].text = "Domain"
        headers[1].text = "# of instances"
        headers[2].text = "# of pages affected"
        
        # Add domain data
        for i, (domain, stats) in enumerate(sorted(tag_domain_details[tag].items()), 1):
            row = domain_table.rows[i].cells
            row[0].text = domain
            row[1].text = str(stats['count'])
            row[2].text = str(len(stats['pages']))
        
        # Format the table text
        format_table_text(domain_table)
        
        # Add space after each tag's breakdown
        doc.add_paragraph()    

    #############################
    # Animation
    #############################

    def parse_duration(duration_str):
        """Convert duration string to milliseconds"""
        if not duration_str or duration_str == '0ms':
            return 0
        
        value = float(duration_str.replace('ms', '').replace('s', ''))
        return value * 1000 if duration_str.endswith('s') else value

    doc.add_page_break()
    h2 = doc.add_heading('Animation', level=2)
    h2.style = doc.styles['Heading 2']

    doc.add_paragraph("The 'prefers-reduced-motion' media query allows websites to respect a user's system-level preference for reduced motion. This accessibility feature is crucial for several user groups:")

    doc.add_paragraph("People with vestibular disorders who can experience dizziness, nausea, and disorientation from animated content", style='List Bullet')
    doc.add_paragraph("Users with attention-related disabilities who may find animations distracting and disruptive", style='List Bullet')
    doc.add_paragraph("People with migraine sensitivity who can be triggered by certain types of motion", style='List Bullet')
    doc.add_paragraph("Users with cognitive disabilities who may find it difficult to focus on content when animations are present", style='List Bullet')

    doc.add_paragraph("""
    When websites don't support reduced motion preferences, users who rely on this setting remain exposed to animations that could affect their ability to use the site or even cause physical discomfort. This is particularly important for essential services and information websites where users need to access content regardless of their motion sensitivity.
    """.strip())
                      
    doc.add_paragraph("""
    Note that "prefers-reduced-andimation" does not mean no animation, but you do need to cosider the impact of each, especially longer ones. In these tests long animations are those over 5 seconds, but in practice, any animation of over 1 second needs to respect the prefers-reduced-motion media query to help neuro-diverse users who may struggle to read a parge with significant animation "mnise".
    """.strip())

    doc.add_paragraph()  # Add space before the tables

    # Query for pages that have animations but lack reduced motion support
    pages_lacking_motion_support = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.animations.animations.pageFlags.hasAnimations": True,
            "results.accessibility.tests.animations.animations.pageFlags.lacksReducedMotionSupport": True
        },
        {
            "url": 1,
            "results.accessibility.tests.animations.animations.details.summary": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Count affected domains
    affected_domains = set()
    for page in pages_lacking_motion_support:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        affected_domains.add(domain)

    # Calculate percentage
    percentage = (len(affected_domains) / len(total_domains)) * 100 if total_domains else 0

    # Query for pages that have animations but lack reduced motion support
    pages_with_animation_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.animations.animations.pageFlags.hasAnimations": True,
            "results.accessibility.tests.animations.animations.pageFlags.lacksReducedMotionSupport": True
        },
        {
            "url": 1,
            "results.accessibility.tests.animations.animations": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Count affected domains and collect statistics
    domain_stats = {}
    for page in pages_with_animation_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        animation_data = page['results']['accessibility']['tests']['animations']['animations']
        summary = animation_data['details']['summary']
        page_flags = animation_data['pageFlags']['details']
        
        if domain not in domain_stats:
            domain_stats[domain] = {
                'pages': 0,
                'total_animations': 0,
                'infinite_animations': 0,
                'long_animations': 0,
                'shortest_animation': None,
                'longest_animation': None
            }
        
        domain_stats[domain]['pages'] += 1
        domain_stats[domain]['total_animations'] += summary['totalAnimations']
        domain_stats[domain]['infinite_animations'] += summary['infiniteAnimations']
        domain_stats[domain]['long_animations'] += summary['longDurationAnimations']
        
        # Update shortest animation
        if 'shortestAnimation' in page_flags and page_flags['shortestAnimation'] != '0ms':
            if domain_stats[domain]['shortest_animation'] is None:
                domain_stats[domain]['shortest_animation'] = page_flags['shortestAnimation']
            else:
                # Compare durations (convert to ms for comparison)
                current = parse_duration(domain_stats[domain]['shortest_animation'])
                new = parse_duration(page_flags['shortestAnimation'])
                if new < current:
                    domain_stats[domain]['shortest_animation'] = page_flags['shortestAnimation']
        
        # Update longest animation
        if 'longestAnimationElement' in page_flags and page_flags['longestAnimationElement'] and page_flags['longestAnimationElement']['duration']:
            if domain_stats[domain]['longest_animation'] is None:
                domain_stats[domain]['longest_animation'] = page_flags['longestAnimationElement']['duration']
            else:
                # Compare durations (convert to ms for comparison)
                current = parse_duration(domain_stats[domain]['longest_animation'])
                new = parse_duration(page_flags['longestAnimationElement']['duration'])
                if new > current:
                    domain_stats[domain]['longest_animation'] = page_flags['longestAnimationElement']['duration']
    
    # Create summary table
    if domain_stats:
        # Add paragraph to keep table with previous content
        last_para = doc.add_paragraph()
        last_para._element.get_or_add_pPr().append(
            parse_xml(r'<w:keepNext xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        )

        # Calculate totals for summary table
        total_pages = sum(stats['pages'] for stats in domain_stats.values())
        affected_domains = len(domain_stats)
        percentage = (affected_domains / len(total_domains)) * 100 if total_domains else 0

        summary_table = doc.add_table(rows=2, cols=4)
        summary_table.style = 'Table Grid'
        
        # Keep table together
        for row in summary_table.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcPr.append(parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>'))
                tcPr.append(parse_xml(r'<w:noWrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue"
        headers[1].text = "# of pages"
        headers[2].text = "# of sites affected"
        headers[3].text = "% of sites"

        # Add data
        row = summary_table.rows[1].cells
        row[0].text = "Pages with animations lacking reduced motion support"
        row[1].text = str(total_pages)
        row[2].text = str(affected_domains)
        row[3].text = f"{percentage:.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add some space after the table
        doc.add_paragraph()

        # Create detailed domain table
        last_para = doc.add_paragraph()
        last_para._element.get_or_add_pPr().append(
            parse_xml(r'<w:keepNext xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        )

        domain_table = doc.add_table(rows=len(domain_stats) + 1, cols=7)  # Updated number of columns
        domain_table.style = 'Table Grid'
        
        # Keep table together
        for row in domain_table.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcPr.append(parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>'))
                tcPr.append(parse_xml(r'<w:noWrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

        # Add headers
        headers = domain_table.rows[0].cells
        headers[0].text = "Domain"
        headers[1].text = "Pages without reduced motion support"
        headers[2].text = "Total animations"
        headers[3].text = "Infinite animations"
        headers[4].text = "Long animations"
        headers[5].text = "Shortest animation"
        headers[6].text = "Longest animation"

        # Add domain data
        for i, (domain, stats) in enumerate(sorted(domain_stats.items()), 1):
            row = domain_table.rows[i].cells
            row[0].text = domain
            row[1].text = str(stats['pages'])
            row[2].text = str(stats['total_animations'])
            row[3].text = str(stats['infinite_animations'])
            row[4].text = str(stats['long_animations'])
            row[5].text = stats['shortest_animation'] or 'N/A'
            row[6].text = stats['longest_animation'] or 'N/A'

        # Format the table text
        format_table_text(domain_table)

    ###################################
    # Colour Contrast
    ###################################

    doc.add_page_break()
    h2 = doc.add_heading('Colour Contrast', level=2)
    h2.style = doc.styles['Heading 2']

    # 1. Text Contrast Issues
    doc.add_paragraph("Text Contrast Issues:").bold = True
    doc.add_paragraph("Checks if text has sufficient contrast with its background", style='List Bullet')
    doc.add_paragraph("Small text needs a contrast ratio of at least 4.5:1", style='List Bullet')
    doc.add_paragraph("Large text (18pt+ or 14pt+ bold) needs at least 3:1", style='List Bullet')
    doc.add_paragraph("Essential for users with low vision or color vision deficiencies", style='List Bullet')
    doc.add_paragraph()

    # 2. Non-Text Contrast Issues
    doc.add_paragraph("Non-Text Contrast Issues:").bold = True
    doc.add_paragraph("Examines contrast of UI components like buttons and horizontal rules", style='List Bullet')
    doc.add_paragraph("Requires a minimum contrast ratio of 3:1 for boundaries and visual information", style='List Bullet')
    doc.add_paragraph("Important for identifying clickable elements and interface boundaries", style='List Bullet')
    doc.add_paragraph()

    # 3. Adjacent Contrast Issues
    doc.add_paragraph("Adjacent Contrast Issues:").bold = True
    doc.add_paragraph("Checks contrast between neighboring content blocks", style='List Bullet')
    doc.add_paragraph("Requires sufficient contrast between adjacent sections", style='List Bullet')
    doc.add_paragraph("Helps users distinguish between different content areas", style='List Bullet')
    doc.add_paragraph()

    # 4. Contrast Preferences Support
    doc.add_paragraph("Contrast Preferences Support:").bold = True
    doc.add_paragraph("Checks if the site responds to system-level contrast preferences", style='List Bullet')
    doc.add_paragraph("Should adapt to user's contrast preference settings", style='List Bullet')
    doc.add_paragraph("Important for users who need specific contrast levels", style='List Bullet')
    doc.add_paragraph()

    # Define the contrast issues to be analyzed
    contrast_issues = [
        {
            'name': 'Text Contrast Issues',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.hasContrastIssues',
            'details_field': 'results.accessibility.tests.colors.colors.details.summary.contrastViolations'
        },
        {
            'name': 'Non-Text Contrast Issues',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.hasNonTextContrastIssues',
            'details_field': 'results.accessibility.tests.colors.colors.details.summary.nonTextContrastViolations'
        },
        {
            'name': 'Adjacent Contrast Issues',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.hasAdjacentContrastIssues',
            'details_field': 'results.accessibility.tests.colors.colors.details.summary.adjacentContrastViolations'
        },
        {
            'name': 'Contrast Preferences Support',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.supportsContrastPreferences',
            'details_field': None  # This is a boolean field, not a count
        }
    ]

    # Gather the data for each issue type
    issue_data = {}

    for issue in contrast_issues:
        # For the Contrast Preferences Support, we want sites that DO support it
        # For other issues, we want sites that have problems
        if issue['name'] == 'Contrast Preferences Support':
            query = {issue['db_field']: True}
        else:
            query = {issue['db_field']: True}
        
        # Prepare projection
        projection = {"url": 1, "_id": 0}
        if issue['details_field']:
            projection[issue['details_field']] = 1
        
        # Query the database to find pages with this issue
        pages_with_issue = list(db_connection.page_results.find(query, projection))
        
        # Count affected domains and total issue instances
        affected_domains = set()
        total_instances = 0
        
        for page in pages_with_issue:
            domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
            affected_domains.add(domain)
            
            # Count instances if applicable
            if issue['details_field']:
                # Navigate the nested structure to get the count
                parts = issue['details_field'].split('.')
                value = page
                try:
                    for part in parts:
                        if part in value:
                            value = value[part]
                        else:
                            value = 0
                            break
                    
                    if isinstance(value, (int, float)):
                        total_instances += value
                except:
                    pass  # Handle any issues with nested access
        
        # Store the data
        issue_data[issue['name']] = {
            'pages': pages_with_issue,
            'domains': affected_domains,
            'instances': total_instances
        }

    # Create summary table
    last_para = doc.add_paragraph()
    last_para._element.get_or_add_pPr().append(
        parse_xml(r'<w:keepNext xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    )

    summary_table = doc.add_table(rows=len(contrast_issues) + 1, cols=4)
    summary_table.style = 'Table Grid'

    # Keep table together
    for row in summary_table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcPr.append(parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>'))
            tcPr.append(parse_xml(r'<w:noWrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

    # Set column headers
    headers = summary_table.rows[0].cells
    headers[0].text = "Color Accessibility Issue"
    headers[1].text = "Pages Affected"
    headers[2].text = "Sites Affected"
    headers[3].text = "% of Total Sites"

    # Add data
    for i, issue in enumerate(contrast_issues, 1):
        row = summary_table.rows[i].cells
        data = issue_data[issue['name']]
        
        row[0].text = issue['name']
        row[1].text = str(len(data['pages']))
        row[2].text = str(len(data['domains']))
        
        percentage = (len(data['domains']) / len(total_domains)) * 100 if total_domains else 0
        row[3].text = f"{percentage:.1f}%"

    # Format the table text
    format_table_text(summary_table)

    # Add some space after the table
    doc.add_paragraph()

    # Add details for each issue type that has occurrences
    for issue in contrast_issues:
        data = issue_data[issue['name']]
        if data['domains']:
            doc.add_paragraph(f"Sites with {issue['name'].lower()}:")
            
            # Create a dictionary to count pages per domain
            domain_counts = {}
            for page in data['pages']:
                domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
                domain_counts[domain] = domain_counts.get(domain, 0) + 1
            
            # Create domain details table
            domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
            domain_table.style = 'Table Grid'

            # Keep table together
            for row in domain_table.rows:
                for cell in row.cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcPr.append(parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>'))
                    tcPr.append(parse_xml(r'<w:noWrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

            # Add headers
            headers = domain_table.rows[0].cells
            headers[0].text = "Domain"
            headers[1].text = "Number of pages"

            # Add domain data
            for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                row = domain_table.rows[i].cells
                row[0].text = domain
                row[1].text = str(count)

            # Format the table text
            format_table_text(domain_table)
            doc.add_paragraph()  # Add space after each table


    ###################################
    # Colour as indicator
    ###################################

    doc.add_page_break()
    h2 = doc.add_heading('Colour as Indicator', level=2)
    h2.style = doc.styles['Heading 2']
    doc.add_paragraph()

    # 5. Color-Only Links
    doc.add_paragraph("Color-Only Links:").bold = True
    doc.add_paragraph("Identifies links that are distinguished only by color", style='List Bullet')
    doc.add_paragraph("Links should have additional indicators (e.g., underlines, roles, or other visual markers)", style='List Bullet')
    doc.add_paragraph("Critical for users who are color blind or using monochrome displays", style='List Bullet')
    doc.add_paragraph()

    # 6. Color References
    doc.add_paragraph("Color References:").bold = True
    doc.add_paragraph("Detects content that relies on color to convey information", style='List Bullet')
    doc.add_paragraph("Looks for phrases like \"click the red button\" or \"items marked in blue\"", style='List Bullet')
    doc.add_paragraph("Must have alternative ways to convey the same information", style='List Bullet')
    doc.add_paragraph()

    # 7. Color Scheme Preferences Support
    doc.add_paragraph("Color Scheme Preferences Support:").bold = True
    doc.add_paragraph("Verifies support for system color scheme preferences (like dark mode)", style='List Bullet')
    doc.add_paragraph("Should respect user's preferred color scheme", style='List Bullet')
    doc.add_paragraph("Helps users who are sensitive to bright displays or need specific color schemes", style='List Bullet')
    doc.add_paragraph()

    # Define the indicator issues to be analyzed
    indicator_issues = [
        {
            'name': 'Color-Only Links',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.hasColorOnlyLinks',
            'details_field': 'results.accessibility.tests.colors.colors.details.summary.colorOnlyLinks'
        },
        {
            'name': 'Color References',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.hasColorReferences',
            'details_field': 'results.accessibility.tests.colors.colors.details.summary.colorReferenceCount'
        },
        {
            'name': 'Color Scheme Preferences Support',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.supportsColorSchemePreferences',
            'details_field': None  # This is a boolean field, not a count
        }
    ]

    # Gather the data for each issue type
    indicator_data = {}

    for issue in indicator_issues:
        # For the Color Scheme Preferences Support, we want sites that DO support it
        # For other issues, we want sites that have problems
        if issue['name'] == 'Color Scheme Preferences Support':
            query = {issue['db_field']: True}
        else:
            query = {issue['db_field']: True}
        
        # Prepare projection
        projection = {"url": 1, "_id": 0}
        if issue['details_field']:
            projection[issue['details_field']] = 1
        
        # Query the database to find pages with this issue
        pages_with_issue = list(db_connection.page_results.find(query, projection))
        
        # Count affected domains and total issue instances
        affected_domains = set()
        total_instances = 0
        
        for page in pages_with_issue:
            domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
            affected_domains.add(domain)
            
            # Count instances if applicable
            if issue['details_field']:
                # Navigate the nested structure to get the count
                parts = issue['details_field'].split('.')
                value = page
                try:
                    for part in parts:
                        if part in value:
                            value = value[part]
                        else:
                            value = 0
                            break
                    
                    if isinstance(value, (int, float)):
                        total_instances += value
                except:
                    pass  # Handle any issues with nested access
        
        # Store the data
        indicator_data[issue['name']] = {
            'pages': pages_with_issue,
            'domains': affected_domains,
            'instances': total_instances
        }

    # Create summary table
    last_para = doc.add_paragraph()
    last_para._element.get_or_add_pPr().append(
        parse_xml(r'<w:keepNext xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    )

    summary_table = doc.add_table(rows=len(indicator_issues) + 1, cols=4)
    summary_table.style = 'Table Grid'

    # Keep table together
    for row in summary_table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcPr.append(parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>'))
            tcPr.append(parse_xml(r'<w:noWrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

    # Set column headers
    headers = summary_table.rows[0].cells
    headers[0].text = "Color Accessibility Issue"
    headers[1].text = "Pages Affected"
    headers[2].text = "Sites Affected"
    headers[3].text = "% of Total Sites"

    # Add data
    for i, issue in enumerate(indicator_issues, 1):
        row = summary_table.rows[i].cells
        data = indicator_data[issue['name']]
        
        row[0].text = issue['name']
        row[1].text = str(len(data['pages']))
        row[2].text = str(len(data['domains']))
        
        percentage = (len(data['domains']) / len(total_domains)) * 100 if total_domains else 0
        row[3].text = f"{percentage:.1f}%"

    # Format the table text
    format_table_text(summary_table)

    # Add some space after the table
    doc.add_paragraph()

    # Add details for each issue type that has occurrences
    for issue in indicator_issues:
        data = indicator_data[issue['name']]
        if data['domains']:
            doc.add_paragraph(f"Sites with {issue['name'].lower()}:")
            
            # Create a dictionary to count pages per domain
            domain_counts = {}
            for page in data['pages']:
                domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
                domain_counts[domain] = domain_counts.get(domain, 0) + 1
            
            # Create domain details table
            domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
            domain_table.style = 'Table Grid'

            # Keep table together
            for row in domain_table.rows:
                for cell in row.cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcPr.append(parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>'))
                    tcPr.append(parse_xml(r'<w:noWrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

            # Add headers
            headers = domain_table.rows[0].cells
            headers[0].text = "Domain"
            headers[1].text = "Number of pages"

            # Add domain data
            for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                row = domain_table.rows[i].cells
                row[0].text = domain
                row[1].text = str(count)

            # Format the table text
            format_table_text(domain_table)
            doc.add_paragraph()  # Add space after each table
            
            # For Color References, provide more detailed information
            if issue['name'] == 'Color References':
                doc.add_paragraph("Color references found in content:", style='Heading 4')
                
                # Query for pages with color references to see the actual references
                color_refs_query = {
                    'results.accessibility.tests.colors.colors.pageFlags.hasColorReferences': True
                }
                color_refs_projection = {
                    'url': 1,
                    'results.accessibility.tests.colors.colors.details.colorReferences.instances': 1,
                    '_id': 0
                }
                
                pages_with_refs = list(db_connection.page_results.find(color_refs_query, color_refs_projection))
                
                # Extract all references and count them
                reference_counts = {}
                
                for page in pages_with_refs:
                    try:
                        # Navigate to get the instances
                        instances = page['results']['accessibility']['tests']['colors']['colors']['details']['colorReferences']['instances']
                        
                        # Handle case where instances might be a string (JSON)
                        if isinstance(instances, str):
                            try:
                                instances = json.loads(instances)
                            except:
                                continue
                        
                        # Process each instance
                        if isinstance(instances, list):
                            for instance in instances:
                                if isinstance(instance, dict) and 'references' in instance:
                                    for ref in instance['references']:
                                        ref = ref.lower()  # Normalize to lowercase
                                        reference_counts[ref] = reference_counts.get(ref, 0) + 1
                    except (KeyError, TypeError):
                        # Handle any issues with accessing the data
                        continue
                
                # Create a table of color references if any were found
                if reference_counts:
                    reference_table = doc.add_table(rows=len(reference_counts) + 1, cols=2)
                    reference_table.style = 'Table Grid'
                    
                    # Add headers
                    headers = reference_table.rows[0].cells
                    headers[0].text = "Color Referenced"
                    headers[1].text = "Number of Occurrences"
                    
                    # Add reference data
                    for i, (ref, count) in enumerate(sorted(reference_counts.items(), key=lambda x: x[1], reverse=True), 1):
                        row = reference_table.rows[i].cells
                        row[0].text = ref.capitalize()
                        row[1].text = str(count)
                    
                    # Format the table text
                    format_table_text(reference_table)
                    doc.add_paragraph()  # Add space after the table
                    
                    # Provide an explanation about the impact
                    doc.add_paragraph(
                        "Color references in content (like 'click the red button' or 'items marked in blue') "
                        "create barriers for users who are color blind or using screen readers. Information conveyed "
                        "through color alone must have alternative indicators that don't rely on color perception."
                    )
                else:
                    doc.add_paragraph("No specific color references were identified in the scanned content.")

    ################################################
    # Dialogs
    ################################################

    doc.add_page_break()
    h2 = doc.add_heading('Dialogs', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
    Modal dialogs present unique accessibility challenges and must be implemented correctly to ensure all users can interact with them effectively. When a modal is open, keyboard focus must be trapped within it, and screen readers must be properly informed of the modal's presence and purpose.
    """.strip())

    doc.add_paragraph("Common accessibility issues with modal dialogs include:", style='Normal')

    doc.add_paragraph("Missing close mechanisms (close button or escape key)", style='List Bullet')
    doc.add_paragraph("Improper focus management when opening and closing", style='List Bullet')
    doc.add_paragraph("Missing or improper heading structure", style='List Bullet')
    doc.add_paragraph("Missing or improper trigger buttons", style='List Bullet')

    # Add recommendations
    doc.add_paragraph()
    doc.add_paragraph("Recommendations for Modal Dialog Implementation:", style='Normal')
    
    doc.add_paragraph("Ensure all modals have a visible close button and respond to the escape key", style='List Bullet')
    doc.add_paragraph("Implement proper focus management - trap focus within the modal when open and return focus when closed", style='List Bullet')
    doc.add_paragraph("Include proper heading structure within modals for clear content hierarchy", style='List Bullet')
    doc.add_paragraph("Use proper trigger buttons with appropriate ARIA attributes and keyboard interaction", style='List Bullet')
    doc.add_paragraph("Test modal interactions with keyboard-only navigation and screen readers", style='List Bullet')

    # Query for pages with modal issues
    pages_with_modal_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.modals.modals.pageFlags.hasModals": True,
            "results.accessibility.tests.modals.modals.pageFlags.hasModalViolations": True
        },
        {
            "url": 1,
            "results.accessibility.tests.modals.modals.pageFlags": 1,
            "results.accessibility.tests.modals.modals.details.summary": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    modal_issues = {
        "modalsWithoutClose": {"name": "Missing close mechanism", "pages": set(), "domains": set()},
        "modalsWithoutFocusManagement": {"name": "Improper focus management", "pages": set(), "domains": set()},
        "modalsWithoutProperHeading": {"name": "Missing/improper heading", "pages": set(), "domains": set()},
        "modalsWithoutTriggers": {"name": "Missing/improper triggers", "pages": set(), "domains": set()}
    }

    # Count issues
    for page in pages_with_modal_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        summary = page['results']['accessibility']['tests']['modals']['modals']['details']['summary']
        
        for flag in modal_issues:
            if summary.get(flag, 0) > 0:
                modal_issues[flag]['pages'].add(page['url'])
                modal_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in modal_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Modal Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)


        # Add domain details for each issue
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)

        # Add statistics about total modals if available
        doc.add_paragraph()
        total_modals = sum(page['results']['accessibility']['tests']['modals']['modals']['details']['summary']['totalModals'] 
                        for page in pages_with_modal_issues)
        doc.add_paragraph(f"Total number of modals detected across all pages: {total_modals}")

    else:
        doc.add_paragraph("No dialog accessibility issues were found.")


    ####################################################
    # Event handling
    ####################################################

    doc.add_page_break()
    h2 = doc.add_heading('Event Handling and Keyboard Interaction', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
    Event handling and keyboard interaction are crucial for accessibility. This analysis examines event types, tab order, interactive elements, and modal dialog behavior. Issues with event handling can significantly impact keyboard and screen reader users.
    """.strip())

    # Add recommendations
    doc.add_paragraph()
    doc.add_heading('Event Handling Recommendations', level=3)
    
    doc.add_paragraph("Event Implementation:", style='List Bullet')
    doc.add_paragraph("Ensure keyboard alternatives for mouse-only interactions", style='List Bullet 2')
    doc.add_paragraph("Add keyboard event handlers alongside mouse events", style='List Bullet 2')
    doc.add_paragraph("Implement proper focus management", style='List Bullet 2')
    
    doc.add_paragraph("Tab Order:", style='List Bullet')
    doc.add_paragraph("Maintain logical tab sequence matching visual layout", style='List Bullet 2')
    doc.add_paragraph("Avoid using tabindex values greater than 0", style='List Bullet 2')
    doc.add_paragraph("Ensure all interactive elements are keyboard accessible", style='List Bullet 2')
    
    doc.add_paragraph("Modal Dialogs:", style='List Bullet')
    doc.add_paragraph("Implement escape key handling for all modals", style='List Bullet 2')
    doc.add_paragraph("Manage focus properly when opening/closing modals", style='List Bullet 2')
    doc.add_paragraph("Ensure modal content is properly contained", style='List Bullet 2')
    
    doc.add_paragraph()

    # Query for pages with event information
    pages_with_events = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.events.events": {"$exists": True}
        },
        {
            "url": 1,
            "results.accessibility.tests.events.events": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize tracking structures
    property_data = {
        # Event Types
        "event_mouse": {"name": "Mouse Events", "pages": set(), "domains": set(), "count": 0},
        "event_keyboard": {"name": "Keyboard Events", "pages": set(), "domains": set(), "count": 0},
        "event_focus": {"name": "Focus Events", "pages": set(), "domains": set(), "count": 0},
        "event_touch": {"name": "Touch Events", "pages": set(), "domains": set(), "count": 0},
        "event_timer": {"name": "Timer Events", "pages": set(), "domains": set(), "count": 0},
        "event_lifecycle": {"name": "Lifecycle Events", "pages": set(), "domains": set(), "count": 0},
        "event_other": {"name": "Other Events", "pages": set(), "domains": set(), "count": 0},
        
        # Tab Order
        "explicit_tabindex": {"name": "Explicit tabindex Usage", "pages": set(), "domains": set(), "count": 0},
        "visual_violations": {"name": "Visual Order Violations", "pages": set(), "domains": set(), "count": 0},
        "column_violations": {"name": "Column Order Violations", "pages": set(), "domains": set(), "count": 0},
        "negative_tabindex": {"name": "Negative Tabindex", "pages": set(), "domains": set(), "count": 0},
        "high_tabindex": {"name": "High Tabindex Values", "pages": set(), "domains": set(), "count": 0},
        
        # Interactive Elements
        "mouse_only": {"name": "Mouse-only Elements", "pages": set(), "domains": set(), "count": 0},
        "missing_tabindex": {"name": "Missing tabindex", "pages": set(), "domains": set(), "count": 0},
        "non_interactive": {"name": "Non-interactive with Handlers", "pages": set(), "domains": set(), "count": 0},
        
        # Modal Support
        "modals_no_escape": {"name": "Modals Missing Escape", "pages": set(), "domains": set(), "count": 0}
    }

    # Create detailed violation tracking organized by domain and URL
    domain_data = {}

    # Process each page
    for page in pages_with_events:
        try:
            url = page['url']
            
            domain = url.replace('http://', '').replace('https://', '').split('/')[0]
            event_data = page['results']['accessibility']['tests']['events']['events']
            
            # Initialize domain and URL tracking if needed
            if domain not in domain_data:
                domain_data[domain] = {
                    'urls': {}
                }
            
            # Initialize URL data
            domain_data[domain]['urls'][url] = {
                'event_types': {},
                'violations': {},
                'handlers_count': 0,
                'focusable_elements': 0,
                'total_violations': 0
            }
            
            # Get pageFlags data for the most reliable summary information
            pageFlags = event_data.get('pageFlags', {})
            details = pageFlags.get('details', {})
            
            # Track total handlers and violations
            total_handlers = details.get('totalHandlers', 0)
            total_violations = details.get('totalViolations', 0)
            domain_data[domain]['urls'][url]['handlers_count'] = total_handlers
            domain_data[domain]['urls'][url]['total_violations'] = total_violations

            # Track total focusable elements
            tab_order_data = details.get('tabOrder', {})
            focusable_elements = tab_order_data.get('totalFocusableElements', 0)
            domain_data[domain]['urls'][url]['focusable_elements'] = focusable_elements
            
            # Process event types using the updated structure
            by_type = details.get('byType', {})
            
            for event_type in ['mouse', 'keyboard', 'focus', 'touch', 'timer', 'lifecycle', 'other']:
                count = by_type.get(event_type, 0)
                
                if isinstance(count, list):
                    count = len(count)
                elif not isinstance(count, (int, float)):
                    try:
                        count = int(count or 0)
                    except (ValueError, TypeError):
                        count = 0
                
                # Track event type for this URL
                domain_data[domain]['urls'][url]['event_types'][event_type] = count
                
                if count > 0:
                    key = f"event_{event_type}"
                    property_data[key]['pages'].add(url)
                    property_data[key]['domains'].add(domain)
                    property_data[key]['count'] += count

            # Process violation counts by type
            violation_counts = details.get('violationCounts', {})
            
            # Tab order violations
            explicit_count = tab_order_data.get('elementsWithExplicitTabIndex', 0)
            visual_violations = violation_counts.get('visual-order', 0) or tab_order_data.get('visualOrderViolations', 0)
            column_violations = violation_counts.get('column-order', 0) or tab_order_data.get('columnOrderViolations', 0)
            
            # Track negative and high tabindex
            negative_tabindex = 1 if pageFlags.get('hasNegativeTabindex', False) else 0
            high_tabindex = 1 if pageFlags.get('hasHighTabindex', False) else 0
            
            # Track violations for this URL
            domain_data[domain]['urls'][url]['violations']['explicit_tabindex'] = explicit_count
            domain_data[domain]['urls'][url]['violations']['visual_order'] = visual_violations
            domain_data[domain]['urls'][url]['violations']['column_order'] = column_violations
            domain_data[domain]['urls'][url]['violations']['negative_tabindex'] = negative_tabindex
            domain_data[domain]['urls'][url]['violations']['high_tabindex'] = high_tabindex
            
            if explicit_count > 0:
                property_data['explicit_tabindex']['pages'].add(url)
                property_data['explicit_tabindex']['domains'].add(domain)
                property_data['explicit_tabindex']['count'] += explicit_count
                
            if visual_violations > 0:
                property_data['visual_violations']['pages'].add(url)
                property_data['visual_violations']['domains'].add(domain)
                property_data['visual_violations']['count'] += visual_violations
                
            if column_violations > 0:
                property_data['column_violations']['pages'].add(url)
                property_data['column_violations']['domains'].add(domain)
                property_data['column_violations']['count'] += column_violations
                
            if negative_tabindex > 0:
                property_data['negative_tabindex']['pages'].add(url)
                property_data['negative_tabindex']['domains'].add(domain)
                property_data['negative_tabindex']['count'] += negative_tabindex
                
            if high_tabindex > 0:
                property_data['high_tabindex']['pages'].add(url)
                property_data['high_tabindex']['domains'].add(domain)
                property_data['high_tabindex']['count'] += high_tabindex

            # Process element violations
            mouse_only = violation_counts.get('mouse-only', 0) or details.get('mouseOnlyElements', {}).get('count', 0)
            missing_tabindex = violation_counts.get('missing-tabindex', 0) or details.get('missingTabindex', 0)
            non_interactive = details.get('nonInteractiveWithHandlers', 0)
            modals_without_escape = violation_counts.get('modal-without-escape', 0)
            
            # Track violations for this URL
            domain_data[domain]['urls'][url]['violations']['mouse_only'] = mouse_only
            domain_data[domain]['urls'][url]['violations']['missing_tabindex'] = missing_tabindex
            domain_data[domain]['urls'][url]['violations']['non_interactive'] = non_interactive
            domain_data[domain]['urls'][url]['violations']['modals_no_escape'] = modals_without_escape
            
            if mouse_only > 0:
                property_data['mouse_only']['pages'].add(url)
                property_data['mouse_only']['domains'].add(domain)
                property_data['mouse_only']['count'] += mouse_only
                
            if missing_tabindex > 0:
                property_data['missing_tabindex']['pages'].add(url)
                property_data['missing_tabindex']['domains'].add(domain)
                property_data['missing_tabindex']['count'] += missing_tabindex
                
            if non_interactive > 0:
                property_data['non_interactive']['pages'].add(url)
                property_data['non_interactive']['domains'].add(domain)
                property_data['non_interactive']['count'] += non_interactive
                
            if modals_without_escape > 0:
                property_data['modals_no_escape']['pages'].add(url)
                property_data['modals_no_escape']['domains'].add(domain)
                property_data['modals_no_escape']['count'] += modals_without_escape

        except Exception as e:
            print(f"Error processing page {url}:")
            print("Exception:", str(e))
            traceback.print_exc()
            continue

    if pages_with_events:
        # Part 1: Overall Summary Table
        doc.add_heading('Event Handling Summary', level=3)
        
        # Calculate number of rows needed
        rows_needed = 1  # Header row
        
        # Event Types section (header + all event types)
        rows_needed += 1  # Section header
        rows_needed += len([k for k in property_data.keys() if k.startswith('event_')])
        
        # Tab Order section (header + 5 items)
        rows_needed += 1  # Section header
        rows_needed += 5  # explicit_tabindex, visual_violations, column_violations, negative_tabindex, high_tabindex
        
        # Interactive Elements section (header + 3 items)
        rows_needed += 1  # Section header
        rows_needed += 3  # mouse_only, missing_tabindex, non_interactive
        
        # Modal Support section (header + 1 item)
        rows_needed += 1  # Section header
        rows_needed += 1  # modals_no_escape

        # Create table with correct number of rows
        table = doc.add_table(rows=rows_needed, cols=4)
        table.style = 'Table Grid'
        
        # Add headers
        headers = table.rows[0].cells
        headers[0].text = "Property"
        headers[1].text = "Occurrences"
        headers[2].text = "Pages Affected"
        headers[3].text = "% of Sites"
        
        current_row = 1
        
        # Add Event Types section
        row = table.rows[current_row].cells
        row[0].text = "Event Types:"
        current_row += 1
        
        for key, data in sorted([(k, v) for k, v in property_data.items() if k.startswith('event_')], 
                            key=lambda x: x[1]['count'], reverse=True):
            row = table.rows[current_row].cells
            row[0].text = "  " + data['name']
            row[1].text = str(data['count'])
            row[2].text = str(len(data['pages']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"
            current_row += 1
        
        # Add Tab Order section
        row = table.rows[current_row].cells
        row[0].text = "Tab Order:"
        current_row += 1
        
        for key in ['explicit_tabindex', 'visual_violations', 'column_violations', 'negative_tabindex', 'high_tabindex']:
            row = table.rows[current_row].cells
            row[0].text = "  " + property_data[key]['name']
            row[1].text = str(property_data[key]['count'])
            row[2].text = str(len(property_data[key]['pages']))
            row[3].text = f"{(len(property_data[key]['domains']) / len(total_domains) * 100):.1f}%"
            current_row += 1
        
        # Add Interactive Elements section
        row = table.rows[current_row].cells
        row[0].text = "Interactive Elements:"
        current_row += 1
        
        for key in ['mouse_only', 'missing_tabindex', 'non_interactive']:
            row = table.rows[current_row].cells
            row[0].text = "  " + property_data[key]['name']
            row[1].text = str(property_data[key]['count'])
            row[2].text = str(len(property_data[key]['pages']))
            row[3].text = f"{(len(property_data[key]['domains']) / len(total_domains) * 100):.1f}%"
            current_row += 1
        
        # Add Modal Support section
        row = table.rows[current_row].cells
        row[0].text = "Modal Support:"
        current_row += 1
        
        key = 'modals_no_escape'
        row = table.rows[current_row].cells
        row[0].text = "  " + property_data[key]['name']
        row[1].text = str(property_data[key]['count'])
        row[2].text = str(len(property_data[key]['pages']))
        row[3].text = f"{(len(property_data[key]['domains']) / len(total_domains) * 100):.1f}%"

        format_table_text(table)
        
        # Part 2: Detailed Per-Site and Per-URL Tables
        doc.add_paragraph()
        doc.add_heading('Event Handling Details by Site', level=3)
        
        # Process each domain
        for domain_name in sorted(domain_data.keys()):
            domain_urls = domain_data[domain_name]['urls']
            if not domain_urls:
                continue
                
            # Add domain heading
            doc.add_paragraph()
            doc.add_heading(f'Domain: {domain_name}', level=4)
            
            # Table 1: Event Types for this domain
            doc.add_paragraph("Event handler distribution:", style='Normal')
            
            # Create event types table
            event_type_headers = ["URL", "Total Handlers", "Mouse", "Keyboard", "Focus", "Touch", "Timer", "Lifecycle", "Other"]
            event_table = doc.add_table(rows=len(domain_urls) + 1, cols=len(event_type_headers))
            event_table.style = 'Table Grid'
            
            # Add headers
            for i, header in enumerate(event_type_headers):
                event_table.cell(0, i).text = header
            
            # Add data rows
            for row_idx, (url, url_data) in enumerate(sorted(domain_urls.items()), 1):
                event_table.cell(row_idx, 0).text = url
                event_table.cell(row_idx, 1).text = str(url_data['handlers_count'])
                
                for col_idx, event_type in enumerate(['mouse', 'keyboard', 'focus', 'touch', 'timer', 'lifecycle', 'other'], 2):
                    event_table.cell(row_idx, col_idx).text = str(url_data['event_types'].get(event_type, 0))
            
            format_table_text(event_table)
            
            # Table 2: Tab Order Violations for this domain
            doc.add_paragraph()
            doc.add_paragraph("Tab order violations:", style='Normal')
            
            # Create tab order table
            tab_order_headers = ["URL", "Focusable Elements", "Total Violations", "Visual Order", "Column Order", 
                                "Explicit tabindex", "Negative tabindex", "High tabindex"]
            tab_order_table = doc.add_table(rows=len(domain_urls) + 1, cols=len(tab_order_headers))
            tab_order_table.style = 'Table Grid'
            
            # Add headers
            for i, header in enumerate(tab_order_headers):
                tab_order_table.cell(0, i).text = header
            
            # Add data rows
            for row_idx, (url, url_data) in enumerate(sorted(domain_urls.items()), 1):
                violations = url_data['violations']
                
                tab_order_table.cell(row_idx, 0).text = url
                tab_order_table.cell(row_idx, 1).text = str(url_data['focusable_elements'])
                tab_order_table.cell(row_idx, 2).text = str(url_data['total_violations'])
                tab_order_table.cell(row_idx, 3).text = str(violations.get('visual_order', 0))
                tab_order_table.cell(row_idx, 4).text = str(violations.get('column_order', 0))
                tab_order_table.cell(row_idx, 5).text = str(violations.get('explicit_tabindex', 0))
                tab_order_table.cell(row_idx, 6).text = str(violations.get('negative_tabindex', 0))
                tab_order_table.cell(row_idx, 7).text = str(violations.get('high_tabindex', 0))
            
            format_table_text(tab_order_table)
            
            # Table 3: Interactive Element Violations for this domain
            doc.add_paragraph()
            doc.add_paragraph("Interactive element violations:", style='Normal')
            
            # Create interactive elements table
            interactive_headers = ["URL", "Mouse-only", "Missing tabindex", "Non-interactive with Handlers", "Modals Missing Escape"]
            interactive_table = doc.add_table(rows=len(domain_urls) + 1, cols=len(interactive_headers))
            interactive_table.style = 'Table Grid'
            
            # Add headers
            for i, header in enumerate(interactive_headers):
                interactive_table.cell(0, i).text = header
            
            # Add data rows
            for row_idx, (url, url_data) in enumerate(sorted(domain_urls.items()), 1):
                violations = url_data['violations']
                
                interactive_table.cell(row_idx, 0).text = url
                interactive_table.cell(row_idx, 1).text = str(violations.get('mouse_only', 0))
                interactive_table.cell(row_idx, 2).text = str(violations.get('missing_tabindex', 0))
                interactive_table.cell(row_idx, 3).text = str(violations.get('non_interactive', 0))
                interactive_table.cell(row_idx, 4).text = str(violations.get('modals_no_escape', 0))
            
            format_table_text(interactive_table)

    else:
        doc.add_paragraph("No event handling data was found.")

    #################################
    # Floating dialogs
    #################################

    doc.add_page_break()
    h3 = doc.add_heading('Floating Dialogs', level=2)
    h3.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
    Floating dialogs (permanent floating dialogs e.g. chatbots, cookie notices, modals) require specific accessibility considerations to ensure they are usable by all. These interface elements can create significant barriers when not implemented properly.
    """.strip())

    # Query for documentation to use in our explanation
    dialog_docs = list(db_connection.page_results.find(
        {"results.accessibility.tests.floating_dialogs.dialogs.documentation": {"$exists": True}},
        {"results.accessibility.tests.floating_dialogs.dialogs.documentation": 1, "_id": 0}
    ).limit(1))

    # If we have documentation, use it to add a rich explanation
    if dialog_docs:
        doc_data = dialog_docs[0]['results']['accessibility']['tests']['floating_dialogs']['dialogs']['documentation']
        
        # Add WCAG criteria references
        wcag_references = []
        for key, criteria in doc_data.get('wcagReferences', {}).items():
            wcag_references.append(f"WCAG {key} ({criteria.get('level', '')}) - {criteria.get('title', '')}")
        
        if wcag_references:
            doc.add_paragraph("This test evaluates the following WCAG criteria:")
            for ref in sorted(wcag_references):
                doc.add_paragraph(ref, style='List Bullet')
        
        # Add common issues based on violation types
        violation_types = doc_data.get('violationTypes', {})
        if violation_types:
            doc.add_paragraph("Common accessibility issues with floating dialogs include:")
            for violation_type, details in violation_types.items():
                doc.add_paragraph(f"{details.get('description', '')}", style='List Bullet')
                
        # Add best practices from documentation
        best_practices = doc_data.get('bestPractices', {})
        if best_practices:
            doc.add_paragraph()
            doc.add_paragraph("Recommendations for making floating dialogs accessible:")
            
            for category, practices in best_practices.items():
                for practice in practices:
                    doc.add_paragraph(practice, style='List Bullet')
    else:
        # Fallback if no documentation is available
        doc.add_paragraph("Common issues include:")
        doc.add_paragraph("Missing or incorrect ARIA roles and attributes", style='List Bullet')
        doc.add_paragraph("Lack of proper focus management when dialogs expand and collapse", style='List Bullet')
        doc.add_paragraph("Missing or inadequate heading structure within dialogs", style='List Bullet')
        doc.add_paragraph("No visible close button or keyboard escape mechanism", style='List Bullet')
        doc.add_paragraph("Interactive content obscured by floating dialogs", style='List Bullet')

        # Add recommendations
        doc.add_paragraph()
        doc.add_paragraph("Recommendations for making floating dialogs accessible:")
        doc.add_paragraph("Use role='dialog' or role='alertdialog' appropriately", style='List Bullet')
        doc.add_paragraph("Manage keyboard focus when dialogs expand and collapse", style='List Bullet')
        doc.add_paragraph("Provide proper heading structure within dialogs", style='List Bullet')
        doc.add_paragraph("Include a visible close button and support Escape key", style='List Bullet')
        doc.add_paragraph("Use aria-modal='true' when appropriate", style='List Bullet')
        doc.add_paragraph("Ensure the dialog doesn't obscure important interactive elements", style='List Bullet')

    doc.add_paragraph()

    # Query for pages with dialog issues - using the consolidated results field
    pages_with_dialog_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.floating_dialogs.dialogs.consolidated": {"$exists": True},
            "results.accessibility.tests.floating_dialogs.dialogs.consolidated.summary.totalIssues": {"$gt": 0}
        },
        {
            "url": 1,
            "results.accessibility.tests.floating_dialogs.dialogs.consolidated": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type by severity
    dialog_issues = {
        "violations": {
            "hiddenInteractiveContent": {"name": "Hidden interactive content", "pages": set(), "domains": set(), "severity": "critical"},
            "incorrectHeadingLevel": {"name": "Incorrect heading structure", "pages": set(), "domains": set(), "severity": "high"},
            "missingCloseButton": {"name": "Missing close button", "pages": set(), "domains": set(), "severity": "high"},
            "improperFocusManagement": {"name": "Improper focus management", "pages": set(), "domains": set(), "severity": "high"}
        },
        "warnings": {
            "contentOverlap": {"name": "Content overlap issues", "pages": set(), "domains": set(), "severity": "moderate"}
        }
    }

    # Count issues and store URLs by domain
    domain_to_urls = {}

    for page in pages_with_dialog_issues:
        url = page['url']
        domain = url.replace('http://', '').replace('https://', '').split('/')[0]
        consolidated = page['results']['accessibility']['tests']['floating_dialogs']['dialogs']['consolidated']
        
        # Initialize domain entry if it doesn't exist
        if domain not in domain_to_urls:
            domain_to_urls[domain] = {}
        
        # Process violations
        if 'issuesByType' in consolidated:
            issues_by_type = consolidated['issuesByType']
            
            # Process violations
            for violation_type, violation_data in issues_by_type.get('violations', {}).items():
                if violation_type in dialog_issues['violations'] and violation_data.get('count', 0) > 0:
                    dialog_issues['violations'][violation_type]['pages'].add(url)
                    dialog_issues['violations'][violation_type]['domains'].add(domain)
                    
                    # Store the severity if available
                    if 'severity' in violation_data:
                        dialog_issues['violations'][violation_type]['severity'] = violation_data['severity']
                    
                    # Store URL by issue type for this domain
                    if violation_type not in domain_to_urls[domain]:
                        domain_to_urls[domain][violation_type] = []
                    domain_to_urls[domain][violation_type].append(url)
            
            # Process warnings
            for warning_type, warning_data in issues_by_type.get('warnings', {}).items():
                if warning_type in dialog_issues['warnings'] and warning_data.get('count', 0) > 0:
                    dialog_issues['warnings'][warning_type]['pages'].add(url)
                    dialog_issues['warnings'][warning_type]['domains'].add(domain)
                    
                    # Store the severity if available
                    if 'severity' in warning_data:
                        dialog_issues['warnings'][warning_type]['severity'] = warning_data['severity']
                    
                    # Store URL by issue type for this domain
                    if warning_type not in domain_to_urls[domain]:
                        domain_to_urls[domain][warning_type] = []
                    domain_to_urls[domain][warning_type].append(url)

    # Create filtered list of issues that have affected pages
    all_active_issues = []

    for category in ['violations', 'warnings']:
        for issue_type, data in dialog_issues[category].items():
            if len(data['pages']) > 0:
                all_active_issues.append({
                    'category': category,
                    'type': issue_type,
                    'name': data['name'],
                    'severity': data['severity'],
                    'pages': data['pages'],
                    'domains': data['domains']
                })

    # Sort issues by severity - critical first, then high, then moderate
    severity_order = {'critical': 0, 'high': 1, 'moderate': 2, 'low': 3}
    all_active_issues.sort(key=lambda x: severity_order.get(x['severity'], 4))

    if all_active_issues:
        # Create summary table
        doc.add_paragraph("The analysis identified the following floating dialog accessibility issues:")
        
        summary_table = doc.add_table(rows=len(all_active_issues) + 1, cols=5)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue"
        headers[1].text = "Severity"
        headers[2].text = "Pages Affected"
        headers[3].text = "Sites Affected"
        headers[4].text = "% of Total Sites"

        # Add data
        for i, issue in enumerate(all_active_issues, 1):
            row = summary_table.rows[i].cells
            row[0].text = issue['name']
            row[1].text = issue['severity'].capitalize()
            row[2].text = str(len(issue['pages']))
            row[3].text = str(len(issue['domains']))
            row[4].text = f"{(len(issue['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add details about each issue by domain with all affected URLs
        doc.add_paragraph()
        heading_para = doc.add_paragraph()
        heading_para.add_run("Detailed List of Affected Pages by Issue and Domain").bold = True
        
        for issue in all_active_issues:
            issue_para = doc.add_paragraph()
            issue_para.add_run(f"{issue['name']} ({issue['severity'].capitalize()})").bold = True
            
            # Get domains that have this specific issue
            issue_domains = {
                domain: urls.get(issue['type'], [])
                for domain, urls in domain_to_urls.items()
                if issue['type'] in urls
            }
            
            # Sort domains by number of affected pages (highest first)
            sorted_domains = sorted(issue_domains.items(), key=lambda x: len(x[1]), reverse=True)
            
            # If we have any domains with this issue
            if sorted_domains:
                for domain, urls in sorted_domains:
                    # Add domain header
                    domain_para = doc.add_paragraph()
                    domain_para.add_run(f"{domain} ({len(urls)} affected pages)").bold = True
                    
                    # Sort URLs for consistent output
                    sorted_urls = sorted(urls)
                    
                    # Add every URL as a bullet point
                    for url in sorted_urls:
                        doc.add_paragraph(url, style='List Bullet')
                    
                    # Add a space between domains
                    doc.add_paragraph()
            else:
                doc.add_paragraph("No domains found with this issue.")

        # Add breakdown by viewport sizes
        responsive_issues = []
        for page in pages_with_dialog_issues:
            if 'consolidated' in page['results']['accessibility']['tests']['floating_dialogs']['dialogs']:
                consolidated = page['results']['accessibility']['tests']['floating_dialogs']['dialogs']['consolidated']
                if 'issuesByType' in consolidated:
                    # Check violations
                    for category in ['violations', 'warnings']:
                        for issue_type, issue_data in consolidated['issuesByType'].get(category, {}).items():
                            if 'elements' in issue_data:
                                for element in issue_data['elements']:
                                    if 'breakpointRange' in element:
                                        responsive_issues.append({
                                            'url': page['url'],
                                            'issue': issue_type,
                                            'name': dialog_issues.get(category, {}).get(issue_type, {}).get('name', issue_type),
                                            'breakpointRange': element['breakpointRange'],
                                            'details': element.get('details', '')
                                        })
        
        if responsive_issues:
            doc.add_paragraph()
            # Use a paragraph with a run for heading formatting instead of a style
            heading_para = doc.add_paragraph()
            heading_para.add_run("Responsive Design Considerations").bold = True
            
            doc.add_paragraph("Many floating dialog issues are specific to certain viewport sizes. The following issues were identified across different breakpoints:")
            
            # Group by breakpoint range
            breakpoint_grouping = {}
            for issue in responsive_issues:
                breakpoint_key = issue['breakpointRange']
                if breakpoint_key not in breakpoint_grouping:
                    breakpoint_grouping[breakpoint_key] = []
                breakpoint_grouping[breakpoint_key].append(issue)
            
            # Create a table for responsive issues
            responsive_table = doc.add_table(rows=len(breakpoint_grouping) + 1, cols=3)
            responsive_table.style = 'Table Grid'
            
            # Set column headers
            headers = responsive_table.rows[0].cells
            headers[0].text = "Viewport Size Range"
            headers[1].text = "Issue Types"
            headers[2].text = "Number of Affected Sites"
            
            # Add data
            row_idx = 1
            for breakpoint_range, issues in breakpoint_grouping.items():
                row = responsive_table.rows[row_idx].cells
                row[0].text = breakpoint_range
                
                # Get unique issue names
                unique_issues = set(issue['name'] for issue in issues)
                row[1].text = ", ".join(unique_issues)
                
                # Count unique domains
                unique_domains = set()
                for issue in issues:
                    domain = issue['url'].replace('http://', '').replace('https://', '').split('/')[0]
                    unique_domains.add(domain)
                
                row[2].text = str(len(unique_domains))
                row_idx += 1
            
            # Format the table text
            format_table_text(responsive_table)

        # Add technical implementation recommendations
        doc.add_paragraph()
        # Use a paragraph with a run for heading formatting instead of a style
        heading_para = doc.add_paragraph()
        heading_para.add_run("Technical Implementation Recommendations").bold = True
        
        if dialog_docs and 'bestPractices' in dialog_docs[0]['results']['accessibility']['tests']['floating_dialogs']['dialogs']['documentation']:
            best_practices = dialog_docs[0]['results']['accessibility']['tests']['floating_dialogs']['dialogs']['documentation']['bestPractices']
            
            for category, practices in best_practices.items():
                # Use individual runs with bold formatting instead of styles
                category_para = doc.add_paragraph()
                category_para.add_run(f"{category.replace('_', ' ').title()}").bold = True
                
                for practice in practices:
                    doc.add_paragraph(practice, style='List Bullet')
        else:
            # Fallback recommendations if documentation is unavailable
            # Use individual runs with bold formatting instead of styles
            dialog_para = doc.add_paragraph()
            dialog_para.add_run("Dialog Structure").bold = True
            
            doc.add_paragraph("Use native <dialog> element with the showModal() method where supported", style='List Bullet')
            doc.add_paragraph("Add role='dialog' or role='alertdialog' for screen reader announcement", style='List Bullet')
            doc.add_paragraph("Use aria-modal='true' for modal dialogs to limit screen reader focus", style='List Bullet')
            
            # Use individual runs with bold formatting instead of styles
            focus_para = doc.add_paragraph()
            focus_para.add_run("Focus Management").bold = True
            
            doc.add_paragraph("Move focus to the dialog when it opens, typically to the first focusable element", style='List Bullet')
            doc.add_paragraph("Trap keyboard focus within modal dialogs", style='List Bullet')
            doc.add_paragraph("When dialog closes, return focus to the element that opened it", style='List Bullet')
            
            # Use individual runs with bold formatting instead of styles
            content_para = doc.add_paragraph()
            content_para.add_run("Dialog Content").bold = True
            
            doc.add_paragraph("Include a descriptive heading (h2) at the start of the dialog content", style='List Bullet')
            doc.add_paragraph("Provide a visible close button with a clear accessible name", style='List Bullet')
            doc.add_paragraph("Support Escape key for closing the dialog", style='List Bullet')

    else:
        doc.add_paragraph("No floating dialog accessibility issues were found.")


    #################################
    # Focus Management (general)
    #################################

    doc.add_page_break()
    h2 = doc.add_heading('Focus Management (General)', level=2)
    h2.style = doc.styles['Heading 2']

    # Query for pages with focus management information
    pages_with_focus = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.focus_management.focus_management": {"$exists": True}
        },
        {
            "url": 1,
            "results.accessibility.tests.focus_management.focus_management": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize tracking
    site_data = {}
    url_data = {}  # Track data by individual URL
    total_interactive_elements = 0
    total_violations = 0
    total_breakpoints_tested = 0

    # First, add a description of the test based on the documentation
    if pages_with_focus:
        first_page = pages_with_focus[0]
        test_doc = first_page.get('results', {}).get('accessibility', {}).get('tests', {}).get(
            'focus_management', {}).get('focus_management', {}).get('test_documentation', {})
        
        if test_doc:
            doc.add_paragraph(f"{test_doc.get('description', '')}", style='Normal')
            wcag_criteria = test_doc.get('wcag_criteria', [])
            if wcag_criteria:
                doc.add_paragraph("WCAG Success Criteria:", style='Normal')
                for criterion in wcag_criteria:
                    doc.add_paragraph(f"• {criterion}", style='List Bullet')
            doc.add_paragraph()

    # Process each page
    for page in pages_with_focus:
        try:
            url = page['url']
            domain = url.replace('http://', '').replace('https://', '').split('/')[0]
            focus_data = page['results']['accessibility']['tests']['focus_management']['focus_management']
            
            # Initialize domain tracking
            if domain not in site_data:
                site_data[domain] = {
                    "total_violations": 0,
                    "breakpoints_tested": 0,
                    "urls": set(),
                    "tests": {
                        "focus_outline_presence": {"violations": 0, "elements": set()},
                        "focus_outline_contrast": {"violations": 0, "elements": set()},
                        "focus_outline_offset": {"violations": 0, "elements": set()},
                        "hover_feedback": {"violations": 0, "elements": set()},
                        "focus_obscurement": {"violations": 0, "elements": set()},
                        "anchor_target_tabindex": {"violations": 0, "elements": set()}
                    }
                }
            
            # Initialize URL tracking
            if url not in url_data:
                url_data[url] = {
                    "domain": domain,
                    "total_violations": 0,
                    "breakpoints_tested": 0,
                    "tests": {
                        "focus_outline_presence": {"violations": 0, "elements": set()},
                        "focus_outline_contrast": {"violations": 0, "elements": set()},
                        "focus_outline_offset": {"violations": 0, "elements": set()},
                        "hover_feedback": {"violations": 0, "elements": set()},
                        "focus_obscurement": {"violations": 0, "elements": set()},
                        "anchor_target_tabindex": {"violations": 0, "elements": set()}
                    }
                }

            # Get metadata
            metadata = focus_data.get('metadata', {})
            url_violations = metadata.get('total_violations_found', 0)
            
            total_violations += url_violations
            site_data[domain]["total_violations"] += url_violations
            url_data[url]["total_violations"] = url_violations
            
            breakpoints_tested = metadata.get('total_breakpoints_tested', 0)
            total_breakpoints_tested += breakpoints_tested
            site_data[domain]["breakpoints_tested"] = max(site_data[domain]["breakpoints_tested"], breakpoints_tested)
            url_data[url]["breakpoints_tested"] = breakpoints_tested
            
            # Add URL to domain list
            site_data[domain]["urls"].add(url)

            # Process each test
            tests = focus_data.get('tests', {})
            for test_name, test_data in tests.items():
                if test_name in site_data[domain]["tests"]:
                    # Get summary data
                    summary = test_data.get('summary', {})
                    violations = summary.get('total_violations', 0)
                    
                    # Update site data
                    site_data[domain]["tests"][test_name]["violations"] += violations
                    
                    # Update URL data
                    url_data[url]["tests"][test_name]["violations"] = violations
                    
                    # Track affected elements
                    elements = test_data.get('elements_affected', [])
                    if isinstance(elements, list):
                        site_data[domain]["tests"][test_name]["elements"].update(elements)
                        url_data[url]["tests"][test_name]["elements"].update(elements)

        except Exception as e:
            print(f"Error processing page {page.get('url', 'unknown')}: {str(e)}")
            continue

    if pages_with_focus:
        # Overall Statistics
        doc.add_paragraph("Focus Management Statistics:", style='Normal')
        stats_table = doc.add_table(rows=3, cols=2)
        stats_table.style = 'Table Grid'
        
        rows = stats_table.rows
        rows[0].cells[0].text = "Pages Tested"
        rows[0].cells[1].text = str(len(pages_with_focus))
        rows[1].cells[0].text = "Total Breakpoints Tested (across all pages)"
        rows[1].cells[1].text = str(total_breakpoints_tested)
        rows[2].cells[0].text = "Total Violations Found"
        rows[2].cells[1].text = str(total_violations)
        
        format_table_text(stats_table)

        # Tests information
        doc.add_paragraph()
        doc.add_paragraph("Tests Performed:", style='Normal')
        
        if first_page:
            tests_performed = test_doc.get('tests_performed', [])
            test_table = doc.add_table(rows=len(tests_performed) + 1, cols=3)
            test_table.style = 'Table Grid'
            
            # Add headers
            headers = test_table.rows[0].cells
            headers[0].text = "Test Name"
            headers[1].text = "Description"
            headers[2].text = "Success Criteria"
            
            # Add data for each test
            for i, test_info in enumerate(tests_performed, 1):
                row = test_table.rows[i].cells
                row[0].text = test_info.get('name', '')
                row[1].text = test_info.get('description', '')
                row[2].text = test_info.get('success_criteria', '')
                
            format_table_text(test_table)

        # Map test IDs to more readable names
        test_name_map = {
            "focus_outline_presence": "Missing Focus Outlines",
            "focus_outline_contrast": "Insufficient Outline Contrast",
            "focus_outline_offset": "Insufficient Outline Offset/Width",
            "hover_feedback": "Insufficient Hover Feedback",
            "focus_obscurement": "Obscured Focus Outlines",
            "anchor_target_tabindex": "Improper Local Target Configuration"
        }

        # Detailed Issues Summary
        doc.add_paragraph()
        doc.add_paragraph("Focus Management Issues by Test Type:", style='Normal')
        
        # Calculate totals for each test type
        test_totals = {}
        for test_id, display_name in test_name_map.items():
            total_violations = sum(site["tests"][test_id]["violations"] for site in site_data.values())
            affected_sites = sum(1 for site in site_data.values() if site["tests"][test_id]["violations"] > 0)
            affected_pages = sum(1 for url in url_data.values() if url["tests"][test_id]["violations"] > 0)
            
            # Count unique elements across all sites
            all_elements = set()
            for site in site_data.values():
                all_elements.update(site["tests"][test_id]["elements"])
            
            test_totals[test_id] = {
                "violations": total_violations,
                "affected_sites": affected_sites,
                "affected_pages": affected_pages,
                "unique_elements": len(all_elements)
            }
        
        # Create table for test summaries
        summary_table = doc.add_table(rows=len(test_name_map) + 1, cols=5)
        summary_table.style = 'Table Grid'
        
        # Add headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue Type"
        headers[1].text = "Total Violations"
        headers[2].text = "Pages Affected"
        headers[3].text = "Sites Affected"
        headers[4].text = "% of Total Sites"
        
        # Add data for each test type
        row_idx = 1
        for test_id, display_name in test_name_map.items():
            totals = test_totals[test_id]
            row = summary_table.rows[row_idx].cells
            row[0].text = display_name
            row[1].text = str(totals["violations"])
            row[2].text = str(totals["affected_pages"])
            row[3].text = str(totals["affected_sites"])
            row[4].text = f"{(totals['affected_sites'] / len(site_data) * 100):.1f}%" if site_data else "0%"
            row_idx += 1
        
        format_table_text(summary_table)
        
        # Detailed Report by Domain/URL
        doc.add_paragraph()
        doc.add_heading("Detailed Focus Management Issues by Domain", level=4)
        
        # Sort domains by total violations
        sorted_domains = sorted(site_data.items(), key=lambda x: x[1]["total_violations"], reverse=True)
        
        for domain, domain_data in sorted_domains:
            # Add domain heading
            doc.add_paragraph()
            doc.add_paragraph(f"Domain: {domain}", style='Heading 5')
            
            # Add domain summary
            domain_summary = doc.add_paragraph(style='Normal')
            domain_summary.add_run(f"Total Violations: {domain_data['total_violations']}").bold = True
            domain_summary.add_run(f" | URLs: {len(domain_data['urls'])} | Breakpoints tested: {domain_data['breakpoints_tested']}")
            
            # Get URLs for this domain and sort by violations
            domain_urls = [url for url, url_info in url_data.items() if url_info["domain"] == domain]
            sorted_urls = sorted(domain_urls, key=lambda u: url_data[u]["total_violations"], reverse=True)
            
            # Create table for URLs in this domain
            url_table = doc.add_table(rows=len(sorted_urls) + 1, cols=7)
            url_table.style = 'Table Grid'
            
            # Add headers
            url_headers = url_table.rows[0].cells
            url_headers[0].text = "URL"
            url_headers[1].text = "Total Violations"
            url_headers[2].text = "Missing Outlines"
            url_headers[3].text = "Hover Issues"
            url_headers[4].text = "Contrast Issues"
            url_headers[5].text = "Offset Issues"
            url_headers[6].text = "Other Issues"
            
            # Add URL rows
            for i, url in enumerate(sorted_urls, 1):
                url_info = url_data[url]
                row = url_table.rows[i].cells
                
                # Display full URL for copy/paste
                row[0].text = url
                row[1].text = str(url_info["total_violations"])
                row[2].text = str(url_info["tests"]["focus_outline_presence"]["violations"])
                row[3].text = str(url_info["tests"]["hover_feedback"]["violations"])
                row[4].text = str(url_info["tests"]["focus_outline_contrast"]["violations"])
                row[5].text = str(url_info["tests"]["focus_outline_offset"]["violations"])
                
                # Combine other issues
                other_issues = (
                    url_info["tests"]["focus_obscurement"]["violations"] + 
                    url_info["tests"]["anchor_target_tabindex"]["violations"]
                )
                row[6].text = str(other_issues)
            
            format_table_text(url_table)
    else:
        doc.add_paragraph("No focus management data available in the database.", style='Normal')


    #################################
    # Fonts
    #################################

    doc.add_page_break()
    h2 = doc.add_heading('Fonts', level=2)
    h2.style = doc.styles['Heading 2']
    doc.add_paragraph()

    # System fonts definition
    SYSTEM_FONTS = [
        "Arial", "Helvetica", "Times New Roman", "Times", "Courier New", 
        "Courier", "Verdana", "Georgia", "Palatino", "Garamond", "Bookman",
        "Tahoma", "Trebuchet MS", "Impact", "Comic Sans MS", "Webdings", 
        "Symbol", "Calibri", "Cambria", "Segoe UI"
    ]



    # Add explanation
    doc.add_paragraph("""
Font choices and typography implementation significantly impact content accessibility. This analysis covers font selection, text size, line height, text alignment, and overall typography. Key factors include heading hierarchy, text readability, and proper use of font styling.
    """.strip())

    doc.add_paragraph("Typography accessibility concerns include:", style='Normal')
    doc.add_paragraph("Text size and readability issues", style='List Bullet')
    doc.add_paragraph("Line height and spacing problems", style='List Bullet')
    doc.add_paragraph("Text alignment affecting readability", style='List Bullet')
    doc.add_paragraph("Improper use of italic and bold text", style='List Bullet')
    doc.add_paragraph("Heading size hierarchy issues", style='List Bullet')
    doc.add_paragraph("Font selection and implementation", style='List Bullet')

    # 6. Recommendations
    doc.add_paragraph()
    doc.add_paragraph("Recommendations:", style='Normal')
    
    doc.add_paragraph("Typography:", style='List Bullet')
    doc.add_paragraph("Ensure body text is at least 16px", style='List Bullet 2')
    doc.add_paragraph("Maintain line height of at least 1.5", style='List Bullet 2')
    doc.add_paragraph("Avoid justified text alignment", style='List Bullet 2')
    doc.add_paragraph("Use left-aligned text for better readability", style='List Bullet 2')
    doc.add_paragraph("Limit use of italics", style='List Bullet 2')
    
    doc.add_paragraph("Font Usage:", style='List Bullet')
    doc.add_paragraph("Limit number of different fonts", style='List Bullet 2')
    doc.add_paragraph("Provide appropriate fallback fonts", style='List Bullet 2')
    doc.add_paragraph("Ensure proper font loading", style='List Bullet 2')
    doc.add_paragraph("Use system fonts where possible", style='List Bullet 2')
    
    doc.add_paragraph("Heading Structure:", style='List Bullet')
    doc.add_paragraph("Maintain clear size hierarchy", style='List Bullet 2')
    doc.add_paragraph("Ensure headings are larger than body text", style='List Bullet 2')
    doc.add_paragraph("Use consistent heading sizes across pages", style='List Bullet 2')
    doc.add_paragraph()
    
    # Query for pages with font information
    pages_with_fonts = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.fonts.font_analysis": {"$exists": True}
        },
        {
            "url": 1,
            "results.accessibility.tests.fonts.font_analysis": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize tracking structures
    font_usage = {}
    site_fonts = {}
    total_font_count = 0
    heading_sizes = []

    # Typography issue tracking
    typography_issues = {
        "small_text": {"name": "Small text", "pages": set(), "domains": set()},
        "small_line_height": {"name": "Small line height", "pages": set(), "domains": set()},
        "justified_text": {"name": "Justified text", "pages": set(), "domains": set()},
        "right_aligned": {"name": "Right-aligned text", "pages": set(), "domains": set()},
        "italic_text": {"name": "Italic text usage", "pages": set(), "domains": set()},
        "bold_larger_than_headings": {"name": "Bold text larger than headings", "pages": set(), "domains": set()}
    }

    # CSS Variable tracking
    css_var_usage = {
        'awb-text-font-family': set(),
        'body_typography-font-family': set(),
        'fontsBaseFamily': set(),
        'h1_typography-font-family': set(),
        'h2_typography-font-family': set(),
        'h3_typography-font-family': set(),
        'homepage-title-font': set()
    }

    # Process each page
    for page in pages_with_fonts:
        try:
            domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
            font_data = page['results']['accessibility']['tests']['fonts']['font_analysis']
            
            # Initialize domain in site_fonts if not present
            if domain not in site_fonts:
                site_fonts[domain] = {
                    'fonts': set(),
                    'css_vars': set(),
                    'system_fonts': set(),
                    'web_fonts': set(),
                    'smallest_heading': None,
                    'typography_issues': set()
                }

            # Process accessibility data
            accessibility_data = font_data.get('accessibility', {})
            tests = accessibility_data.get('tests', {})
            
            # Track smallest heading size
            smallest_heading = accessibility_data.get('smallestHeadingSize')
            if smallest_heading:
                heading_sizes.append(smallest_heading)
                site_fonts[domain]['smallest_heading'] = smallest_heading

            # Check typography issues
            if tests.get('hasSmallText'):
                typography_issues['small_text']['pages'].add(page['url'])
                typography_issues['small_text']['domains'].add(domain)
                site_fonts[domain]['typography_issues'].add('small_text')

            if tests.get('hasSmallLineHeight'):
                typography_issues['small_line_height']['pages'].add(page['url'])
                typography_issues['small_line_height']['domains'].add(domain)
                site_fonts[domain]['typography_issues'].add('small_line_height')

            if tests.get('hasJustifiedText'):
                typography_issues['justified_text']['pages'].add(page['url'])
                typography_issues['justified_text']['domains'].add(domain)
                site_fonts[domain]['typography_issues'].add('justified_text')

            if tests.get('hasRightAlignedText'):
                typography_issues['right_aligned']['pages'].add(page['url'])
                typography_issues['right_aligned']['domains'].add(domain)
                site_fonts[domain]['typography_issues'].add('right_aligned')

            if tests.get('hasItalicText'):
                typography_issues['italic_text']['pages'].add(page['url'])
                typography_issues['italic_text']['domains'].add(domain)
                site_fonts[domain]['typography_issues'].add('italic_text')

            if tests.get('hasBoldNonHeadingLargerThanHeadings'):
                typography_issues['bold_larger_than_headings']['pages'].add(page['url'])
                typography_issues['bold_larger_than_headings']['domains'].add(domain)
                site_fonts[domain]['typography_issues'].add('bold_larger_than_headings')

            # Process fonts
            fonts = font_data.get('fonts', {})
            for font_name in fonts.keys():
                # Skip generic families
                if font_name.lower() in ['inherit', 'sans-serif', 'monospace']:
                    continue
                
                # Handle CSS variables
                if font_name.startswith('var('):
                    site_fonts[domain]['css_vars'].add(font_name)
                    for var_name in css_var_usage.keys():
                        if var_name in font_name:
                            css_var_usage[var_name].add(domain)
                    continue
                
                # Track regular fonts
                if font_name not in font_usage:
                    font_usage[font_name] = {"domains": set()}
                
                font_usage[font_name]["domains"].add(domain)
                site_fonts[domain]['fonts'].add(font_name)
                
                # Categorize as system or web font
                if any(sf.lower() in font_name.lower() for sf in SYSTEM_FONTS):
                    site_fonts[domain]['system_fonts'].add(font_name)
                else:
                    site_fonts[domain]['web_fonts'].add(font_name)

            total_font_count += font_data.get('totalFonts', 0)
                
        except Exception as e:
            print(f"Error processing page {page.get('url', 'unknown')}: {str(e)}")
            continue

    if pages_with_fonts:
        # 1. Typography Issues Section
        doc.add_paragraph()
        doc.add_paragraph("Typography Accessibility Issues:", style='Normal')
        
        active_issues = {flag: data for flag, data in typography_issues.items() 
                        if len(data['pages']) > 0}
        
        if active_issues:
            issues_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
            issues_table.style = 'Table Grid'

            headers = issues_table.rows[0].cells
            headers[0].text = "Issue"
            headers[1].text = "Pages Affected"
            headers[2].text = "Sites Affected"
            headers[3].text = "% of Total Sites"

            for i, (flag, data) in enumerate(active_issues.items(), 1):
                row = issues_table.rows[i].cells
                row[0].text = data['name']
                row[1].text = str(len(data['pages']))
                row[2].text = str(len(data['domains']))
                row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

            format_table_text(issues_table)

        # 2. Heading Size Analysis
        doc.add_paragraph()
        doc.add_paragraph("Heading Size Analysis:", style='Normal')
        if heading_sizes:
            doc.add_paragraph(f"Smallest heading size detected: {min(heading_sizes)}px")
            doc.add_paragraph(f"Average smallest heading size: {sum(heading_sizes) / len(heading_sizes):.1f}px")
        
        # 3. Font Usage Summary
        doc.add_paragraph()
        doc.add_paragraph("Font Usage Summary:", style='Normal')
        
        sorted_fonts = sorted(font_usage.items(), 
                            key=lambda x: (len(x[1]["domains"]), x[0].lower()), 
                            reverse=True)
        
        if sorted_fonts:
            fonts_table = doc.add_table(rows=len(sorted_fonts) + 1, cols=3)
            fonts_table.style = 'Table Grid'

            headers = fonts_table.rows[0].cells
            headers[0].text = "Font Name"
            headers[1].text = "Sites Using"
            headers[2].text = "% of Total Sites"

            for i, (font_name, data) in enumerate(sorted_fonts, 1):
                row = fonts_table.rows[i].cells
                row[0].text = font_name
                row[1].text = str(len(data["domains"]))
                row[2].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

            format_table_text(fonts_table)

        # 4. CSS Variable Usage
        doc.add_paragraph()
        doc.add_paragraph("CSS Variable Font Usage:", style='Normal')
        
        active_vars = {name: domains for name, domains in css_var_usage.items() if domains}
        if active_vars:
            var_table = doc.add_table(rows=len(active_vars) + 1, cols=3)
            var_table.style = 'Table Grid'
            
            headers = var_table.rows[0].cells
            headers[0].text = "CSS Variable"
            headers[1].text = "Sites Using"
            headers[2].text = "% of Total Sites"
            
            for i, (var_name, domains) in enumerate(sorted(active_vars.items()), 1):
                row = var_table.rows[i].cells
                row[0].text = var_name
                row[1].text = str(len(domains))
                row[2].text = f"{(len(domains) / len(total_domains) * 100):.1f}%"
            
            format_table_text(var_table)

        # 5. Site-Specific Analysis
        doc.add_paragraph()
        doc.add_paragraph("Analysis by Site:", style='Normal')
        
        for domain, data in sorted(site_fonts.items()):
            if data['fonts'] or data['css_vars'] or data['typography_issues']:
                doc.add_paragraph()
                doc.add_paragraph(f"{domain}:", style='List Bullet')
                
                # Create detailed table for this site
                details = []
                if data['system_fonts']:
                    details.append(("System Fonts", f"{len(data['system_fonts'])} fonts"))
                if data['web_fonts']:
                    details.append(("Web Fonts", f"{len(data['web_fonts'])} fonts"))
                if data['css_vars']:
                    details.append(("CSS Variables", f"{len(data['css_vars'])} variables"))
                if data['typography_issues']:
                    details.append(("Typography Issues", f"{len(data['typography_issues'])} issues"))
                if data['smallest_heading']:
                    details.append(("Smallest Heading", f"{data['smallest_heading']}px"))

                site_table = doc.add_table(rows=len(details) + 1, cols=2)
                site_table.style = 'Table Grid'
                
                headers = site_table.rows[0].cells
                headers[0].text = "Category"
                headers[1].text = "Details"
                
                for i, (category, detail) in enumerate(details, 1):
                    row = site_table.rows[i].cells
                    row[0].text = category
                    row[1].text = detail
                
                format_table_text(site_table)
    else:
        doc.add_paragraph("No font usage data was found.")


    #################################
    # Forms
    #################################

    doc.add_page_break()
    h2 = doc.add_heading('Forms', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
    Forms must be accessible to all users, including those using screen readers or keyboard navigation. Proper labeling, structure, and organization are essential for form accessibility. Forms should have clear instructions, properly associated labels, and appropriate error handling.
    """.strip())

    doc.add_paragraph("Common form accessibility issues include:", style='Normal')

    doc.add_paragraph("Form inputs without proper labels", style='List Bullet')
    doc.add_paragraph("Reliance on placeholders instead of labels", style='List Bullet')
    doc.add_paragraph("Forms without proper heading structure", style='List Bullet')
    doc.add_paragraph("Forms placed outside landmark regions", style='List Bullet')
    doc.add_paragraph("Input fields with insufficient contrast", style='List Bullet')
    doc.add_paragraph("Layout issues affecting form usability", style='List Bullet')

    # Add recommendations
    doc.add_paragraph()
    doc.add_paragraph("Recommendations for Form Accessibility:", style='Normal')
    
    doc.add_paragraph("Ensure all form controls have properly associated labels", style='List Bullet')
    doc.add_paragraph("Use labels instead of relying solely on placeholders", style='List Bullet')
    doc.add_paragraph("Include proper heading structure for forms", style='List Bullet')
    doc.add_paragraph("Place forms within appropriate landmark regions", style='List Bullet')
    doc.add_paragraph("Maintain sufficient contrast for all form elements", style='List Bullet')
    doc.add_paragraph("Ensure proper spacing and layout for form controls", style='List Bullet')
    doc.add_paragraph("Provide clear error messages and validation feedback", style='List Bullet')
    doc.add_paragraph("Ensure forms are keyboard accessible", style='List Bullet')

    doc.add_paragraph("NOTE: There are many more accessibility issues surrounding form error handling, form submission, and pagination not tested here", style='Normal')


    doc.add_paragraph()

    # Query for pages with form issues
    pages_with_form_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.forms.forms.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.forms.forms.pageFlags.hasInputsWithoutLabels": True},
                {"results.accessibility.tests.forms.forms.pageFlags.hasPlaceholderOnlyInputs": True},
                {"results.accessibility.tests.forms.forms.pageFlags.hasFormsWithoutHeadings": True},
                {"results.accessibility.tests.forms.forms.pageFlags.hasFormsOutsideLandmarks": True},
                {"results.accessibility.tests.forms.forms.pageFlags.hasContrastIssues": True},
                {"results.accessibility.tests.forms.forms.pageFlags.hasLayoutIssues": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.forms.forms": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for different form issues
    form_issues = {
        "missing_labels": {
            "name": "Inputs without labels",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "placeholder_only": {
            "name": "Placeholder-only inputs",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "no_headings": {
            "name": "Forms without headings",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "outside_landmarks": {
            "name": "Forms outside landmarks",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "contrast_issues": {
            "name": "Input contrast issues",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "layout_issues": {
            "name": "Form layout issues",
            "pages": set(),
            "domains": set(),
            "count": 0
        }
    }

    # Process each page
    total_forms = 0
    for page in pages_with_form_issues:
        try:
            domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
            form_data = page['results']['accessibility']['tests']['forms']['forms']
            flags = form_data.get('pageFlags', {})
            summary = form_data.get('details', {}).get('summary', {})
            
            # Update total forms count
            total_forms += summary.get('totalForms', 0)
            
            # Check inputs without labels
            if flags.get('hasInputsWithoutLabels'):
                form_issues['missing_labels']['pages'].add(page['url'])
                form_issues['missing_labels']['domains'].add(domain)
                form_issues['missing_labels']['count'] += summary.get('inputsWithoutLabels', 0)
            
            # Check placeholder-only inputs
            if flags.get('hasPlaceholderOnlyInputs'):
                form_issues['placeholder_only']['pages'].add(page['url'])
                form_issues['placeholder_only']['domains'].add(domain)
                form_issues['placeholder_only']['count'] += summary.get('inputsWithPlaceholderOnly', 0)
            
            # Check forms without headings
            if flags.get('hasFormsWithoutHeadings'):
                form_issues['no_headings']['pages'].add(page['url'])
                form_issues['no_headings']['domains'].add(domain)
                form_issues['no_headings']['count'] += summary.get('formsWithoutHeadings', 0)
            
            # Check forms outside landmarks
            if flags.get('hasFormsOutsideLandmarks'):
                form_issues['outside_landmarks']['pages'].add(page['url'])
                form_issues['outside_landmarks']['domains'].add(domain)
                form_issues['outside_landmarks']['count'] += summary.get('formsOutsideLandmarks', 0)
            
            # Check contrast issues
            if flags.get('hasContrastIssues'):
                form_issues['contrast_issues']['pages'].add(page['url'])
                form_issues['contrast_issues']['domains'].add(domain)
                form_issues['contrast_issues']['count'] += summary.get('inputsWithContrastIssues', 0)
            
            # Check layout issues
            if flags.get('hasLayoutIssues'):
                form_issues['layout_issues']['pages'].add(page['url'])
                form_issues['layout_issues']['domains'].add(domain)
                form_issues['layout_issues']['count'] += summary.get('inputsWithLayoutIssues', 0)
                
        except Exception as e:
            print(f"Error processing page {page.get('url', 'unknown')}: {str(e)}")
            continue

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in form_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=5)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue Type"
        headers[1].text = "Number of Occurrences"
        headers[2].text = "Pages Affected"
        headers[3].text = "Sites Affected"
        headers[4].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(data['count'])
            row[2].text = str(len(data['pages']))
            row[3].text = str(len(data['domains']))
            row[4].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add statistics
        doc.add_paragraph()
        doc.add_paragraph("Form Statistics:", style='Normal')
        doc.add_paragraph(f"Total number of forms across all pages: {total_forms}")

 
        # Add domain details for each issue type
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)

    else:
        doc.add_paragraph("No form accessibility issues were found.")

    #######################################
    # Headings
    #######################################

    doc.add_paragraph()
    h2 = doc.add_heading('Headings', level=2)
    h2.style = doc.styles['Heading 2']
 

    # Add explanation
    doc.add_paragraph("""
    Proper heading structure is essential for accessibility as it helps users understand the organization of content and navigate pages effectively. Headings should follow a logical hierarchy and accurately reflect the content structure. Screen reader users often navigate by headings, making proper structure crucial.
    """.strip())

    doc.add_paragraph("Common heading structure issues include:", style='Normal')

    doc.add_paragraph("Missing or multiple main headings (H1)", style='List Bullet')
    doc.add_paragraph("Skipped heading levels creating hierarchy gaps", style='List Bullet')
    doc.add_paragraph("Headings placed before the main content", style='List Bullet')
    doc.add_paragraph("Visual styling that doesn't match heading levels", style='List Bullet')

    doc.add_paragraph()

    # Query for pages with heading issues
    pages_with_heading_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.headings.headings.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.headings.headings.pageFlags.missingH1": True},
                {"results.accessibility.tests.headings.headings.pageFlags.multipleH1s": True},
                {"results.accessibility.tests.headings.headings.pageFlags.hasHierarchyGaps": True},
                {"results.accessibility.tests.headings.headings.pageFlags.hasHeadingsBeforeMain": True},
                {"results.accessibility.tests.headings.headings.pageFlags.hasVisualHierarchyIssues": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.headings.headings": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for different heading issues
    heading_issues = {
        "missing_h1": {"name": "Missing H1", "pages": set(), "domains": set()},
        "multiple_h1": {"name": "Multiple H1s", "pages": set(), "domains": set()},
        "hierarchy_gaps": {"name": "Hierarchy gaps", "pages": set(), "domains": set(), "count": 0},
        "headings_before_main": {"name": "Headings before main", "pages": set(), "domains": set(), "count": 0},
        "visual_hierarchy": {"name": "Visual hierarchy issues", "pages": set(), "domains": set(), "count": 0}
    }

    # Process each page
    total_headings = 0
    for page in pages_with_heading_issues:
        try:
            domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
            heading_data = page['results']['accessibility']['tests']['headings']['headings']
            flags = heading_data.get('pageFlags', {})
            
            # The issue is here - details might be in a different place or structure than expected
            details = flags.get('details', {})
            if not details:  # If details not found in flags, try the main heading_data
                details = heading_data.get('details', {})
            
            summary = heading_data.get('details', {}).get('summary', {})
            
            # Update total headings count
            headings_count = summary.get('totalHeadings', 0)
            if isinstance(headings_count, (int, float)):
                total_headings += headings_count
            
            # Check missing H1
            if flags.get('missingH1'):
                heading_issues['missing_h1']['pages'].add(page['url'])
                heading_issues['missing_h1']['domains'].add(domain)
            
            # Check multiple H1s
            if flags.get('multipleH1s'):
                heading_issues['multiple_h1']['pages'].add(page['url'])
                heading_issues['multiple_h1']['domains'].add(domain)
            
            # Check hierarchy gaps
            if flags.get('hasHierarchyGaps'):
                heading_issues['hierarchy_gaps']['pages'].add(page['url'])
                heading_issues['hierarchy_gaps']['domains'].add(domain)
                
                # Fix for the potential list issue
                hierarchy_gaps = details.get('hierarchyGaps', 0)
                if isinstance(hierarchy_gaps, list):
                    heading_issues['hierarchy_gaps']['count'] += len(hierarchy_gaps)
                elif isinstance(hierarchy_gaps, (int, float)):
                    heading_issues['hierarchy_gaps']['count'] += hierarchy_gaps
            
            # Check headings before main
            if flags.get('hasHeadingsBeforeMain'):
                heading_issues['headings_before_main']['pages'].add(page['url'])
                heading_issues['headings_before_main']['domains'].add(domain)
                
                # Fix for the potential list issue
                headings_before_main = details.get('headingsBeforeMain', 0)
                if isinstance(headings_before_main, list):
                    heading_issues['headings_before_main']['count'] += len(headings_before_main)
                elif isinstance(headings_before_main, (int, float)):
                    heading_issues['headings_before_main']['count'] += headings_before_main
            
            # Check visual hierarchy issues
            if flags.get('hasVisualHierarchyIssues'):
                heading_issues['visual_hierarchy']['pages'].add(page['url'])
                heading_issues['visual_hierarchy']['domains'].add(domain)
                
                # Fix for the potential list issue
                visual_hierarchy_issues = details.get('visualHierarchyIssues', 0)
                if isinstance(visual_hierarchy_issues, list):
                    heading_issues['visual_hierarchy']['count'] += len(visual_hierarchy_issues)
                elif isinstance(visual_hierarchy_issues, (int, float)):
                    heading_issues['visual_hierarchy']['count'] += visual_hierarchy_issues
                
        except Exception as e:
            print(f"Error processing page {page.get('url', 'unknown')}: {str(e)}")
            continue

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in heading_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=5)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue Type"
        headers[1].text = "Number of Occurrences"
        headers[2].text = "Pages Affected"
        headers[3].text = "Sites Affected"
        headers[4].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(data.get('count', len(data['pages'])))
            row[2].text = str(len(data['pages']))
            row[3].text = str(len(data['domains']))
            row[4].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add statistics
        doc.add_paragraph()
        doc.add_paragraph("Heading Statistics:", style='Normal')
        doc.add_paragraph(f"Total number of headings across all pages: {total_headings}")

        # Add recommendations
        doc.add_paragraph()
        doc.add_paragraph("Recommendations for Heading Structure:", style='Normal')
        
        doc.add_paragraph("Use exactly one H1 heading per page as the main title", style='List Bullet')
        doc.add_paragraph("Maintain proper heading hierarchy without skipping levels", style='List Bullet')
        doc.add_paragraph("Ensure heading levels match their visual presentation", style='List Bullet')
        doc.add_paragraph("Place meaningful headings in the main content area", style='List Bullet')
        doc.add_paragraph("Use headings to create a clear content outline", style='List Bullet')
        doc.add_paragraph("Make heading text descriptive and meaningful", style='List Bullet')

        # Add example of proper heading structure
        doc.add_paragraph()
        doc.add_paragraph("Example of proper heading structure:", style='Normal')
        doc.add_paragraph("H1: Main page title", style='List Bullet')
        doc.add_paragraph("    H2: Major section", style='List Bullet 2')
        doc.add_paragraph("        H3: Subsection", style='List Bullet 3')
        doc.add_paragraph("        H3: Another subsection", style='List Bullet 3')
        doc.add_paragraph("    H2: Another major section", style='List Bullet 2')

        # Add domain details for each issue type
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)

    else:
        doc.add_paragraph("No heading structure issues were found.")


    ##################################
    # Images
    ##################################
    
    doc.add_page_break()
    h2 = doc.add_heading('Images', level=2)
    h2.style = doc.styles['Heading 2']
    findings = doc.add_paragraph()


    # Add explanation
    doc.add_paragraph("""
    Images must be accessible to all users, including those using screen readers. This requires proper alternative text descriptions and appropriate ARIA roles. Images that convey information need descriptive alt text, while decorative images should be properly marked as such.
    """.strip())

    doc.add_paragraph("Common image accessibility issues include:", style='Normal')

    doc.add_paragraph("Missing alternative text for informative images", style='List Bullet')
    doc.add_paragraph("Invalid or uninformative alt text", style='List Bullet')
    doc.add_paragraph("Missing ARIA roles for SVG elements", style='List Bullet')
    doc.add_paragraph("Decorative images not properly marked", style='List Bullet')

    # Add recommendations
    doc.add_paragraph()
    doc.add_paragraph("Recommendations for Image Accessibility:", style='Normal')
    
    doc.add_paragraph("Provide meaningful alt text for all informative images", style='List Bullet')
    doc.add_paragraph("Use empty alt text (alt=\"\") for decorative images", style='List Bullet')
    doc.add_paragraph("Ensure SVG elements have appropriate ARIA roles", style='List Bullet')
    doc.add_paragraph("Make sure alt text is descriptive and conveys the image's purpose", style='List Bullet')
    doc.add_paragraph("Avoid using generic text like 'image' or 'photo' in alt attributes", style='List Bullet')
    doc.add_paragraph("Include text alternatives for complex images, charts, and graphs", style='List Bullet')

    doc.add_paragraph()
 
    # Query for pages with image issues
    pages_with_image_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.images.images.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.images.images.pageFlags.hasImagesWithoutAlt": True},
                {"results.accessibility.tests.images.images.pageFlags.hasImagesWithInvalidAlt": True},
                {"results.accessibility.tests.images.images.pageFlags.hasSVGWithoutRole": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.images.images": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for different image issues
    image_issues = {
        "missing_alt": {
            "name": "Missing alt text",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "invalid_alt": {
            "name": "Invalid alt text",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "missing_role": {
            "name": "SVGs missing role",
            "pages": set(),
            "domains": set(),
            "count": 0
        }
    }

    # Process each page
    total_images = 0
    total_decorative = 0

    for page in pages_with_image_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        image_data = page['results']['accessibility']['tests']['images']['images']
        flags = image_data['pageFlags']
        details = flags['details']
        
        # Count total and decorative images
        total_images += details.get('totalImages', 0)
        total_decorative += details.get('decorativeImages', 0)
        
        # Check missing alt text
        if flags.get('hasImagesWithoutAlt'):
            image_issues['missing_alt']['pages'].add(page['url'])
            image_issues['missing_alt']['domains'].add(domain)
            image_issues['missing_alt']['count'] += details.get('missingAlt', 0)
        
        # Check invalid alt text
        if flags.get('hasImagesWithInvalidAlt'):
            image_issues['invalid_alt']['pages'].add(page['url'])
            image_issues['invalid_alt']['domains'].add(domain)
            image_issues['invalid_alt']['count'] += details.get('invalidAlt', 0)
        
        # Check missing SVG roles
        if flags.get('hasSVGWithoutRole'):
            image_issues['missing_role']['pages'].add(page['url'])
            image_issues['missing_role']['domains'].add(domain)
            image_issues['missing_role']['count'] += details.get('missingRole', 0)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in image_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=5)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue Type"
        headers[1].text = "Number of Images"
        headers[2].text = "Pages Affected"
        headers[3].text = "Sites Affected"
        headers[4].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(data['count'])
            row[2].text = str(len(data['pages']))
            row[3].text = str(len(data['domains']))
            row[4].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add image statistics
        doc.add_paragraph()
        doc.add_paragraph("Image Statistics:", style='Normal')
        stats_table = doc.add_table(rows=3, cols=2)
        stats_table.style = 'Table Grid'

        # Add statistics data
        rows = stats_table.rows
        rows[0].cells[0].text = "Total Images"
        rows[0].cells[1].text = str(total_images)
        rows[1].cells[0].text = "Decorative Images"
        rows[1].cells[1].text = str(total_decorative)
        rows[2].cells[0].text = "Informative Images"
        rows[2].cells[1].text = str(total_images - total_decorative)

        # Format the table text
        format_table_text(stats_table)

       # Add domain details for each issue type
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)

    else:
        doc.add_paragraph("No image accessibility issues were found.")

    #######################################
    # Landmarks
    #######################################

    doc.add_page_break()
    h2 = doc.add_heading('Landmarks', level=2)
    h2.style = doc.styles['Heading 2']


    # Add explanation
    doc.add_paragraph("""
    HTML landmarks provide a navigational structure that helps screen reader users understand the organization of a page's content. Properly implemented landmarks are crucial for efficient navigation and orientation. Each landmark role serves a specific purpose and should be used appropriately.
    """.strip())

    doc.add_paragraph("Common landmark roles include:", style='Normal')

    doc.add_paragraph("banner - Header content", style='List Bullet')
    doc.add_paragraph("main - Primary content area", style='List Bullet')
    doc.add_paragraph("navigation - Navigation sections", style='List Bullet')
    doc.add_paragraph("complementary - Supporting content", style='List Bullet')
    doc.add_paragraph("contentinfo - Footer content", style='List Bullet')
    doc.add_paragraph("search - Search functionality", style='List Bullet')
    doc.add_paragraph("form - Form sections", style='List Bullet')
    doc.add_paragraph("region - Distinct sections requiring labels", style='List Bullet')

    # Add recommendations
    doc.add_paragraph()
    doc.add_paragraph("Recommendations for Landmark Implementation:", style='Normal')
    
    doc.add_paragraph("Ensure all pages have the required landmarks (banner, main, contentinfo)", style='List Bullet')
    doc.add_paragraph("Provide unique names for duplicate landmarks using aria-label or aria-labelledby", style='List Bullet')
    doc.add_paragraph("Avoid nesting top-level landmarks", style='List Bullet')
    doc.add_paragraph("Ensure all content is contained within appropriate landmarks", style='List Bullet')
    doc.add_paragraph("Use semantic HTML elements with implicit landmark roles where possible", style='List Bullet')

    # Add statistics
    doc.add_paragraph()
    doc.add_paragraph(f"Total number of landmarks detected across all pages: {total_landmarks}")

    # Query for pages with landmark issues
    pages_with_landmark_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.landmarks.landmarks.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.landmarks.landmarks.pageFlags.missingRequiredLandmarks": True},
                {"results.accessibility.tests.landmarks.landmarks.pageFlags.hasDuplicateLandmarksWithoutNames": True},
                {"results.accessibility.tests.landmarks.landmarks.pageFlags.hasNestedTopLevelLandmarks": True},
                {"results.accessibility.tests.landmarks.landmarks.pageFlags.hasContentOutsideLandmarks": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.landmarks.landmarks": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for different landmark issues
    landmark_issues = {
        "missing": {
            "name": "Missing required landmarks",
            "pages": set(),
            "domains": set(),
            "details": {
                "banner": 0,
                "main": 0,
                "contentinfo": 0,
                "search": 0
            }
        },
        "duplicate": {
            "name": "Duplicate landmarks without unique names",
            "pages": set(),
            "domains": set(),
            "details": {
                "banner": 0,
                "main": 0,
                "navigation": 0,
                "complementary": 0,
                "contentinfo": 0,
                "search": 0,
                "form": 0,
                "region": 0
            }
        },
        "nested": {
            "name": "Nested top-level landmarks",
            "pages": set(),
            "domains": set()
        },
        "outside": {
            "name": "Content outside landmarks",
            "pages": set(),
            "domains": set(),
            "count": 0
        }
    }

    # Process each page
    total_landmarks = 0
    for page in pages_with_landmark_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        landmark_data = page['results']['accessibility']['tests']['landmarks']['landmarks']
        flags = landmark_data['pageFlags']
        details = flags['details']
        
        # Count total landmarks
        if 'totalLandmarks' in landmark_data.get('details', {}).get('summary', {}):
            total_landmarks += landmark_data['details']['summary']['totalLandmarks']
        
        # Check missing landmarks
        if flags.get('missingRequiredLandmarks'):
            landmark_issues['missing']['pages'].add(page['url'])
            landmark_issues['missing']['domains'].add(domain)
            missing = details.get('missingLandmarks', {})
            for landmark in ['banner', 'main', 'contentinfo', 'search']:
                if missing.get(landmark):
                    landmark_issues['missing']['details'][landmark] += 1

        # Check duplicate landmarks
        if flags.get('hasDuplicateLandmarksWithoutNames'):
            landmark_issues['duplicate']['pages'].add(page['url'])
            landmark_issues['duplicate']['domains'].add(domain)
            duplicates = details.get('duplicateLandmarks', {})
            for landmark in landmark_issues['duplicate']['details'].keys():
                if landmark in duplicates:
                    landmark_issues['duplicate']['details'][landmark] += duplicates[landmark].get('count', 0)

        # Check nested landmarks
        if flags.get('hasNestedTopLevelLandmarks'):
            landmark_issues['nested']['pages'].add(page['url'])
            landmark_issues['nested']['domains'].add(domain)

        # Check content outside landmarks
        if flags.get('hasContentOutsideLandmarks'):
            landmark_issues['outside']['pages'].add(page['url'])
            landmark_issues['outside']['domains'].add(domain)
            landmark_issues['outside']['count'] += details.get('contentOutsideLandmarksCount', 0)

    # Create summary table
    if any(len(issue['pages']) > 0 for issue in landmark_issues.values()):
        # Create main issues summary table
        summary_table = doc.add_table(rows=len(landmark_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue Type"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        row_idx = 1
        for issue_type, data in landmark_issues.items():
            if len(data['pages']) > 0:
                row = summary_table.rows[row_idx].cells
                row[0].text = data['name']
                row[1].text = str(len(data['pages']))
                row[2].text = str(len(data['domains']))
                row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"
                row_idx += 1

        # Format the table text
        format_table_text(summary_table)

        # Add specific details for missing landmarks
        if landmark_issues['missing']['pages']:
            doc.add_paragraph()
            doc.add_paragraph("Missing Required Landmarks Breakdown:", style='Normal')
            
            missing_table = doc.add_table(rows=5, cols=2)
            missing_table.style = 'Table Grid'
            
            headers = missing_table.rows[0].cells
            headers[0].text = "Landmark Type"
            headers[1].text = "Number of Pages Missing"
            
            landmarks = [("Banner", "banner"), ("Main", "main"), 
                        ("Footer", "contentinfo"), ("Search", "search")]
            
            for idx, (name, key) in enumerate(landmarks, 1):
                row = missing_table.rows[idx].cells
                row[0].text = name
                row[1].text = str(landmark_issues['missing']['details'][key])
            
            format_table_text(missing_table)

 
        # Add domain details for each issue type
        for issue_type, data in landmark_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)

    else:
        doc.add_paragraph("No landmark structure issues were found.")


    #######################################################
    # Language of page
    #######################################################

    doc.add_page_break()
    h2 = doc.add_heading('Language of Page', level=2)
    h2.style = doc.styles['Heading 2']    

    doc.add_paragraph("""
        A "lang" attribute is required in the <html> tag at the beginning of a web page so that assitive tehnology knows what the dominant language of the page is and can choose an appropriate pronounciation model for the text it finds. Even when sites are monlingual this attribute is required  as you cannot assume what the default language of a screen-reader or other assstive tech is. Without the "lang" attribute, it is likely to be this default value of the screen-reader that will determine how text is announced.
                            """.strip())
        
    findings = doc.add_paragraph()

    # If there are pages without lang attribute, list them
    pages_without_lang = list(db_connection.page_results.find(
        {"results.accessibility.tests.html_structure.html_structure.tests.hasValidLang": False},
        {"url": 1, "_id": 0}
    ).sort("url", 1))

    if pages_without_lang:
        doc.add_paragraph(f"{len(pages_without_lang)} pages found without valid language attribute:".strip())

        for page in pages_without_lang:
            doc.add_paragraph(page['url'], style='List Bullet')

    else:
        doc.add_paragraph("""
All pages have a valid lang attribute.
                        """.strip())

 
    #################################
    # Lists
    #################################

    doc.add_page_break()
    h2 = doc.add_heading('Lists', level=2)
    h2.style = doc.styles['Heading 2']
    
    # Add explanation
    doc.add_paragraph("""
Proper semantic list markup is important for accessibility. Lists should use appropriate HTML elements (ul, ol, li) rather than visual formatting to create list-like structures. Common issues include:
""".strip())

    doc.add_paragraph("Using DIVs with bullets or numbers instead of proper list elements", style='List Bullet')
    doc.add_paragraph("Empty lists that serve no semantic purpose", style='List Bullet')
    doc.add_paragraph("Custom bullet implementations that may not be accessible", style='List Bullet')
    doc.add_paragraph("Excessively deep nesting of lists", style='List Bullet')

    doc.add_paragraph()

    # Query for pages with list issues
    pages_with_list_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.lists.lists.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.lists.lists.pageFlags.hasEmptyLists": True},
                {"results.accessibility.tests.lists.lists.pageFlags.hasFakeLists": True},
                {"results.accessibility.tests.lists.lists.pageFlags.hasCustomBullets": True},
                {"results.accessibility.tests.lists.lists.pageFlags.hasDeepNesting": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.lists.lists.pageFlags": 1,
            "results.accessibility.tests.lists.lists.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    list_issues = {
        "hasEmptyLists": {"name": "Empty lists", "pages": set(), "domains": set()},
        "hasFakeLists": {"name": "Fake lists (not using proper HTML)", "pages": set(), "domains": set()},
        "hasCustomBullets": {"name": "Custom bullet implementations", "pages": set(), "domains": set()},
        "hasDeepNesting": {"name": "Excessively nested lists", "pages": set(), "domains": set()}
    }

    # Count issues
    for page in pages_with_list_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        flags = page['results']['accessibility']['tests']['lists']['lists']['pageFlags']
        
        for flag in list_issues:
            if flags.get(flag, False):
                list_issues[flag]['pages'].add(page['url'])
                list_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in list_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "List Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add domain details for each issue
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)

    else:
        doc.add_paragraph("No list markup issues were found.")


    #####################################
    # Maps
    #####################################

    doc.add_page_break()
    h2 = doc.add_heading('Maps', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
Interactive maps present particular accessibility challenges. Common issues include:
""".strip())

    doc.add_paragraph("Maps without text alternatives for the information they convey", style='List Bullet')
    doc.add_paragraph("Interactive maps that can't be operated by keyboard", style='List Bullet')
    doc.add_paragraph("Map features that aren't properly labeled for screen readers", style='List Bullet')
    doc.add_paragraph("Missing alternative ways to access location information", style='List Bullet')

    # Add recommendations
    doc.add_paragraph()
    doc.add_paragraph("Recommendations for making maps accessible:")
    doc.add_paragraph("Provide text alternatives that describe the key information the map conveys", style='List Bullet')
    doc.add_paragraph("Ensure all map controls can be operated by keyboard", style='List Bullet')
    doc.add_paragraph("Include proper ARIA labels and roles for map features", style='List Bullet')
    doc.add_paragraph("Offer alternative formats (e.g., text list of locations, address lookup)", style='List Bullet')
    doc.add_paragraph("Ensure interactive elements within maps are properly labeled", style='List Bullet')

    doc.add_paragraph()

    # Query for pages with map issues
    pages_with_map_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.maps.maps.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.maps.maps.pageFlags.hasMaps": True},
                {"results.accessibility.tests.maps.maps.pageFlags.hasMapsWithoutTitle": True},
                {"results.accessibility.tests.maps.maps.pageFlags.hasMapsWithAriaHidden": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.maps.maps.pageFlags": 1,
            "results.accessibility.tests.maps.maps.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    map_issues = {
        "hasMaps": {"name": "Pages containing maps", "pages": set(), "domains": set()},
        "hasMapsWithoutTitle": {"name": "Maps without proper titles", "pages": set(), "domains": set()},
        "hasMapsWithAriaHidden": {"name": "Maps hidden from screen readers", "pages": set(), "domains": set()}
    }

    # Count issues
    for page in pages_with_map_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        flags = page['results']['accessibility']['tests']['maps']['maps']['pageFlags']
        
        for flag in map_issues:
            if flags.get(flag, False):
                map_issues[flag]['pages'].add(page['url'])
                map_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in map_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add domain details for each issue
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)

        # Add examples if available
        for page in pages_with_map_issues:
            details = page['results']['accessibility']['tests']['maps']['maps']['details']
            if 'violations' in details and details['violations']:
                doc.add_paragraph()
                doc.add_paragraph("Examples of map accessibility issues found:")
                for violation in details['violations'][:5]:  # Show up to 5 examples
                    doc.add_paragraph(violation, style='List Bullet')
                break  # Only show examples from first page with violations

    else:
        doc.add_paragraph("No map-related accessibility issues were found.")


    #####################################
    # Menus
    #####################################

    doc.add_page_break()
    h2 = doc.add_heading('Menus', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
Navigation menus are critical components for website accessibility. They must be properly structured, labeled, and implement correct ARIA roles and attributes to ensure all users can navigate effectively. Screen reader users particularly rely on well-implemented navigation menus.
    """.strip())

    doc.add_paragraph("Common accessibility issues with navigation menus include:", style='Normal')

    doc.add_paragraph("Missing or invalid ARIA roles for navigation elements", style='List Bullet')
    doc.add_paragraph("Missing current page indicators", style='List Bullet')
    doc.add_paragraph("Missing or improper menu labels and names", style='List Bullet')
    doc.add_paragraph("Duplicate menu names causing confusion", style='List Bullet')

    # Add recommendations
    doc.add_paragraph()
    doc.add_paragraph("Recommendations for Navigation Menu Implementation:", style='Normal')
    
    doc.add_paragraph("Use proper ARIA roles (e.g., navigation, menubar, menu) for navigation elements", style='List Bullet')
    doc.add_paragraph("Implement clear current page indicators using aria-current", style='List Bullet')
    doc.add_paragraph("Ensure all navigation menus have unique, descriptive labels", style='List Bullet')
    doc.add_paragraph("Use appropriate heading levels for menu sections", style='List Bullet')
    doc.add_paragraph("Ensure keyboard navigation works properly within menus", style='List Bullet')
    doc.add_paragraph("Test menu functionality with screen readers", style='List Bullet')

    # Query for pages with menu issues
    pages_with_menu_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.menus.menus.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.menus.menus.pageFlags.hasInvalidMenuRoles": True},
                {"results.accessibility.tests.menus.menus.pageFlags.hasMenusWithoutCurrent": True},
                {"results.accessibility.tests.menus.menus.pageFlags.hasUnnamedMenus": True},
                {"results.accessibility.tests.menus.menus.pageFlags.hasDuplicateMenuNames": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.menus.menus": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    menu_issues = {
        "invalidRoles": {"name": "Invalid menu roles", "pages": set(), "domains": set(), "count": 0},
        "menusWithoutCurrent": {"name": "Missing current page indicators", "pages": set(), "domains": set(), "count": 0},
        "unnamedMenus": {"name": "Unnamed menus", "pages": set(), "domains": set(), "count": 0},
        "duplicateNames": {"name": "Duplicate menu names", "pages": set(), "domains": set(), "count": 0}
    }

    # Count issues
    total_menus = 0
    for page in pages_with_menu_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        menu_data = page['results']['accessibility']['tests']['menus']['menus']
        flags = menu_data['pageFlags']
        details = menu_data['pageFlags']['details']
        
        total_menus += details.get('totalMenus', 0)
        
        # Check each type of issue
        if flags.get('hasInvalidMenuRoles'):
            menu_issues['invalidRoles']['pages'].add(page['url'])
            menu_issues['invalidRoles']['domains'].add(domain)
            menu_issues['invalidRoles']['count'] += details.get('invalidRoles', 0)
            
        if flags.get('hasMenusWithoutCurrent'):
            menu_issues['menusWithoutCurrent']['pages'].add(page['url'])
            menu_issues['menusWithoutCurrent']['domains'].add(domain)
            menu_issues['menusWithoutCurrent']['count'] += details.get('menusWithoutCurrent', 0)
            
        if flags.get('hasUnnamedMenus'):
            menu_issues['unnamedMenus']['pages'].add(page['url'])
            menu_issues['unnamedMenus']['domains'].add(domain)
            menu_issues['unnamedMenus']['count'] += details.get('unnamedMenus', 0)
            
        if flags.get('hasDuplicateMenuNames'):
            menu_issues['duplicateNames']['pages'].add(page['url'])
            menu_issues['duplicateNames']['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in menu_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=5)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Menu Issue"
        headers[1].text = "Number of Occurrences"
        headers[2].text = "Pages Affected"
        headers[3].text = "Sites Affected"
        headers[4].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(data['count']) if flag != 'duplicateNames' else 'N/A'
            row[2].text = str(len(data['pages']))
            row[3].text = str(len(data['domains']))
            row[4].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add statistics about total menus
        doc.add_paragraph()
        doc.add_paragraph(f"Total number of navigation menus detected across all pages: {total_menus}")

        # Add domain details for each issue
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)

    else:
        doc.add_paragraph("No navigation menu accessibility issues were found.")

    #########################################
    # 'More' Controls
    #########################################
    
    doc.add_page_break()
    h2 = doc.add_heading('"More" Controls', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
Generic "Read More" or "Learn More" style links can create barriers for screen reader users who rely on link and button text to understand where a link/button will take them. When they are taken out of context:
""".strip())

    doc.add_paragraph("Users can't determine the link's purpose from the link text alone", style='List Bullet')
    doc.add_paragraph("Screen reader users may get a list of identical 'read more' links", style='List Bullet')
    doc.add_paragraph("The destination of the link isn't clear without surrounding context slowing down reading for screen-reader and screen-magnifier users", style='List Bullet')

    # Add recommendations
    doc.add_paragraph()
    doc.add_paragraph("Recommendations for improving link text:")
    doc.add_paragraph("Make link and button text descriptive of its destination or purpose", style='List Bullet')
    doc.add_paragraph("Use aria-label or visually hidden text if additional context is needed", style='List Bullet')
    doc.add_paragraph("Ensure link text makes sense when read out of context", style='List Bullet')

    doc.add_paragraph()

    # Query for pages with read more link issues
    pages_with_readmore_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.read_more_links.read_more_links.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.read_more_links.read_more_links.pageFlags.hasGenericReadMoreLinks": True},
                {"results.accessibility.tests.read_more_links.read_more_links.pageFlags.hasInvalidReadMoreLinks": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.read_more_links.read_more_links.pageFlags": 1,
            "results.accessibility.tests.read_more_links.read_more_links.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    readmore_issues = {
        "hasGenericReadMoreLinks": {"name": "Generic 'Read More' links", "pages": set(), "domains": set()},
        "hasInvalidReadMoreLinks": {"name": "Invalid implementation of 'Read More' links", "pages": set(), "domains": set()}
    }

    # Count issues
    for page in pages_with_readmore_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        flags = page['results']['accessibility']['tests']['read_more_links']['read_more_links']['pageFlags']
        
        for flag in readmore_issues:
            if flags.get(flag, False):
                readmore_issues[flag]['pages'].add(page['url'])
                readmore_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in readmore_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add domain details for each issue
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)

        # Add examples if available
        for page in pages_with_readmore_issues:
            details = page['results']['accessibility']['tests']['read_more_links']['read_more_links']['details']
            if 'items' in details and details['items']:
                doc.add_paragraph()
                doc.add_paragraph("Examples of problematic link text found:")
                for item in details['items'][:5]:  # Show up to 5 examples
                    doc.add_paragraph(item, style='List Bullet')
                break  # Only show examples from first page with issues

    else:
        doc.add_paragraph("No issues with generic 'Read More' links were found.")


    ################################################
    # Tabindex
    ################################################

    doc.add_page_break()
    h2 = doc.add_heading('Tabindex', level=2)
    h2.style = doc.styles['Heading 2']

    # Query for pages with tabindex issues
    pages_with_tabindex_issues = list(db_connection.page_results.find(
        {"results.accessibility.tests.tabindex.tabindex.pageFlags": {"$exists": True}},
        {
            "url": 1,
            "results.accessibility.tests.tabindex.tabindex.pageFlags": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    tabindex_issues = {
        "hasPositiveTabindex": {"name": "Elements with positive tabindex", "pages": set(), "domains": set()},
        "hasNonInteractiveZeroTabindex": {"name": "Non-interactive elements with tabindex=0", "pages": set(), "domains": set()},
        "hasMissingRequiredTabindex": {"name": "Interactive elements missing required tabindex", "pages": set(), "domains": set()},
        "hasSvgTabindexWarnings": {"name": "SVG elements with tabindex warnings", "pages": set(), "domains": set()}
    }    

    # Add explanation
    doc.add_paragraph("""
The tabindex attribute controls whether, and in what order elements can be focused using the keyboard. Improper use of tabindex can disrupt the natural tab order and create accessibility barriers:
""".strip())

    # Add bullet points using Word's built-in bullet style
    doc.add_paragraph("Positive tabindex values force elements into a specific tab order, which can be confusing and unpredictable", style='List Bullet')
    doc.add_paragraph("Non-interactive elements with tabindex=0 create unnecessary tab stops", style='List Bullet')
    doc.add_paragraph("Interactive elements without proper tabindex may be unreachable by keyboard", style='List Bullet')
    doc.add_paragraph("SVG elements need special consideration for keyboard accessibility", style='List Bullet')

    doc.add_paragraph()  # Add space after the list

    # Count issues
    for page in pages_with_tabindex_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        flags = page['results']['accessibility']['tests']['tabindex']['tabindex']['pageFlags']
        
        for flag in tabindex_issues:
            if flags.get(flag, False):  # If issue exists (True)
                tabindex_issues[flag]['pages'].add(page['url'])
                tabindex_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in tabindex_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add domain details for each issue
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)
    else:
        doc.add_paragraph("No issues were found.")

    ##########################################
    # Title Attribute
    ##########################################

    doc.add_page_break()
    h2 = doc.add_heading('Title Attribute', level=2)
    h2.style = doc.styles['Heading 2']
 
    # Add explanation
    doc.add_paragraph("""
The title attribute is often misused as a tooltip or to provide additional information. However, it has several accessibility limitations:
""".strip())

    doc.add_paragraph("Not consistently exposed by screen readers or available on mobile devices", style='List Bullet')
    doc.add_paragraph("Cannot be accessed by keyboard-only users", style='List Bullet')
    doc.add_paragraph("Cannot be reliably accessed by screen-magnifier users as the title attribute may be unreachable as the user moves the mouse to read it", style='List Bullet')
    doc.add_paragraph("Content is not visible until hover, which some users cannot do", style='List Bullet')
    doc.add_paragraph("Should not be used as the only way to convey important information", style='List Bullet')
    doc.add_paragraph("""
There is one case when the title attribute must be used, and that is fir <iframe> as it is the only way to give a name ot and embedded element. Typically that is the title of a YouTube or Vimeo video.
""".strip())

    doc.add_paragraph()

    # Query for pages with title attribute issues
    pages_with_title_issues = list(db_connection.page_results.find(
        {"results.accessibility.tests.title.titleAttribute.pageFlags.hasImproperTitleAttributes": True},
        {
            "url": 1,
            "results.accessibility.tests.title.titleAttribute.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Count affected domains
    affected_domains = set()
    total_improper_uses = 0
    domain_counts = {}

    for page in pages_with_title_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        affected_domains.add(domain)
        
        # Count improper uses from the details
        improper_uses = len(page['results']['accessibility']['tests']['title']['titleAttribute']['details']['improperUse'])
        total_improper_uses += improper_uses
        
        # Track counts by domain
        if domain not in domain_counts:
            domain_counts[domain] = 0
        domain_counts[domain] += improper_uses

    # Calculate percentage
    percentage = (len(affected_domains) / len(total_domains)) * 100 if total_domains else 0

    # Create summary table
    summary_table = doc.add_table(rows=2, cols=4)
    summary_table.style = 'Table Grid'

    # Set column headers
    headers = summary_table.rows[0].cells
    headers[0].text = "Issue"
    headers[1].text = "Total Occurrences"
    headers[2].text = "Sites Affected"
    headers[3].text = "% of Total Sites"

    # Add data
    row = summary_table.rows[1].cells
    row[0].text = "Improper use of title attribute"
    row[1].text = str(total_improper_uses)
    row[2].text = str(len(affected_domains))
    row[3].text = f"{percentage:.1f}%"

    # Format the table text
    format_table_text(summary_table)

    # Add some space after the table
    doc.add_paragraph()

    if domain_counts:
        # Create domain details table
        doc.add_paragraph("Breakdown by site:")
        domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
        domain_table.style = 'Table Grid'

        # Add headers
        headers = domain_table.rows[0].cells
        headers[0].text = "Domain"
        headers[1].text = "Number of improper title attributes"

        # Add domain data
        for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
            row = domain_table.rows[i].cells
            row[0].text = domain
            row[1].text = str(count)

        # Format the table text
        format_table_text(domain_table)

    ########################################
    # Tables
    ########################################
            
    doc.add_page_break()
    h2 = doc.add_heading('Tables', level=2)
    h2.style = doc.styles['Heading 2']
 
    # Add explanation
    doc.add_paragraph("""
Tables should be used for presenting tabular data, not for layout purposes. Proper table markup with appropriate headers and structure is crucial for screen reader users. Common issues include:
""".strip())

    doc.add_paragraph("Missing table headers (th elements)", style='List Bullet')
    doc.add_paragraph("Lack of proper scope attributes on header cells", style='List Bullet')
    doc.add_paragraph("Missing caption or summary for complex tables", style='List Bullet')
    doc.add_paragraph("Tables used for layout purposes instead of CSS", style='List Bullet')
    doc.add_paragraph("Complex tables without proper row/column headers", style='List Bullet')

    doc.add_paragraph()

    # Query for pages with table issues
    pages_with_table_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.tables.tables.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.tables.tables.pageFlags.hasMissingHeaders": True},
                {"results.accessibility.tests.tables.tables.pageFlags.hasNoScope": True},
                {"results.accessibility.tests.tables.tables.pageFlags.hasMissingCaption": True},
                {"results.accessibility.tests.tables.tables.pageFlags.hasLayoutTables": True},
                {"results.accessibility.tests.tables.tables.pageFlags.hasComplexTables": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.tables.tables.pageFlags": 1,
            "results.accessibility.tests.tables.tables.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    table_issues = {
        "hasMissingHeaders": {"name": "Missing table headers", "pages": set(), "domains": set()},
        "hasNoScope": {"name": "Missing scope attributes", "pages": set(), "domains": set()},
        "hasMissingCaption": {"name": "Missing table captions", "pages": set(), "domains": set()},
        "hasLayoutTables": {"name": "Layout tables", "pages": set(), "domains": set()},
        "hasComplexTables": {"name": "Complex tables without proper structure", "pages": set(), "domains": set()}
    }

    # Count issues
    for page in pages_with_table_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        flags = page['results']['accessibility']['tests']['tables']['tables']['pageFlags']
        
        for flag in table_issues:
            if flags.get(flag, False):
                table_issues[flag]['pages'].add(page['url'])
                table_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in table_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Table Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add domain details for each issue
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)

    else:
        doc.add_paragraph("No table markup issues were found.")

    #########################################
    # Timers
    #########################################

    doc.add_page_break()
    h2 = doc.add_heading('Timers', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
Time limits and automatic updates can create significant barriers for users who need more time to read content or complete tasks. Common issues with timers include:
""".strip())

    doc.add_paragraph("Auto-starting timers that begin without user initiation", style='List Bullet')
    doc.add_paragraph("Timers without proper controls to pause, stop, or extend time", style='List Bullet')
    doc.add_paragraph("Session timeouts without adequate warning or ability to extend", style='List Bullet')
    doc.add_paragraph("Content that updates automatically without user control", style='List Bullet')

    # Add recommendations
    doc.add_paragraph()
    doc.add_paragraph("Recommendations for addressing timer issues:")
    doc.add_paragraph("Provide options to turn off, adjust, or extend time limits", style='List Bullet')
    doc.add_paragraph("Ensure all auto-updating content can be paused", style='List Bullet')
    doc.add_paragraph("Give adequate warning before session timeouts", style='List Bullet')
    doc.add_paragraph("Provide mechanisms to request more time", style='List Bullet')

    doc.add_paragraph()

    # Query for pages with timer issues
    pages_with_timer_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.timers.timers.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.timers.timers.pageFlags.hasTimers": True},
                {"results.accessibility.tests.timers.timers.pageFlags.hasAutoStartTimers": True},
                {"results.accessibility.tests.timers.timers.pageFlags.hasTimersWithoutControls": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.timers.timers.pageFlags": 1,
            "results.accessibility.tests.timers.timers.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    timer_issues = {
        "hasTimers": {"name": "Pages with timers", "pages": set(), "domains": set()},
        "hasAutoStartTimers": {"name": "Auto-starting timers", "pages": set(), "domains": set()},
        "hasTimersWithoutControls": {"name": "Timers without adequate controls", "pages": set(), "domains": set()}
    }

    # Count issues
    for page in pages_with_timer_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        flags = page['results']['accessibility']['tests']['timers']['timers']['pageFlags']
        
        for flag in timer_issues:
            if flags.get(flag, False):
                timer_issues[flag]['pages'].add(page['url'])
                timer_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in timer_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Timer Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add domain details for each issue
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)

    else:
        doc.add_paragraph("No timer-related issues were found.")

    ##################################
    # Videos
    ##################################

    doc.add_page_break()
    h2 = doc.add_heading('Videos', level=2)
    h2.style = doc.styles['Heading 2']
    findings = doc.add_paragraph()

    # Add explanation
    doc.add_paragraph("""
    Video content must be accessible to all users, including those with visual or hearing impairments. Videos should include appropriate alternatives and controls. Common accessibility issues with video content include:
    """.strip())

    doc.add_paragraph("Missing closed captions for audio content", style='List Bullet')
    doc.add_paragraph("Lack of audio descriptions for visual information", style='List Bullet')
    doc.add_paragraph("Inaccessible video controls", style='List Bullet')
    doc.add_paragraph("Missing transcripts", style='List Bullet')
    doc.add_paragraph("Autoplay videos without user control", style='List Bullet')
    doc.add_paragraph("Videos without proper labels or titles", style='List Bullet')

    # Add recommendations paragraph
    doc.add_paragraph()
    doc.add_paragraph("Recommendations for Video Accessibility:")
    doc.add_paragraph("Ensure all videos have accurate closed captions that include both speech and important sound effects.", style='List Bullet')
    doc.add_paragraph("Provide audio descriptions for important visual information when necessary.", style='List Bullet')
    doc.add_paragraph("Include complete transcripts for all video content.", style='List Bullet')
    doc.add_paragraph("Ensure video players have keyboard-accessible controls.", style='List Bullet')
    doc.add_paragraph("Avoid autoplay or provide easy controls to stop playback.", style='List Bullet')
    doc.add_paragraph("Include clear, descriptive titles and labels for all video content.", style='List Bullet')

    doc.add_paragraph()

    # Query for pages with video issues
    pages_with_video_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.video.video.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.video.video.pageFlags.missingCaptions": True},
                {"results.accessibility.tests.video.video.pageFlags.missingAudioDescription": True},
                {"results.accessibility.tests.video.video.pageFlags.inaccessibleControls": True},
                {"results.accessibility.tests.video.video.pageFlags.missingTranscript": True},
                {"results.accessibility.tests.video.video.pageFlags.hasAutoplay": True},
                {"results.accessibility.tests.video.video.pageFlags.missingLabels": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.video.video.pageFlags": 1,
            "results.accessibility.tests.video.video.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    video_issues = {
        "missingCaptions": {"name": "Missing closed captions", "pages": set(), "domains": set()},
        "missingAudioDescription": {"name": "Missing audio descriptions", "pages": set(), "domains": set()},
        "inaccessibleControls": {"name": "Inaccessible video controls", "pages": set(), "domains": set()},
        "missingTranscript": {"name": "Missing transcripts", "pages": set(), "domains": set()},
        "hasAutoplay": {"name": "Autoplay without user control", "pages": set(), "domains": set()},
        "missingLabels": {"name": "Missing video labels/titles", "pages": set(), "domains": set()}
    }

    # Count issues
    if (len(pages_with_video_issues) > 0):

        for page in pages_with_video_issues:
            domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
            flags = page['results']['accessibility']['tests']['video']['video']['pageFlags']
            
            for flag in video_issues:
                if flags.get(flag, False):
                    video_issues[flag]['pages'].add(page['url'])
                    video_issues[flag]['domains'].add(domain)

        # Create filtered list of issues that have affected pages
        active_issues = {flag: data for flag, data in video_issues.items() 
                        if len(data['pages']) > 0}

        if active_issues:
            # Create summary table
            summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
            summary_table.style = 'Table Grid'

            # Set column headers
            headers = summary_table.rows[0].cells
            headers[0].text = "Video Issue"
            headers[1].text = "Pages Affected"
            headers[2].text = "Sites Affected"
            headers[3].text = "% of Total Sites"

            # Add data
            for i, (flag, data) in enumerate(active_issues.items(), 1):
                row = summary_table.rows[i].cells
                row[0].text = data['name']
                row[1].text = str(len(data['pages']))
                row[2].text = str(len(data['domains']))
                row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

            # Format the table text
            format_table_text(summary_table)

            # Add domain details for each issue
            for flag, data in active_issues.items():
                if data['domains']:
                    doc.add_paragraph()
                    doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                    
                    # Group by domain and count occurrences
                    domain_counts = {}
                    for page in data['pages']:
                        domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                        domain_counts[domain] = domain_counts.get(domain, 0) + 1

                    # Create domain details table
                    domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                    domain_table.style = 'Table Grid'

                    # Add headers
                    headers = domain_table.rows[0].cells
                    headers[0].text = "Domain"
                    headers[1].text = "Number of pages"

                    # Add domain data
                    for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                        row = domain_table.rows[i].cells
                        row[0].text = domain
                        row[1].text = str(count)

                    # Format the table text
                    format_table_text(domain_table)

            # Add extra details about video findings if available
            if any('details' in page.get('results', {}).get('accessibility', {}).get('tests', {}).get('video', {}).get('video', {}) 
                for page in pages_with_video_issues):
                doc.add_paragraph()
                doc.add_paragraph("Additional Details:", style='Intense Quote')
                for page in pages_with_video_issues:
                    details = page.get('results', {}).get('accessibility', {}).get('tests', {}).get('video', {}).get('video', {}).get('details', [])
                    if details:
                        doc.add_paragraph(f"URL: {page['url']}", style='Subtle Reference')
                        for detail in details:
                            doc.add_paragraph(detail, style='List Bullet')

        else:
            doc.add_paragraph("No video accessibility issues were found.")
    else:
        doc.add_paragraph("No videos were found.")

    ############################################
    # Appendices
    ############################################

    h2 = doc.add_heading('APPENDICES', level=1)
    h2.style = doc.styles['Heading 1']

    doc.add_paragraph()
    h2 = doc.add_heading('Test Coverage', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
    This section provides an overview of the sites and pages included in the accessibility analysis. Understanding the scope
    of testing is important for evaluating the comprehensiveness of the assessment and identifying areas that may need
    additional coverage.
    """.strip())

    # Add notes about coverage
    doc.add_paragraph()
    doc.add_paragraph("Notes about Coverage:", style='Normal')
    doc.add_paragraph("The pages tested represent a sample of each site's content", style='List Bullet')
    doc.add_paragraph("Testing includes various page types (home pages, content pages, forms, etc.)", style='List Bullet')
    doc.add_paragraph("Coverage may vary by site based on site structure and complexity", style='List Bullet')

    doc.add_paragraph()

    # Query for all tested pages
    tested_pages = list(db_connection.page_results.find(
        {},
        {
            "url": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Process the pages
    sites_data = {}
    for page in tested_pages:
        url = page['url']
        domain = url.replace('http://', '').replace('https://', '').split('/')[0]
        
        if domain not in sites_data:
            sites_data[domain] = {
                'pages': set(),
                'count': 0
            }
        
        sites_data[domain]['pages'].add(url)
        sites_data[domain]['count'] += 1

    # Create summary statistics
    total_sites = len(sites_data)
    total_pages = sum(site['count'] for site in sites_data.values())
    avg_pages_per_site = total_pages / total_sites if total_sites > 0 else 0

    # Add summary statistics
    doc.add_paragraph("Summary Statistics:", style='Normal')
    summary_table = doc.add_table(rows=3, cols=2)
    summary_table.style = 'Table Grid'

    # Add summary data
    rows = summary_table.rows
    rows[0].cells[0].text = "Total Sites Tested"
    rows[0].cells[1].text = str(total_sites)
    rows[1].cells[0].text = "Total Pages Tested"
    rows[1].cells[1].text = str(total_pages)
    rows[2].cells[0].text = "Average Pages per Site"
    rows[2].cells[1].text = f"{avg_pages_per_site:.1f}"

    format_table_text(summary_table)

    # Add site-by-site breakdown
    doc.add_paragraph()
    doc.add_paragraph("Coverage by Site:", style='Normal')

    # Create sites overview table
    sites_table = doc.add_table(rows=len(sites_data) + 1, cols=2)
    sites_table.style = 'Table Grid'

    # Add headers
    headers = sites_table.rows[0].cells
    headers[0].text = "Site"
    headers[1].text = "Pages Tested"

    # Add site data
    for i, (domain, data) in enumerate(sorted(sites_data.items()), 1):
        row = sites_table.rows[i].cells
        row[0].text = domain
        row[1].text = str(data['count'])

    format_table_text(sites_table)

    # Add detailed page listings
    doc.add_paragraph()
    doc.add_paragraph("Detailed Page Listings:", style='Normal')

    for domain, data in sorted(sites_data.items()):
        doc.add_paragraph()
        doc.add_paragraph(f"{domain}:", style='List Bullet')
        
        # Sort pages for consistent display
        sorted_pages = sorted(data['pages'])
        
        # Create page list table
        page_table = doc.add_table(rows=len(sorted_pages) + 1, cols=2)
        page_table.style = 'Table Grid'
        
        # Add headers
        headers = page_table.rows[0].cells
        headers[0].text = "#"
        headers[1].text = "Page URL"
        
        # Add pages
        for i, page in enumerate(sorted_pages, 1):
            row = page_table.rows[i].cells
            row[0].text = str(i)
            row[1].text = page
        
        format_table_text(page_table)


    #################################
    # Electronic Documents
    #################################

    doc.add_page_break()
    h2 = doc.add_heading('Electronic documents found', level=2)
    h2.style = doc.styles['Heading 2']


    # Add explanation
    doc.add_paragraph("""
    This section lists the electronic documents found across all tested pages.
    """.strip())

    doc.add_paragraph()

    # Query for pages with document information
    pages_with_documents = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.documents.document_links": {"$exists": True}
        },
        {
            "url": 1,
            "results.accessibility.tests.documents.document_links": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Create a list of all documents
    all_documents = []
    for page in pages_with_documents:
        try:
            doc_links = page.get('results', {}).get('accessibility', {}).get('tests', {}).get('documents', {}).get('document_links', {})
            if 'documents' in doc_links:
                for document in doc_links['documents']:
                    all_documents.append({
                        'page_url': page['url'],
                        'doc_url': document.get('url'),
                        'type': document.get('type')
                    })
        except Exception as e:
            print(f"Error processing page {page.get('url')}: {str(e)}")
            continue

    # Count documents by type
    type_counts = {}
    for document in all_documents:
        doc_type = document.get('type', 'unknown').upper()
        type_counts[doc_type] = type_counts.get(doc_type, 0) + 1

    # Create summary table
    summary_table = doc.add_table(rows=len(type_counts) + 1, cols=2)
    summary_table.style = 'Table Grid'

    # Add summary headers
    headers = summary_table.rows[0].cells
    headers[0].text = "Document Type"
    headers[1].text = "Count"

    # Add summary data
    for i, (doc_type, count) in enumerate(sorted(type_counts.items()), 1):
        row = summary_table.rows[i].cells
        row[0].text = doc_type
        row[1].text = str(count)

    format_table_text(summary_table)

    doc.add_paragraph()
    doc.add_paragraph("Document Listing:", style='Normal')

    # Create document listing table
    table = doc.add_table(rows=len(all_documents) + 1, cols=3)
    table.style = 'Table Grid'

    # Add headers
    headers = table.rows[0].cells
    headers[0].text = "Type"
    headers[1].text = "Document URL"
    headers[2].text = "Found On Page"

    # Add documents
    for i, document in enumerate(sorted(all_documents, key=lambda x: x['type']), 1):
        row = table.rows[i].cells
        row[0].text = document.get('type', 'unknown').upper()
        row[1].text = document.get('doc_url', 'No URL')
        row[2].text = document.get('page_url', 'Unknown page')

    format_table_text(table)

    # Add total count
    doc.add_paragraph()
    doc.add_paragraph(f"Total Documents Found: {len(all_documents)}")


    #############################################
    # Return the created Word Document
    #############################################

    return doc

################################################
# Main Report generation fn
################################################

def generate_report(db_connection, title, author, date):
    try:
        doc = create_report_template(db_connection, title, author, date)
        output_filename = f'accessibility_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        doc.save(output_filename)
        return output_filename
    except Exception as e:
        print(f"Error generating report: {e}")
        return None

#################################################
# Command Line Interface
#################################################

@click.command()
@click.option('--title', '-t', 
              default='Accessibility Test Report', 
              help='Title of the report')
@click.option('--author', '-a', 
              required=True, 
              help='Author of the report')
@click.option('--date', '-d', 
              default=datetime.now().strftime("%Y-%m-%d"),
              help='Date of the report (YYYY-MM-DD)')
def main(title, author, date):
    """Generate an accessibility test report with specified parameters."""
    try:
        datetime.strptime(date, "%Y-%m-%d")
    except ValueError:
        click.echo("Error: Date must be in YYYY-MM-DD format")
        return

    click.echo(f"Generating report with the following parameters:")
    click.echo(f"Title: {title}")
    click.echo(f"Author: {author}")
    click.echo(f"Date: {date}")
    
    db = AccessibilityDB()
    report_file = generate_report(db, title, author, date)
    
    if report_file:
        click.echo(f"\nReport generated successfully: {report_file}")
        click.echo("\nIMPORTANT: To complete the report formatting:")
        click.echo("1. Open the document in Microsoft Word")
        click.echo("2. Right-click anywhere in the table of contents")
        click.echo("3. Select 'Update Field'")
        click.echo("4. Choose 'Update entire table'")
    else:
        click.echo("Failed to generate report", err=True)

if __name__ == "__main__":
    main()

