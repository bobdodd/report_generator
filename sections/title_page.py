import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from report_styling import format_table_text

def add_title_page(doc, db_connection, title, author, date):
    """Create the title page of the report"""
    title_heading = doc.add_heading(title, level=0)
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add image - calculate width based on page width
    doc.add_paragraph()  # Add space before image
    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Get page width (accounting for margins)
    section = doc.sections[0]
    page_width = section.page_width - section.left_margin - section.right_margin
    
    if os.path.exists('access labs.jpg'):
        image_run = image_paragraph.add_run()
        image_run.add_picture('access labs.jpg', width=page_width)
    else:
        print("Warning: 'access labs.jpg' not found in the current directory")
    
    doc.add_paragraph()  # Add space after image
    
    # Add subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("CNIB Access Labs")
    subtitle_run.font.size = Pt(24)  # Large text
    subtitle_run.font.bold = True

    doc.add_paragraph()  # Add space after image
        
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_para.style.font.size = Pt(14)
    author_para.add_run(f"Author: {author}")

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"Date: {date}")
    
    # Get the latest test run data
    latest_test_run = db_connection.get_latest_test_run()

    if latest_test_run:
        doc.add_paragraph()
        test_run_para = doc.add_paragraph()
        test_run_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        test_run_para.add_run(f"Test Run Date: {latest_test_run['timestamp_start']}")
        status_para = doc.add_paragraph()
        status_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        status_para.add_run(f"Status: {latest_test_run['status']}")