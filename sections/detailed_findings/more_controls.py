# sections/detailed_findings/more_controls.py
from report_styling import format_table_text
from docx.shared import Pt

def add_detailed_more_controls(doc, db_connection, total_domains):
    """Add the detailed 'More' Controls section"""
    doc.add_page_break()
    h2 = doc.add_heading('"More" Controls', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
Generic "Read More" or "Learn More" style links can create barriers for screen reader users who rely on link and button text to understand where a link/button will take them. When they are taken out of context:
""".strip())

    doc.add_paragraph("Users can't determine the link's purpose from the link text alone", style='List Bullet')
    doc.add_paragraph("Screen reader users may get a list of identical 'read more' links", style='List Bullet')
    doc.add_paragraph("The destination of the link isn't clear without surrounding context slowing down reading for screen-reader and screen-magnifier users", style='List Bullet')

    # Add recommendations
    doc.add_paragraph()
    doc.add_paragraph("Recommendations for improving link text:")
    doc.add_paragraph("Make link and button text descriptive of its destination or purpose", style='List Bullet')
    doc.add_paragraph("Use aria-label or visually hidden text if additional context is needed", style='List Bullet')
    doc.add_paragraph("Ensure link text makes sense when read out of context", style='List Bullet')

    doc.add_paragraph()

    # Query for pages with read more link issues
    pages_with_readmore_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.read_more_links.read_more_links.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.read_more_links.read_more_links.pageFlags.hasGenericReadMoreLinks": True},
                {"results.accessibility.tests.read_more_links.read_more_links.pageFlags.hasInvalidReadMoreLinks": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.read_more_links.read_more_links.pageFlags": 1,
            "results.accessibility.tests.read_more_links.read_more_links.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    readmore_issues = {
        "hasGenericReadMoreLinks": {"name": "Generic 'Read More' links", "pages": set(), "domains": set()},
        "hasInvalidReadMoreLinks": {"name": "Invalid implementation of 'Read More' links", "pages": set(), "domains": set()}
    }

    # Count issues
    for page in pages_with_readmore_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        flags = page['results']['accessibility']['tests']['read_more_links']['read_more_links']['pageFlags']
        
        for flag in readmore_issues:
            if flags.get(flag, False):
                readmore_issues[flag]['pages'].add(page['url'])
                readmore_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in readmore_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add domain details for each issue
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)

        # Add examples if available
        for page in pages_with_readmore_issues:
            details = page['results']['accessibility']['tests']['read_more_links']['read_more_links']['details']
            if 'items' in details and details['items']:
                doc.add_paragraph()
                doc.add_paragraph("Examples of problematic link text found:")
                for item in details['items'][:5]:  # Show up to 5 examples
                    doc.add_paragraph(item, style='List Bullet')
                break  # Only show examples from first page with issues
                
        # Add technical implementation section
        doc.add_paragraph()
        doc.add_heading('Technical Implementation Guidelines', level=3)
        
        # Problematic examples
        doc.add_paragraph("Problematic 'Read More' Implementations:", style='Normal').bold = True
        
        bad_examples = doc.add_paragraph("""
<!-- Problematic: Generic "read more" link -->
<div class="news-item">
  <h3>Company Launches New Product</h3>
  <p>Our company has launched an exciting new product that will revolutionize...</p>
  <a href="/news/new-product">Read more</a>
</div>

<!-- Problematic: Multiple identical "learn more" links -->
<div class="services">
  <div class="service">
    <h3>Web Development</h3>
    <p>Custom website development services...</p>
    <a href="/services/web">Learn more</a>
  </div>
  <div class="service">
    <h3>App Development</h3>
    <p>Mobile application development services...</p>
    <a href="/services/app">Learn more</a>
  </div>
  <div class="service">
    <h3>Design Services</h3>
    <p>UI/UX design services for digital products...</p>
    <a href="/services/design">Learn more</a>
  </div>
</div>
        """)
        bad_examples.style = doc.styles['Normal']
        bad_examples.paragraph_format.left_indent = Pt(36)
        
        # Accessible alternatives
        doc.add_paragraph("Accessible Alternatives:", style='Normal').bold = True
        
        good_examples = doc.add_paragraph("""
<!-- Good: Descriptive link text -->
<div class="news-item">
  <h3>Company Launches New Product</h3>
  <p>Our company has launched an exciting new product that will revolutionize...</p>
  <a href="/news/new-product">Read more about our new product launch</a>
</div>

<!-- Good: Using aria-label for additional context -->
<div class="news-item">
  <h3>Company Launches New Product</h3>
  <p>Our company has launched an exciting new product that will revolutionize...</p>
  <a href="/news/new-product" aria-label="Read more about our new product launch">Read more</a>
</div>

<!-- Good: Using visually hidden text -->
<div class="news-item">
  <h3>Company Launches New Product</h3>
  <p>Our company has launched an exciting new product that will revolutionize...</p>
  <a href="/news/new-product">
    Read more
    <span class="visually-hidden">about our new product launch</span>
  </a>
</div>

<!-- Good: Descriptive link text in service cards -->
<div class="services">
  <div class="service">
    <h3>Web Development</h3>
    <p>Custom website development services...</p>
    <a href="/services/web">Learn more about web development</a>
  </div>
  <div class="service">
    <h3>App Development</h3>
    <p>Mobile application development services...</p>
    <a href="/services/app">Learn more about app development</a>
  </div>
  <div class="service">
    <h3>Design Services</h3>
    <p>UI/UX design services for digital products...</p>
    <a href="/services/design">Learn more about design services</a>
  </div>
</div>
        """)
        good_examples.style = doc.styles['Normal']
        good_examples.paragraph_format.left_indent = Pt(36)
        
        # CSS for visually hidden text
        doc.add_paragraph("CSS for Visually Hidden Text:", style='Normal').bold = True
        
        css_example = doc.add_paragraph("""
/* Visually hidden text - accessible to screen readers but not visible on screen */
.visually-hidden {
  position: absolute;
  width: 1px;
  height: 1px;
  margin: -1px;
  padding: 0;
  overflow: hidden;
  clip: rect(0, 0, 0, 0);
  border: 0;
}
        """)
        css_example.style = doc.styles['Normal']
        css_example.paragraph_format.left_indent = Pt(36)
        
        # JavaScript enhancement for consistent context
        doc.add_paragraph("JavaScript Enhancement for Consistent Context:", style='Normal').bold = True
        
        js_example = doc.add_paragraph("""
// JavaScript to enhance generic "Read More" links with their context
document.addEventListener('DOMContentLoaded', function() {
  // Find all "Read more" or "Learn more" links
  const genericLinks = document.querySelectorAll('a');
  
  genericLinks.forEach(link => {
    const linkText = link.textContent.trim().toLowerCase();
    
    // Check if this is a generic "read more" type link
    if (linkText === 'read more' || linkText === 'learn more') {
      // Look for a heading within the same container
      const container = link.closest('article, section, div');
      if (container) {
        const heading = container.querySelector('h1, h2, h3, h4, h5, h6');
        
        if (heading) {
          const headingText = heading.textContent.trim();
          
          // Add visually hidden text with context
          const hiddenText = document.createElement('span');
          hiddenText.className = 'visually-hidden';
          hiddenText.textContent = ` about ${headingText}`;
          link.appendChild(hiddenText);
        }
      }
    }
  });
});
        """)
        js_example.style = doc.styles['Normal']
        js_example.paragraph_format.left_indent = Pt(36)

    else:
        doc.add_paragraph("No issues with generic 'Read More' links were found.")
        