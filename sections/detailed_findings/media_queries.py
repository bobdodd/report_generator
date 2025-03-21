# sections/detailed_findings/media_queries.py
from report_styling import format_table_text
from docx.shared import Pt

def add_detailed_media_queries(doc, db_connection, total_domains):
    """Add the detailed Media Queries section"""
    doc.add_page_break()
    h2 = doc.add_heading('Media Queries and Responsive Design', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
CSS Media queries are essential for creating accessible websites that adapt to different devices, screen sizes, and user preferences. Key accessibility concerns include:
""".strip())

    doc.add_paragraph("Responsive layouts that adapt to different screen sizes and zoom levels", style='List Bullet')
    doc.add_paragraph("Print stylesheets that ensure content is accessible when printed", style='List Bullet')
    doc.add_paragraph("Support for reduced motion preferences for users with vestibular disorders", style='List Bullet')
    doc.add_paragraph("Dark mode support for users with light sensitivity", style='List Bullet')
    doc.add_paragraph("Orientation-specific layouts for different device orientations", style='List Bullet')

    # Add recommendations
    doc.add_paragraph()
    doc.add_paragraph("Recommendations for implementing accessible media queries:")
    doc.add_paragraph("Implement responsive breakpoints for common device sizes (320px, 768px, 1024px, etc.)", style='List Bullet')
    doc.add_paragraph("Add print stylesheets to optimize content for printing", style='List Bullet')
    doc.add_paragraph("Support reduced motion preferences with @media (prefers-reduced-motion: reduce)", style='List Bullet')
    doc.add_paragraph("Support dark mode with @media (prefers-color-scheme: dark)", style='List Bullet')
    doc.add_paragraph("Test layouts in both portrait and landscape orientations", style='List Bullet')

    doc.add_paragraph()

    # Get breakpoint data from all pages, using the new dedicated structure
    breakpoint_data = {}
    pages_with_breakpoints = list(db_connection.page_results.find(
        {"results.accessibility.tests.media_queries.media_queries.responsiveBreakpoints": {"$exists": True}},
        {
            "url": 1,
            "results.accessibility.tests.media_queries.media_queries.responsiveBreakpoints": 1,
            "_id": 0
        }
    ))

    # Collect all breakpoints across pages
    all_breakpoints = set()
    breakpoint_by_category = {
        'mobile': set(),
        'tablet': set(),
        'desktop': set(),
        'largeScreen': set()
    }
    breakpoint_histogram = {}
    
    for page in pages_with_breakpoints:
        if 'responsiveBreakpoints' in page['results']['accessibility']['tests']['media_queries']['media_queries']:
            breakpoints_data = page['results']['accessibility']['tests']['media_queries']['media_queries']['responsiveBreakpoints']
            
            # Add to the full list of breakpoints
            if 'allBreakpoints' in breakpoints_data:
                for bp in breakpoints_data['allBreakpoints']:
                    all_breakpoints.add(bp)
                    breakpoint_histogram[bp] = breakpoint_histogram.get(bp, 0) + 1
            
            # Add to category-specific sets
            if 'byCategory' in breakpoints_data:
                for category, bps in breakpoints_data['byCategory'].items():
                    if category in breakpoint_by_category:
                        for bp in bps:
                            breakpoint_by_category[category].add(bp)
    
    # Add common breakpoints section if we have data
    if breakpoint_histogram:
        doc.add_heading('Common Responsive Breakpoints', level=3)
        doc.add_paragraph("The following breakpoints (in pixels) were detected across the site:")
        
        # Create a table for breakpoints by category
        doc.add_heading('Breakpoints by Device Category', level=4)
        category_table = doc.add_table(rows=5, cols=2)
        category_table.style = 'Table Grid'
        
        # Add headers
        headers = category_table.rows[0].cells
        headers[0].text = "Device Category"
        headers[1].text = "Breakpoints (px)"
        
        # Add category data
        categories = [
            ("Mobile (≤480px)", sorted(breakpoint_by_category['mobile'])),
            ("Tablet (481-768px)", sorted(breakpoint_by_category['tablet'])),
            ("Desktop (769-1200px)", sorted(breakpoint_by_category['desktop'])),
            ("Large Screen (>1200px)", sorted(breakpoint_by_category['largeScreen']))
        ]
        
        for i, (category, bps) in enumerate(categories, 1):
            row = category_table.rows[i].cells
            row[0].text = category
            row[1].text = ", ".join(str(bp) for bp in bps) if bps else "None detected"
        
        # Format the table
        format_table_text(category_table)
        
        doc.add_paragraph()
        
        # Create table for breakpoint histogram
        doc.add_heading('Breakpoint Frequency', level=4)
        doc.add_paragraph("This table shows how frequently each breakpoint appears across all pages:")
        
        sorted_breakpoints = sorted(breakpoint_histogram.items())
        bp_table = doc.add_table(rows=len(sorted_breakpoints) + 1, cols=2)
        bp_table.style = 'Table Grid'
        
        # Add headers
        headers = bp_table.rows[0].cells
        headers[0].text = "Breakpoint (px)"
        headers[1].text = "Frequency"
        
        # Add data
        for i, (breakpoint, count) in enumerate(sorted_breakpoints, 1):
            row = bp_table.rows[i].cells
            row[0].text = str(breakpoint)
            row[1].text = str(count)
        
        # Format table
        format_table_text(bp_table)
        
        doc.add_paragraph()

    # Query for pages with media query issues
    pages_with_media_query_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.media_queries.media_queries.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.media_queries.media_queries.pageFlags.hasResponsiveBreakpoints": False},
                {"results.accessibility.tests.media_queries.media_queries.pageFlags.hasPrintStyles": False},
                {"results.accessibility.tests.media_queries.media_queries.pageFlags.hasReducedMotionSupport": False},
                {"results.accessibility.tests.media_queries.media_queries.pageFlags.hasDarkModeSupport": False},
                {"results.accessibility.tests.media_queries.media_queries.pageFlags.hasOrientationStyles": False}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.media_queries.media_queries.pageFlags": 1,
            "results.accessibility.tests.media_queries.media_queries.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    media_query_issues = {
        "no_responsive": {"name": "No responsive breakpoints", "pages": set(), "domains": set()},
        "no_print": {"name": "No print stylesheets", "pages": set(), "domains": set()},
        "no_reduced_motion": {"name": "No reduced motion support", "pages": set(), "domains": set()},
        "no_dark_mode": {"name": "No dark mode support", "pages": set(), "domains": set()},
        "no_orientation": {"name": "No orientation-specific styles", "pages": set(), "domains": set()}
    }

    # Count issues
    for page in pages_with_media_query_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        flags = page['results']['accessibility']['tests']['media_queries']['media_queries']['pageFlags']
        
        if not flags.get('hasResponsiveBreakpoints', True):
            media_query_issues["no_responsive"]["pages"].add(page['url'])
            media_query_issues["no_responsive"]["domains"].add(domain)
            
        if not flags.get('hasPrintStyles', True):
            media_query_issues["no_print"]["pages"].add(page['url'])
            media_query_issues["no_print"]["domains"].add(domain)
            
        if not flags.get('hasReducedMotionSupport', True):
            media_query_issues["no_reduced_motion"]["pages"].add(page['url'])
            media_query_issues["no_reduced_motion"]["domains"].add(domain)
            
        if not flags.get('hasDarkModeSupport', True):
            media_query_issues["no_dark_mode"]["pages"].add(page['url'])
            media_query_issues["no_dark_mode"]["domains"].add(domain)
            
        if not flags.get('hasOrientationStyles', True):
            media_query_issues["no_orientation"]["pages"].add(page['url'])
            media_query_issues["no_orientation"]["domains"].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in media_query_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        doc.add_heading('Issues Summary', level=3)
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add domain details for critical issues (reduced motion is most critical)
        if "no_reduced_motion" in active_issues and active_issues["no_reduced_motion"]["domains"]:
            doc.add_paragraph()
            doc.add_heading('Sites Lacking Reduced Motion Support', level=3)
            doc.add_paragraph("The following sites do not implement the prefers-reduced-motion media query, which is essential for users with vestibular disorders:")
            
            # Group by domain and count occurrences
            domain_counts = {}
            for page in active_issues["no_reduced_motion"]["pages"]:
                domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                domain_counts[domain] = domain_counts.get(domain, 0) + 1

            # Create domain details table
            domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
            domain_table.style = 'Table Grid'

            # Add headers
            headers = domain_table.rows[0].cells
            headers[0].text = "Domain"
            headers[1].text = "Number of pages"

            # Add domain data
            for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                row = domain_table.rows[i].cells
                row[0].text = domain
                row[1].text = str(count)

            # Format the table text
            format_table_text(domain_table)

        # Add recommendations from test results if available
        recommendations_added = False
        for page in pages_with_media_query_issues:
            details = page['results']['accessibility']['tests']['media_queries']['media_queries']['details']
            if 'recommendations' in details and details['recommendations']:
                if not recommendations_added:
                    doc.add_paragraph()
                    doc.add_heading('Specific Recommendations', level=3)
                    recommendations_added = True
                    
                for rec in details['recommendations']:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"{rec['issue']} ").bold = True
                    p.add_run(f"(WCAG {rec['wcag']}) - {rec['recommendation']}")
                
                break  # Only show recommendations from first page
                
        # Add technical implementation guidance
        doc.add_paragraph()
        doc.add_heading('Technical Implementation Guidelines', level=3)
        
        # Responsive Design
        doc.add_paragraph("Responsive breakpoints for common device sizes:", style='Normal').bold = True
        
        responsive_code = doc.add_paragraph("""
/* Base styles for mobile devices */
body {
  font-size: 16px;
  line-height: 1.5;
}

/* Small tablets (portrait) and large phones */
@media (min-width: 600px) {
  body {
    font-size: 17px;
  }
}

/* Tablets and small desktops */
@media (min-width: 768px) {
  body {
    font-size: 18px;
  }
}

/* Large tablets and desktops */
@media (min-width: 992px) {
  body {
    width: 992px;
    margin: 0 auto;
  }
}

/* Ensure content reflows for zoom or small viewports */
@media (max-width: 320px), (forced-colors: active) {
  /* Simplified layout for very small screens or when zoom is applied */
  .multi-column {
    display: block;
  }
}
        """)
        responsive_code.style = doc.styles['Normal']
        responsive_code.paragraph_format.left_indent = Pt(36)
        
        # Print styles
        doc.add_paragraph("Print stylesheets for better printed output:", style='Normal').bold = True
        
        print_code = doc.add_paragraph("""
@media print {
  /* Hide navigation, sidebars, ads, and other non-essential content */
  nav, .sidebar, .ads, .comments, footer {
    display: none !important;
  }
  
  /* Ensure text is readable when printed */
  body {
    font-size: 12pt;
    line-height: 1.5;
    color: #000 !important;
    background: #fff !important;
  }
  
  /* Make links more useful in printed format */
  a[href]:after {
    content: " (" attr(href) ")";
    font-size: 10pt;
  }
  
  /* Avoid page breaks inside important elements */
  h1, h2, h3, img, table {
    page-break-inside: avoid;
    page-break-after: avoid;
  }
  
  /* Ensure tables print properly */
  table {
    border-collapse: collapse;
  }
  
  table, th, td {
    border: 1px solid #000;
  }
}
        """)
        print_code.style = doc.styles['Normal']
        print_code.paragraph_format.left_indent = Pt(36)
        
        # Reduced motion
        doc.add_paragraph("Respecting user preferences for reduced motion:", style='Normal').bold = True
        
        motion_code = doc.add_paragraph("""
/* Default animations */
.card {
  transition: transform 0.3s ease, box-shadow 0.3s ease;
}

.card:hover {
  transform: translateY(-5px);
  box-shadow: 0 10px 20px rgba(0,0,0,0.1);
}

/* Respect user preference for reduced motion */
@media (prefers-reduced-motion: reduce) {
  /* Disable non-essential animations and transitions */
  * {
    animation-duration: 0.001ms !important;
    animation-iteration-count: 1 !important;
    transition-duration: 0.001ms !important;
    scroll-behavior: auto !important;
  }
  
  /* Alternative hover effect without motion */
  .card:hover {
    transform: none;
    box-shadow: 0 0 0 2px #0078d7;
  }
  
  /* For essential animations, use significantly reduced motion */
  .progress-indicator {
    animation: none !important;
    /* Use opacity or color changes instead of movement */
    transition: opacity 0.5s linear !important;
  }
}
        """)
        motion_code.style = doc.styles['Normal']
        motion_code.paragraph_format.left_indent = Pt(36)
        
        # Dark mode
        doc.add_paragraph("Supporting dark mode for users with light sensitivity:", style='Normal').bold = True
        
        dark_code = doc.add_paragraph("""
:root {
  /* Light mode variables (default) */
  --background-color: #ffffff;
  --text-color: #222222;
  --link-color: #0066cc;
  --heading-color: #333333;
  --border-color: #dddddd;
}

/* Dark mode styles */
@media (prefers-color-scheme: dark) {
  :root {
    --background-color: #222222;
    --text-color: #f0f0f0;
    --link-color: #4d9cf6;
    --heading-color: #ffffff;
    --border-color: #444444;
  }
  
  /* Improve contrast for dark mode */
  img, video {
    opacity: 0.8;
  }
  
  /* Reduce eye strain with softer white text */
  body {
    background-color: var(--background-color);
    color: var(--text-color);
  }
  
  /* Ensure form controls have sufficient contrast */
  input, textarea, select {
    background-color: #333333;
    color: #ffffff;
    border: 1px solid #555555;
  }
}
        """)
        dark_code.style = doc.styles['Normal']
        dark_code.paragraph_format.left_indent = Pt(36)

    else:
        doc.add_paragraph("No media query-related accessibility issues were found.")