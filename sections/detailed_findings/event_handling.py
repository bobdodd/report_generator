from docx.shared import Pt
import traceback
from report_styling import format_table_text

def add_detailed_event_handling(doc, db_connection, total_domains):
    """Add the detailed Event Handling section"""
    doc.add_page_break()
    h2 = doc.add_heading('Event Handling and Keyboard Interaction', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
    Event handling and keyboard interaction are crucial for accessibility. This analysis examines event types, tab order, interactive elements, and modal dialog behavior. Issues with event handling can significantly impact keyboard and screen reader users.
    """.strip())

    # Add recommendations
    doc.add_paragraph()
    doc.add_heading('Event Handling Recommendations', level=3)
    
    doc.add_paragraph("Event Implementation:", style='List Bullet')
    doc.add_paragraph("Ensure keyboard alternatives for mouse-only interactions", style='List Bullet 2')
    doc.add_paragraph("Add keyboard event handlers alongside mouse events", style='List Bullet 2')
    doc.add_paragraph("Implement proper focus management", style='List Bullet 2')
    
    doc.add_paragraph("Tab Order:", style='List Bullet')
    doc.add_paragraph("Maintain logical tab sequence matching visual layout", style='List Bullet 2')
    doc.add_paragraph("Avoid using tabindex values greater than 0", style='List Bullet 2')
    doc.add_paragraph("Ensure all interactive elements are keyboard accessible", style='List Bullet 2')
    
    doc.add_paragraph("Modal Dialogs:", style='List Bullet')
    doc.add_paragraph("Implement escape key handling for all modals", style='List Bullet 2')
    doc.add_paragraph("Manage focus properly when opening/closing modals", style='List Bullet 2')
    doc.add_paragraph("Ensure modal content is properly contained", style='List Bullet 2')
    
    doc.add_paragraph()

    # Query for pages with event information
    pages_with_events = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.events.events": {"$exists": True}
        },
        {
            "url": 1,
            "results.accessibility.tests.events.events": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize tracking structures
    property_data = {
        # Event Types
        "event_mouse": {"name": "Mouse Events", "pages": set(), "domains": set(), "count": 0},
        "event_keyboard": {"name": "Keyboard Events", "pages": set(), "domains": set(), "count": 0},
        "event_focus": {"name": "Focus Events", "pages": set(), "domains": set(), "count": 0},
        "event_touch": {"name": "Touch Events", "pages": set(), "domains": set(), "count": 0},
        "event_timer": {"name": "Timer Events", "pages": set(), "domains": set(), "count": 0},
        "event_lifecycle": {"name": "Lifecycle Events", "pages": set(), "domains": set(), "count": 0},
        "event_other": {"name": "Other Events", "pages": set(), "domains": set(), "count": 0},
        
        # Tab Order
        "explicit_tabindex": {"name": "Explicit tabindex Usage", "pages": set(), "domains": set(), "count": 0},
        "visual_violations": {"name": "Visual Order Violations", "pages": set(), "domains": set(), "count": 0},
        "column_violations": {"name": "Column Order Violations", "pages": set(), "domains": set(), "count": 0},
        "negative_tabindex": {"name": "Negative Tabindex", "pages": set(), "domains": set(), "count": 0},
        "high_tabindex": {"name": "High Tabindex Values", "pages": set(), "domains": set(), "count": 0},
        
        # Interactive Elements
        "mouse_only": {"name": "Mouse-only Elements", "pages": set(), "domains": set(), "count": 0},
        "missing_tabindex": {"name": "Missing tabindex", "pages": set(), "domains": set(), "count": 0},
        "non_interactive": {"name": "Non-interactive with Handlers", "pages": set(), "domains": set(), "count": 0},
        
        # Modal Support
        "modals_no_escape": {"name": "Modals Missing Escape", "pages": set(), "domains": set(), "count": 0}
    }

    # Create detailed violation tracking organized by domain and URL
    domain_data = {}

    # Process each page
    for page in pages_with_events:
        try:
            url = page['url']
            
            domain = url.replace('http://', '').replace('https://', '').split('/')[0]
            event_data = page['results']['accessibility']['tests']['events']['events']
            
            # Initialize domain and URL tracking if needed
            if domain not in domain_data:
                domain_data[domain] = {
                    'urls': {}
                }
            
            # Initialize URL data
            domain_data[domain]['urls'][url] = {
                'event_types': {},
                'violations': {},
                'handlers_count': 0,
                'focusable_elements': 0,
                'total_violations': 0
            }
            
            # Get pageFlags data for the most reliable summary information
            pageFlags = event_data.get('pageFlags', {})
            details = pageFlags.get('details', {})
            
            # Track total handlers and violations
            total_handlers = details.get('totalHandlers', 0)
            total_violations = details.get('totalViolations', 0)
            domain_data[domain]['urls'][url]['handlers_count'] = total_handlers
            domain_data[domain]['urls'][url]['total_violations'] = total_violations

            # Track total focusable elements
            tab_order_data = details.get('tabOrder', {})
            focusable_elements = tab_order_data.get('totalFocusableElements', 0)
            domain_data[domain]['urls'][url]['focusable_elements'] = focusable_elements
            
            # Process event types using the updated structure
            by_type = details.get('byType', {})
            
            for event_type in ['mouse', 'keyboard', 'focus', 'touch', 'timer', 'lifecycle', 'other']:
                count = by_type.get(event_type, 0)
                
                if isinstance(count, list):
                    count = len(count)
                elif not isinstance(count, (int, float)):
                    try:
                        count = int(count or 0)
                    except (ValueError, TypeError):
                        count = 0
                
                # Track event type for this URL
                domain_data[domain]['urls'][url]['event_types'][event_type] = count
                
                if count > 0:
                    key = f"event_{event_type}"
                    property_data[key]['pages'].add(url)
                    property_data[key]['domains'].add(domain)
                    property_data[key]['count'] += count

            # Process violation counts by type
            violation_counts = details.get('violationCounts', {})
            
            # Tab order violations
            explicit_count = tab_order_data.get('elementsWithExplicitTabIndex', 0)
            visual_violations = violation_counts.get('visual-order', 0) or tab_order_data.get('visualOrderViolations', 0)
            column_violations = violation_counts.get('column-order', 0) or tab_order_data.get('columnOrderViolations', 0)
            
            # Track negative and high tabindex
            negative_tabindex = 1 if pageFlags.get('hasNegativeTabindex', False) else 0
            high_tabindex = 1 if pageFlags.get('hasHighTabindex', False) else 0
            
            # Track violations for this URL
            domain_data[domain]['urls'][url]['violations']['explicit_tabindex'] = explicit_count
            domain_data[domain]['urls'][url]['violations']['visual_order'] = visual_violations
            domain_data[domain]['urls'][url]['violations']['column_order'] = column_violations
            domain_data[domain]['urls'][url]['violations']['negative_tabindex'] = negative_tabindex
            domain_data[domain]['urls'][url]['violations']['high_tabindex'] = high_tabindex
            
            if explicit_count > 0:
                property_data['explicit_tabindex']['pages'].add(url)
                property_data['explicit_tabindex']['domains'].add(domain)
                property_data['explicit_tabindex']['count'] += explicit_count
                
            if visual_violations > 0:
                property_data['visual_violations']['pages'].add(url)
                property_data['visual_violations']['domains'].add(domain)
                property_data['visual_violations']['count'] += visual_violations
                
            if column_violations > 0:
                property_data['column_violations']['pages'].add(url)
                property_data['column_violations']['domains'].add(domain)
                property_data['column_violations']['count'] += column_violations
                
            if negative_tabindex > 0:
                property_data['negative_tabindex']['pages'].add(url)
                property_data['negative_tabindex']['domains'].add(domain)
                property_data['negative_tabindex']['count'] += negative_tabindex
                
            if high_tabindex > 0:
                property_data['high_tabindex']['pages'].add(url)
                property_data['high_tabindex']['domains'].add(domain)
                property_data['high_tabindex']['count'] += high_tabindex

            # Process element violations
            mouse_only = violation_counts.get('mouse-only', 0) or details.get('mouseOnlyElements', {}).get('count', 0)
            missing_tabindex = violation_counts.get('missing-tabindex', 0) or details.get('missingTabindex', 0)
            non_interactive = details.get('nonInteractiveWithHandlers', 0)
            modals_without_escape = violation_counts.get('modal-without-escape', 0)
            
            # Track violations for this URL
            domain_data[domain]['urls'][url]['violations']['mouse_only'] = mouse_only
            domain_data[domain]['urls'][url]['violations']['missing_tabindex'] = missing_tabindex
            domain_data[domain]['urls'][url]['violations']['non_interactive'] = non_interactive
            domain_data[domain]['urls'][url]['violations']['modals_no_escape'] = modals_without_escape
            
            if mouse_only > 0:
                property_data['mouse_only']['pages'].add(url)
                property_data['mouse_only']['domains'].add(domain)
                property_data['mouse_only']['count'] += mouse_only
                
            if missing_tabindex > 0:
                property_data['missing_tabindex']['pages'].add(url)
                property_data['missing_tabindex']['domains'].add(domain)
                property_data['missing_tabindex']['count'] += missing_tabindex
                
            if non_interactive > 0:
                property_data['non_interactive']['pages'].add(url)
                property_data['non_interactive']['domains'].add(domain)
                property_data['non_interactive']['count'] += non_interactive
                
            if modals_without_escape > 0:
                property_data['modals_no_escape']['pages'].add(url)
                property_data['modals_no_escape']['domains'].add(domain)
                property_data['modals_no_escape']['count'] += modals_without_escape

        except Exception as e:
            print(f"Error processing page {url}:")
            print("Exception:", str(e))
            traceback.print_exc()
            continue

    if pages_with_events:
        # Overall Summary section
        doc.add_heading('Event Handling Summary', level=3)
        
        # Create summary tables for different categories
        
        # 1. Event Types
        doc.add_paragraph("Event Types Distribution:", style='Normal').bold = True
        
        event_table = doc.add_table(rows=len([k for k in property_data.keys() if k.startswith('event_')]) + 1, cols=4)
        event_table.style = 'Table Grid'
        
        # Add headers
        headers = event_table.rows[0].cells
        headers[0].text = "Event Type"
        headers[1].text = "Total Count"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Sites"
        
        # Add data
        row_idx = 1
        for key, data in sorted([(k, v) for k, v in property_data.items() if k.startswith('event_')], 
                           key=lambda x: x[1]['count'], reverse=True):
            row = event_table.rows[row_idx].cells
            row[0].text = data['name']
            row[1].text = str(data['count'])
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"
            row_idx += 1
        
        format_table_text(event_table)
        
        # 2. Tab Order Issues
        doc.add_paragraph()
        doc.add_paragraph("Tab Order Issues:", style='Normal').bold = True
        
        tab_order_table = doc.add_table(rows=5 + 1, cols=4)
        tab_order_table.style = 'Table Grid'
        
        # Add headers
        headers = tab_order_table.rows[0].cells
        headers[0].text = "Issue Type"
        headers[1].text = "Total Count"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Sites"
        
        # Add data
        tab_order_keys = ['explicit_tabindex', 'visual_violations', 'column_violations', 'negative_tabindex', 'high_tabindex']
        for i, key in enumerate(tab_order_keys, 1):
            row = tab_order_table.rows[i].cells
            data = property_data[key]
            row[0].text = data['name']
            row[1].text = str(data['count'])
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"
        
        format_table_text(tab_order_table)
        
        # 3. Interactive Element Issues
        doc.add_paragraph()
        doc.add_paragraph("Interactive Element Issues:", style='Normal').bold = True
        
        interactive_table = doc.add_table(rows=3 + 1, cols=4)
        interactive_table.style = 'Table Grid'
        
        # Add headers
        headers = interactive_table.rows[0].cells
        headers[0].text = "Issue Type"
        headers[1].text = "Total Count"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Sites"
        
        # Add data
        interactive_keys = ['mouse_only', 'missing_tabindex', 'non_interactive']
        for i, key in enumerate(interactive_keys, 1):
            row = interactive_table.rows[i].cells
            data = property_data[key]
            row[0].text = data['name']
            row[1].text = str(data['count'])
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"
        
        format_table_text(interactive_table)
        
        # Technical implementation recommendations
        doc.add_paragraph()
        doc.add_heading('Technical Implementation Recommendations', level=3)
        
        # Mouse-only elements
        doc.add_paragraph("Fixing Mouse-only Elements:", style='Normal').bold = True
        doc.add_paragraph("Add keyboard event handlers for all mouse interactions:", style='List Bullet')
        
        mouse_code = doc.add_paragraph("""
// Instead of this:
element.addEventListener('click', handleAction);

// Do this:
element.addEventListener('click', handleAction);
element.addEventListener('keydown', function(e) {
  if (e.key === 'Enter' || e.key === ' ') {
    handleAction(e);
  }
});
        """)
        mouse_code.style = doc.styles['Normal']
        mouse_code.paragraph_format.left_indent = Pt(36)
        
        # tabindex issues
        doc.add_paragraph("Proper tabindex Usage:", style='Normal').bold = True
        doc.add_paragraph("Use tabindex='0' to add elements to the natural tab order", style='List Bullet')
        doc.add_paragraph("Use tabindex='-1' for elements that should be focusable by script but not tab", style='List Bullet')
        doc.add_paragraph("Avoid positive tabindex values which override the natural tab order", style='List Bullet')
        
        tabindex_code = doc.add_paragraph("""
<!-- Good usage: Adding a div to the natural tab order -->
<div tabindex="0" role="button" aria-label="Action">Interactive element</div>

<!-- Good usage: Element that can receive focus via script but not tab -->
<div tabindex="-1" id="focusTarget">Script-focusable element</div>

<!-- Bad usage: Creating custom tab order -->
<button tabindex="1">First</button>
<button tabindex="3">Third</button>
<button tabindex="2">Second</button>
        """)
        tabindex_code.style = doc.styles['Normal']
        tabindex_code.paragraph_format.left_indent = Pt(36)
        
        # Case study: Domain with most issues
        if domain_data:
            worst_domain = max(domain_data.items(), 
                              key=lambda x: sum(url_data['total_violations'] for url_data in x[1]['urls'].values()))
            
            doc.add_paragraph()
            doc.add_heading(f'Case Study: {worst_domain[0]}', level=3)
            
            # Calculate total violations for this domain
            total_domain_violations = sum(url_data['total_violations'] for url_data in worst_domain[1]['urls'].values())
            doc.add_paragraph(f"This site has a total of {total_domain_violations} event handling violations across {len(worst_domain[1]['urls'])} pages.")
            
            # Create a breakdown of the top issues
            issue_counts = {
                'Mouse-only Elements': sum(url_data['violations'].get('mouse_only', 0) for url_data in worst_domain[1]['urls'].values()),
                'Visual Order Violations': sum(url_data['violations'].get('visual_order', 0) for url_data in worst_domain[1]['urls'].values()),
                'Explicit tabindex Usage': sum(url_data['violations'].get('explicit_tabindex', 0) for url_data in worst_domain[1]['urls'].values()),
                'Non-interactive with Handlers': sum(url_data['violations'].get('non_interactive', 0) for url_data in worst_domain[1]['urls'].values())
            }
            
            # Create a table with the top issues
            issue_table = doc.add_table(rows=len(issue_counts) + 1, cols=2)
            issue_table.style = 'Table Grid'
            
            # Add headers
            headers = issue_table.rows[0].cells
            headers[0].text = "Issue Type"
            headers[1].text = "Count"
            
            # Add data
            for i, (issue, count) in enumerate(sorted(issue_counts.items(), key=lambda x: x[1], reverse=True), 1):
                row = issue_table.rows[i].cells
                row[0].text = issue
                row[1].text = str(count)
            
            format_table_text(issue_table)
    else:
        doc.add_paragraph("No event handling data was found.")
        