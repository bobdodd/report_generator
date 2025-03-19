from report_styling import format_table_text
from docx.shared import Pt

def add_detailed_language(doc, db_connection, total_domains):
    """Add the detailed Language of Page section"""
    doc.add_page_break()
    h2 = doc.add_heading('Language of Page', level=2)
    h2.style = doc.styles['Heading 2']
    
    # Add explanation
    doc.add_paragraph("""
    The lang attribute on the <html> element is crucial for accessibility as it allows assistive technologies to determine the correct pronunciation rules for content. Without a properly specified language, screen readers may use incorrect pronunciation, making content difficult to understand for users who rely on text-to-speech.
    """.strip())
    
    doc.add_paragraph("The lang attribute also affects:", style='Normal')
    doc.add_paragraph("How screen readers pronounce content", style='List Bullet')
    doc.add_paragraph("How browsers apply language-specific typography and text processing", style='List Bullet')  
    doc.add_paragraph("Spell checking functionality", style='List Bullet')
    doc.add_paragraph("Hyphenation and other language-specific features", style='List Bullet')

    # If there are pages without lang attribute, list them
    pages_without_lang = list(db_connection.page_results.find(
        {"results.accessibility.tests.html_structure.html_structure.tests.hasValidLang": False},
        {"url": 1, "_id": 0}
    ).sort("url", 1))

    # Count affected domains
    affected_domains = set()
    for page in pages_without_lang:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        affected_domains.add(domain)

    # Calculate percentage
    percentage = (len(affected_domains) / len(total_domains)) * 100 if total_domains else 0

    if pages_without_lang:
        doc.add_paragraph(f"{len(pages_without_lang)} pages found without valid language attribute ({percentage:.1f}% of sites).")
        
        # Group pages by domain
        pages_by_domain = {}
        for page in pages_without_lang:
            domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
            if domain not in pages_by_domain:
                pages_by_domain[domain] = []
            pages_by_domain[domain].append(page['url'])
        
        # Create a summary table
        domain_table = doc.add_table(rows=len(pages_by_domain) + 1, cols=2)
        domain_table.style = 'Table Grid'
        
        # Set headers
        headers = domain_table.rows[0].cells
        headers[0].text = "Domain"
        headers[1].text = "Pages without language attribute"
        
        # Add data
        for i, (domain, pages) in enumerate(sorted(pages_by_domain.items()), 1):
            row = domain_table.rows[i].cells
            row[0].text = domain
            row[1].text = str(len(pages))
        
        format_table_text(domain_table)
        
        # Add example implementation
        doc.add_paragraph()
        doc.add_heading('Implementation Guidelines', level=3)
        
        doc.add_paragraph("Proper implementation of the lang attribute:", style='Normal').bold = True
        doc.add_paragraph("The lang attribute should be added to the html element:", style='List Bullet')
        
        lang_code = doc.add_paragraph("""
<!DOCTYPE html>
<html lang="en">
  <head>
    <meta charset="UTF-8">
    <title>Page Title</title>
    ...
  </head>
  <body>
    ...
  </body>
</html>
        """)
        lang_code.style = doc.styles['Normal']
        lang_code.paragraph_format.left_indent = Pt(36)
        
        # Add examples for other languages
        doc.add_paragraph("Examples for other languages:", style='List Bullet')
        
        examples_code = doc.add_paragraph("""
<html lang="fr"> <!-- French -->
<html lang="es"> <!-- Spanish -->
<html lang="zh-CN"> <!-- Simplified Chinese -->
<html lang="ar"> <!-- Arabic -->
<html lang="ru"> <!-- Russian -->
        """)
        examples_code.style = doc.styles['Normal']
        examples_code.paragraph_format.left_indent = Pt(36)
        
        # Content in multiple languages
        doc.add_paragraph("For content in multiple languages:", style='List Bullet')
        
        multi_lang_code = doc.add_paragraph("""
<html lang="en">
  <body>
    <p>This is English text.</p>
    <p lang="fr">Ceci est du texte français.</p>
    <p lang="es">Este es texto en español.</p>
  </body>
</html>
        """)
        multi_lang_code.style = doc.styles['Normal']
        multi_lang_code.paragraph_format.left_indent = Pt(36)
        
        # Common issues
        doc.add_paragraph("Common issues to avoid:", style='Normal').bold = True
        
        issues_code = doc.add_paragraph("""
<!-- Missing lang attribute -->
<html>
  <head>...</head>
  <body>...</body>
</html>

<!-- Incorrect format -->
<html language="en"> <!-- Wrong attribute name -->
<html lang=en> <!-- Missing quotes -->
<html lang="english"> <!-- Not using ISO language code -->

<!-- Inconsistent language -->
<html lang="en">
  <body>
    <!-- Content is actually in French but marked as English -->
    <p>Ceci est du texte français sans attribut lang correct.</p>
  </body>
</html>
        """)
        issues_code.style = doc.styles['Normal']
        issues_code.paragraph_format.left_indent = Pt(36)

    else:
        doc.add_paragraph("All pages have a valid lang attribute. This is excellent practice for accessibility.")
        