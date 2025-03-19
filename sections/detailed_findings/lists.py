from report_styling import format_table_text
from docx.shared import Pt

def add_detailed_lists(doc, db_connection, total_domains):
    """Add the detailed Lists section"""
    doc.add_page_break()
    h2 = doc.add_heading('Lists', level=2)
    h2.style = doc.styles['Heading 2']
    
    # Add explanation
    doc.add_paragraph("""
Proper semantic list markup is important for accessibility. Lists should use appropriate HTML elements (ul, ol, li) rather than visual formatting to create list-like structures. This helps screen reader users understand content structure and navigate through items more efficiently.
""".strip())

    doc.add_paragraph("Common list accessibility issues include:", style='Normal')
    doc.add_paragraph("Using DIVs with bullets or numbers instead of proper list elements", style='List Bullet')
    doc.add_paragraph("Empty lists that serve no semantic purpose", style='List Bullet')
    doc.add_paragraph("Custom bullet implementations that may not be accessible", style='List Bullet')
    doc.add_paragraph("Excessively deep nesting of lists", style='List Bullet')

    doc.add_paragraph()

    # Query for pages with list issues
    pages_with_list_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.lists.lists.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.lists.lists.pageFlags.hasEmptyLists": True},
                {"results.accessibility.tests.lists.lists.pageFlags.hasFakeLists": True},
                {"results.accessibility.tests.lists.lists.pageFlags.hasCustomBullets": True},
                {"results.accessibility.tests.lists.lists.pageFlags.hasDeepNesting": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.lists.lists.pageFlags": 1,
            "results.accessibility.tests.lists.lists.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    list_issues = {
        "hasEmptyLists": {"name": "Empty lists", "pages": set(), "domains": set()},
        "hasFakeLists": {"name": "Fake lists (not using proper HTML)", "pages": set(), "domains": set()},
        "hasCustomBullets": {"name": "Custom bullet implementations", "pages": set(), "domains": set()},
        "hasDeepNesting": {"name": "Excessively nested lists", "pages": set(), "domains": set()}
    }

    # Count issues
    for page in pages_with_list_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        flags = page['results']['accessibility']['tests']['lists']['lists']['pageFlags']
        
        for flag in list_issues:
            if flags.get(flag, False):
                list_issues[flag]['pages'].add(page['url'])
                list_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in list_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "List Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add domain details for each issue
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)
        
        # Add technical implementation section
        doc.add_paragraph()
        doc.add_heading('Technical Implementation Guidelines', level=3)
        
        # Proper list examples
        doc.add_paragraph("Proper semantic list implementations:", style='Normal').bold = True
        
        proper_list_code = doc.add_paragraph("""
<!-- Unordered list -->
<ul>
  <li>First item</li>
  <li>Second item</li>
  <li>Third item</li>
</ul>

<!-- Ordered list -->
<ol>
  <li>Step 1</li>
  <li>Step 2</li>
  <li>Step 3</li>
</ol>

<!-- Description list -->
<dl>
  <dt>Term 1</dt>
  <dd>Definition 1</dd>
  <dt>Term 2</dt>
  <dd>Definition 2</dd>
</dl>
        """)
        proper_list_code.style = doc.styles['Normal']
        proper_list_code.paragraph_format.left_indent = Pt(36)
        
        # Examples of issues to avoid
        doc.add_paragraph("Common issues to avoid:", style='Normal').bold = True
        
        issues_code = doc.add_paragraph("""
<!-- Empty list -->
<ul></ul>

<!-- Fake list using divs and CSS -->
<div class="fake-list">
  <div class="list-item">• First item</div>
  <div class="list-item">• Second item</div>
</div>

<!-- Custom bullets with accessibility issues -->
<ul style="list-style: none;">
  <li><span class="custom-bullet">→</span> First item</li>
  <li><span class="custom-bullet">→</span> Second item</li>
</ul>

<!-- Excessively nested lists -->
<ul>
  <li>Level 1
    <ul>
      <li>Level 2
        <ul>
          <li>Level 3
            <ul>
              <li>Level 4 (getting excessive)
                <ul>
                  <li>Level 5 (too deep)</li>
                </ul>
              </li>
            </ul>
          </li>
        </ul>
      </li>
    </ul>
  </li>
</ul>
        """)
        issues_code.style = doc.styles['Normal']
        issues_code.paragraph_format.left_indent = Pt(36)
        
        # Accessible custom bullets
        doc.add_paragraph("Accessible custom bullets (if needed):", style='Normal').bold = True
        
        custom_bullets_code = doc.add_paragraph("""
<!-- Using CSS for custom bullets -->
<ul class="custom-bullets">
  <li>First item</li>
  <li>Second item</li>
</ul>

<style>
  .custom-bullets {
    list-style: none;
    padding-left: 1.5em;
  }
  
  .custom-bullets li {
    position: relative;
  }
  
  .custom-bullets li::before {
    content: "→";  /* Custom bullet character */
    position: absolute;
    left: -1.2em;
  }
</style>
        """)
        custom_bullets_code.style = doc.styles['Normal']
        custom_bullets_code.paragraph_format.left_indent = Pt(36)
        
        # Best practices for list accessibility
        doc.add_paragraph("Best practices for list accessibility:", style='Normal').bold = True
        
        doc.add_paragraph("Keep list structure simple and avoid excessive nesting", style='List Bullet')
        doc.add_paragraph("Use semantic HTML elements (ul, ol, li) rather than CSS for visual styling", style='List Bullet')
        doc.add_paragraph("When custom styling is needed, apply it through CSS while preserving semantic structure", style='List Bullet')
        doc.add_paragraph("Avoid empty lists or lists with only one item (consider using a paragraph instead)", style='List Bullet')
        doc.add_paragraph("Ensure list items are concise and organized logically", style='List Bullet')
        doc.add_paragraph("For complex nested lists, consider breaking them into separate lists with descriptive headings", style='List Bullet')

    else:
        doc.add_paragraph("No list markup issues were found.")
        