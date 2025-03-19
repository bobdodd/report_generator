# sections/detailed_findings/tabindex.py
from report_styling import format_table_text
from docx.shared import Pt

def add_detailed_tabindex(doc, db_connection, total_domains):
    """Add the detailed Tabindex section"""
    doc.add_page_break()
    h2 = doc.add_heading('Tabindex', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
The tabindex attribute controls whether, and in what order elements can be focused using the keyboard. Improper use of tabindex can disrupt the natural tab order and create accessibility barriers:
""".strip())

    # Add bullet points
    doc.add_paragraph("Positive tabindex values force elements into a specific tab order, which can be confusing and unpredictable", style='List Bullet')
    doc.add_paragraph("Non-interactive elements with tabindex=0 create unnecessary tab stops", style='List Bullet')
    doc.add_paragraph("Interactive elements without proper tabindex may be unreachable by keyboard", style='List Bullet')
    doc.add_paragraph("SVG elements need special consideration for keyboard accessibility", style='List Bullet')

    doc.add_paragraph()

    # Query for pages with tabindex issues
    pages_with_tabindex_issues = list(db_connection.page_results.find(
        {"results.accessibility.tests.tabindex.tabindex.pageFlags": {"$exists": True}},
        {
            "url": 1,
            "results.accessibility.tests.tabindex.tabindex.pageFlags": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    tabindex_issues = {
        "hasPositiveTabindex": {"name": "Elements with positive tabindex", "pages": set(), "domains": set()},
        "hasNonInteractiveZeroTabindex": {"name": "Non-interactive elements with tabindex=0", "pages": set(), "domains": set()},
        "hasMissingRequiredTabindex": {"name": "Interactive elements missing required tabindex", "pages": set(), "domains": set()},
        "hasSvgTabindexWarnings": {"name": "SVG elements with tabindex warnings", "pages": set(), "domains": set()}
    }

    # Count issues
    for page in pages_with_tabindex_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        flags = page['results']['accessibility']['tests']['tabindex']['tabindex']['pageFlags']
        
        for flag in tabindex_issues:
            if flags.get(flag, False):  # If issue exists (True)
                tabindex_issues[flag]['pages'].add(page['url'])
                tabindex_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in tabindex_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add domain details for each issue
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)
                
        # Add technical implementation section
        doc.add_paragraph()
        doc.add_heading('Technical Implementation Guidelines', level=3)
        
        # Problematic tabindex examples
        doc.add_paragraph("Problematic tabindex usage:", style='Normal').bold = True
        
        bad_examples = doc.add_paragraph("""
<!-- Problematic: Positive tabindex -->
<div>
  <button tabindex="1">First button</button>
  <button tabindex="3">Third button</button>
  <button tabindex="2">Second button</button>
  <button>Regular button</button>
</div>

<!-- Problematic: tabindex=0 on non-interactive element -->
<p tabindex="0">This paragraph can receive focus but has no interaction.</p>

<!-- Problematic: div made focusable without proper keyboard handling -->
<div tabindex="0" onclick="handleClick()">
  Click me
</div>
        """)
        bad_examples.style = doc.styles['Normal']
        bad_examples.paragraph_format.left_indent = Pt(36)
        
        # Proper tabindex usage
        doc.add_paragraph("Proper tabindex usage:", style='Normal').bold = True
        
        good_examples = doc.add_paragraph("""
<!-- Good: tabindex="-1" to programmatically focus -->
<div id="notification" role="alert" tabindex="-1">
  Your settings have been saved
</div>

<!-- Good: Using tabindex="0" with interactive role -->
<div tabindex="0" role="button" aria-pressed="false"
     onkeydown="if(event.key === ' ' || event.key === 'Enter') this.click()"
     onclick="toggleState(this)">
  Toggle Feature
</div>

<!-- Good: Natural tab order without tabindex -->
<div>
  <button>First button</button>
  <button>Second button</button>
  <button>Third button</button>
</div>

<!-- Good: Accessibility-enhanced custom widget -->
<div class="custom-select" role="combobox" tabindex="0"
     aria-expanded="false" aria-haspopup="listbox" aria-labelledby="label">
  <span id="label">Choose an option</span>
  <span class="dropdown-arrow" aria-hidden="true">▼</span>
  <div class="options" role="listbox" tabindex="-1">
    <div role="option" tabindex="-1">Option 1</div>
    <div role="option" tabindex="-1">Option 2</div>
    <div role="option" tabindex="-1">Option 3</div>
  </div>
</div>
        """)
        good_examples.style = doc.styles['Normal']
        good_examples.paragraph_format.left_indent = Pt(36)
        
        # SVG accessibility
        doc.add_paragraph("SVG accessibility and tabindex:", style='Normal').bold = True
        
        svg_examples = doc.add_paragraph("""
<!-- SVG as a button -->
<svg role="button" tabindex="0" aria-label="Close dialog"
     onclick="closeDialog()" 
     onkeydown="if(event.key === ' ' || event.key === 'Enter') closeDialog()">
  <!-- SVG contents (X icon) -->
  <path d="M10 10 L20 20 M10 20 L20 10" stroke="black" stroke-width="2"/>
</svg>

<!-- Decorative SVG (not focusable) -->
<svg aria-hidden="true" focusable="false">
  <!-- Decorative SVG content -->
</svg>

<!-- Interactive SVG with accessible description -->
<svg role="img" aria-labelledby="chart-title chart-desc" tabindex="0">
  <title id="chart-title">Company Growth 2023</title>
  <desc id="chart-desc">Bar chart showing 24% growth in Q1, 18% in Q2, 32% in Q3, and 26% in Q4.</desc>
  <!-- SVG chart content -->
</svg>
        """)
        svg_examples.style = doc.styles['Normal']
        svg_examples.paragraph_format.left_indent = Pt(36)
        
        # JavaScript for managing focus
        doc.add_paragraph("JavaScript for focus management:", style='Normal').bold = True
        
        js_examples = doc.add_paragraph("""
// Show dialog and trap focus
function openDialog() {
  const dialog = document.getElementById('my-dialog');
  
  // Make dialog visible
  dialog.classList.remove('hidden');
  
  // Store previously focused element to restore later
  const previousFocus = document.activeElement;
  dialog.dataset.previousFocus = previousFocus;
  
  // Move focus to the first focusable element in the dialog
  const firstFocusable = dialog.querySelector('button, [href], input, select, textarea, [tabindex]:not([tabindex="-1"])');
  if (firstFocusable) {
    firstFocusable.focus();
  } else {
    // If no focusable elements, focus the dialog itself
    dialog.tabIndex = -1; 
    dialog.focus();
  }
  
  // Set up focus trap
  dialog.addEventListener('keydown', trapFocus);
}

// Close dialog and restore focus
function closeDialog() {
  const dialog = document.getElementById('my-dialog');
  
  // Hide dialog
  dialog.classList.add('hidden');
  
  // Remove focus trap
  dialog.removeEventListener('keydown', trapFocus);
  
  // Restore focus to previous element
  const previousFocus = dialog.dataset.previousFocus;
  if (previousFocus) {
    document.getElementById(previousFocus).focus();
  }
}

// Trap focus within the dialog
function trapFocus(e) {
  if (e.key !== 'Tab') return;
  
  const dialog = document.getElementById('my-dialog');
  const focusableElements = dialog.querySelectorAll(
    'button, [href], input, select, textarea, [tabindex]:not([tabindex="-1"])'
  );
  
  const firstFocusable = focusableElements[0];
  const lastFocusable = focusableElements[focusableElements.length - 1];
  
  // If Shift+Tab on first element, move to last element
  if (e.shiftKey && document.activeElement === firstFocusable) {
    e.preventDefault();
    lastFocusable.focus();
  } 
  // If Tab on last element, move to first element
  else if (!e.shiftKey && document.activeElement === lastFocusable) {
    e.preventDefault();
    firstFocusable.focus();
  }
}
        """)
        js_examples.style = doc.styles['Normal']
        js_examples.paragraph_format.left_indent = Pt(36)

    else:
        doc.add_paragraph("No tabindex accessibility issues were found.")
        