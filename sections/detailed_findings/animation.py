from docx.oxml import parse_xml
from report_styling import format_table_text

def parse_duration(duration_str):
    """Convert duration string to milliseconds"""
    if not duration_str or duration_str == '0ms':
        return 0
    
    value = float(duration_str.replace('ms', '').replace('s', ''))
    return value * 1000 if duration_str.endswith('s') else value

def add_detailed_animation(doc, db_connection, total_domains):
    """Add the detailed Animation section"""
    doc.add_page_break()
    h2 = doc.add_heading('Animation', level=2)
    h2.style = doc.styles['Heading 2']

    doc.add_paragraph("The 'prefers-reduced-motion' media query allows websites to respect a user's system-level preference for reduced motion. This accessibility feature is crucial for several user groups:")

    doc.add_paragraph("People with vestibular disorders who can experience dizziness, nausea, and disorientation from animated content", style='List Bullet')
    doc.add_paragraph("Users with attention-related disabilities who may find animations distracting and disruptive", style='List Bullet')
    doc.add_paragraph("People with migraine sensitivity who can be triggered by certain types of motion", style='List Bullet')
    doc.add_paragraph("Users with cognitive disabilities who may find it difficult to focus on content when animations are present", style='List Bullet')

    doc.add_paragraph("""
    When websites don't support reduced motion preferences, users who rely on this setting remain exposed to animations that could affect their ability to use the site or even cause physical discomfort. This is particularly important for essential services and information websites where users need to access content regardless of their motion sensitivity.
    """.strip())
                      
    doc.add_paragraph("""
    Note that "prefers-reduced-animation" does not mean no animation, but you do need to consider the impact of each, especially longer ones. In these tests long animations are those over 5 seconds, but in practice, any animation of over 1 second needs to respect the prefers-reduced-motion media query to help neuro-diverse users who may struggle to read a page with significant animation "noise".
    """.strip())

    doc.add_paragraph()  # Add space before the tables

    # Query for pages that have animations but lack reduced motion support
    pages_with_animation_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.animations.animations.pageFlags.hasAnimations": True,
            "results.accessibility.tests.animations.animations.pageFlags.lacksReducedMotionSupport": True
        },
        {
            "url": 1,
            "results.accessibility.tests.animations.animations": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Count affected domains and collect statistics
    domain_stats = {}
    for page in pages_with_animation_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        animation_data = page['results']['accessibility']['tests']['animations']['animations']
        summary = animation_data['details']['summary']
        page_flags = animation_data['pageFlags']['details']
        
        if domain not in domain_stats:
            domain_stats[domain] = {
                'pages': 0,
                'total_animations': 0,
                'infinite_animations': 0,
                'long_animations': 0,
                'shortest_animation': None,
                'longest_animation': None
            }
        
        domain_stats[domain]['pages'] += 1
        domain_stats[domain]['total_animations'] += summary['totalAnimations']
        domain_stats[domain]['infinite_animations'] += summary['infiniteAnimations']
        domain_stats[domain]['long_animations'] += summary['longDurationAnimations']
        
        # Update shortest animation
        if 'shortestAnimation' in page_flags and page_flags['shortestAnimation'] != '0ms':
            if domain_stats[domain]['shortest_animation'] is None:
                domain_stats[domain]['shortest_animation'] = page_flags['shortestAnimation']
            else:
                # Compare durations (convert to ms for comparison)
                current = parse_duration(domain_stats[domain]['shortest_animation'])
                new = parse_duration(page_flags['shortestAnimation'])
                if new < current:
                    domain_stats[domain]['shortest_animation'] = page_flags['shortestAnimation']
        
        # Update longest animation
        if 'longestAnimationElement' in page_flags and page_flags['longestAnimationElement'] and page_flags['longestAnimationElement']['duration']:
            if domain_stats[domain]['longest_animation'] is None:
                domain_stats[domain]['longest_animation'] = page_flags['longestAnimationElement']['duration']
            else:
                # Compare durations (convert to ms for comparison)
                current = parse_duration(domain_stats[domain]['longest_animation'])
                new = parse_duration(page_flags['longestAnimationElement']['duration'])
                if new > current:
                    domain_stats[domain]['longest_animation'] = page_flags['longestAnimationElement']['duration']
    
    # Create summary table
    if domain_stats:
        # Add paragraph to keep table with previous content
        last_para = doc.add_paragraph()
        last_para._element.get_or_add_pPr().append(
            parse_xml(r'<w:keepNext xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        )

        # Calculate totals for summary table
        total_pages = sum(stats['pages'] for stats in domain_stats.values())
        affected_domains = len(domain_stats)
        percentage = (affected_domains / len(total_domains)) * 100 if total_domains else 0

        summary_table = doc.add_table(rows=2, cols=4)
        summary_table.style = 'Table Grid'
        
        # Keep table together
        for row in summary_table.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcPr.append(parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>'))
                tcPr.append(parse_xml(r'<w:noWrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue"
        headers[1].text = "# of pages"
        headers[2].text = "# of sites affected"
        headers[3].text = "% of sites"

        # Add data
        row = summary_table.rows[1].cells
        row[0].text = "Pages with animations lacking reduced motion support"
        row[1].text = str(total_pages)
        row[2].text = str(affected_domains)
        row[3].text = f"{percentage:.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add some space after the table
        doc.add_paragraph()

        # Create detailed domain table
        last_para = doc.add_paragraph()
        last_para._element.get_or_add_pPr().append(
            parse_xml(r'<w:keepNext xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        )

        domain_table = doc.add_table(rows=len(domain_stats) + 1, cols=7)  # Updated number of columns
        domain_table.style = 'Table Grid'
        
        # Keep table together
        for row in domain_table.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcPr.append(parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>'))
                tcPr.append(parse_xml(r'<w:noWrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

        # Add headers
        headers = domain_table.rows[0].cells
        headers[0].text = "Domain"
        headers[1].text = "Pages without reduced motion support"
        headers[2].text = "Total animations"
        headers[3].text = "Infinite animations"
        headers[4].text = "Long animations"
        headers[5].text = "Shortest animation"
        headers[6].text = "Longest animation"

        # Add domain data
        for i, (domain, stats) in enumerate(sorted(domain_stats.items()), 1):
            row = domain_table.rows[i].cells
            row[0].text = domain
            row[1].text = str(stats['pages'])
            row[2].text = str(stats['total_animations'])
            row[3].text = str(stats['infinite_animations'])
            row[4].text = str(stats['long_animations'])
            row[5].text = stats['shortest_animation'] or 'N/A'
            row[6].text = stats['longest_animation'] or 'N/A'

        # Format the table text
        format_table_text(domain_table)