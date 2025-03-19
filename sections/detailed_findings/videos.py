# sections/detailed_findings/videos.py
from report_styling import format_table_text
from docx.shared import Pt

def add_detailed_videos(doc, db_connection, total_domains):
    """Add the detailed Videos section"""
    doc.add_page_break()
    h2 = doc.add_heading('Videos', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
    Video content must be accessible to all users, including those with visual or hearing impairments. Videos should include appropriate alternatives and controls. Common accessibility issues with video content include:
    """.strip())

    doc.add_paragraph("Missing closed captions for audio content", style='List Bullet')
    doc.add_paragraph("Lack of audio descriptions for visual information", style='List Bullet')
    doc.add_paragraph("Inaccessible video controls", style='List Bullet')
    doc.add_paragraph("Missing transcripts", style='List Bullet')
    doc.add_paragraph("Autoplay videos without user control", style='List Bullet')
    doc.add_paragraph("Videos without proper labels or titles", style='List Bullet')

    # Add recommendations paragraph
    doc.add_paragraph()
    doc.add_paragraph("Recommendations for Video Accessibility:")
    doc.add_paragraph("Ensure all videos have accurate closed captions that include both speech and important sound effects.", style='List Bullet')
    doc.add_paragraph("Provide audio descriptions for important visual information when necessary.", style='List Bullet')
    doc.add_paragraph("Include complete transcripts for all video content.", style='List Bullet')
    doc.add_paragraph("Ensure video players have keyboard-accessible controls.", style='List Bullet')
    doc.add_paragraph("Avoid autoplay or provide easy controls to stop playback.", style='List Bullet')
    doc.add_paragraph("Include clear, descriptive titles and labels for all video content.", style='List Bullet')

    doc.add_paragraph()

    # Query for pages with video issues
    pages_with_video_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.video.video.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.video.video.pageFlags.missingCaptions": True},
                {"results.accessibility.tests.video.video.pageFlags.missingAudioDescription": True},
                {"results.accessibility.tests.video.video.pageFlags.inaccessibleControls": True},
                {"results.accessibility.tests.video.video.pageFlags.missingTranscript": True},
                {"results.accessibility.tests.video.video.pageFlags.hasAutoplay": True},
                {"results.accessibility.tests.video.video.pageFlags.missingLabels": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.video.video.pageFlags": 1,
            "results.accessibility.tests.video.video.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    video_issues = {
        "missingCaptions": {"name": "Missing closed captions", "pages": set(), "domains": set()},
        "missingAudioDescription": {"name": "Missing audio descriptions", "pages": set(), "domains": set()},
        "inaccessibleControls": {"name": "Inaccessible video controls", "pages": set(), "domains": set()},
        "missingTranscript": {"name": "Missing transcripts", "pages": set(), "domains": set()},
        "hasAutoplay": {"name": "Autoplay without user control", "pages": set(), "domains": set()},
        "missingLabels": {"name": "Missing video labels/titles", "pages": set(), "domains": set()}
    }

    # Count issues
    if len(pages_with_video_issues) > 0:
        for page in pages_with_video_issues:
            domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
            flags = page['results']['accessibility']['tests']['video']['video']['pageFlags']
            
            for flag in video_issues:
                if flags.get(flag, False):
                    video_issues[flag]['pages'].add(page['url'])
                    video_issues[flag]['domains'].add(domain)

        # Create filtered list of issues that have affected pages
        active_issues = {flag: data for flag, data in video_issues.items() 
                        if len(data['pages']) > 0}

        if active_issues:
            # Create summary table
            summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
            summary_table.style = 'Table Grid'

            # Set column headers
            headers = summary_table.rows[0].cells
            headers[0].text = "Video Issue"
            headers[1].text = "Pages Affected"
            headers[2].text = "Sites Affected"
            headers[3].text = "% of Total Sites"

            # Add data
            for i, (flag, data) in enumerate(active_issues.items(), 1):
                row = summary_table.rows[i].cells
                row[0].text = data['name']
                row[1].text = str(len(data['pages']))
                row[2].text = str(len(data['domains']))
                row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

            # Format the table text
            format_table_text(summary_table)

            # Add domain details for each issue
            for flag, data in active_issues.items():
                if data['domains']:
                    doc.add_paragraph()
                    doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                    
                    # Group by domain and count occurrences
                    domain_counts = {}
                    for page in data['pages']:
                        domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                        domain_counts[domain] = domain_counts.get(domain, 0) + 1

                    # Create domain details table
                    domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                    domain_table.style = 'Table Grid'

                    # Add headers
                    headers = domain_table.rows[0].cells
                    headers[0].text = "Domain"
                    headers[1].text = "Number of pages"

                    # Add domain data
                    for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                        row = domain_table.rows[i].cells
                        row[0].text = domain
                        row[1].text = str(count)

                    # Format the table text
                    format_table_text(domain_table)

            # Add technical implementation section
            doc.add_paragraph()
            doc.add_heading('Technical Implementation Guidelines', level=3)
            
            # YouTube embed with captions and controls
            doc.add_paragraph("Accessible YouTube Embed:", style='Normal').bold = True
            
            youtube_example = doc.add_paragraph("""
<!-- Accessible YouTube embed with captions enabled -->
<div class="video-container">
  <h3 id="video-title">Company Overview Video</h3>
  <iframe 
    width="560" 
    height="315" 
    src="https://www.youtube.com/embed/VIDEO_ID?cc_load_policy=1&cc_lang_pref=en&rel=0" 
    title="Company Overview Video"
    allow="accelerometer; encrypted-media; gyroscope; picture-in-picture" 
    allowfullscreen>
  </iframe>
  
  <!-- Transcript toggle button -->
  <button type="button" aria-expanded="false" aria-controls="video-transcript">
    Show Transcript
  </button>
  
  <!-- Transcript with proper structure -->
  <div id="video-transcript" hidden>
    <h4>Video Transcript</h4>
    <p><strong>[Music]</strong></p>
    <p><strong>Narrator:</strong> Welcome to our company overview.</p>
    <p><strong>CEO:</strong> Since our founding in 2005, we've been committed to...</p>
    <!-- More transcript content -->
  </div>
</div>
            """)
            youtube_example.style = doc.styles['Normal']
            youtube_example.paragraph_format.left_indent = Pt(36)
            
            # Video player with audio description
            doc.add_paragraph("HTML5 Video with Audio Description:", style='Normal').bold = True
            
            html5_example = doc.add_paragraph("""
<!-- HTML5 video with audio description track -->
<div class="video-container">
  <h3 id="product-demo">Product Demonstration</h3>
  <video 
    controls 
    width="560" 
    height="315"
    poster="video-thumbnail.jpg"
    aria-describedby="video-description">
    
    <!-- Video sources for different formats -->
    <source src="product-demo.mp4" type="video/mp4">
    <source src="product-demo.webm" type="video/webm">
    
    <!-- Caption tracks -->
    <track 
      kind="subtitles" 
      label="English" 
      src="captions-en.vtt" 
      srclang="en" 
      default>
    <track 
      kind="subtitles" 
      label="Spanish" 
      src="captions-es.vtt" 
      srclang="es">
    
    <!-- Audio description track -->
    <track 
      kind="descriptions" 
      label="Audio Descriptions" 
      src="descriptions-en.vtt" 
      srclang="en">
    
    <!-- Fallback for browsers that don't support video element -->
    <p>Your browser doesn't support HTML5 video. 
       Here is a <a href="product-demo.mp4">link to the video</a> instead.</p>
  </video>
  
  <!-- Description for screen readers -->
  <p id="video-description" class="sr-only">
    This video demonstrates how to set up and use our product, 
    showing step-by-step instructions for installation and configuration.
  </p>
  
  <!-- Transcript section -->
  <details>
    <summary>Video Transcript</summary>
    <div class="transcript">
      <p>[0:00] <strong>Narrator:</strong> Welcome to our product demonstration.</p>
      <p>[0:05] To begin installation, first download the package from our website.</p>
      <p>[0:12] <strong>[Visual: Screen shows download button highlighted]</strong></p>
      <!-- Full transcript content -->
    </div>
  </details>
</div>

<!-- Additional CSS for accessibility -->
<style>
  .video-container {
    max-width: 100%;
    margin: 2em 0;
  }
  
  video {
    max-width: 100%;
    height: auto;
  }
  
  .sr-only {
    position: absolute;
    width: 1px;
    height: 1px;
    margin: -1px;
    padding: 0;
    overflow: hidden;
    clip: rect(0, 0, 0, 0);
    border: 0;
  }
  
  .transcript {
    max-height: 300px;
    overflow-y: auto;
    padding: 1em;
    border: 1px solid #ddd;
    margin-top: 1em;
  }
</style>
            """)
            html5_example.style = doc.styles['Normal']
            html5_example.paragraph_format.left_indent = Pt(36)
            
            # JavaScript for transcript toggle
            doc.add_paragraph("JavaScript for Transcript Toggle:", style='Normal').bold = True
            
            js_example = doc.add_paragraph("""
// JavaScript to handle transcript toggle
document.addEventListener('DOMContentLoaded', function() {
  const transcriptButtons = document.querySelectorAll('button[aria-controls]');
  
  transcriptButtons.forEach(button => {
    const transcriptId = button.getAttribute('aria-controls');
    const transcript = document.getElementById(transcriptId);
    
    button.addEventListener('click', function() {
      // Get current state
      const expanded = button.getAttribute('aria-expanded') === 'true';
      
      // Toggle state
      button.setAttribute('aria-expanded', !expanded);
      
      if (expanded) {
        // Hide transcript
        transcript.hidden = true;
        button.textContent = 'Show Transcript';
      } else {
        // Show transcript
        transcript.hidden = false;
        button.textContent = 'Hide Transcript';
      }
    });
  });
  
  // Automatically pause videos when description track becomes active
  const videos = document.querySelectorAll('video');
  videos.forEach(video => {
    const descTrack = Array.from(video.textTracks)
                        .find(track => track.kind === 'descriptions');
    
    if (descTrack) {
      descTrack.addEventListener('cuechange', function() {
        // If a description cue is active and video is playing, pause it
        // to give time to listen to the description
        if (this.activeCues.length > 0 && !video.paused) {
          video.pause();
          
          // Resume play after the description finishes
          const cue = this.activeCues[0];
          const duration = cue.endTime - cue.startTime;
          
          setTimeout(() => {
            video.play();
          }, duration * 1000);
        }
      });
    }
  });
});
            """)
            js_example.style = doc.styles['Normal']
            js_example.paragraph_format.left_indent = Pt(36)
            
            # WebVTT caption file example
            doc.add_paragraph("WebVTT Caption File Example:", style='Normal').bold = True
            
            webvtt_example = doc.add_paragraph("""
WEBVTT

00:00:01.000 --> 00:00:05.000
Welcome to our product demonstration.

00:00:05.500 --> 00:00:10.000
To begin installation, first download the package from our website.

00:00:10.500 --> 00:00:15.000
Click the Download button highlighted in green on the top right corner.

00:00:15.500 --> 00:00:20.000
[Upbeat music plays while download starts]

00:00:20.500 --> 00:00:25.000
Once downloaded, open the installer package.
            """)
            webvtt_example.style = doc.styles['Normal']
            webvtt_example.paragraph_format.left_indent = Pt(36)
            
            # Audio description VTT file
            doc.add_paragraph("WebVTT Audio Description File Example:", style='Normal').bold = True
            
            desc_example = doc.add_paragraph("""
WEBVTT

00:00:03.000 --> 00:00:05.000
The screen shows the company's homepage with a prominent banner.

00:00:12.000 --> 00:00:15.000
A green download button appears in the top right corner.

00:00:22.000 --> 00:00:26.000
The installer package opens, showing a welcome screen with the company logo.

00:00:35.000 --> 00:00:40.000
A progress bar appears, filling from left to right as the installation proceeds.
            """)
            desc_example.style = doc.styles['Normal']
            desc_example.paragraph_format.left_indent = Pt(36)

        else:
            doc.add_paragraph("No video accessibility issues were found.")
    else:
        doc.add_paragraph("No videos were found.")