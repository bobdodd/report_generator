# sections/detailed_findings/timers.py
from report_styling import format_table_text
from docx.shared import Pt

def add_detailed_timers(doc, db_connection, total_domains):
    """Add the detailed Timers section"""
    doc.add_page_break()
    h2 = doc.add_heading('Timers', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
Time limits and automatic updates can create significant barriers for users who need more time to read content or complete tasks. Common issues with timers include:
""".strip())

    doc.add_paragraph("Auto-starting timers that begin without user initiation", style='List Bullet')
    doc.add_paragraph("Timers without proper controls to pause, stop, or extend time", style='List Bullet')
    doc.add_paragraph("Session timeouts without adequate warning or ability to extend", style='List Bullet')
    doc.add_paragraph("Content that updates automatically without user control", style='List Bullet')

    # Add recommendations
    doc.add_paragraph()
    doc.add_paragraph("Recommendations for addressing timer issues:")
    doc.add_paragraph("Provide options to turn off, adjust, or extend time limits", style='List Bullet')
    doc.add_paragraph("Ensure all auto-updating content can be paused", style='List Bullet')
    doc.add_paragraph("Give adequate warning before session timeouts", style='List Bullet')
    doc.add_paragraph("Provide mechanisms to request more time", style='List Bullet')

    doc.add_paragraph()

    # Query for pages with timer issues
    pages_with_timer_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.timers.timers.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.timers.timers.pageFlags.hasTimers": True},
                {"results.accessibility.tests.timers.timers.pageFlags.hasAutoStartTimers": True},
                {"results.accessibility.tests.timers.timers.pageFlags.hasTimersWithoutControls": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.timers.timers.pageFlags": 1,
            "results.accessibility.tests.timers.timers.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    timer_issues = {
        "hasTimers": {"name": "Pages with timers", "pages": set(), "domains": set()},
        "hasAutoStartTimers": {"name": "Auto-starting timers", "pages": set(), "domains": set()},
        "hasTimersWithoutControls": {"name": "Timers without adequate controls", "pages": set(), "domains": set()}
    }

    # Count issues
    for page in pages_with_timer_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        flags = page['results']['accessibility']['tests']['timers']['timers']['pageFlags']
        
        for flag in timer_issues:
            if flags.get(flag, False):
                timer_issues[flag]['pages'].add(page['url'])
                timer_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in timer_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Timer Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add domain details for each issue
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)
                
        # Add technical implementation section
        doc.add_paragraph()
        doc.add_heading('Technical Implementation Guidelines', level=3)
        
        # Session timeout example
        doc.add_paragraph("Accessible Session Timeout Implementation:", style='Normal').bold = True
        
        session_timeout = doc.add_paragraph("""
// JavaScript for accessible session timeout warning
let sessionTimeoutMinutes = 20;
let warningTimeMinutes = 5;
let sessionTimer;
let warningTimer;
let sessionDialogVisible = false;

// Initialize session timeout on page load
function initSessionTimeout() {
  // Reset timers whenever there's user activity
  document.addEventListener('click', resetSessionTimers);
  document.addEventListener('keypress', resetSessionTimers);
  
  // Set initial timers
  resetSessionTimers();
}

// Reset the session timers
function resetSessionTimers() {
  // Clear existing timers
  clearTimeout(sessionTimer);
  clearTimeout(warningTimer);
  
  // Hide warning dialog if it's visible
  if (sessionDialogVisible) {
    hideSessionWarningDialog();
  }
  
  // Set new timers
  warningTimer = setTimeout(showSessionWarningDialog, 
                          (sessionTimeoutMinutes - warningTimeMinutes) * 60 * 1000);
  
  sessionTimer = setTimeout(handleSessionTimeout, 
                           sessionTimeoutMinutes * 60 * 1000);
}

// Show the warning dialog when session is about to expire
function showSessionWarningDialog() {
  const warningDialog = document.getElementById('session-warning-dialog');
  const minutesSpan = document.getElementById('minutes-remaining');
  
  // Update the minutes remaining
  minutesSpan.textContent = warningTimeMinutes;
  
  // Show the dialog
  warningDialog.classList.remove('hidden');
  warningDialog.setAttribute('aria-hidden', 'false');
  
  // Move focus to the dialog for screen reader announcement
  const firstButton = warningDialog.querySelector('button');
  if (firstButton) {
    firstButton.focus();
  }
  
  sessionDialogVisible = true;
  
  // Start countdown timer
  startCountdown(warningTimeMinutes * 60);
}

// Hide the session warning dialog
function hideSessionWarningDialog() {
  const warningDialog = document.getElementById('session-warning-dialog');
  warningDialog.classList.add('hidden');
  warningDialog.setAttribute('aria-hidden', 'true');
  sessionDialogVisible = false;
}

// Handle session timeout
function handleSessionTimeout() {
  // Redirect to logout page
  window.location.href = '/logout?reason=timeout';
}

// Extend the session
function extendSession() {
  // Make AJAX call to extend session on server
  fetch('/api/extend-session', { method: 'POST' })
    .then(response => response.json())
    .then(data => {
      if (data.success) {
        // Announce success to screen readers
        document.getElementById('session-status').textContent = 
          'Your session has been extended.';
        
        // Reset timers
        resetSessionTimers();
      }
    });
}

// Start countdown timer
function startCountdown(seconds) {
  const countdownElement = document.getElementById('countdown-timer');
  const countdownInterval = setInterval(() => {
    seconds--;
    
    // Format time as MM:SS
    const minutes = Math.floor(seconds / 60);
    const remainingSeconds = seconds % 60;
    countdownElement.textContent = 
      `${minutes}:${remainingSeconds < 10 ? '0' : ''}${remainingSeconds}`;
    
    // Update ARIA live region every 15 seconds or at important thresholds
    if (seconds % 15 === 0 || seconds === 30 || seconds === 10) {
      document.getElementById('countdown-announcement').textContent = 
        `${minutes} minute${minutes !== 1 ? 's' : ''} and ${remainingSeconds} second${remainingSeconds !== 1 ? 's' : ''} remaining before session timeout`;
    }
    
    if (seconds <= 0) {
      clearInterval(countdownInterval);
    }
  }, 1000);
}

// Call initialization on page load
window.addEventListener('DOMContentLoaded', initSessionTimeout);
        """)
        session_timeout.style = doc.styles['Normal']
        session_timeout.paragraph_format.left_indent = Pt(36)
        
        # Session timeout HTML example
        doc.add_paragraph("HTML for Session Timeout Dialog:", style='Normal').bold = True
        
        html_example = doc.add_paragraph("""
<!-- Session timeout warning dialog -->
<div id="session-warning-dialog" role="dialog" aria-modal="true" 
     aria-labelledby="dialog-title" aria-describedby="dialog-desc" class="hidden">
  <h2 id="dialog-title">Session Timeout Warning</h2>
  
  <div id="dialog-desc">
    <p>Your session will expire in <span id="minutes-remaining">5</span> minutes due to inactivity.</p>
    <p>Time remaining: <span id="countdown-timer">5:00</span></p>
    <p>Would you like to continue your session?</p>
  </div>
  
  <div class="dialog-buttons">
    <button type="button" onclick="extendSession()">Yes, Continue Session</button>
    <button type="button" onclick="handleSessionTimeout()">No, Log Out Now</button>
  </div>
  
  <!-- ARIA live regions for announcements -->
  <div aria-live="polite" id="countdown-announcement" class="sr-only"></div>
  <div aria-live="assertive" id="session-status" class="sr-only"></div>
</div>

<!-- CSS for dialog styling -->
<style>
  .hidden {
    display: none;
  }
  
  .sr-only {
    position: absolute;
    width: 1px;
    height: 1px;
    margin: -1px;
    padding: 0;
    overflow: hidden;
    clip: rect(0, 0, 0, 0);
    border: 0;
  }
  
  #session-warning-dialog {
    position: fixed;
    top: 50%;
    left: 50%;
    transform: translate(-50%, -50%);
    width: 400px;
    background: white;
    padding: 20px;
    border: 2px solid #333;
    border-radius: 5px;
    box-shadow: 0 0 10px rgba(0, 0, 0, 0.3);
    z-index: 1000;
  }
</style>
        """)
        html_example.style = doc.styles['Normal']
        html_example.paragraph_format.left_indent = Pt(36)
        
        # Carousel with accessible controls
        doc.add_paragraph("Accessible Carousel with Pause Controls:", style='Normal').bold = True
        
        carousel_example = doc.add_paragraph("""
<!-- HTML for accessible carousel -->
<div class="carousel" aria-roledescription="carousel" aria-label="Featured content">
  <!-- Carousel slides container -->
  <div class="carousel-slides" aria-live="polite">
    <div id="slide1" class="carousel-slide active" role="group" aria-roledescription="slide" aria-label="1 of 3">
      <img src="slide1.jpg" alt="Description of slide 1 image">
      <h3>Slide 1 Title</h3>
      <p>Slide 1 description text.</p>
    </div>
    <div id="slide2" class="carousel-slide" role="group" aria-roledescription="slide" aria-label="2 of 3">
      <img src="slide2.jpg" alt="Description of slide 2 image">
      <h3>Slide 2 Title</h3>
      <p>Slide 2 description text.</p>
    </div>
    <div id="slide3" class="carousel-slide" role="group" aria-roledescription="slide" aria-label="3 of 3">
      <img src="slide3.jpg" alt="Description of slide 3 image">
      <h3>Slide 3 Title</h3>
      <p>Slide 3 description text.</p>
    </div>
  </div>
  
  <!-- Carousel controls -->
  <div class="carousel-controls">
    <button type="button" class="carousel-prev" aria-label="Previous slide">
      <span aria-hidden="true">←</span>
    </button>
    
    <!-- Play/pause button -->
    <button type="button" class="carousel-pause" aria-pressed="false" aria-label="Pause automatic slide rotation">
      <span class="icon-pause" aria-hidden="true">⏸</span>
      <span class="icon-play hidden" aria-hidden="true">▶️</span>
    </button>
    
    <button type="button" class="carousel-next" aria-label="Next slide">
      <span aria-hidden="true">→</span>
    </button>
  </div>
  
  <!-- Slide indicators -->
  <div class="carousel-indicators" role="tablist" aria-label="Slide select">
    <button type="button" id="indicator1" class="active" aria-selected="true" 
            aria-label="Show slide 1" aria-controls="slide1" role="tab">1</button>
    <button type="button" id="indicator2" aria-selected="false" 
            aria-label="Show slide 2" aria-controls="slide2" role="tab">2</button>
    <button type="button" id="indicator3" aria-selected="false" 
            aria-label="Show slide 3" aria-controls="slide3" role="tab">3</button>
  </div>
  
  <!-- ARIA live region for announcements -->
  <div aria-live="polite" class="sr-only carousel-status"></div>
</div>

<!-- JavaScript for accessible carousel -->
<script>
  document.addEventListener('DOMContentLoaded', function() {
    const carousel = document.querySelector('.carousel');
    const slides = carousel.querySelectorAll('.carousel-slide');
    const pauseButton = carousel.querySelector('.carousel-pause');
    const prevButton = carousel.querySelector('.carousel-prev');
    const nextButton = carousel.querySelector('.carousel-next');
    const indicators = carousel.querySelectorAll('.carousel-indicators button');
    const statusRegion = carousel.querySelector('.carousel-status');
    
    // Respect user's preference for reduced motion
    let prefersReducedMotion = window.matchMedia('(prefers-reduced-motion: reduce)').matches;
    
    // Initial state
    let currentSlide = 0;
    let isPlaying = !prefersReducedMotion; // Don't autoplay if user prefers reduced motion
    let slideInterval;
    
    // Set pause button to reflect initial state
    updatePlayPauseState();
    
    // Start or stop autoplay based on initial preference
    if (isPlaying) {
      startSlideInterval();
    } else {
      pauseButton.setAttribute('aria-pressed', 'true');
    }
    
    // Toggle play/pause
    pauseButton.addEventListener('click', function() {
      isPlaying = !isPlaying;
      updatePlayPauseState();
      
      if (isPlaying) {
        startSlideInterval();
        statusRegion.textContent = "Slideshow playing automatically";
      } else {
        clearInterval(slideInterval);
        statusRegion.textContent = "Slideshow paused";
      }
    });
    
    // Previous slide
    prevButton.addEventListener('click', function() {
      goToSlide(currentSlide - 1);
    });
    
    // Next slide
    nextButton.addEventListener('click', function() {
      goToSlide(currentSlide + 1);
    });
    
    // Indicator buttons
    indicators.forEach((indicator, index) => {
      indicator.addEventListener('click', function() {
        goToSlide(index);
      });
    });
    
    // Pause on keyboard focus inside the carousel
    carousel.addEventListener('focusin', function() {
      if (isPlaying) {
        clearInterval(slideInterval);
      }
    });
    
    // Resume when focus leaves (if autoplay was on)
    carousel.addEventListener('focusout', function(e) {
      if (!carousel.contains(e.relatedTarget) && isPlaying) {
        startSlideInterval();
      }
    });
    
    // Pause on mouse hover
    carousel.addEventListener('mouseenter', function() {
      if (isPlaying) {
        clearInterval(slideInterval);
      }
    });
    
    // Resume on mouse leave (if autoplay was on)
    carousel.addEventListener('mouseleave', function() {
      if (isPlaying) {
        startSlideInterval();
      }
    });
    
    // Update play/pause button state
    function updatePlayPauseState() {
      pauseButton.setAttribute('aria-pressed', !isPlaying);
      
      const pauseIcon = pauseButton.querySelector('.icon-pause');
      const playIcon = pauseButton.querySelector('.icon-play');
      
      if (isPlaying) {
        pauseIcon.classList.remove('hidden');
        playIcon.classList.add('hidden');
        pauseButton.setAttribute('aria-label', 'Pause automatic slide rotation');
      } else {
        pauseIcon.classList.add('hidden');
        playIcon.classList.remove('hidden');
        pauseButton.setAttribute('aria-label', 'Start automatic slide rotation');
      }
    }
    
    // Go to a specific slide
    function goToSlide(index) {
      // Reset autoplay if it was on
      if (isPlaying) {
        clearInterval(slideInterval);
      }
      
      // Handle wrapping around
      if (index < 0) {
        index = slides.length - 1;
      } else if (index >= slides.length) {
        index = 0;
      }
      
      // Update active slide
      slides[currentSlide].classList.remove('active');
      slides[index].classList.add('active');
      
      // Update ARIA attributes for slides
      slides[currentSlide].setAttribute('aria-hidden', 'true');
      slides[index].setAttribute('aria-hidden', 'false');
      
      // Update indicators
      indicators[currentSlide].classList.remove('active');
      indicators[currentSlide].setAttribute('aria-selected', 'false');
      indicators[index].classList.add('active');
      indicators[index].setAttribute('aria-selected', 'true');
      
      // Announce slide change
      statusRegion.textContent = `Slide ${index + 1} of ${slides.length}`;
      
      // Update current slide index
      currentSlide = index;
      
      // Restart autoplay if it was on
      if (isPlaying) {
        startSlideInterval();
      }
    }
    
    // Start autoplay interval
    function startSlideInterval() {
      clearInterval(slideInterval);
      slideInterval = setInterval(() => {
        goToSlide(currentSlide + 1);
      }, 5000); // 5 second interval
    }
    
    // Listen for changes to the user's motion preference
    window.matchMedia('(prefers-reduced-motion: reduce)').addEventListener('change', (e) => {
      prefersReducedMotion = e.matches;
      if (prefersReducedMotion && isPlaying) {
        // Stop autoplay when user switches to preferring reduced motion
        isPlaying = false;
        updatePlayPauseState();
        clearInterval(slideInterval);
        statusRegion.textContent = "Slideshow paused due to reduced motion preference";
      }
    });
  });
</script>
        """)
        carousel_example.style = doc.styles['Normal']
        carousel_example.paragraph_format.left_indent = Pt(36)

    else:
        doc.add_paragraph("No timer-related issues were found.")