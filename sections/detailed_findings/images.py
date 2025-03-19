from report_styling import format_table_text
from docx.shared import Pt

def add_detailed_images(doc, db_connection, total_domains):
    """Add the detailed Images section"""
    doc.add_page_break()
    h2 = doc.add_heading('Images', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
    Images must be accessible to all users, including those using screen readers. This requires proper alternative text descriptions and appropriate ARIA roles. Images that convey information need descriptive alt text, while decorative images should be properly marked as such.
    """.strip())

    doc.add_paragraph("Common image accessibility issues include:", style='Normal')

    doc.add_paragraph("Missing alternative text for informative images", style='List Bullet')
    doc.add_paragraph("Invalid or uninformative alt text", style='List Bullet')
    doc.add_paragraph("Missing ARIA roles for SVG elements", style='List Bullet')
    doc.add_paragraph("Decorative images not properly marked", style='List Bullet')

    # Add recommendations
    doc.add_paragraph()
    doc.add_paragraph("Recommendations for Image Accessibility:", style='Normal')
    
    doc.add_paragraph("Provide meaningful alt text for all informative images", style='List Bullet')
    doc.add_paragraph("Use empty alt text (alt=\"\") for decorative images", style='List Bullet')
    doc.add_paragraph("Ensure SVG elements have appropriate ARIA roles", style='List Bullet')
    doc.add_paragraph("Make sure alt text is descriptive and conveys the image's purpose", style='List Bullet')
    doc.add_paragraph("Avoid using generic text like 'image' or 'photo' in alt attributes", style='List Bullet')
    doc.add_paragraph("Include text alternatives for complex images, charts, and graphs", style='List Bullet')

    doc.add_paragraph()

    # Query for pages with image issues
    pages_with_image_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.images.images.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.images.images.pageFlags.hasImagesWithoutAlt": True},
                {"results.accessibility.tests.images.images.pageFlags.hasImagesWithInvalidAlt": True},
                {"results.accessibility.tests.images.images.pageFlags.hasSVGWithoutRole": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.images.images": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for different image issues
    image_issues = {
        "missing_alt": {
            "name": "Missing alt text",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "invalid_alt": {
            "name": "Invalid alt text",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "missing_role": {
            "name": "SVGs missing role",
            "pages": set(),
            "domains": set(),
            "count": 0
        }
    }

    # Process each page
    total_images = 0
    total_decorative = 0

    for page in pages_with_image_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        image_data = page['results']['accessibility']['tests']['images']['images']
        flags = image_data['pageFlags']
        details = flags['details']
        
        # Count total and decorative images
        total_images += details.get('totalImages', 0)
        total_decorative += details.get('decorativeImages', 0)
        
        # Check missing alt text
        if flags.get('hasImagesWithoutAlt'):
            image_issues['missing_alt']['pages'].add(page['url'])
            image_issues['missing_alt']['domains'].add(domain)
            image_issues['missing_alt']['count'] += details.get('missingAlt', 0)
        
        # Check invalid alt text
        if flags.get('hasImagesWithInvalidAlt'):
            image_issues['invalid_alt']['pages'].add(page['url'])
            image_issues['invalid_alt']['domains'].add(domain)
            image_issues['invalid_alt']['count'] += details.get('invalidAlt', 0)
        
        # Check missing SVG roles
        if flags.get('hasSVGWithoutRole'):
            image_issues['missing_role']['pages'].add(page['url'])
            image_issues['missing_role']['domains'].add(domain)
            image_issues['missing_role']['count'] += details.get('missingRole', 0)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in image_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=5)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue Type"
        headers[1].text = "Number of Images"
        headers[2].text = "Pages Affected"
        headers[3].text = "Sites Affected"
        headers[4].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(data['count'])
            row[2].text = str(len(data['pages']))
            row[3].text = str(len(data['domains']))
            row[4].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add image statistics
        doc.add_paragraph()
        doc.add_paragraph("Image Statistics:", style='Normal')
        stats_table = doc.add_table(rows=3, cols=2)
        stats_table.style = 'Table Grid'

        # Add statistics data
        rows = stats_table.rows
        rows[0].cells[0].text = "Total Images"
        rows[0].cells[1].text = str(total_images)
        rows[1].cells[0].text = "Decorative Images"
        rows[1].cells[1].text = str(total_decorative)
        rows[2].cells[0].text = "Informative Images"
        rows[2].cells[1].text = str(total_images - total_decorative)

        # Format the table text
        format_table_text(stats_table)

        # Add domain details for each issue type
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)
                
        # Add technical implementation section
        doc.add_paragraph()
        doc.add_heading('Technical Implementation Guidelines', level=3)
        
        # Alt text examples
        doc.add_paragraph("Proper alt text implementation:", style='Normal').bold = True
        doc.add_paragraph("Descriptive alt text for informative images:", style='List Bullet')
        
        good_alt_code = doc.add_paragraph("""
<!-- Good practice: Descriptive alt text -->
<img src="map-to-office.png" alt="Map showing directions to our office from the train station">

<!-- Good practice: Empty alt for decorative images -->
<img src="decorative-swirl.png" alt="">

<!-- Good practice: For logos that link -->
<a href="/">
  <img src="company-logo.png" alt="Company Name - return to homepage">
</a>
        """)
        good_alt_code.style = doc.styles['Normal']
        good_alt_code.paragraph_format.left_indent = Pt(36)
        
        # Bad alt text examples
        doc.add_paragraph("Examples of poor alt text:", style='List Bullet')
        
        bad_alt_code = doc.add_paragraph("""
<!-- Bad practice: Missing alt attribute -->
<img src="important-chart.png">

<!-- Bad practice: Uninformative alt text -->
<img src="sales-graph.png" alt="image">

<!-- Bad practice: Redundant alt text -->
<img src="photo.jpg" alt="photo">

<!-- Bad practice: Filename as alt text -->
<img src="DSC1234.jpg" alt="DSC1234.jpg">
        """)
        bad_alt_code.style = doc.styles['Normal']
        bad_alt_code.paragraph_format.left_indent = Pt(36)
        
        # SVG accessibility
        doc.add_paragraph("SVG accessibility implementation:", style='Normal').bold = True
        
        svg_code = doc.add_paragraph("""
<!-- Decorative SVG -->
<svg aria-hidden="true" focusable="false">
  <!-- SVG content -->
</svg>

<!-- Informative SVG -->
<svg role="img" aria-labelledby="svg-title">
  <title id="svg-title">Description of what the SVG shows</title>
  <!-- SVG content -->
</svg>

<!-- Interactive SVG -->
<svg role="button" aria-label="Open menu" tabindex="0">
  <!-- SVG content -->
</svg>
        """)
        svg_code.style = doc.styles['Normal']
        svg_code.paragraph_format.left_indent = Pt(36)
        
        # Complex images
        doc.add_paragraph("Complex images (charts, infographics, maps):", style='Normal').bold = True
        
        complex_code = doc.add_paragraph("""
<!-- Chart with longer description -->
<figure>
  <img src="quarterly-sales.png" 
       alt="Q3 sales chart showing a 15% increase over Q2">
  <figcaption>
    <details>
      <summary>Detailed chart description</summary>
      <p>The chart shows quarterly sales for 2023. Q1 sales were $1.2M, 
         Q2 sales were $1.5M, and Q3 sales reached $1.725M, representing 
         a 15% increase over Q2 and a 43.75% increase year-to-date.</p>
    </details>
  </figcaption>
</figure>
        """)
        complex_code.style = doc.styles['Normal']
        complex_code.paragraph_format.left_indent = Pt(36)

    else:
        doc.add_paragraph("No image accessibility issues were found.")
        