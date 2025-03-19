import json
from docx.oxml import parse_xml
from report_styling import format_table_text

def add_detailed_color_as_indicator(doc, db_connection, total_domains):
    """Add the detailed Color as Indicator section"""
    doc.add_page_break()
    h2 = doc.add_heading('Colour as Indicator', level=2)
    h2.style = doc.styles['Heading 2']

    # Introduction
    doc.add_paragraph("""
Color should never be used as the only visual means of conveying information, indicating an action, prompting a response, or distinguishing a visual element. According to WCAG Success Criterion 1.4.1, color should not be the only visual means of conveying information.
    """.strip())
    
    doc.add_paragraph("This is critical for users with color blindness or other visual impairments, who may not be able to perceive color differences. It's also important for users who have monochrome displays or who print content in black and white.")
    
    doc.add_paragraph()

    # 5. Color-Only Links
    doc.add_paragraph("Color-Only Links:").bold = True
    doc.add_paragraph("Identifies links that are distinguished only by color", style='List Bullet')
    doc.add_paragraph("Links should have additional indicators (e.g., underlines, roles, or other visual markers)", style='List Bullet')
    doc.add_paragraph("Critical for users who are color blind or using monochrome displays", style='List Bullet')
    doc.add_paragraph()

    # 6. Color References
    doc.add_paragraph("Color References:").bold = True
    doc.add_paragraph("Detects content that relies on color to convey information", style='List Bullet')
    doc.add_paragraph("Looks for phrases like \"click the red button\" or \"items marked in blue\"", style='List Bullet')
    doc.add_paragraph("Must have alternative ways to convey the same information", style='List Bullet')
    doc.add_paragraph()

    # 7. Color Scheme Preferences Support
    doc.add_paragraph("Color Scheme Preferences Support:").bold = True
    doc.add_paragraph("Verifies support for system color scheme preferences (like dark mode)", style='List Bullet')
    doc.add_paragraph("Should respect user's preferred color scheme", style='List Bullet')
    doc.add_paragraph("Helps users who are sensitive to bright displays or need specific color schemes", style='List Bullet')
    doc.add_paragraph()

    # Define the indicator issues to be analyzed
    indicator_issues = [
        {
            'name': 'Color-Only Links',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.hasColorOnlyLinks',
            'details_field': 'results.accessibility.tests.colors.colors.details.summary.colorOnlyLinks'
        },
        {
            'name': 'Color References',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.hasColorReferences',
            'details_field': 'results.accessibility.tests.colors.colors.details.summary.colorReferenceCount'
        },
        {
            'name': 'Color Scheme Preferences Support',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.supportsColorSchemePreferences',
            'details_field': None  # This is a boolean field, not a count
        }
    ]

    # Gather the data for each issue type
    indicator_data = {}

    for issue in indicator_issues:
        # For the Color Scheme Preferences Support, we want sites that DO support it
        # For other issues, we want sites that have problems
        if issue['name'] == 'Color Scheme Preferences Support':
            query = {issue['db_field']: True}
        else:
            query = {issue['db_field']: True}
        
        # Prepare projection
        projection = {"url": 1, "_id": 0}
        if issue['details_field']:
            projection[issue['details_field']] = 1
        
        # Query the database to find pages with this issue
        pages_with_issue = list(db_connection.page_results.find(query, projection))
        
        # Count affected domains and total issue instances
        affected_domains = set()
        total_instances = 0
        
        for page in pages_with_issue:
            domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
            affected_domains.add(domain)
            
            # Count instances if applicable
            if issue['details_field']:
                # Navigate the nested structure to get the count
                parts = issue['details_field'].split('.')
                value = page
                try:
                    for part in parts:
                        if part in value:
                            value = value[part]
                        else:
                            value = 0
                            break
                    
                    if isinstance(value, (int, float)):
                        total_instances += value
                except:
                    pass  # Handle any issues with nested access
        
        # Store the data
        indicator_data[issue['name']] = {
            'pages': pages_with_issue,
            'domains': affected_domains,
            'instances': total_instances
        }

    # Create summary table
    last_para = doc.add_paragraph()
    last_para._element.get_or_add_pPr().append(
        parse_xml(r'<w:keepNext xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    )

    summary_table = doc.add_table(rows=len(indicator_issues) + 1, cols=4)
    summary_table.style = 'Table Grid'

    # Keep table together
    for row in summary_table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcPr.append(parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>'))
            tcPr.append(parse_xml(r'<w:noWrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

    # Set column headers
    headers = summary_table.rows[0].cells
    headers[0].text = "Color Accessibility Issue"
    headers[1].text = "Pages Affected"
    headers[2].text = "Sites Affected"
    headers[3].text = "% of Total Sites"

    # Add data
    for i, issue in enumerate(indicator_issues, 1):
        row = summary_table.rows[i].cells
        data = indicator_data[issue['name']]
        
        row[0].text = issue['name']
        row[1].text = str(len(data['pages']))
        row[2].text = str(len(data['domains']))
        
        percentage = (len(data['domains']) / len(total_domains)) * 100 if total_domains else 0
        row[3].text = f"{percentage:.1f}%"

    # Format the table text
    format_table_text(summary_table)

    # Add some space after the table
    doc.add_paragraph()

    # Add details for each issue type that has occurrences
    for issue in indicator_issues:
        data = indicator_data[issue['name']]
        if data['domains']:
            doc.add_paragraph(f"Sites with {issue['name'].lower()}:")
            
            # Create a dictionary to count pages per domain
            domain_counts = {}
            for page in data['pages']:
                domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
                domain_counts[domain] = domain_counts.get(domain, 0) + 1
            
            # Create domain details table
            domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
            domain_table.style = 'Table Grid'

            # Keep table together
            for row in domain_table.rows:
                for cell in row.cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcPr.append(parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>'))
                    tcPr.append(parse_xml(r'<w:noWrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

            # Add headers
            headers = domain_table.rows[0].cells
            headers[0].text = "Domain"
            headers[1].text = "Number of pages"

            # Add domain data
            for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                row = domain_table.rows[i].cells
                row[0].text = domain
                row[1].text = str(count)

            # Format the table text
            format_table_text(domain_table)
            doc.add_paragraph()  # Add space after each table
            
            # For Color References, provide more detailed information
            if issue['name'] == 'Color References':
                doc.add_paragraph("Color references found in content:", style='Heading 4')
                
                # Query for pages with color references to see the actual references
                color_refs_query = {
                    'results.accessibility.tests.colors.colors.pageFlags.hasColorReferences': True
                }
                color_refs_projection = {
                    'url': 1,
                    'results.accessibility.tests.colors.colors.details.colorReferences.instances': 1,
                    '_id': 0
                }
                
                pages_with_refs = list(db_connection.page_results.find(color_refs_query, color_refs_projection))
                
                # Extract all references and count them
                reference_counts = {}
                
                for page in pages_with_refs:
                    try:
                        # Navigate to get the instances
                        instances = page['results']['accessibility']['tests']['colors']['colors']['details']['colorReferences']['instances']
                        
                        # Handle case where instances might be a string (JSON)
                        if isinstance(instances, str):
                            try:
                                instances = json.loads(instances)
                            except:
                                continue
                        
                        # Process each instance
                        if isinstance(instances, list):
                            for instance in instances:
                                if isinstance(instance, dict) and 'references' in instance:
                                    for ref in instance['references']:
                                        ref = ref.lower()  # Normalize to lowercase
                                        reference_counts[ref] = reference_counts.get(ref, 0) + 1
                    except (KeyError, TypeError):
                        # Handle any issues with accessing the data
                        continue
                
                # Create a table of color references if any were found
                if reference_counts:
                    reference_table = doc.add_table(rows=len(reference_counts) + 1, cols=2)
                    reference_table.style = 'Table Grid'
                    
                    # Add headers
                    headers = reference_table.rows[0].cells
                    headers[0].text = "Color Referenced"
                    headers[1].text = "Number of Occurrences"
                    
                    # Add reference data
                    for i, (ref, count) in enumerate(sorted(reference_counts.items(), key=lambda x: x[1], reverse=True), 1):
                        row = reference_table.rows[i].cells
                        row[0].text = ref.capitalize()
                        row[1].text = str(count)
                    
                    # Format the table text
                    format_table_text(reference_table)
                    doc.add_paragraph()  # Add space after the table
                    
                    # Provide an explanation about the impact
                    doc.add_paragraph(
                        "Color references in content (like 'click the red button' or 'items marked in blue') "
                        "create barriers for users who are color blind or using screen readers. Information conveyed "
                        "through color alone must have alternative indicators that don't rely on color perception."
                    )
                else:
                    doc.add_paragraph("No specific color references were identified in the scanned content.")
                    
    # Add recommendations
    doc.add_paragraph()
    doc.add_heading('Recommendations for Addressing Color as Indicator Issues', level=3)
    
    doc.add_paragraph("For Color-Only Links:", style='Normal').bold = True
    doc.add_paragraph("Always include an underline or other non-color indicator for links", style='List Bullet')
    doc.add_paragraph("Ensure sufficient contrast between link and non-link text", style='List Bullet')
    doc.add_paragraph("Consider adding icons or other visual cues alongside links", style='List Bullet')
    
    doc.add_paragraph("For Color References in Content:", style='Normal').bold = True
    doc.add_paragraph("Avoid phrases like 'click the red button' or 'items in green'", style='List Bullet')
    doc.add_paragraph("Add additional descriptors like position, shape, or labels", style='List Bullet')
    doc.add_paragraph("Use patterns, icons, or text labels in addition to color", style='List Bullet')
    
    doc.add_paragraph("For Color Scheme Preferences:", style='Normal').bold = True
    doc.add_paragraph("Implement support for prefers-color-scheme media query", style='List Bullet')
    doc.add_paragraph("Design with both light and dark modes in mind", style='List Bullet')
    doc.add_paragraph("Ensure adequate contrast in all color schemes", style='List Bullet')
    