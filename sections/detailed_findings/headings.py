from report_styling import format_table_text
from docx.shared import Pt

def add_detailed_headings(doc, db_connection, total_domains):
    """Add the detailed Headings section"""
    doc.add_page_break()
    h2 = doc.add_heading('Headings', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
    Proper heading structure is essential for accessibility as it helps users understand the organization of content and navigate pages effectively. Headings should follow a logical hierarchy and accurately reflect the content structure. Screen reader users often navigate by headings, making proper structure crucial.
    """.strip())

    doc.add_paragraph("Common heading structure issues include:", style='Normal')

    doc.add_paragraph("Missing or multiple main headings (H1)", style='List Bullet')
    doc.add_paragraph("Skipped heading levels creating hierarchy gaps", style='List Bullet')
    doc.add_paragraph("Headings placed before the main content", style='List Bullet')
    doc.add_paragraph("Visual styling that doesn't match heading levels", style='List Bullet')

    doc.add_paragraph()

    # Query for pages with heading issues
    pages_with_heading_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.headings.headings.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.headings.headings.pageFlags.missingH1": True},
                {"results.accessibility.tests.headings.headings.pageFlags.multipleH1s": True},
                {"results.accessibility.tests.headings.headings.pageFlags.hasHierarchyGaps": True},
                {"results.accessibility.tests.headings.headings.pageFlags.hasHeadingsBeforeMain": True},
                {"results.accessibility.tests.headings.headings.pageFlags.hasVisualHierarchyIssues": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.headings.headings": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for different heading issues
    heading_issues = {
        "missing_h1": {"name": "Missing H1", "pages": set(), "domains": set()},
        "multiple_h1": {"name": "Multiple H1s", "pages": set(), "domains": set()},
        "hierarchy_gaps": {"name": "Hierarchy gaps", "pages": set(), "domains": set(), "count": 0},
        "headings_before_main": {"name": "Headings before main", "pages": set(), "domains": set(), "count": 0},
        "visual_hierarchy": {"name": "Visual hierarchy issues", "pages": set(), "domains": set(), "count": 0}
    }

    # Process each page
    total_headings = 0
    for page in pages_with_heading_issues:
        try:
            domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
            heading_data = page['results']['accessibility']['tests']['headings']['headings']
            flags = heading_data.get('pageFlags', {})
            
            # The issue is here - details might be in a different place or structure than expected
            details = flags.get('details', {})
            if not details:  # If details not found in flags, try the main heading_data
                details = heading_data.get('details', {})
            
            summary = heading_data.get('details', {}).get('summary', {})
            
            # Update total headings count
            headings_count = summary.get('totalHeadings', 0)
            if isinstance(headings_count, (int, float)):
                total_headings += headings_count
            
            # Check missing H1
            if flags.get('missingH1'):
                heading_issues['missing_h1']['pages'].add(page['url'])
                heading_issues['missing_h1']['domains'].add(domain)
            
            # Check multiple H1s
            if flags.get('multipleH1s'):
                heading_issues['multiple_h1']['pages'].add(page['url'])
                heading_issues['multiple_h1']['domains'].add(domain)
            
            # Check hierarchy gaps
            if flags.get('hasHierarchyGaps'):
                heading_issues['hierarchy_gaps']['pages'].add(page['url'])
                heading_issues['hierarchy_gaps']['domains'].add(domain)
                
                # Fix for the potential list issue
                hierarchy_gaps = details.get('hierarchyGaps', 0)
                if isinstance(hierarchy_gaps, list):
                    heading_issues['hierarchy_gaps']['count'] += len(hierarchy_gaps)
                elif isinstance(hierarchy_gaps, (int, float)):
                    heading_issues['hierarchy_gaps']['count'] += hierarchy_gaps
            
            # Check headings before main
            if flags.get('hasHeadingsBeforeMain'):
                heading_issues['headings_before_main']['pages'].add(page['url'])
                heading_issues['headings_before_main']['domains'].add(domain)
                
                # Fix for the potential list issue
                headings_before_main = details.get('headingsBeforeMain', 0)
                if isinstance(headings_before_main, list):
                    heading_issues['headings_before_main']['count'] += len(headings_before_main)
                elif isinstance(headings_before_main, (int, float)):
                    heading_issues['headings_before_main']['count'] += headings_before_main
            
            # Check visual hierarchy issues
            if flags.get('hasVisualHierarchyIssues'):
                heading_issues['visual_hierarchy']['pages'].add(page['url'])
                heading_issues['visual_hierarchy']['domains'].add(domain)
                
                # Fix for the potential list issue
                visual_hierarchy_issues = details.get('visualHierarchyIssues', 0)
                if isinstance(visual_hierarchy_issues, list):
                    heading_issues['visual_hierarchy']['count'] += len(visual_hierarchy_issues)
                elif isinstance(visual_hierarchy_issues, (int, float)):
                    heading_issues['visual_hierarchy']['count'] += visual_hierarchy_issues
                
        except Exception as e:
            print(f"Error processing page {page.get('url', 'unknown')}: {str(e)}")
            continue

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in heading_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=5)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue Type"
        headers[1].text = "Number of Occurrences"
        headers[2].text = "Pages Affected"
        headers[3].text = "Sites Affected"
        headers[4].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(data.get('count', len(data['pages'])))
            row[2].text = str(len(data['pages']))
            row[3].text = str(len(data['domains']))
            row[4].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add statistics
        doc.add_paragraph()
        doc.add_paragraph("Heading Statistics:", style='Normal')
        doc.add_paragraph(f"Total number of headings across all pages: {total_headings}")

        # Add recommendations
        doc.add_paragraph()
        doc.add_heading('Recommendations for Heading Structure', level=3)
        
        doc.add_paragraph("Use exactly one H1 heading per page as the main title", style='List Bullet')
        doc.add_paragraph("Maintain proper heading hierarchy without skipping levels", style='List Bullet')
        doc.add_paragraph("Ensure heading levels match their visual presentation", style='List Bullet')
        doc.add_paragraph("Place meaningful headings in the main content area", style='List Bullet')
        doc.add_paragraph("Use headings to create a clear content outline", style='List Bullet')
        doc.add_paragraph("Make heading text descriptive and meaningful", style='List Bullet')

        # Add example of proper heading structure
        doc.add_paragraph()
        doc.add_paragraph("Example of proper heading structure:", style='Normal')
        doc.add_paragraph("H1: Main page title", style='List Bullet')
        doc.add_paragraph("    H2: Major section", style='List Bullet 2')
        doc.add_paragraph("        H3: Subsection", style='List Bullet 3')
        doc.add_paragraph("        H3: Another subsection", style='List Bullet 3')
        doc.add_paragraph("    H2: Another major section", style='List Bullet 2')
        
        # Add code samples for proper heading implementation
        doc.add_paragraph()
        doc.add_paragraph("Technical Implementation of Proper Heading Structure:", style='Normal').bold = True
        
        # Example of good heading structure
        doc.add_paragraph("Good Example - Proper Heading Hierarchy:", style='List Bullet')
        good_code = doc.add_paragraph("""
<h1>Website Title</h1>
<main>
  <h2>Major Section Title</h2>
  <p>Some introductory content...</p>
  
  <h3>Subsection Title</h3>
  <p>More detailed content...</p>
  
  <h3>Another Subsection</h3>
  <p>Additional content...</p>
  
  <h2>Another Major Section</h2>
  <p>More content...</p>
</main>
        """)
        good_code.style = doc.styles['Normal']
        good_code.paragraph_format.left_indent = Pt(36)
        
        # Example of bad heading structure
        doc.add_paragraph("Bad Example - Skipped Hierarchy Levels:", style='List Bullet')
        bad_code = doc.add_paragraph("""
<h1>Website Title</h1>
<main>
  <h2>Major Section Title</h2>
  <p>Some introductory content...</p>
  
  <!-- Bad practice: Skipping from H2 to H4 -->
  <h4>Subsection Title</h4>
  <p>More detailed content...</p>
  
  <h2>Another Major Section</h2>
  <p>More content...</p>
</main>
        """)
        bad_code.style = doc.styles['Normal']
        bad_code.paragraph_format.left_indent = Pt(36)
        
        # Visual styling example
        doc.add_paragraph("Example - Visual vs. Semantic Heading Levels:", style='List Bullet')
        styling_code = doc.add_paragraph("""
<!-- Bad practice: Using styling instead of semantic headings -->
<div class="looks-like-h2">This looks like a heading but isn't</div>

<!-- Bad practice: Using lower-level heading with larger styling -->
<h3 style="font-size: 24px;">This looks like H2 but is H3</h3>

<!-- Good practice: Heading level matches visual importance -->
<h2>This is properly an H2 and looks like one</h2>

<!-- Good practice: If styling is needed, keep semantic level correct -->
<h3 class="special-style">Styled but still correctly an H3</h3>
        """)
        styling_code.style = doc.styles['Normal']
        styling_code.paragraph_format.left_indent = Pt(36)

        # Add domain details for each issue type
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)

    else:
        doc.add_paragraph("No heading structure issues were found.")