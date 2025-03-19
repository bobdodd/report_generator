# sections/detailed_findings/tables.py
from report_styling import format_table_text
from docx.shared import Pt

def add_detailed_tables(doc, db_connection, total_domains):
    """Add the detailed Tables section"""
    doc.add_page_break()
    h2 = doc.add_heading('Tables', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
Tables should be used for presenting tabular data, not for layout purposes. Proper table markup with appropriate headers and structure is crucial for screen reader users. Common issues include:
""".strip())

    doc.add_paragraph("Missing table headers (th elements)", style='List Bullet')
    doc.add_paragraph("Lack of proper scope attributes on header cells", style='List Bullet')
    doc.add_paragraph("Missing caption or summary for complex tables", style='List Bullet')
    doc.add_paragraph("Tables used for layout purposes instead of CSS", style='List Bullet')
    doc.add_paragraph("Complex tables without proper row/column headers", style='List Bullet')

    doc.add_paragraph()

    # Query for pages with table issues
    pages_with_table_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.tables.tables.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.tables.tables.pageFlags.hasMissingHeaders": True},
                {"results.accessibility.tests.tables.tables.pageFlags.hasNoScope": True},
                {"results.accessibility.tests.tables.tables.pageFlags.hasMissingCaption": True},
                {"results.accessibility.tests.tables.tables.pageFlags.hasLayoutTables": True},
                {"results.accessibility.tests.tables.tables.pageFlags.hasComplexTables": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.tables.tables.pageFlags": 1,
            "results.accessibility.tests.tables.tables.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    table_issues = {
        "hasMissingHeaders": {"name": "Missing table headers", "pages": set(), "domains": set()},
        "hasNoScope": {"name": "Missing scope attributes", "pages": set(), "domains": set()},
        "hasMissingCaption": {"name": "Missing table captions", "pages": set(), "domains": set()},
        "hasLayoutTables": {"name": "Layout tables", "pages": set(), "domains": set()},
        "hasComplexTables": {"name": "Complex tables without proper structure", "pages": set(), "domains": set()}
    }

    # Count issues
    for page in pages_with_table_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        flags = page['results']['accessibility']['tests']['tables']['tables']['pageFlags']
        
        for flag in table_issues:
            if flags.get(flag, False):
                table_issues[flag]['pages'].add(page['url'])
                table_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in table_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Table Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add domain details for each issue
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)
                
        # Add technical implementation section
        doc.add_paragraph()
        doc.add_heading('Technical Implementation Guidelines', level=3)
        
        # Simple table example
        doc.add_paragraph("Simple Table with Headers:", style='Normal').bold = True
        
        simple_table = doc.add_paragraph("""
<!-- Simple table with proper headers -->
<table>
  <caption>Monthly Sales by Region</caption>
  <thead>
    <tr>
      <th scope="col">Region</th>
      <th scope="col">January</th>
      <th scope="col">February</th>
      <th scope="col">March</th>
    </tr>
  </thead>
  <tbody>
    <tr>
      <th scope="row">North</th>
      <td>$10,000</td>
      <td>$12,000</td>
      <td>$15,000</td>
    </tr>
    <tr>
      <th scope="row">South</th>
      <td>$8,000</td>
      <td>$9,500</td>
      <td>$11,000</td>
    </tr>
    <tr>
      <th scope="row">East</th>
      <td>$9,000</td>
      <td>$8,500</td>
      <td>$12,500</td>
    </tr>
  </tbody>
  <tfoot>
    <tr>
      <th scope="row">Total</th>
      <td>$27,000</td>
      <td>$30,000</td>
      <td>$38,500</td>
    </tr>
  </tfoot>
</table>
        """)
        simple_table.style = doc.styles['Normal']
        simple_table.paragraph_format.left_indent = Pt(36)
        
        # Complex table example
        doc.add_paragraph("Complex Table with Row and Column Headers:", style='Normal').bold = True
        
        complex_table = doc.add_paragraph("""
<!-- Complex table with both row and column headers -->
<table>
  <caption>Quarterly Product Sales by Division</caption>
  <thead>
    <tr>
      <th scope="col">Division</th>
      <th scope="col" colspan="3">Q1</th>
      <th scope="col" colspan="3">Q2</th>
    </tr>
    <tr>
      <th scope="col"></th>
      <th scope="col">Product A</th>
      <th scope="col">Product B</th>
      <th scope="col">Product C</th>
      <th scope="col">Product A</th>
      <th scope="col">Product B</th>
      <th scope="col">Product C</th>
    </tr>
  </thead>
  <tbody>
    <tr>
      <th scope="row">North</th>
      <td>500</td>
      <td>600</td>
      <td>400</td>
      <td>550</td>
      <td>650</td>
      <td>450</td>
    </tr>
    <tr>
      <th scope="row">South</th>
      <td>300</td>
      <td>400</td>
      <td>350</td>
      <td>325</td>
      <td>425</td>
      <td>400</td>
    </tr>
  </tbody>
</table>
        """)
        complex_table.style = doc.styles['Normal']
        complex_table.paragraph_format.left_indent = Pt(36)
        
        # Table with headers in first column and row
        doc.add_paragraph("Table with Headers in First Column and Row:", style='Normal').bold = True
        
        dual_headers_table = doc.add_paragraph("""
<!-- Table with headers in first column and row (using id and headers) -->
<table>
  <caption>Employee Schedule</caption>
  <tr>
    <th id="empty-cell"></th>
    <th id="monday">Monday</th>
    <th id="tuesday">Tuesday</th>
    <th id="wednesday">Wednesday</th>
    <th id="thursday">Thursday</th>
    <th id="friday">Friday</th>
  </tr>
  <tr>
    <th id="alice">Alice</th>
    <td headers="alice monday">Meeting</td>
    <td headers="alice tuesday">Training</td>
    <td headers="alice wednesday">Project Work</td>
    <td headers="alice thursday">Client Visit</td>
    <td headers="alice friday">Remote Work</td>
  </tr>
  <tr>
    <th id="bob">Bob</th>
    <td headers="bob monday">Project Work</td>
    <td headers="bob tuesday">Meeting</td>
    <td headers="bob wednesday">Remote Work</td>
    <td headers="bob thursday">Training</td>
    <td headers="bob friday">Client Visit</td>
  </tr>
</table>
        """)
        dual_headers_table.style = doc.styles['Normal']
        dual_headers_table.paragraph_format.left_indent = Pt(36)
        
        # Table with responsive design
        doc.add_paragraph("Responsive Table Design:", style='Normal').bold = True
        
        responsive_table = doc.add_paragraph("""
<!-- CSS for responsive tables -->
<style>
@media screen and (max-width: 600px) {
  /* Convert table to non-tabular display for small screens */
  table.responsive {
    display: block;
    width: 100%;
  }
  
  table.responsive thead {
    display: none; /* Hide the header row */
  }
  
  table.responsive tbody, 
  table.responsive tr, 
  table.responsive th, 
  table.responsive td {
    display: block;
    width: 100%;
  }
  
  /* Make each row look like a card */
  table.responsive tr {
    margin-bottom: 15px;
    border: 1px solid #ddd;
  }
  
  /* Add data labels using pseudo-elements and data attributes */
  table.responsive td::before {
    content: attr(data-label);
    font-weight: bold;
    display: inline-block;
    width: 40%;
  }
}
</style>

<!-- HTML for responsive table -->
<table class="responsive">
  <caption>Product Comparison</caption>
  <thead>
    <tr>
      <th scope="col">Product</th>
      <th scope="col">Price</th>
      <th scope="col">Features</th>
      <th scope="col">Rating</th>
    </tr>
  </thead>
  <tbody>
    <tr>
      <th scope="row">Basic Model</th>
      <td data-label="Price">$99</td>
      <td data-label="Features">Standard features only</td>
      <td data-label="Rating">3.5/5</td>
    </tr>
    <tr>
      <th scope="row">Pro Model</th>
      <td data-label="Price">$199</td>
      <td data-label="Features">Advanced features included</td>
      <td data-label="Rating">4.5/5</td>
    </tr>
  </tbody>
</table>
        """)
        responsive_table.style = doc.styles['Normal']
        responsive_table.paragraph_format.left_indent = Pt(36)
        
        # Accessible alternative to layout tables
        doc.add_paragraph("Accessible Alternatives to Layout Tables:", style='Normal').bold = True
        
        layout_alternatives = doc.add_paragraph("""
<!-- Instead of using layout tables, use CSS Grid or Flexbox -->

<!-- CSS Grid Example -->
<style>
.grid-layout {
  display: grid;
  grid-template-columns: 1fr 2fr;
  grid-gap: 20px;
}

.grid-layout > div {
  padding: 20px;
}
</style>

<div class="grid-layout">
  <div>
    <h2>Side Content</h2>
    <p>This would be navigation or sidebar content.</p>
  </div>
  <div>
    <h2>Main Content</h2>
    <p>This is the main content area where most of the information would appear.</p>
  </div>
</div>

<!-- Flexbox Example -->
<style>
.flex-layout {
  display: flex;
  flex-wrap: wrap;
}

.flex-sidebar {
  flex: 1;
  min-width: 200px;
  padding: 20px;
}

.flex-main {
  flex: 2;
  min-width: 300px;
  padding: 20px;
}
</style>

<div class="flex-layout">
  <div class="flex-sidebar">
    <h2>Side Content</h2>
    <p>This would be navigation or sidebar content.</p>
  </div>
  <div class="flex-main">
    <h2>Main Content</h2>
    <p>This is the main content area where most of the information would appear.</p>
  </div>
</div>
        """)
        layout_alternatives.style = doc.styles['Normal']
        layout_alternatives.paragraph_format.left_indent = Pt(36)

    else:
        doc.add_paragraph("No table markup issues were found.")