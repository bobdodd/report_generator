# sections/detailed_findings/maps.py
from report_styling import format_table_text
from docx.shared import Pt

def add_detailed_maps(doc, db_connection, total_domains):
    """Add the detailed Maps section"""
    doc.add_page_break()
    h2 = doc.add_heading('Maps', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
Interactive maps present particular accessibility challenges. Common issues include:
""".strip())

    doc.add_paragraph("Maps without text alternatives for the information they convey", style='List Bullet')
    doc.add_paragraph("Interactive maps that can't be operated by keyboard", style='List Bullet')
    doc.add_paragraph("Map features that aren't properly labeled for screen readers", style='List Bullet')
    doc.add_paragraph("Missing alternative ways to access location information", style='List Bullet')

    # Add recommendations
    doc.add_paragraph()
    doc.add_paragraph("Recommendations for making maps accessible:")
    doc.add_paragraph("Provide text alternatives that describe the key information the map conveys", style='List Bullet')
    doc.add_paragraph("Ensure all map controls can be operated by keyboard", style='List Bullet')
    doc.add_paragraph("Include proper ARIA labels and roles for map features", style='List Bullet')
    doc.add_paragraph("Offer alternative formats (e.g., text list of locations, address lookup)", style='List Bullet')
    doc.add_paragraph("Ensure interactive elements within maps are properly labeled", style='List Bullet')

    doc.add_paragraph()

    # Query for pages with map issues
    pages_with_map_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.maps.maps.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.maps.maps.pageFlags.hasMaps": True},
                {"results.accessibility.tests.maps.maps.pageFlags.hasMapsWithoutTitle": True},
                {"results.accessibility.tests.maps.maps.pageFlags.hasMapsWithAriaHidden": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.maps.maps.pageFlags": 1,
            "results.accessibility.tests.maps.maps.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    map_issues = {
        "hasMaps": {"name": "Pages containing maps", "pages": set(), "domains": set()},
        "hasMapsWithoutTitle": {"name": "Maps without proper titles", "pages": set(), "domains": set()},
        "hasMapsWithAriaHidden": {"name": "Maps hidden from screen readers", "pages": set(), "domains": set()}
    }

    # Count issues
    for page in pages_with_map_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        flags = page['results']['accessibility']['tests']['maps']['maps']['pageFlags']
        
        for flag in map_issues:
            if flags.get(flag, False):
                map_issues[flag]['pages'].add(page['url'])
                map_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in map_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add domain details for each issue
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)

        # Add examples if available
        for page in pages_with_map_issues:
            details = page['results']['accessibility']['tests']['maps']['maps']['details']
            if 'violations' in details and details['violations']:
                doc.add_paragraph()
                doc.add_paragraph("Examples of map accessibility issues found:")
                for violation in details['violations'][:5]:  # Show up to 5 examples
                    doc.add_paragraph(violation, style='List Bullet')
                break  # Only show examples from first page with violations
                
        # Add technical implementation guidance
        doc.add_paragraph()
        doc.add_heading('Technical Implementation Guidelines', level=3)
        
        # Code examples
        doc.add_paragraph("Example of an accessible map implementation:", style='Normal').bold = True
        
        map_code = doc.add_paragraph("""
<!-- Accessible map implementation -->
<div role="region" aria-label="Location map">
  <!-- Map container -->
  <div id="map" aria-hidden="false" tabindex="0" 
       aria-label="Interactive map showing our location at 123 Main Street">
    <!-- Map content loaded by JavaScript -->
  </div>
  
  <!-- Accessible alternative -->
  <div class="map-alternatives">
    <h3>Location Information</h3>
    <p>Our office is located at:</p>
    <address>
      123 Main Street<br>
      Springfield, IL 62701<br>
      <a href="https://goo.gl/maps/example" target="_blank">View on Google Maps</a>
    </address>
    
    <details>
      <summary>Directions from major landmarks</summary>
      <ul>
        <li>From the train station: Walk east on Railway Ave for 2 blocks, then north on Main St.</li>
        <li>From the bus terminal: Take bus #42 or #56 to Main & Oak stop.</li>
        <li>From the airport: Take the airport shuttle to downtown, then walk 3 blocks east.</li>
      </ul>
    </details>
  </div>
</div>
        """)
        map_code.style = doc.styles['Normal']
        map_code.paragraph_format.left_indent = Pt(36)
        
        # JavaScript considerations
        doc.add_paragraph("JavaScript considerations for map accessibility:", style='Normal').bold = True
        
        js_code = doc.add_paragraph("""
// Initialize map with keyboard support
function initAccessibleMap() {
  const map = new MapLibrary('map-container');
  
  // Enable keyboard navigation
  map.enableKeyboardSupport();
  
  // Ensure focus is managed properly
  map.on('open-popup', function(e) {
    // When popup opens, move focus to it
    document.getElementById('map-popup').focus();
  });
  
  // Ensure Escape key closes popups
  map.on('keydown', function(e) {
    if (e.key === 'Escape' && map.hasOpenPopups()) {
      map.closeAllPopups();
      // Return focus to the map
      document.getElementById('map-container').focus();
    }
  });
  
  // Announce screen reader updates when map changes
  map.on('zoom', function() {
    updateAriaLiveRegion('Map zoom level changed');
  });
  
  map.on('move', function() {
    updateAriaLiveRegion('Map view has moved');
  });
}

function updateAriaLiveRegion(message) {
  document.getElementById('map-announcer').textContent = message;
}
        """)
        js_code.style = doc.styles['Normal']
        js_code.paragraph_format.left_indent = Pt(36)
        
        # ARIA live region example
        doc.add_paragraph("Example of ARIA live region for map updates:", style='Normal').bold = True
        
        aria_code = doc.add_paragraph("""
<!-- ARIA live region for map announcements -->
<div id="map-announcer" class="sr-only" aria-live="polite"></div>
        """)
        aria_code.style = doc.styles['Normal']
        aria_code.paragraph_format.left_indent = Pt(36)

    else:
        doc.add_paragraph("No map-related accessibility issues were found.")
        