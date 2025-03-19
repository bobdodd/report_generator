from report_styling import format_table_text
from docx.shared import Pt

def add_detailed_forms(doc, db_connection, total_domains):
    """Add the detailed Forms section"""
    doc.add_page_break()
    h2 = doc.add_heading('Forms', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
    Forms must be accessible to all users, including those using screen readers or keyboard navigation. Proper labeling, structure, and organization are essential for form accessibility. Forms should have clear instructions, properly associated labels, and appropriate error handling.
    """.strip())

    doc.add_paragraph("Common form accessibility issues include:", style='Normal')

    doc.add_paragraph("Form inputs without proper labels", style='List Bullet')
    doc.add_paragraph("Reliance on placeholders instead of labels", style='List Bullet')
    doc.add_paragraph("Forms without proper heading structure", style='List Bullet')
    doc.add_paragraph("Forms placed outside landmark regions", style='List Bullet')
    doc.add_paragraph("Input fields with insufficient contrast", style='List Bullet')
    doc.add_paragraph("Layout issues affecting form usability", style='List Bullet')

    # Add recommendations
    doc.add_paragraph()
    doc.add_paragraph("Recommendations for Form Accessibility:", style='Normal')
    
    doc.add_paragraph("Ensure all form controls have properly associated labels", style='List Bullet')
    doc.add_paragraph("Use labels instead of relying solely on placeholders", style='List Bullet')
    doc.add_paragraph("Include proper heading structure for forms", style='List Bullet')
    doc.add_paragraph("Place forms within appropriate landmark regions", style='List Bullet')
    doc.add_paragraph("Maintain sufficient contrast for all form elements", style='List Bullet')
    doc.add_paragraph("Ensure proper spacing and layout for form controls", style='List Bullet')
    doc.add_paragraph("Provide clear error messages and validation feedback", style='List Bullet')
    doc.add_paragraph("Ensure forms are keyboard accessible", style='List Bullet')

    doc.add_paragraph()

    # Query for pages with form issues
    pages_with_form_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.forms.forms.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.forms.forms.pageFlags.hasInputsWithoutLabels": True},
                {"results.accessibility.tests.forms.forms.pageFlags.hasPlaceholderOnlyInputs": True},
                {"results.accessibility.tests.forms.forms.pageFlags.hasFormsWithoutHeadings": True},
                {"results.accessibility.tests.forms.forms.pageFlags.hasFormsOutsideLandmarks": True},
                {"results.accessibility.tests.forms.forms.pageFlags.hasContrastIssues": True},
                {"results.accessibility.tests.forms.forms.pageFlags.hasLayoutIssues": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.forms.forms": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for different form issues
    form_issues = {
        "missing_labels": {
            "name": "Inputs without labels",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "placeholder_only": {
            "name": "Placeholder-only inputs",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "no_headings": {
            "name": "Forms without headings",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "outside_landmarks": {
            "name": "Forms outside landmarks",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "contrast_issues": {
            "name": "Input contrast issues",
            "pages": set(),
            "domains": set(),
            "count": 0
        },
        "layout_issues": {
            "name": "Form layout issues",
            "pages": set(),
            "domains": set(),
            "count": 0
        }
    }

    # Process each page
    total_forms = 0
    for page in pages_with_form_issues:
        try:
            domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
            form_data = page['results']['accessibility']['tests']['forms']['forms']
            flags = form_data.get('pageFlags', {})
            summary = form_data.get('details', {}).get('summary', {})
            
            # Update total forms count
            total_forms += summary.get('totalForms', 0)
            
            # Check inputs without labels
            if flags.get('hasInputsWithoutLabels'):
                form_issues['missing_labels']['pages'].add(page['url'])
                form_issues['missing_labels']['domains'].add(domain)
                form_issues['missing_labels']['count'] += summary.get('inputsWithoutLabels', 0)
            
            # Check placeholder-only inputs
            if flags.get('hasPlaceholderOnlyInputs'):
                form_issues['placeholder_only']['pages'].add(page['url'])
                form_issues['placeholder_only']['domains'].add(domain)
                form_issues['placeholder_only']['count'] += summary.get('inputsWithPlaceholderOnly', 0)
            
            # Check forms without headings
            if flags.get('hasFormsWithoutHeadings'):
                form_issues['no_headings']['pages'].add(page['url'])
                form_issues['no_headings']['domains'].add(domain)
                form_issues['no_headings']['count'] += summary.get('formsWithoutHeadings', 0)
            
            # Check forms outside landmarks
            if flags.get('hasFormsOutsideLandmarks'):
                form_issues['outside_landmarks']['pages'].add(page['url'])
                form_issues['outside_landmarks']['domains'].add(domain)
                form_issues['outside_landmarks']['count'] += summary.get('formsOutsideLandmarks', 0)
            
            # Check contrast issues
            if flags.get('hasContrastIssues'):
                form_issues['contrast_issues']['pages'].add(page['url'])
                form_issues['contrast_issues']['domains'].add(domain)
                form_issues['contrast_issues']['count'] += summary.get('inputsWithContrastIssues', 0)
            
            # Check layout issues
            if flags.get('hasLayoutIssues'):
                form_issues['layout_issues']['pages'].add(page['url'])
                form_issues['layout_issues']['domains'].add(domain)
                form_issues['layout_issues']['count'] += summary.get('inputsWithLayoutIssues', 0)
                
        except Exception as e:
            print(f"Error processing page {page.get('url', 'unknown')}: {str(e)}")
            continue

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in form_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=5)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue Type"
        headers[1].text = "Number of Occurrences"
        headers[2].text = "Pages Affected"
        headers[3].text = "Sites Affected"
        headers[4].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(data['count'])
            row[2].text = str(len(data['pages']))
            row[3].text = str(len(data['domains']))
            row[4].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add statistics
        doc.add_paragraph()
        doc.add_paragraph("Form Statistics:", style='Normal')
        doc.add_paragraph(f"Total number of forms across all pages: {total_forms}")

        # Add domain details for each issue type
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)

        # Add detailed technical implementation guidelines
        doc.add_paragraph()
        doc.add_heading('Technical Implementation Guidelines for Accessible Forms', level=3)
        
        # Proper Labeling
        doc.add_paragraph("Proper form control labeling:", style='Normal').bold = True
        doc.add_paragraph("Use explicit <label> elements with 'for' attributes matching input IDs", style='List Bullet')
        doc.add_paragraph("Ensure labels are descriptive and concise", style='List Bullet')
        doc.add_paragraph("Place labels consistently (usually above or to the left of inputs)", style='List Bullet')
        doc.add_paragraph("Example code:")
        
        label_code = doc.add_paragraph("""
<!-- Good practice: Explicit label with for attribute -->
<label for="user-email">Email Address</label>
<input type="email" id="user-email" name="email">

<!-- Bad practice: Input without label -->
<input type="email" name="email" placeholder="Email Address">

<!-- Bad practice: Using placeholder instead of label -->
<input type="email" name="email" placeholder="Email Address">

<!-- Alternative if visible label isn't possible: Use aria-label -->
<input type="email" name="email" aria-label="Email Address">
        """)
        label_code.style = doc.styles['Normal']
        label_code.paragraph_format.left_indent = Pt(36)
        
        # Placeholders
        doc.add_paragraph("Proper use of placeholders:", style='Normal').bold = True
        doc.add_paragraph("Use placeholders for hints or examples, not as replacements for labels", style='List Bullet')
        doc.add_paragraph("Ensure placeholder text has sufficient contrast", style='List Bullet')
        doc.add_paragraph("Example code:")
        
        placeholder_code = doc.add_paragraph("""
<!-- Good practice: Using placeholder as hint with proper label -->
<label for="phone">Phone Number</label>
<input type="tel" id="phone" name="phone" placeholder="Example: 555-123-4567">

<!-- Bad practice: Placeholder as label -->
<input type="tel" name="phone" placeholder="Phone Number">
        """)
        placeholder_code.style = doc.styles['Normal']
        placeholder_code.paragraph_format.left_indent = Pt(36)
        
        # Form Structure
        doc.add_paragraph("Proper form structure:", style='Normal').bold = True
        doc.add_paragraph("Use appropriate fieldset and legend for grouping related inputs", style='List Bullet')
        doc.add_paragraph("Include proper heading for the form", style='List Bullet')
        doc.add_paragraph("Place forms inside a region with role='form' or use the <form> element", style='List Bullet')
        doc.add_paragraph("Example code:")
        
        structure_code = doc.add_paragraph("""
<main>
  <section>
    <h2>Contact Form</h2>
    <form role="form" aria-labelledby="form-heading">
      <h3 id="form-heading">Send us a message</h3>
      
      <fieldset>
        <legend>Personal Information</legend>
        <div>
          <label for="name">Full Name</label>
          <input type="text" id="name" name="name">
        </div>
        <div>
          <label for="email">Email Address</label>
          <input type="email" id="email" name="email">
        </div>
      </fieldset>
      
      <fieldset>
        <legend>Your Message</legend>
        <div>
          <label for="subject">Subject</label>
          <input type="text" id="subject" name="subject">
        </div>
        <div>
          <label for="message">Message</label>
          <textarea id="message" name="message"></textarea>
        </div>
      </fieldset>
      
      <button type="submit">Send Message</button>
    </form>
  </section>
</main>
        """)
        structure_code.style = doc.styles['Normal']
        structure_code.paragraph_format.left_indent = Pt(36)
        
        # Error Handling
        doc.add_paragraph("Accessible error handling:", style='Normal').bold = True
        doc.add_paragraph("Provide clear error messages that identify the specific issues", style='List Bullet')
        doc.add_paragraph("Associate error messages with their corresponding inputs using aria-describedby", style='List Bullet')
        doc.add_paragraph("Use both color and text to indicate errors", style='List Bullet')
        doc.add_paragraph("Example code:")
        
        error_code = doc.add_paragraph("""
<!-- When an error occurs -->
<div>
  <label for="password">Password</label>
  <input type="password" id="password" name="password" 
         aria-invalid="true" aria-describedby="password-error">
  <p id="password-error" class="error-message">
    Password must be at least 8 characters and include a number and symbol
  </p>
</div>

<script>
  // Focus on the first field with an error when form submission fails
  document.querySelector('[aria-invalid="true"]').focus();
</script>
        """)
        error_code.style = doc.styles['Normal']
        error_code.paragraph_format.left_indent = Pt(36)

    else:
        doc.add_paragraph("No form accessibility issues were found.")
        