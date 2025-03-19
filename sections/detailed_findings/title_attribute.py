# sections/detailed_findings/title_attribute.py
from report_styling import format_table_text
from docx.shared import Pt

def add_detailed_title_attribute(doc, db_connection, total_domains):
    """Add the detailed Title Attribute section"""
    doc.add_page_break()
    h2 = doc.add_heading('Title Attribute', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
The title attribute is often misused as a tooltip or to provide additional information. However, it has several accessibility limitations:
""".strip())

    doc.add_paragraph("Not consistently exposed by screen readers or available on mobile devices", style='List Bullet')
    doc.add_paragraph("Cannot be accessed by keyboard-only users", style='List Bullet')
    doc.add_paragraph("Cannot be reliably accessed by screen-magnifier users as the title attribute may be unreachable as the user moves the mouse to read it", style='List Bullet')
    doc.add_paragraph("Content is not visible until hover, which some users cannot do", style='List Bullet')
    doc.add_paragraph("Should not be used as the only way to convey important information", style='List Bullet')

    doc.add_paragraph("""
There is one case when the title attribute must be used, and that is for <iframe> as it is the only way to give a name to an embedded element. Typically that is the title of a YouTube or Vimeo video.
""".strip())

    doc.add_paragraph()

    # Query for pages with title attribute issues
    pages_with_title_issues = list(db_connection.page_results.find(
        {"results.accessibility.tests.title.titleAttribute.pageFlags.hasImproperTitleAttributes": True},
        {
            "url": 1,
            "results.accessibility.tests.title.titleAttribute.details": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Count affected domains
    affected_domains = set()
    total_improper_uses = 0
    domain_counts = {}

    for page in pages_with_title_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        affected_domains.add(domain)
        
        # Count improper uses from the details
        improper_uses = len(page['results']['accessibility']['tests']['title']['titleAttribute']['details']['improperUse'])
        total_improper_uses += improper_uses
        
        # Track counts by domain
        if domain not in domain_counts:
            domain_counts[domain] = 0
        domain_counts[domain] += improper_uses

    # Calculate percentage
    percentage = (len(affected_domains) / len(total_domains)) * 100 if total_domains else 0

    # Create summary table
    summary_table = doc.add_table(rows=2, cols=4)
    summary_table.style = 'Table Grid'

    # Set column headers
    headers = summary_table.rows[0].cells
    headers[0].text = "Issue"
    headers[1].text = "Total Occurrences"
    headers[2].text = "Sites Affected"
    headers[3].text = "% of Total Sites"

    # Add data
    row = summary_table.rows[1].cells
    row[0].text = "Improper use of title attribute"
    row[1].text = str(total_improper_uses)
    row[2].text = str(len(affected_domains))
    row[3].text = f"{percentage:.1f}%"

    # Format the table text
    format_table_text(summary_table)

    # Add some space after the table
    doc.add_paragraph()

    if domain_counts:
        # Create domain details table
        doc.add_paragraph("Breakdown by site:")
        domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
        domain_table.style = 'Table Grid'

        # Add headers
        headers = domain_table.rows[0].cells
        headers[0].text = "Domain"
        headers[1].text = "Number of improper title attributes"

        # Add domain data
        for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
            row = domain_table.rows[i].cells
            row[0].text = domain
            row[1].text = str(count)

        # Format the table text
        format_table_text(domain_table)
        
        # Add technical implementation section
        doc.add_paragraph()
        doc.add_heading('Technical Implementation Guidelines', level=3)
        
        # When title attribute should be used
        doc.add_paragraph("When the title attribute should be used:", style='Normal').bold = True
        
        good_title_uses = doc.add_paragraph("""
<!-- Good use: title on iframe -->
<iframe 
  src="https://www.youtube.com/embed/abc123" 
  title="Introduction to Web Accessibility"
  width="560" 
  height="315" 
  allowfullscreen
></iframe>

<!-- Good use: title on abbr -->
<abbr title="Web Content Accessibility Guidelines">WCAG</abbr>
        """)
        good_title_uses.style = doc.styles['Normal']
        good_title_uses.paragraph_format.left_indent = Pt(36)
        
        # Improper title attribute usage
        doc.add_paragraph("Common improper uses of title attribute:", style='Normal').bold = True
        
        improper_title_uses = doc.add_paragraph("""
<!-- Improper: title as sole descriptor for link -->
<a href="/contact" title="Contact us">Contact</a>

<!-- Improper: title on image instead of alt text -->
<img src="logo.png" title="Company Logo">

<!-- Improper: title to provide critical instructions -->
<button title="Click to submit form">Submit</button>

<!-- Improper: title for tooltip-like functionality -->
<span title="This shows up on hover only">Hover me for more info</span>
        """)
        improper_title_uses.style = doc.styles['Normal']
        improper_title_uses.paragraph_format.left_indent = Pt(36)
        
        # Better alternatives
        doc.add_paragraph("Better alternatives to the title attribute:", style='Normal').bold = True
        
        better_alternatives = doc.add_paragraph("""
<!-- Better: Descriptive link text -->
<a href="/contact">Contact our support team</a>

<!-- Better: Proper alt text -->
<img src="logo.png" alt="Company Logo">

<!-- Better: Visible instructions -->
<button>Submit <span class="instruction">(Completes your application)</span></button>

<!-- Better: Accessible tooltip with ARIA -->
<span aria-describedby="tooltip1">Hover or focus for more info</span>
<div id="tooltip1" role="tooltip" class="tooltip">
  This information is accessible to all users including keyboard and screen reader users
</div>
        """)
        better_alternatives.style = doc.styles['Normal']
        better_alternatives.paragraph_format.left_indent = Pt(36)
        
        # CSS and JavaScript for accessible tooltips
        doc.add_paragraph("CSS and JavaScript for accessible tooltips:", style='Normal').bold = True
        
        tooltip_code = doc.add_paragraph("""
/* CSS for accessible tooltips */
.tooltip {
  display: none;
  position: absolute;
  background: #333;
  color: white;
  padding: 10px;
  border-radius: 4px;
  z-index: 100;
}

.tooltip.visible {
  display: block;
}

/* JavaScript for accessible tooltips */
document.addEventListener('DOMContentLoaded', function() {
  // Find all elements that have tooltip descriptions
  const triggers = document.querySelectorAll('[aria-describedby]');
  
  triggers.forEach(trigger => {
    const tooltipId = trigger.getAttribute('aria-describedby');
    const tooltip = document.getElementById(tooltipId);
    
    if (!tooltip) return;
    
    // Show tooltip on hover for mouse users
    trigger.addEventListener('mouseenter', function() {
      positionTooltip(tooltip, trigger);
      tooltip.classList.add('visible');
    });
    
    trigger.addEventListener('mouseleave', function() {
      tooltip.classList.remove('visible');
    });
    
    // Show tooltip on focus for keyboard users
    trigger.addEventListener('focus', function() {
      positionTooltip(tooltip, trigger);
      tooltip.classList.add('visible');
    });
    
    trigger.addEventListener('blur', function() {
      tooltip.classList.remove('visible');
    });
  });
  
  // Position tooltip near its trigger
  function positionTooltip(tooltip, trigger) {
    const triggerRect = trigger.getBoundingClientRect();
    tooltip.style.top = (triggerRect.bottom + window.scrollY + 10) + 'px';
    tooltip.style.left = (triggerRect.left + window.scrollX) + 'px';
  }
});
        """)
        tooltip_code.style = doc.styles['Normal']
        tooltip_code.paragraph_format.left_indent = Pt(36)

    else:
        doc.add_paragraph("No improper uses of the title attribute were found.")
        