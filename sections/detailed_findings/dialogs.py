from docx.shared import Pt
from report_styling import format_table_text

def add_detailed_dialogs(doc, db_connection, total_domains):
    """Add the detailed Dialogs section"""
    doc.add_page_break()
    h2 = doc.add_heading('Dialogs', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
    Modal dialogs present unique accessibility challenges and must be implemented correctly to ensure all users can interact with them effectively. When a modal is open, keyboard focus must be trapped within it, and screen readers must be properly informed of the modal's presence and purpose.
    """.strip())

    doc.add_paragraph("Common accessibility issues with modal dialogs include:", style='Normal')

    doc.add_paragraph("Missing close mechanisms (close button or escape key)", style='List Bullet')
    doc.add_paragraph("Improper focus management when opening and closing", style='List Bullet')
    doc.add_paragraph("Missing or improper heading structure", style='List Bullet')
    doc.add_paragraph("Missing or improper trigger buttons", style='List Bullet')

    # Add recommendations
    doc.add_paragraph()
    doc.add_paragraph("Recommendations for Modal Dialog Implementation:", style='Normal')
    
    doc.add_paragraph("Ensure all modals have a visible close button and respond to the escape key", style='List Bullet')
    doc.add_paragraph("Implement proper focus management - trap focus within the modal when open and return focus when closed", style='List Bullet')
    doc.add_paragraph("Include proper heading structure within modals for clear content hierarchy", style='List Bullet')
    doc.add_paragraph("Use proper trigger buttons with appropriate ARIA attributes and keyboard interaction", style='List Bullet')
    doc.add_paragraph("Test modal interactions with keyboard-only navigation and screen readers", style='List Bullet')

    # Query for pages with modal issues
    pages_with_modal_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.modals.modals.pageFlags.hasModals": True,
            "results.accessibility.tests.modals.modals.pageFlags.hasModalViolations": True
        },
        {
            "url": 1,
            "results.accessibility.tests.modals.modals.pageFlags": 1,
            "results.accessibility.tests.modals.modals.details.summary": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    modal_issues = {
        "modalsWithoutClose": {"name": "Missing close mechanism", "pages": set(), "domains": set()},
        "modalsWithoutFocusManagement": {"name": "Improper focus management", "pages": set(), "domains": set()},
        "modalsWithoutProperHeading": {"name": "Missing/improper heading", "pages": set(), "domains": set()},
        "modalsWithoutTriggers": {"name": "Missing/improper triggers", "pages": set(), "domains": set()}
    }

    # Count issues
    for page in pages_with_modal_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        summary = page['results']['accessibility']['tests']['modals']['modals']['details']['summary']
        
        for flag in modal_issues:
            if summary.get(flag, 0) > 0:
                modal_issues[flag]['pages'].add(page['url'])
                modal_issues[flag]['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in modal_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Modal Issue"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(len(data['pages']))
            row[2].text = str(len(data['domains']))
            row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add domain details for each issue
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)

        # Add statistics about total modals if available
        doc.add_paragraph()
        total_modals = sum(page['results']['accessibility']['tests']['modals']['modals']['details']['summary']['totalModals'] 
                        for page in pages_with_modal_issues)
        doc.add_paragraph(f"Total number of modals detected across all pages: {total_modals}")

        # Add a detailed technical implementation section
        doc.add_paragraph()
        doc.add_heading('Technical Implementation Guidelines for Accessible Modals', level=3)
        
        # Close mechanism section
        doc.add_paragraph("Implementing proper close mechanisms:", style='Normal').bold = True
        doc.add_paragraph("Include a visible close button with a clear accessible name (e.g., 'Close', 'Close dialog')", style='List Bullet')
        doc.add_paragraph("Add an event listener for the Escape key to close the modal", style='List Bullet')
        doc.add_paragraph("Ensure the close button is keyboard focusable and has sufficient contrast", style='List Bullet')
        doc.add_paragraph("Example code:")
        
        code_close = doc.add_paragraph("""
// Example close button implementation
const closeButton = document.createElement('button');
closeButton.textContent = 'Close';
closeButton.setAttribute('aria-label', 'Close dialog');
closeButton.classList.add('modal-close');
closeButton.addEventListener('click', closeModal);

// Example Escape key handler
document.addEventListener('keydown', function(event) {
  if (event.key === 'Escape' && modalIsOpen) {
    closeModal();
  }
});
        """)
        code_close.style = doc.styles['Normal']
        code_close.paragraph_format.left_indent = Pt(36)
        
        # Focus management section
        doc.add_paragraph("Proper focus management:", style='Normal').bold = True
        doc.add_paragraph("Store the element that opened the modal to return focus later", style='List Bullet')
        doc.add_paragraph("Set focus to the first focusable element in the modal when it opens", style='List Bullet')
        doc.add_paragraph("Trap keyboard focus within the modal using a focus trap", style='List Bullet')
        doc.add_paragraph("Return focus to the triggering element when the modal closes", style='List Bullet')
        doc.add_paragraph("Example code:")
        
        code_focus = doc.add_paragraph("""
// Store the element that triggered the modal
let lastFocusedElement;

function openModal() {
  lastFocusedElement = document.activeElement;
  
  // Show modal
  modal.style.display = 'block';
  
  // Focus the first focusable element
  const focusableElements = modal.querySelectorAll('button, [href], input, select, textarea, [tabindex]:not([tabindex="-1"])');
  if (focusableElements.length > 0) {
    focusableElements[0].focus();
  }
  
  // Set up focus trap
  modal.addEventListener('keydown', trapFocus);
}

function closeModal() {
  // Hide modal
  modal.style.display = 'none';
  
  // Return focus to triggering element
  if (lastFocusedElement) {
    lastFocusedElement.focus();
  }
  
  // Remove focus trap
  modal.removeEventListener('keydown', trapFocus);
}

function trapFocus(e) {
  // Implementation of focus trapping...
}
        """)
        code_focus.style = doc.styles['Normal']
        code_focus.paragraph_format.left_indent = Pt(36)
        
        # Heading structure section
        doc.add_paragraph("Proper heading structure:", style='Normal').bold = True
        doc.add_paragraph("Include a descriptive heading as the first element in the modal", style='List Bullet')
        doc.add_paragraph("Use appropriate heading level (usually h2) for the modal title", style='List Bullet')
        doc.add_paragraph("Connect the modal with its heading using aria-labelledby", style='List Bullet')
        doc.add_paragraph("Example code:")
        
        code_heading = doc.add_paragraph("""
<div role="dialog" aria-modal="true" aria-labelledby="dialog-title">
  <h2 id="dialog-title">Modal Dialog Title</h2>
  <div class="dialog-content">
    <!-- Modal content here -->
  </div>
  <button aria-label="Close dialog">Close</button>
</div>
        """)
        code_heading.style = doc.styles['Normal']
        code_heading.paragraph_format.left_indent = Pt(36)

    else:
        doc.add_paragraph("No dialog accessibility issues were found.")
