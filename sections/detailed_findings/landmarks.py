from report_styling import format_table_text
from docx.shared import Pt

def add_detailed_landmarks(doc, db_connection, total_domains):
    """Add the detailed Landmarks section"""
    doc.add_page_break()
    h2 = doc.add_heading('Landmarks', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
    HTML landmarks provide a navigational structure that helps screen reader users understand the organization of a page's content. Properly implemented landmarks are crucial for efficient navigation and orientation. Each landmark role serves a specific purpose and should be used appropriately.
    """.strip())

    doc.add_paragraph("Common landmark roles include:", style='Normal')

    doc.add_paragraph("banner - Header content", style='List Bullet')
    doc.add_paragraph("main - Primary content area", style='List Bullet')
    doc.add_paragraph("navigation - Navigation sections", style='List Bullet')
    doc.add_paragraph("complementary - Supporting content", style='List Bullet')
    doc.add_paragraph("contentinfo - Footer content", style='List Bullet')
    doc.add_paragraph("search - Search functionality", style='List Bullet')
    doc.add_paragraph("form - Form sections", style='List Bullet')
    doc.add_paragraph("region - Distinct sections requiring labels", style='List Bullet')

    # Add recommendations
    doc.add_paragraph()
    doc.add_paragraph("Recommendations for Landmark Implementation:", style='Normal')
    
    doc.add_paragraph("Ensure all pages have the required landmarks (banner, main, contentinfo)", style='List Bullet')
    doc.add_paragraph("Provide unique names for duplicate landmarks using aria-label or aria-labelledby", style='List Bullet')
    doc.add_paragraph("Avoid nesting top-level landmarks", style='List Bullet')
    doc.add_paragraph("Ensure all content is contained within appropriate landmarks", style='List Bullet')
    doc.add_paragraph("Use semantic HTML elements with implicit landmark roles where possible", style='List Bullet')

    # Query for pages with landmark issues
    pages_with_landmark_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.landmarks.landmarks.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.landmarks.landmarks.pageFlags.missingRequiredLandmarks": True},
                {"results.accessibility.tests.landmarks.landmarks.pageFlags.hasDuplicateLandmarksWithoutNames": True},
                {"results.accessibility.tests.landmarks.landmarks.pageFlags.hasNestedTopLevelLandmarks": True},
                {"results.accessibility.tests.landmarks.landmarks.pageFlags.hasContentOutsideLandmarks": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.landmarks.landmarks": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for different landmark issues
    landmark_issues = {
        "missing": {
            "name": "Missing required landmarks",
            "pages": set(),
            "domains": set(),
            "details": {
                "banner": 0,
                "main": 0,
                "contentinfo": 0,
                "search": 0
            }
        },
        "duplicate": {
            "name": "Duplicate landmarks without unique names",
            "pages": set(),
            "domains": set(),
            "details": {
                "banner": 0,
                "main": 0,
                "navigation": 0,
                "complementary": 0,
                "contentinfo": 0,
                "search": 0,
                "form": 0,
                "region": 0
            }
        },
        "nested": {
            "name": "Nested top-level landmarks",
            "pages": set(),
            "domains": set()
        },
        "outside": {
            "name": "Content outside landmarks",
            "pages": set(),
            "domains": set(),
            "count": 0
        }
    }

    # Process each page
    total_landmarks = 0
    for page in pages_with_landmark_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        landmark_data = page['results']['accessibility']['tests']['landmarks']['landmarks']
        flags = landmark_data['pageFlags']
        details = flags['details']
        
        # Count total landmarks
        if 'totalLandmarks' in landmark_data.get('details', {}).get('summary', {}):
            total_landmarks += landmark_data['details']['summary']['totalLandmarks']
        
        # Check missing landmarks
        if flags.get('missingRequiredLandmarks'):
            landmark_issues['missing']['pages'].add(page['url'])
            landmark_issues['missing']['domains'].add(domain)
            missing = details.get('missingLandmarks', {})
            for landmark in ['banner', 'main', 'contentinfo', 'search']:
                if missing.get(landmark):
                    landmark_issues['missing']['details'][landmark] += 1

        # Check duplicate landmarks
        if flags.get('hasDuplicateLandmarksWithoutNames'):
            landmark_issues['duplicate']['pages'].add(page['url'])
            landmark_issues['duplicate']['domains'].add(domain)
            duplicates = details.get('duplicateLandmarks', {})
            for landmark in landmark_issues['duplicate']['details'].keys():
                if landmark in duplicates:
                    landmark_issues['duplicate']['details'][landmark] += duplicates[landmark].get('count', 0)

        # Check nested landmarks
        if flags.get('hasNestedTopLevelLandmarks'):
            landmark_issues['nested']['pages'].add(page['url'])
            landmark_issues['nested']['domains'].add(domain)

        # Check content outside landmarks
        if flags.get('hasContentOutsideLandmarks'):
            landmark_issues['outside']['pages'].add(page['url'])
            landmark_issues['outside']['domains'].add(domain)
            landmark_issues['outside']['count'] += details.get('contentOutsideLandmarksCount', 0)

    # Add statistics
    doc.add_paragraph()
    doc.add_paragraph("Landmark Statistics:", style='Normal')
    doc.add_paragraph(f"Total number of landmarks detected across all pages: {total_landmarks}")

    # Create summary table
    if any(len(issue['pages']) > 0 for issue in landmark_issues.values()):
        # Create main issues summary table
        summary_table = doc.add_table(rows=len(landmark_issues) + 1, cols=4)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Issue Type"
        headers[1].text = "Pages Affected"
        headers[2].text = "Sites Affected"
        headers[3].text = "% of Total Sites"

        # Add data
        row_idx = 1
        for issue_type, data in landmark_issues.items():
            if len(data['pages']) > 0:
                row = summary_table.rows[row_idx].cells
                row[0].text = data['name']
                row[1].text = str(len(data['pages']))
                row[2].text = str(len(data['domains']))
                row[3].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"
                row_idx += 1

        # Format the table text
        format_table_text(summary_table)

        # Add specific details for missing landmarks
        if landmark_issues['missing']['pages']:
            doc.add_paragraph()
            doc.add_paragraph("Missing Required Landmarks Breakdown:", style='Normal')
            
            missing_table = doc.add_table(rows=5, cols=2)
            missing_table.style = 'Table Grid'
            
            headers = missing_table.rows[0].cells
            headers[0].text = "Landmark Type"
            headers[1].text = "Number of Pages Missing"
            
            landmarks = [("Banner", "banner"), ("Main", "main"), 
                        ("Footer", "contentinfo"), ("Search", "search")]
            
            for idx, (name, key) in enumerate(landmarks, 1):
                row = missing_table.rows[idx].cells
                row[0].text = name
                row[1].text = str(landmark_issues['missing']['details'][key])
            
            format_table_text(missing_table)

        # Technical implementation examples
        doc.add_paragraph()
        doc.add_heading('Technical Implementation Guidelines', level=3)
        
        # HTML5 semantic elements example
        doc.add_paragraph("Using HTML5 semantic elements with implicit landmark roles:", style='Normal').bold = True
        
        html5_code = doc.add_paragraph("""
<!-- Good practice: Using HTML5 semantic elements -->
<header>Site header content</header>  <!-- implicit role="banner" -->

<nav>
  <ul>
    <li><a href="/">Home</a></li>
    <li><a href="/products">Products</a></li>
  </ul>
</nav>  <!-- implicit role="navigation" -->

<main>
  <h1>Page Title</h1>
  <article>Main content...</article>
</main>  <!-- implicit role="main" -->

<aside>Related content...</aside>  <!-- implicit role="complementary" -->

<footer>Site footer content</footer>  <!-- implicit role="contentinfo" -->
        """)
        html5_code.style = doc.styles['Normal']
        html5_code.paragraph_format.left_indent = Pt(36)
        
        # ARIA landmark roles example
        doc.add_paragraph("Using ARIA landmark roles (when semantic HTML isn't possible):", style='Normal').bold = True
        
        aria_code = doc.add_paragraph("""
<!-- Alternative using ARIA roles -->
<div role="banner">Site header content</div>
<div role="navigation">Navigation menu...</div>
<div role="main">Main content...</div>
<div role="complementary">Related content...</div>
<div role="contentinfo">Site footer content</div>
        """)
        aria_code.style = doc.styles['Normal']
        aria_code.paragraph_format.left_indent = Pt(36)
        
        # Naming landmarks example
        doc.add_paragraph("Naming landmarks when you have multiple of the same type:", style='Normal').bold = True
        
        naming_code = doc.add_paragraph("""
<!-- When you have multiple navigation landmarks -->
<nav aria-label="Main">Primary navigation...</nav>
<nav aria-label="Footer">Footer links...</nav>

<!-- When you have multiple complementary landmarks -->
<aside aria-label="Related articles">Related articles...</aside>
<aside aria-label="Advertisements">Advertisements...</aside>

<!-- Using aria-labelledby instead -->
<h2 id="side-nav-heading">Product Categories</h2>
<nav aria-labelledby="side-nav-heading">
  <!-- Navigation content -->
</nav>
        """)
        naming_code.style = doc.styles['Normal']
        naming_code.paragraph_format.left_indent = Pt(36)
        
        # Example page structure with all required landmarks
        doc.add_paragraph("Complete page example with proper landmark structure:", style='Normal').bold = True
        
        complete_code = doc.add_paragraph("""
<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <title>Accessible Page Example</title>
</head>
<body>
  <a href="#main-content" class="skip-link">Skip to main content</a>
  
  <header>
    <div class="logo">Site Logo</div>
    <nav aria-label="Main">
      <ul>
        <li><a href="/">Home</a></li>
        <li><a href="/about">About</a></li>
        <li><a href="/contact">Contact</a></li>
      </ul>
    </nav>
    
    <div role="search">
      <form>
        <label for="search">Search</label>
        <input type="search" id="search">
        <button type="submit">Search</button>
      </form>
    </div>
  </header>
  
  <main id="main-content">
    <h1>Page Title</h1>
    <section>
      <h2>Section Title</h2>
      <p>Content goes here...</p>
    </section>
  </main>
  
  <aside>
    <h2>Related Information</h2>
    <ul>
      <li><a href="#">Related Link 1</a></li>
      <li><a href="#">Related Link 2</a></li>
    </ul>
  </aside>
  
  <footer>
    <nav aria-label="Footer">
      <ul>
        <li><a href="/privacy">Privacy Policy</a></li>
        <li><a href="/terms">Terms of Service</a></li>
      </ul>
    </nav>
    <p>&copy; 2023 Company Name</p>
  </footer>
</body>
</html>
        """)
        complete_code.style = doc.styles['Normal']
        complete_code.paragraph_format.left_indent = Pt(36)

        # Add domain details for each issue type
        for issue_type, data in landmark_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)

    else:
        doc.add_paragraph("No landmark structure issues were found.")
        