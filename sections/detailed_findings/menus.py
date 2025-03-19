# sections/detailed_findings/menus.py
from report_styling import format_table_text
from docx.shared import Pt

def add_detailed_menus(doc, db_connection, total_domains):
    """Add the detailed Menus section"""
    doc.add_page_break()
    h2 = doc.add_heading('Menus', level=2)
    h2.style = doc.styles['Heading 2']

    # Add explanation
    doc.add_paragraph("""
Navigation menus are critical components for website accessibility. They must be properly structured, labeled, and implement correct ARIA roles and attributes to ensure all users can navigate effectively. Screen reader users particularly rely on well-implemented navigation menus.
    """.strip())

    doc.add_paragraph("Common accessibility issues with navigation menus include:", style='Normal')

    doc.add_paragraph("Missing or invalid ARIA roles for navigation elements", style='List Bullet')
    doc.add_paragraph("Missing current page indicators", style='List Bullet')
    doc.add_paragraph("Missing or improper menu labels and names", style='List Bullet')
    doc.add_paragraph("Duplicate menu names causing confusion", style='List Bullet')

    # Add recommendations
    doc.add_paragraph()
    doc.add_paragraph("Recommendations for Navigation Menu Implementation:", style='Normal')
    
    doc.add_paragraph("Use proper ARIA roles (e.g., navigation, menubar, menu) for navigation elements", style='List Bullet')
    doc.add_paragraph("Implement clear current page indicators using aria-current", style='List Bullet')
    doc.add_paragraph("Ensure all navigation menus have unique, descriptive labels", style='List Bullet')
    doc.add_paragraph("Use appropriate heading levels for menu sections", style='List Bullet')
    doc.add_paragraph("Ensure keyboard navigation works properly within menus", style='List Bullet')
    doc.add_paragraph("Test menu functionality with screen readers", style='List Bullet')

    # Query for pages with menu issues
    pages_with_menu_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.menus.menus.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.menus.menus.pageFlags.hasInvalidMenuRoles": True},
                {"results.accessibility.tests.menus.menus.pageFlags.hasMenusWithoutCurrent": True},
                {"results.accessibility.tests.menus.menus.pageFlags.hasUnnamedMenus": True},
                {"results.accessibility.tests.menus.menus.pageFlags.hasDuplicateMenuNames": True}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.menus.menus": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Initialize counters for each issue type
    menu_issues = {
        "invalidRoles": {"name": "Invalid menu roles", "pages": set(), "domains": set(), "count": 0},
        "menusWithoutCurrent": {"name": "Missing current page indicators", "pages": set(), "domains": set(), "count": 0},
        "unnamedMenus": {"name": "Unnamed menus", "pages": set(), "domains": set(), "count": 0},
        "duplicateNames": {"name": "Duplicate menu names", "pages": set(), "domains": set(), "count": 0}
    }

    # Count issues
    total_menus = 0
    for page in pages_with_menu_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        menu_data = page['results']['accessibility']['tests']['menus']['menus']
        flags = menu_data['pageFlags']
        details = menu_data['pageFlags']['details']
        
        total_menus += details.get('totalMenus', 0)
        
        # Check each type of issue
        if flags.get('hasInvalidMenuRoles'):
            menu_issues['invalidRoles']['pages'].add(page['url'])
            menu_issues['invalidRoles']['domains'].add(domain)
            menu_issues['invalidRoles']['count'] += details.get('invalidRoles', 0)
            
        if flags.get('hasMenusWithoutCurrent'):
            menu_issues['menusWithoutCurrent']['pages'].add(page['url'])
            menu_issues['menusWithoutCurrent']['domains'].add(domain)
            menu_issues['menusWithoutCurrent']['count'] += details.get('menusWithoutCurrent', 0)
            
        if flags.get('hasUnnamedMenus'):
            menu_issues['unnamedMenus']['pages'].add(page['url'])
            menu_issues['unnamedMenus']['domains'].add(domain)
            menu_issues['unnamedMenus']['count'] += details.get('unnamedMenus', 0)
            
        if flags.get('hasDuplicateMenuNames'):
            menu_issues['duplicateNames']['pages'].add(page['url'])
            menu_issues['duplicateNames']['domains'].add(domain)

    # Create filtered list of issues that have affected pages
    active_issues = {flag: data for flag, data in menu_issues.items() 
                    if len(data['pages']) > 0}

    if active_issues:
        # Create summary table
        summary_table = doc.add_table(rows=len(active_issues) + 1, cols=5)
        summary_table.style = 'Table Grid'

        # Set column headers
        headers = summary_table.rows[0].cells
        headers[0].text = "Menu Issue"
        headers[1].text = "Number of Occurrences"
        headers[2].text = "Pages Affected"
        headers[3].text = "Sites Affected"
        headers[4].text = "% of Total Sites"

        # Add data
        for i, (flag, data) in enumerate(active_issues.items(), 1):
            row = summary_table.rows[i].cells
            row[0].text = data['name']
            row[1].text = str(data['count']) if flag != 'duplicateNames' else 'N/A'
            row[2].text = str(len(data['pages']))
            row[3].text = str(len(data['domains']))
            row[4].text = f"{(len(data['domains']) / len(total_domains) * 100):.1f}%"

        # Format the table text
        format_table_text(summary_table)

        # Add statistics about total menus
        doc.add_paragraph()
        doc.add_paragraph(f"Total number of navigation menus detected across all pages: {total_menus}")

        # Add domain details for each issue
        for flag, data in active_issues.items():
            if data['domains']:
                doc.add_paragraph()
                doc.add_paragraph(f"Sites with {data['name'].lower()}:")
                
                # Group by domain and count occurrences
                domain_counts = {}
                for page in data['pages']:
                    domain = page.replace('http://', '').replace('https://', '').split('/')[0]
                    domain_counts[domain] = domain_counts.get(domain, 0) + 1

                # Create domain details table
                domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
                domain_table.style = 'Table Grid'

                # Add headers
                headers = domain_table.rows[0].cells
                headers[0].text = "Domain"
                headers[1].text = "Number of pages"

                # Add domain data
                for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                    row = domain_table.rows[i].cells
                    row[0].text = domain
                    row[1].text = str(count)

                # Format the table text
                format_table_text(domain_table)
                
        # Add technical implementation section
        doc.add_paragraph()
        doc.add_heading('Technical Implementation Guidelines', level=3)
        
        # Basic navigation example
        doc.add_paragraph("Primary navigation menu best practices:", style='Normal').bold = True
        
        nav_code = doc.add_paragraph("""
<!-- Primary navigation with proper structure -->
<nav aria-label="Main navigation">
  <ul>
    <li><a href="/" aria-current="page">Home</a></li>
    <li><a href="/about">About Us</a></li>
    <li><a href="/services">Services</a></li>
    <li><a href="/contact">Contact</a></li>
  </ul>
</nav>
        """)
        nav_code.style = doc.styles['Normal']
        nav_code.paragraph_format.left_indent = Pt(36)
        
        # Dropdown menu example
        doc.add_paragraph("Dropdown menu implementation:", style='Normal').bold = True
        
        dropdown_code = doc.add_paragraph("""
<!-- Dropdown menu with ARIA -->
<nav aria-label="Main navigation">
  <ul role="menubar">
    <li role="none">
      <a href="/" role="menuitem" aria-current="page">Home</a>
    </li>
    <li role="none">
      <button role="menuitem" aria-haspopup="true" aria-expanded="false">
        Products
        <span class="dropdown-icon" aria-hidden="true">▼</span>
      </button>
      <ul role="menu">
        <li role="none">
          <a href="/products/software" role="menuitem">Software</a>
        </li>
        <li role="none">
          <a href="/products/hardware" role="menuitem">Hardware</a>
        </li>
        <li role="none">
          <a href="/products/services" role="menuitem">Services</a>
        </li>
      </ul>
    </li>
    <li role="none">
      <a href="/about" role="menuitem">About</a>
    </li>
    <li role="none">
      <a href="/contact" role="menuitem">Contact</a>
    </li>
  </ul>
</nav>
        """)
        dropdown_code.style = doc.styles['Normal']
        dropdown_code.paragraph_format.left_indent = Pt(36)
        
        # JavaScript for keyboard interaction
        doc.add_paragraph("JavaScript for keyboard navigation:", style='Normal').bold = True
        
        js_code = doc.add_paragraph("""
// Handle keyboard navigation for dropdown menus
document.addEventListener('DOMContentLoaded', () => {
  const menuButtons = document.querySelectorAll('button[aria-haspopup="true"]');
  
  menuButtons.forEach(button => {
    // Toggle dropdown
    button.addEventListener('click', () => {
      const expanded = button.getAttribute('aria-expanded') === 'true';
      button.setAttribute('aria-expanded', !expanded);
      
      // Get the dropdown menu
      const menu = button.nextElementSibling;
      if (menu) {
        menu.style.display = expanded ? 'none' : 'block';
        
        // If opening menu, focus first item
        if (!expanded) {
          const firstItem = menu.querySelector('[role="menuitem"]');
          if (firstItem) firstItem.focus();
        }
      }
    });
    
    // Open dropdown on key press
    button.addEventListener('keydown', (e) => {
      // Enter, Space, Down Arrow
      if (e.key === 'Enter' || e.key === ' ' || e.key === 'ArrowDown') {
        e.preventDefault();
        button.setAttribute('aria-expanded', 'true');
        
        const menu = button.nextElementSibling;
        if (menu) {
          menu.style.display = 'block';
          const firstItem = menu.querySelector('[role="menuitem"]');
          if (firstItem) firstItem.focus();
        }
      }
    });
  });
  
  // Handle keyboard navigation within menus
  document.querySelectorAll('[role="menu"]').forEach(menu => {
    menu.addEventListener('keydown', (e) => {
      const items = Array.from(menu.querySelectorAll('[role="menuitem"]'));
      const currentIndex = items.indexOf(document.activeElement);
      
      switch (e.key) {
        case 'ArrowDown':
          e.preventDefault();
          if (currentIndex < items.length - 1) {
            items[currentIndex + 1].focus();
          } else {
            items[0].focus();  // Wrap to first item
          }
          break;
          
        case 'ArrowUp':
          e.preventDefault();
          if (currentIndex > 0) {
            items[currentIndex - 1].focus();
          } else {
            items[items.length - 1].focus();  // Wrap to last item
          }
          break;
          
        case 'Escape':
          e.preventDefault();
          const parentButton = menu.previousElementSibling;
          if (parentButton && parentButton.hasAttribute('aria-haspopup')) {
            parentButton.setAttribute('aria-expanded', 'false');
            menu.style.display = 'none';
            parentButton.focus();
          }
          break;
      }
    });
  });
});
        """)
        js_code.style = doc.styles['Normal']
        js_code.paragraph_format.left_indent = Pt(36)
        
        # Multiple navigation areas
        doc.add_paragraph("Multiple navigation landmarks:", style='Normal').bold = True
        
        multiple_nav_code = doc.add_paragraph("""
<!-- Multiple navigation areas with unique labels -->
<header>
  <nav aria-label="Primary navigation">
    <!-- Primary menu content -->
  </nav>
</header>

<aside>
  <nav aria-label="Section navigation">
    <!-- Section menu content -->
  </nav>
</aside>

<footer>
  <nav aria-label="Footer navigation">
    <!-- Footer links -->
  </nav>
</footer>
        """)
        multiple_nav_code.style = doc.styles['Normal']
        multiple_nav_code.paragraph_format.left_indent = Pt(36)

    else:
        doc.add_paragraph("No navigation menu accessibility issues were found.")
        