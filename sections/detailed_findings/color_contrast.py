import json
from docx.oxml import parse_xml
from report_styling import format_table_text

def add_detailed_color_contrast(doc, db_connection, total_domains):
    """Add the detailed Color Contrast section"""
    doc.add_page_break()
    h2 = doc.add_heading('Colour Contrast', level=2)
    h2.style = doc.styles['Heading 2']

    # 1. Text Contrast Issues
    doc.add_paragraph("Text Contrast Issues:").bold = True
    doc.add_paragraph("Checks if text has sufficient contrast with its background", style='List Bullet')
    doc.add_paragraph("Small text needs a contrast ratio of at least 4.5:1", style='List Bullet')
    doc.add_paragraph("Large text (18pt+ or 14pt+ bold) needs at least 3:1", style='List Bullet')
    doc.add_paragraph("Essential for users with low vision or color vision deficiencies", style='List Bullet')
    doc.add_paragraph()

    # 2. Non-Text Contrast Issues
    doc.add_paragraph("Non-Text Contrast Issues:").bold = True
    doc.add_paragraph("Examines contrast of UI components like buttons and horizontal rules", style='List Bullet')
    doc.add_paragraph("Requires a minimum contrast ratio of 3:1 for boundaries and visual information", style='List Bullet')
    doc.add_paragraph("Important for identifying clickable elements and interface boundaries", style='List Bullet')
    doc.add_paragraph()

    # 3. Adjacent Contrast Issues
    doc.add_paragraph("Adjacent Contrast Issues:").bold = True
    doc.add_paragraph("Checks contrast between neighboring content blocks", style='List Bullet')
    doc.add_paragraph("Requires sufficient contrast between adjacent sections", style='List Bullet')
    doc.add_paragraph("Helps users distinguish between different content areas", style='List Bullet')
    doc.add_paragraph()

    # 4. Contrast Preferences Support
    doc.add_paragraph("Contrast Preferences Support:").bold = True
    doc.add_paragraph("Checks if the site responds to system-level contrast preferences", style='List Bullet')
    doc.add_paragraph("Should adapt to user's contrast preference settings", style='List Bullet')
    doc.add_paragraph("Important for users who need specific contrast levels", style='List Bullet')
    doc.add_paragraph()

    # Define the contrast issues to be analyzed
    contrast_issues = [
        {
            'name': 'Text Contrast Issues',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.hasContrastIssues',
            'details_field': 'results.accessibility.tests.colors.colors.details.summary.contrastViolations'
        },
        {
            'name': 'Non-Text Contrast Issues',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.hasNonTextContrastIssues',
            'details_field': 'results.accessibility.tests.colors.colors.details.summary.nonTextContrastViolations'
        },
        {
            'name': 'Adjacent Contrast Issues',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.hasAdjacentContrastIssues',
            'details_field': 'results.accessibility.tests.colors.colors.details.summary.adjacentContrastViolations'
        },
        {
            'name': 'Contrast Preferences Support',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.supportsContrastPreferences',
            'details_field': None  # This is a boolean field, not a count
        }
    ]

    # Gather the data for each issue type
    issue_data = {}

    for issue in contrast_issues:
        # For the Contrast Preferences Support, we want sites that DO support it
        # For other issues, we want sites that have problems
        if issue['name'] == 'Contrast Preferences Support':
            query = {issue['db_field']: True}
        else:
            query = {issue['db_field']: True}
        
        # Prepare projection
        projection = {"url": 1, "_id": 0}
        if issue['details_field']:
            projection[issue['details_field']] = 1
        
        # Query the database to find pages with this issue
        pages_with_issue = list(db_connection.page_results.find(query, projection))
        
        # Count affected domains and total issue instances
        affected_domains = set()
        total_instances = 0
        
        for page in pages_with_issue:
            domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
            affected_domains.add(domain)
            
            # Count instances if applicable
            if issue['details_field']:
                # Navigate the nested structure to get the count
                parts = issue['details_field'].split('.')
                value = page
                try:
                    for part in parts:
                        if part in value:
                            value = value[part]
                        else:
                            value = 0
                            break
                    
                    if isinstance(value, (int, float)):
                        total_instances += value
                except:
                    pass  # Handle any issues with nested access
        
        # Store the data
        issue_data[issue['name']] = {
            'pages': pages_with_issue,
            'domains': affected_domains,
            'instances': total_instances
        }

    # Create summary table
    last_para = doc.add_paragraph()
    last_para._element.get_or_add_pPr().append(
        parse_xml(r'<w:keepNext xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    )

    summary_table = doc.add_table(rows=len(contrast_issues) + 1, cols=4)
    summary_table.style = 'Table Grid'

    # Keep table together
    for row in summary_table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcPr.append(parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>'))
            tcPr.append(parse_xml(r'<w:noWrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

    # Set column headers
    headers = summary_table.rows[0].cells
    headers[0].text = "Color Accessibility Issue"
    headers[1].text = "Pages Affected"
    headers[2].text = "Sites Affected"
    headers[3].text = "% of Total Sites"

    # Add data
    for i, issue in enumerate(contrast_issues, 1):
        row = summary_table.rows[i].cells
        data = issue_data[issue['name']]
        
        row[0].text = issue['name']
        row[1].text = str(len(data['pages']))
        row[2].text = str(len(data['domains']))
        
        percentage = (len(data['domains']) / len(total_domains)) * 100 if total_domains else 0
        row[3].text = f"{percentage:.1f}%"

    # Format the table text
    format_table_text(summary_table)

    # Add some space after the table
    doc.add_paragraph()

    # Add details for each issue type that has occurrences
    for issue in contrast_issues:
        data = issue_data[issue['name']]
        if data['domains']:
            doc.add_paragraph(f"Sites with {issue['name'].lower()}:")
            
            # Create a dictionary to count pages per domain
            domain_counts = {}
            for page in data['pages']:
                domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
                domain_counts[domain] = domain_counts.get(domain, 0) + 1
            
            # Create domain details table
            domain_table = doc.add_table(rows=len(domain_counts) + 1, cols=2)
            domain_table.style = 'Table Grid'

            # Keep table together
            for row in domain_table.rows:
                for cell in row.cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcPr.append(parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>'))
                    tcPr.append(parse_xml(r'<w:noWrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

            # Add headers
            headers = domain_table.rows[0].cells
            headers[0].text = "Domain"
            headers[1].text = "Number of pages"

            # Add domain data
            for i, (domain, count) in enumerate(sorted(domain_counts.items()), 1):
                row = domain_table.rows[i].cells
                row[0].text = domain
                row[1].text = str(count)

            # Format the table text
            format_table_text(domain_table)
            doc.add_paragraph()  # Add space after each table
            
    # Add recommendations section
    doc.add_paragraph()
    doc.add_heading('Recommendations for Addressing Contrast Issues', level=3)
    
    doc.add_paragraph("For Text Contrast Issues:", style='Normal').bold = True
    doc.add_paragraph("Ensure all text meets minimum contrast ratios (4.5:1 for regular text, 3:1 for large text)", style='List Bullet')
    doc.add_paragraph("Use contrast checking tools during design and development", style='List Bullet')
    doc.add_paragraph("Consider offering high contrast modes or user-selectable themes", style='List Bullet')
    
    doc.add_paragraph("For Non-Text Contrast Issues:", style='Normal').bold = True
    doc.add_paragraph("Ensure UI components and meaningful graphics have at least 3:1 contrast", style='List Bullet')
    doc.add_paragraph("Add borders or additional visual indicators to low-contrast elements", style='List Bullet')
    doc.add_paragraph("Test interface elements against various backgrounds", style='List Bullet')
    
    doc.add_paragraph("For Adjacent Contrast Issues:", style='Normal').bold = True
    doc.add_paragraph("Use borders, spacing, or other visual separators between content blocks", style='List Bullet')
    doc.add_paragraph("Test interfaces with color perception simulators", style='List Bullet')
    
    doc.add_paragraph("For Contrast Preferences Support:", style='Normal').bold = True
    doc.add_paragraph("Implement support for prefers-contrast media query", style='List Bullet')
    doc.add_paragraph("Test with system-level high contrast settings enabled", style='List Bullet')
    doc.add_paragraph("Provide user controls for adjusting contrast when possible", style='List Bullet')
    