from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from report_styling import add_page_number

def setup_document_header_footer(doc, title):
    """Set up the document header and footer"""
    # Get the first section
    section = doc.sections[0]
    
    # Set different first page for header/footer
    section.different_first_page_header_footer = True
    
    # Header (will only appear from second page onwards)
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.text = title
    for run in header_para.runs:
        run.font.size = Pt(14)

    # Footer - Create a table for better alignment control
    footer = section.footer
    footer_table = footer.add_table(1, 3, width=Inches(6.5))
    footer_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set equal widths for all columns
    for cell in footer_table.columns:
        cell.width = Inches(2.17)
    
    # Set table properties
    footer_table.style = 'Table Grid'
    footer_table.autofit = False
    
    # Remove all borders and set font size
    for cell in footer_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(14)
                run.font.name = 'Arial'
        cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>'))
    
    # Footer content
    space_paragraph = footer_table.rows[0].cells[0].paragraphs[0]
    space_paragraph.add_run(" ")
    
    page_num_paragraph = footer_table.rows[0].cells[1].paragraphs[0]
    page_num_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(page_num_paragraph)
    
    logo_paragraph = footer_table.rows[0].cells[2].paragraphs[0]
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    import os
    if os.path.exists('logo.png'):
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture('logo.png', width=Inches(1))
    
    # Set table cell vertical alignment to center
    for cell in footer_table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:vAlign')
        tcVAlign.set(qn('w:val'), "center")
        tcPr.append(tcVAlign)
        