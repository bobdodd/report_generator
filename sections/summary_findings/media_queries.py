from docx.oxml import parse_xml
from report_styling import format_table_text

def add_media_queries_section(doc, db_connection, total_domains):
    """Add the Media Queries section to the summary findings"""
    doc.add_paragraph()
    h2 = doc.add_heading('Media Queries Summary', level=2)
    h2.style = doc.styles['Heading 2']

    # Collect responsive breakpoints across all pages
    breakpoint_data = {}
    pages_with_breakpoints = list(db_connection.page_results.find(
        {"results.accessibility.tests.media_queries.media_queries.responsiveBreakpoints": {"$exists": True}},
        {
            "url": 1,
            "results.accessibility.tests.media_queries.media_queries.responsiveBreakpoints": 1,
            "_id": 0
        }
    ))

    # Collect all breakpoints and organize by category
    all_breakpoints = set()
    breakpoint_by_category = {
        'mobile': set(),
        'tablet': set(),
        'desktop': set(),
        'largeScreen': set()
    }
    
    for page in pages_with_breakpoints:
        if 'responsiveBreakpoints' in page['results']['accessibility']['tests']['media_queries']['media_queries']:
            breakpoints_data = page['results']['accessibility']['tests']['media_queries']['media_queries']['responsiveBreakpoints']
            
            # Add to the full list of breakpoints
            if 'allBreakpoints' in breakpoints_data:
                for bp in breakpoints_data['allBreakpoints']:
                    all_breakpoints.add(bp)
            
            # Add to category-specific sets
            if 'byCategory' in breakpoints_data:
                for category, bps in breakpoints_data['byCategory'].items():
                    if category in breakpoint_by_category:
                        for bp in bps:
                            breakpoint_by_category[category].add(bp)

    # Add responsive breakpoints summary if available
    if all_breakpoints:
        doc.add_heading('Responsive Breakpoints Summary', level=3)
        doc.add_paragraph(f"A total of {len(all_breakpoints)} unique responsive breakpoints were detected across the site.")
        
        # Create a table for breakpoints by category
        category_table = doc.add_table(rows=5, cols=2)
        category_table.style = 'Table Grid'
        
        # Add headers
        headers = category_table.rows[0].cells
        headers[0].text = "Device Category"
        headers[1].text = "Breakpoints Detected"
        
        # Add category data
        categories = [
            ("Mobile (≤480px)", sorted(breakpoint_by_category['mobile'])),
            ("Tablet (481-768px)", sorted(breakpoint_by_category['tablet'])),
            ("Desktop (769-1200px)", sorted(breakpoint_by_category['desktop'])),
            ("Large Screen (>1200px)", sorted(breakpoint_by_category['largeScreen']))
        ]
        
        for i, (category, bps) in enumerate(categories, 1):
            row = category_table.rows[i].cells
            row[0].text = category
            if bps:
                count_text = f"{len(bps)} breakpoints"
                if len(bps) <= 5:  # Show all breakpoints if there are 5 or fewer
                    count_text += f": {', '.join(str(bp) for bp in bps)}"
                row[1].text = count_text
            else:
                row[1].text = "None detected"
        
        # Format the table
        format_table_text(category_table)
        
        doc.add_paragraph()

    # Query for pages with responsive design issues
    pages_with_media_query_issues = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.media_queries.media_queries.pageFlags": {"$exists": True},
            "$or": [
                {"results.accessibility.tests.media_queries.media_queries.pageFlags.hasResponsiveBreakpoints": False},
                {"results.accessibility.tests.media_queries.media_queries.pageFlags.hasPrintStyles": False},
                {"results.accessibility.tests.media_queries.media_queries.pageFlags.hasReducedMotionSupport": False},
                {"results.accessibility.tests.media_queries.media_queries.pageFlags.hasDarkModeSupport": False}
            ]
        },
        {
            "url": 1,
            "results.accessibility.tests.media_queries.media_queries.pageFlags": 1,
            "results.accessibility.tests.media_queries.media_queries.details.summary": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Count affected domains for each issue
    issue_counts = {
        "no_responsive": {"name": "No responsive breakpoints", "pages": set(), "domains": set()},
        "no_print": {"name": "No print stylesheets", "pages": set(), "domains": set()},
        "no_reduced_motion": {"name": "No reduced motion support", "pages": set(), "domains": set()},
        "no_dark_mode": {"name": "No dark mode support", "pages": set(), "domains": set()}
    }

    for page in pages_with_media_query_issues:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        flags = page['results']['accessibility']['tests']['media_queries']['media_queries']['pageFlags']
        
        if not flags.get('hasResponsiveBreakpoints', True):
            issue_counts["no_responsive"]["pages"].add(page['url'])
            issue_counts["no_responsive"]["domains"].add(domain)
            
        if not flags.get('hasPrintStyles', True):
            issue_counts["no_print"]["pages"].add(page['url'])
            issue_counts["no_print"]["domains"].add(domain)
            
        if not flags.get('hasReducedMotionSupport', True):
            issue_counts["no_reduced_motion"]["pages"].add(page['url'])
            issue_counts["no_reduced_motion"]["domains"].add(domain)
            
        if not flags.get('hasDarkModeSupport', True):
            issue_counts["no_dark_mode"]["pages"].add(page['url'])
            issue_counts["no_dark_mode"]["domains"].add(domain)

    # Create issues summary table
    doc.add_heading('Media Query Issues', level=3)
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    # Set column headers
    headers = table.rows[0].cells
    headers[0].text = "Issue"
    headers[1].text = "# of pages"
    headers[2].text = "# of sites affected"
    headers[3].text = "% of sites"

    # Add data for each issue
    for i, (issue_key, issue_data) in enumerate(issue_counts.items(), 1):
        row = table.rows[i].cells
        row[0].text = issue_data["name"]
        row[1].text = str(len(issue_data["pages"]))
        row[2].text = str(len(issue_data["domains"]))
        percentage = (len(issue_data["domains"]) / len(total_domains)) * 100 if total_domains else 0
        row[3].text = f"{percentage:.1f}%"

    # Format the table text
    format_table_text(table)

    # Removed explanatory text to keep the summary concise
    
    # Add space after the section
    doc.add_paragraph()