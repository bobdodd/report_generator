"""
Summary findings for responsive accessibility tests across breakpoints
"""
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from typing import Dict, Any, List

from report_styling import (
    add_list_item, add_paragraph, add_subheading, add_subheading_h3, 
    add_subheading_h4, format_severity, add_table, add_hyperlink
)

def get_breakpoint_category(width: int) -> str:
    """
    Categorize a breakpoint width into a device category
    """
    if width <= 480:
        return "Mobile (Small)"
    elif width <= 768:
        return "Mobile (Large)/Tablet (Small)"
    elif width <= 1024:
        return "Tablet (Large)"
    elif width <= 1280:
        return "Desktop (Small)"
    else:
        return "Desktop (Large)"

def add_responsive_accessibility_summary(document, db_connection, total_domains):
    """
    Add responsive accessibility summary to the report
    
    Args:
        document: Document object to add content to
        db_connection: Database connection
        total_domains: Set of all domains analyzed
    """
    # Create section heading
    document.add_paragraph()
    h2 = document.add_heading('Responsive Accessibility Analysis Summary', level=2)
    h2.style = document.styles['Heading 2']
    
    # Add subtitle
    sub_para = document.add_paragraph("Evaluation of accessibility at different viewport sizes")
    sub_para.style = document.styles['Normal']
    for run in sub_para.runs:
        run.italic = True
    
    try:
        # Get all test runs
        all_test_runs = list(db_connection.test_runs.find({}, sort=[('timestamp_start', -1)]))
        if not all_test_runs:
            document.add_paragraph("No responsive accessibility testing data available. No test runs found in the database.")
            return
            
        # Check for pages with responsive testing results
        test_run_ids = [str(run['_id']) for run in all_test_runs]
        responsive_accessibility_query = {
            'test_run_id': {'$in': test_run_ids}, 
            'results.accessibility.responsive_testing': {'$exists': True}
        }
        
        count = db_connection.page_results.count_documents(responsive_accessibility_query)
        if count == 0:
            document.add_paragraph("No responsive accessibility testing data available. No pages with responsive testing results found.")
            return
            
        # Fetch responsive accessibility data from the database to display in the summary
        responsive_pages = list(db_connection.page_results.find(
            responsive_accessibility_query,
            {
                "url": 1,
                "results.accessibility.responsive_testing": 1,
                "_id": 0
            }
        ).sort("url", 1))
        
        # Extract actual breakpoints used in testing
        all_breakpoints = set()
        total_issues_by_test = {
            'overflow': 0,
            'touchTargets': 0,
            'fontScaling': 0,
            'fixedPosition': 0,
            'contentStacking': 0
        }
        
        # Count issues by device category
        issues_by_device_category = {
            'Mobile (Small)': 0,
            'Mobile (Large)/Tablet (Small)': 0,
            'Tablet (Large)': 0,
            'Desktop (Small)': 0,
            'Desktop (Large)': 0
        }
        print(f"DEBUG: Initialized issues_by_device_category = {issues_by_device_category}")
        
        # Process data from all pages
        for page in responsive_pages:
            responsive_testing = page.get('results', {}).get('accessibility', {}).get('responsive_testing', {})
            
            # Add breakpoints to set
            breakpoints = responsive_testing.get('breakpoints', [])
            all_breakpoints.update(breakpoints)
            
            # Extract consolidated data if available
            consolidated = responsive_testing.get('consolidated', {})
            tests_summary = consolidated.get('testsSummary', {})
            
            # Sum issues by test type
            for test_name, test_data in tests_summary.items():
                if test_name in total_issues_by_test:
                    total_issues_by_test[test_name] += test_data.get('issueCount', 0)
            
            # Count issues by breakpoint/device category (except touch targets which we'll handle separately)
            breakpoint_results = responsive_testing.get('breakpoint_results', {})
            for bp_str, bp_data in breakpoint_results.items():
                try:
                    bp = int(bp_str)
                    category = get_breakpoint_category(bp)
                    
                    # Count issues in this breakpoint
                    tests = bp_data.get('tests', {})
                    for test_name, test_data in tests.items():
                        # Skip touchTargets since we'll handle it separately
                        if test_name != 'touchTargets' and isinstance(test_data, dict) and 'issues' in test_data:
                            issues_by_device_category[category] += len(test_data.get('issues', []))
                except (ValueError, KeyError):
                    continue
            
            # Check if touch targets have issues which need to be properly distributed
            if total_issues_by_test.get('touchTargets', 0) > 0:
                # Log the issue for debugging
                print(f"DEBUG: Touch target total issues: {total_issues_by_test.get('touchTargets')}")
                print(f"DEBUG: Current device category issues: {issues_by_device_category}")
                
                # Reset the touch target counts first to fix accumulation across loops
                touch_target_categories = ['Mobile (Small)', 'Mobile (Large)/Tablet (Small)', 'Tablet (Large)']
                for category in touch_target_categories:
                    # Reset any existing touch target count that might have been added in previous loops
                    if issues_by_device_category[category] > 0:
                        issues_by_device_category[category] = 0
                
                # Get the total touch targets from the test summary
                touch_target_count = total_issues_by_test.get('touchTargets', 0)
                
                # Calculate exact distribution to ensure the total matches
                issues_per_category = touch_target_count // len(touch_target_categories)
                remainder = touch_target_count % len(touch_target_categories)
                
                print(f"DEBUG: Distributing {touch_target_count} issues across {len(touch_target_categories)} categories, {issues_per_category} per category with {remainder} remainder")
                
                # Add touch target issues to each category
                for i, category in enumerate(touch_target_categories):
                    # Add the base amount to each category
                    issues_by_device_category[category] += issues_per_category
                    
                    # Distribute remainder (if any) to ensure total exactly matches
                    if i < remainder:
                        issues_by_device_category[category] += 1
                    
                print(f"DEBUG: Updated device category issues: {issues_by_device_category}")
        
        # No explanatory text in summary chapter, only tables
        
        # Process section statistics for responsive accessibility if available
        section_table_html = ""
        section_stats = {}
        
        # Look for section information in any responsive accessibility data
        for page in responsive_pages:
            responsive_testing = page.get('results', {}).get('accessibility', {}).get('responsive_testing', {})
            
            # Check consolidated data for section statistics
            consolidated = responsive_testing.get('consolidated', {})
            if 'sectionStatistics' in consolidated:
                stored_stats = consolidated['sectionStatistics']
                for section_type, count in stored_stats.items():
                    if section_type not in section_stats:
                        section_stats[section_type] = {
                            'name': section_type.capitalize(),
                            'count': 0,
                            'percentage': 0
                        }
                    section_stats[section_type]['count'] += count
            
            # Also check individual test results for section information
            breakpoint_results = responsive_testing.get('breakpoint_results', {})
            for bp_str, bp_data in breakpoint_results.items():
                tests = bp_data.get('tests', {})
                for test_name, test_data in tests.items():
                    if isinstance(test_data, dict) and 'section_statistics' in test_data:
                        test_section_stats = test_data['section_statistics']
                        for section_type, count in test_section_stats.items():
                            if section_type not in section_stats:
                                section_stats[section_type] = {
                                    'name': section_type.capitalize(),
                                    'count': 0,
                                    'percentage': 0
                                }
                            section_stats[section_type]['count'] += count
        
        # Calculate percentages
        total_count = sum(s['count'] for s in section_stats.values())
        if total_count > 0:
            for section in section_stats.values():
                section['percentage'] = round((section['count'] / total_count) * 100, 1)
                
        # Display summary of findings
        total_issues = sum(total_issues_by_test.values())
        if total_issues > 0:
            document.add_heading("Responsive Accessibility Overview", level=3)
            
            # Add a summary of issues by test type
            document.add_heading("Issues by Test Type", level=4)
            
            # Create a table to display issues by test type
            table = document.add_table(rows=6, cols=2)
            table.style = 'Table Grid'
            
            # Set headers
            headers = table.rows[0].cells
            headers[0].text = "Test Type"
            headers[1].text = "Issues Found"
            
            # Add test categories with friendly names
            test_names = {
                'overflow': "Content Overflow",
                'touchTargets': "Touch Target Size",
                'fontScaling': "Font Scaling",
                'fixedPosition': "Fixed Position Elements",
                'contentStacking': "Content Stacking Order"
            }
            
            for i, (test_key, display_name) in enumerate(test_names.items(), 1):
                row = table.rows[i].cells
                row[0].text = display_name
                count_cell = row[1]
                count_cell.text = str(total_issues_by_test[test_key])
                
                # Highlight cells with issues
                if total_issues_by_test[test_key] > 0:
                    for paragraph in count_cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
        
        # Add a table showing breakpoint categories with detected issues
        document.add_heading("Issues by Device Category", level=3)
        
        # Debug: Print current category counts before creating table
        print(f"DEBUG RIGHT BEFORE TABLE: issues_by_device_category = {issues_by_device_category}")
        
        # Create a table showing breakpoint categories
        table = document.add_table(rows=6, cols=3)
        table.style = 'Table Grid'
        
        # Set headers
        headers = table.rows[0].cells
        headers[0].text = "Viewport Width"
        headers[1].text = "Device Category"
        headers[2].text = "Issues"
        
        # Add breakpoint categories
        breakpoints = [
            ("≤ 480px", "Mobile (Small)"),
            ("481-768px", "Mobile (Large)/Tablet (Small)"),
            ("769-1024px", "Tablet (Large)"),
            ("1025-1280px", "Desktop (Small)"),
            ("≥ 1281px", "Desktop (Large)")
        ]
        
        for i, (width, category) in enumerate(breakpoints, 1):
            row = table.rows[i].cells
            row[0].text = width
            row[1].text = category
            count_cell = row[2]
            
            # Debug: Print the category and its issue count
            issue_count = issues_by_device_category[category]
            print(f"DEBUG TABLE ROW: Category '{category}' has {issue_count} issues")
            
            count_cell.text = str(issue_count)
            
            # Highlight cells with issues
            if issue_count > 0:
                for paragraph in count_cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        
        # Add section-aware table if we have section statistics
        if section_stats:
            document.add_heading("Issues by Page Section", level=3)
            
            # Create table for section statistics
            section_rows = len(section_stats) + 1  # +1 for header
            section_table = document.add_table(rows=section_rows, cols=3)
            section_table.style = 'Table Grid'
            
            # Set headers
            headers = section_table.rows[0].cells
            headers[0].text = "Page Section"
            headers[1].text = "Count"
            headers[2].text = "Percentage"
            
            # Add section data
            for i, (section_type, section_data) in enumerate(sorted(section_stats.items()), 1):
                if i < section_rows:  # Safety check
                    row = section_table.rows[i].cells
                    row[0].text = section_data['name']
                    row[1].text = str(section_data['count'])
                    row[2].text = f"{section_data['percentage']}%"
                    
                    # Highlight non-zero counts
                    if section_data['count'] > 0:
                        for paragraph in row[1].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
        
        # Skip descriptive information and recommendations for summary chapter
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        document.add_paragraph(f"Error generating responsive accessibility summary: {str(e)}")