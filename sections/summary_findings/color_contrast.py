from report_styling import format_table_text

def add_color_contrast_section(doc, db_connection, total_domains):
    """Add the Color Contrast section to the summary findings"""
    h2 = doc.add_heading('Colour Contrast', level=2)
    h2.style = doc.styles['Heading 2']

    # Define the contrast issues to be analyzed
    contrast_issues = [
        {
            'name': 'Text Contrast Issues',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.hasContrastIssues',
            'details_field': 'results.accessibility.tests.colors.colors.details.summary.contrastViolations'
        },
        {
            'name': 'Non-Text Contrast Issues',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.hasNonTextContrastIssues',
            'details_field': 'results.accessibility.tests.colors.colors.details.summary.nonTextContrastViolations'
        },
        {
            'name': 'Adjacent Contrast Issues',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.hasAdjacentContrastIssues',
            'details_field': 'results.accessibility.tests.colors.colors.details.summary.adjacentContrastViolations'
        },
        {
            'name': 'Contrast Preferences Support',
            'db_field': 'results.accessibility.tests.colors.colors.pageFlags.supportsContrastPreferences',
            'details_field': None  # This is a boolean field, not a count
        }
    ]

    # Gather the data for each issue type
    issue_data = {}

    for issue in contrast_issues:
        # For the Contrast Preferences Support, we want sites that DO support it
        # For other issues, we want sites that have problems
        if issue['name'] == 'Contrast Preferences Support':
            query = {issue['db_field']: True}
        else:
            query = {issue['db_field']: True}
        
        # Prepare projection
        projection = {"url": 1, "_id": 0}
        if issue['details_field']:
            projection[issue['details_field']] = 1
        
        # Query the database to find pages with this issue
        pages_with_issue = list(db_connection.page_results.find(query, projection))
        
        # Count affected domains and total issue instances
        affected_domains = set()
        total_instances = 0
        
        for page in pages_with_issue:
            domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
            affected_domains.add(domain)
            
            # Count instances if applicable
            if issue['details_field']:
                # Navigate the nested structure to get the count
                parts = issue['details_field'].split('.')
                value = page
                try:
                    for part in parts:
                        if part in value:
                            value = value[part]
                        else:
                            value = 0
                            break
                    
                    if isinstance(value, (int, float)):
                        total_instances += value
                except:
                    pass  # Handle any issues with nested access
        
        # Store the data
        issue_data[issue['name']] = {
            'pages': pages_with_issue,
            'domains': affected_domains,
            'instances': total_instances
        }

    # Create summary table
    from docx.oxml import parse_xml
    last_para = doc.add_paragraph()
    last_para._element.get_or_add_pPr().append(
        parse_xml(r'<w:keepNext xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    )

    summary_table = doc.add_table(rows=len(contrast_issues) + 1, cols=4)
    summary_table.style = 'Table Grid'

    # Keep table together
    for row in summary_table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcPr.append(parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>'))
            tcPr.append(parse_xml(r'<w:noWrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

    # Set column headers
    headers = summary_table.rows[0].cells
    headers[0].text = "Color Accessibility Issue"
    headers[1].text = "Pages Affected"
    headers[2].text = "Sites Affected"
    headers[3].text = "% of Total Sites"

    # Add data
    for i, issue in enumerate(contrast_issues, 1):
        row = summary_table.rows[i].cells
        data = issue_data[issue['name']]
        
        row[0].text = issue['name']
        row[1].text = str(len(data['pages']))
        row[2].text = str(len(data['domains']))
        
        percentage = (len(data['domains']) / len(total_domains)) * 100 if total_domains else 0
        row[3].text = f"{percentage:.1f}%"

    # Format the table text
    format_table_text(summary_table)

    # Add some space after the table
    doc.add_paragraph()