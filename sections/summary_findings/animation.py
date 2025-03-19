from docx.oxml import parse_xml
from report_styling import format_table_text

def add_animation_section(doc, db_connection, total_domains):
    """Add the Animation section to the summary findings"""
    doc.add_paragraph()
    h2 = doc.add_heading('Animation', level=2)
    h2.style = doc.styles['Heading 2']

    # Query for pages that have animations but lack reduced motion support
    pages_lacking_motion_support = list(db_connection.page_results.find(
        {
            "results.accessibility.tests.animations.animations.pageFlags.hasAnimations": True,
            "results.accessibility.tests.animations.animations.pageFlags.lacksReducedMotionSupport": True
        },
        {
            "url": 1,
            "results.accessibility.tests.animations.animations.details.summary": 1,
            "_id": 0
        }
    ).sort("url", 1))

    # Count affected domains
    affected_domains = set()
    for page in pages_lacking_motion_support:
        domain = page['url'].replace('http://', '').replace('https://', '').split('/')[0]
        affected_domains.add(domain)

    # Calculate percentage
    percentage = (len(affected_domains) / len(total_domains)) * 100 if total_domains else 0

    # Create summary table
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'

    # Set column headers
    headers = table.rows[0].cells
    headers[0].text = "Issue"
    headers[1].text = "# of pages"
    headers[2].text = "# of sites affected"
    headers[3].text = "% of sites"

    # Add data
    row = table.rows[1].cells
    row[0].text = "No reduced motion media query"
    row[1].text = str(len(pages_lacking_motion_support))
    row[2].text = str(len(affected_domains))
    row[3].text = f"{percentage:.1f}%"

    # Format the table text
    format_table_text(table)

    # Add some space after the table
    doc.add_paragraph()
    