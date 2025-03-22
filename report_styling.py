# report_styling.py
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn

def set_document_styles(doc):
    """Set up document styles with specified font sizes and spacing"""
    # Default paragraph text
    style = doc.styles['Normal']
    style.font.size = Pt(14)
    style.font.name = 'Arial'
    
    # Heading Styles
    title_style = doc.styles['Title']  # Level 0
    title_style.font.size = Pt(26)
    title_style.font.name = 'Arial'
    title_style.font.bold = True
    title_style.paragraph_format.space_after = Pt(14)  # One line height
    
    h1_style = doc.styles['Heading 1']
    h1_style.font.size = Pt(22)
    h1_style.font.name = 'Arial'
    h1_style.font.bold = True
    h1_style.paragraph_format.space_after = Pt(14)  # One line height
    
    h2_style = doc.styles['Heading 2']
    h2_style.font.size = Pt(18)
    h2_style.font.name = 'Arial'
    h2_style.font.bold = True
    h2_style.paragraph_format.space_after = Pt(14)  # One line height
    
    h3_style = doc.styles['Heading 3']
    h3_style.font.size = Pt(16)
    h3_style.font.name = 'Arial'
    h3_style.font.bold = True
    h3_style.paragraph_format.space_after = Pt(14)  # One line height
    
    # List styles
    list_style = doc.styles['List Bullet']
    list_style.font.size = Pt(14)
    list_style.font.name = 'Arial'
    
    # Caption style
    caption_style = doc.styles['Caption']
    caption_style.font.size = Pt(14)
    caption_style.font.name = 'Arial'
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # TOC styles are automatically created when the TOC is inserted

def add_table_of_contents(doc):
    """Add a table of contents to the document"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    # Add hyperlink and web options to the TOC field code
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u \\w'  # \h adds hyperlinks, \w preserves tab spacing

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def format_toc_styles(doc):
    """Format TOC styles to properly handle numbered headings"""
    for i in range(1, 4):  # TOC levels 1-3
        style_name = f'TOC {i}'
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.paragraph_format.tab_stops.clear_all()
            
            # Calculate indentation based on level
            base_indent = Inches(0.25)
            level_indent = Inches(0.25 * (i-1))
            
            # Set paragraph indentation
            style.paragraph_format.left_indent = level_indent
            style.paragraph_format.first_line_indent = -base_indent
            
            # Add tab stops:
            # First tab for after the number
            style.paragraph_format.tab_stops.add_tab_stop(level_indent, WD_TAB_ALIGNMENT.LEFT)
            # Final tab for page number
            style.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            
            # Set font
            style.font.name = 'Arial'
            style.font.size = Pt(14)

def format_table_text(table):
    """Format text within tables to ensure consistent font and size"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(14)
                    run.font.name = 'Arial'

def create_element(name):
    """Create an XML element with the specified name"""
    return OxmlElement(name)

def create_attribute(element, name, value):
    """Add an attribute to an XML element"""
    element.set(qn(name), value)

def add_page_number(paragraph):
    """Add a page number field to a paragraph"""
    run = paragraph.add_run()
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_paragraph(document, text=None):
    """Add a paragraph with the specified text"""
    paragraph = document.add_paragraph()
    if text:
        paragraph.add_run(text)
    return paragraph

def add_list_item(document, text):
    """Add a bulleted list item"""
    paragraph = document.add_paragraph(text, style='List Bullet')
    return paragraph

def add_subheading(document, text):
    """Add a subheading (H2)"""
    heading = document.add_heading(text, level=2)
    heading.style = document.styles['Heading 2']
    return heading

def add_subheading_h3(document, text):
    """Add a subheading (H3)"""
    heading = document.add_heading(text, level=3)
    heading.style = document.styles['Heading 3']
    return heading

def add_subheading_h4(document, text):
    """Add a subheading (H4)"""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    return paragraph

def format_severity(run, severity):
    """Format text with severity color"""
    if severity.lower() == 'high':
        run.font.color.rgb = RGBColor(192, 0, 0)  # Red
    elif severity.lower() == 'medium':
        run.font.color.rgb = RGBColor(255, 127, 0)  # Orange
    elif severity.lower() == 'low':
        run.font.color.rgb = RGBColor(150, 150, 0)  # Yellow-ish
    return run

def add_table(document, headers, rows):
    """Add a table with headers and rows"""
    if not rows:
        return None
        
    num_cols = len(headers)
    num_rows = len(rows) + 1  # +1 for header row
    
    table = document.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    
    # Add headers
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(14)
                run.font.name = 'Arial'
    
    # Add data rows
    for i, row_data in enumerate(rows):
        row = table.rows[i + 1]  # +1 to skip header
        for j, cell_data in enumerate(row_data):
            if j < len(row.cells):  # Ensure we don't exceed column count
                cell = row.cells[j]
                cell.text = str(cell_data)
                
                # Apply formatting
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(14)
                        run.font.name = 'Arial'
    
    return table

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph"""
    # Create the hyperlink
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color
    run.font.underline = True
    
    # Create the hyperlink relationship
    r_id = paragraph.part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Append the run to the hyperlink
    hyperlink.append(run._element)
    
    # Replace the run with the hyperlink
    run._element.getparent().replace(run._element, hyperlink)
    
    return hyperlink

def add_code_block(document, code_text):
    """Add a formatted code block with monospace font and gray background"""
    paragraph = document.add_paragraph()
    paragraph.style = document.styles['Normal']
    
    # Set custom spacing and formatting
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.right_indent = Inches(0.5)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(12)
    
    run = paragraph.add_run(code_text)
    run.font.name = 'Courier New'  # Monospace font
    run.font.size = Pt(12)  # Slightly smaller than normal text
    
    # We can't directly set background color in python-docx, 
    # so we'd need to use XML manipulation for that
    
    return paragraph

def add_image_if_exists(document, image_path, width=None, caption=None):
    """Add an image if the file exists, with optional caption and width"""
    import os
    
    if not os.path.exists(image_path):
        return None
        
    try:
        if width:
            document.add_picture(image_path, width=width)
        else:
            document.add_picture(image_path)
            
        if caption:
            cap_para = document.add_paragraph(caption, style='Caption')
            
        return True
    except Exception as e:
        print(f"Error adding image {image_path}: {str(e)}")
        return False