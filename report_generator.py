# report_generator.py
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from datetime import datetime
import os
import sys

# Add parent directory to path to allow relative imports
sys.path.insert(0, os.path.abspath(os.path.dirname(__file__)))

# Import styling utilities
from report_styling import (
    set_document_styles, 
    format_table_text
)

# Import section generators
from sections.sections_header import setup_document_header_footer
from sections.title_page import add_title_page
from sections.table_of_contents import add_toc_section
from sections.executive_summary import add_executive_summary

# Import only the summary findings sections we need
from sections.summary_findings.media_queries import add_media_queries_section
from sections.summary_findings.responsive_accessibility import add_responsive_accessibility_summary

# Import only the detailed findings sections we need
from sections.detailed_findings.media_queries import add_detailed_media_queries
from sections.detailed_findings.responsive_accessibility import add_responsive_accessibility_detailed

# Import appendices
from sections.appendices import add_appendices

def create_report_template(db_connection, title, author, date):
    print("Starting report creation...")

    ####################################################
    # Get list of URLs and domains used by the reporting
    # Gets used everywhere
    ####################################################

    all_test_runs = db_connection.get_all_test_runs()
    if not all_test_runs:
        print("Warning: No test runs found in the database. Creating an empty report template.")
        test_run_ids = []
        all_urls = []
    else:
        test_run_ids = [str(run['_id']) for run in all_test_runs]
        all_urls = db_connection.page_results.distinct('url', {'test_run_id': {'$in': test_run_ids}})
        
        if not all_urls:
            print("Warning: No page results found for the test runs in the database.")

    total_domains = set()
    for url in all_urls:
        domain = url.replace('http://', '').replace('https://', '').split('/')[0]
        total_domains.add(domain)

    doc = Document()
    
    # Set up document styles
    set_document_styles(doc)

    # Set up header and footer
    setup_document_header_footer(doc, title)

    # Add title page
    add_title_page(doc, db_connection, title, author, date)

    # Add table of contents
    doc.add_page_break()
    add_toc_section(doc)

    # Add executive summary
    doc.add_page_break()
    add_executive_summary(doc, db_connection, total_domains)

    #############################################
    # Summary Findings
    #############################################
    h1 = doc.add_heading('Summary findings', level=1)
    h1.style = doc.styles['Heading 1']

    # Add Media Queries Section (first, as it affects overall responsiveness)
    add_media_queries_section(doc, db_connection, total_domains)
    
    # Add Responsive Accessibility Section (right after media queries)
    add_responsive_accessibility_summary(doc, db_connection, total_domains)

    #############################################
    # Detailed Findings
    #############################################
    doc.add_page_break()
    h1 = doc.add_heading('Detailed findings', level=1)
    h1.style = doc.styles['Heading 1']

    # Add Detailed Media Queries Section (first, as it affects overall responsiveness)
    add_detailed_media_queries(doc, db_connection, total_domains)
    
    # Add Detailed Responsive Accessibility Section (right after media queries)
    add_responsive_accessibility_detailed(doc, db_connection, total_domains)

    #############################################
    # Appendices
    #############################################
    doc.add_page_break()
    add_appendices(doc, db_connection)

    return doc

def generate_report(db_connection, title, author, date, output_folder):
    try:
        doc = create_report_template(db_connection, title, author, date)
        output_filename = f'{output_folder}/accessibility_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        doc.save(output_filename)
        return output_filename
    except Exception as e:
        print(f"Error generating report: {e}")
        return None